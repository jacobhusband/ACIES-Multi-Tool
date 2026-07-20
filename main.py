from google.genai import types
from google import genai
import webview
import json
import os
import subprocess
import sys
import shutil
import csv
import io
import logging
import tempfile
import time
import stat as stat_module
import sqlite3
import http.server
import html
from html.parser import HTMLParser
from pathlib import Path
from copy import copy, deepcopy
from dotenv import load_dotenv
import datetime
import threading
import secrets
import hashlib
import requests  # Added for GitHub API calls
import zipfile   # Added for extracting bundles
import math
import random
import re
import base64
from typing import List
import importlib
import uuid
from contextlib import closing
from urllib.parse import parse_qs, urlencode, urlparse, unquote

# Work around NumPy 2.x removing scalar aliases used by openpyxl.
_NUMPY_ALIAS_MAP = {
    "short": "int16",
    "ushort": "uint16",
    "intc": "int32",
    "uintc": "uint32",
    "int_": "int64",
    "uint": "uint64",
    "longlong": "int64",
    "ulonglong": "uint64",
    "half": "float16",
    "single": "float32",
    "double": "float64",
}
_NUMPY_ATTR_FALLBACKS = {
    "int8": int,
    "int16": int,
    "int32": int,
    "int64": int,
    "uint8": int,
    "uint16": int,
    "uint32": int,
    "uint64": int,
    "intp": int,
    "uintp": int,
    "float16": float,
    "float32": float,
    "float64": float,
    "longdouble": float,
    "bool_": bool,
    "floating": float,
    "integer": int,
}


def _patch_numpy_aliases():
    try:
        import numpy as _np  # noqa: F401
    except Exception:
        return False
    changed = False
    for _name, _fallback in _NUMPY_ATTR_FALLBACKS.items():
        if not hasattr(_np, _name):
            setattr(_np, _name, _fallback)
            changed = True
    for _alias, _target in _NUMPY_ALIAS_MAP.items():
        if not hasattr(_np, _alias):
            if hasattr(_np, _target):
                setattr(_np, _alias, getattr(_np, _target))
            else:
                setattr(_np, _alias, _NUMPY_ATTR_FALLBACKS.get(_target, int))
            changed = True
    return changed


_patch_numpy_aliases()

try:
    import openpyxl
except AttributeError as _exc:
    _msg = str(_exc)
    if "numpy" in _msg:
        _patch_numpy_aliases()
        for _name in ("openpyxl.compat.numbers", "openpyxl.compat", "openpyxl"):
            sys.modules.pop(_name, None)
        importlib.invalidate_caches()
        import openpyxl
    else:
        raise
from openpyxl.worksheet.copier import WorksheetCopy
from pydantic import BaseModel, Field
from PIL import Image as PILImage, ImageOps, UnidentifiedImageError
from lighting_plan import LightingPlanValidationError, analyze_lighting_plan


_PROJECT_PAGE_PDF_VOID_TAGS = frozenset({
    "area", "base", "br", "col", "embed", "hr", "img", "input",
    "link", "meta", "param", "source", "track", "wbr",
})


class _ProjectPagePdfHtmlRenderer(HTMLParser):
    """Normalizes editor HTML into the subset supported by PyMuPDF Story."""

    def __init__(self, image_resolver):
        super().__init__(convert_charrefs=False)
        self._image_resolver = image_resolver
        self._output = []
        self._stack = []
        self._blocked_depth = 0

    @staticmethod
    def _attribute_map(attrs):
        return {
            str(name or "").lower(): "" if value is None else str(value)
            for name, value in attrs
        }

    @staticmethod
    def _attribute_text(attrs):
        parts = []
        for name, value in attrs:
            safe_name = html.escape(str(name or ""), quote=True)
            if not safe_name:
                continue
            if value is None or value == "":
                parts.append(f" {safe_name}")
            else:
                parts.append(
                    f' {safe_name}="{html.escape(str(value), quote=True)}"'
                )
        return "".join(parts)

    def _push(self, source_tag, close_html="", kind="", **extra):
        self._stack.append({
            "source": source_tag,
            "close": close_html,
            "kind": kind,
            **extra,
        })

    def _start_tag(self, tag, attrs, self_closing=False):
        source_tag = str(tag or "").lower()
        attr_map = self._attribute_map(attrs)

        if self._blocked_depth:
            if source_tag not in _PROJECT_PAGE_PDF_VOID_TAGS and not self_closing:
                self._blocked_depth += 1
            return
        if source_tag in {"script", "style"}:
            if not self_closing:
                self._blocked_depth = 1
            return

        if (
            source_tag == "div"
            and "page-workbook" in attr_map.get("class", "").split()
        ):
            file_name = str(attr_map.get("data-file-name") or "").strip()
            if not file_name:
                file_ref = str(attr_map.get("data-page-file") or "").replace("\\", "/")
                file_name = file_ref.rsplit("/", 1)[-1] if file_ref else "Workbook.xlsx"
            self._output.append(
                '<div class="pdf-workbook"><b>Excel workbook:</b> '
                f'{html.escape(file_name, quote=False)}</div>'
            )
            if not self_closing:
                self._blocked_depth = 1
            return

        if source_tag == "input":
            return
        if source_tag in {"button", "label"}:
            if not self_closing:
                self._push(source_tag)
            return

        if source_tag == "ul" and attr_map.get("data-type", "").lower() == "tasklist":
            self._output.append('<div class="pdf-task-list">')
            if not self_closing:
                self._push(source_tag, "</div>", "task-list")
            else:
                self._output.append("</div>")
            return

        if source_tag == "li" and attr_map.get("data-type", "").lower() == "taskitem":
            checked = (
                attr_map.get("data-checked", "").lower() == "true"
                or "checked" in attr_map
            )
            css_class = "pdf-task-item is-complete" if checked else "pdf-task-item"
            marker = "[x]" if checked else "[ ]"
            self._output.append(
                f'<table class="{css_class}"><tr>'
                f'<td class="pdf-task-state">{marker}</td>'
                '<td class="pdf-task-content">'
            )
            if not self_closing:
                self._push(
                    source_tag,
                    "</td></tr></table>",
                    "task-item",
                    content_wrapper_skipped=False,
                )
            else:
                self._output.append("</td></tr></table>")
            return

        if (
            source_tag == "div"
            and self._stack
            and self._stack[-1].get("kind") == "task-item"
            and not self._stack[-1].get("content_wrapper_skipped")
        ):
            self._stack[-1]["content_wrapper_skipped"] = True
            if not self_closing:
                self._push(source_tag)
            return

        if source_tag == "details":
            self._output.append('<div class="page-toggle">')
            if not self_closing:
                self._push(source_tag, "</div>", "toggle")
            else:
                self._output.append("</div>")
            return

        if source_tag == "summary":
            self._output.append('<div class="page-toggle-summary">')
            if not self_closing:
                self._push(source_tag, "</div>", "toggle-summary")
            else:
                self._output.append("</div>")
            return

        if source_tag == "img":
            image = self._image_resolver(attr_map)
            if not image:
                self._output.append(
                    '<div class="pdf-image-missing">[Image unavailable]</div>'
                )
                return
            image_source, width_percent = image
            image_attrs = [
                (name, value)
                for name, value in attrs
                if str(name or "").lower() not in {
                    "src", "data-asset", "data-width-percent",
                    "contenteditable", "spellcheck", "style",
                }
            ]
            image_attrs.extend([
                ("src", image_source),
                (
                    "style",
                    f"width:{width_percent}%; max-width:100%; height:auto;",
                ),
            ])
            self._output.append(
                f'<div class="pdf-image-wrap"><img'
                f'{self._attribute_text(image_attrs)}></div>'
            )
            return

        output_tag = source_tag
        if (
            source_tag == "mark"
            and "page-find-match" in attr_map.get("class", "").split()
        ):
            output_tag = "span"

        filtered_attrs = [
            (name, value)
            for name, value in attrs
            if str(name or "").lower() not in {"contenteditable", "spellcheck"}
        ]
        if source_tag == "a" and attr_map.get("href") == "#":
            filtered_attrs = [
                (name, value)
                for name, value in filtered_attrs
                if str(name or "").lower() != "href"
            ]

        self._output.append(
            f"<{output_tag}{self._attribute_text(filtered_attrs)}>"
        )
        close_html = f"</{output_tag}>"
        if output_tag not in _PROJECT_PAGE_PDF_VOID_TAGS:
            if self_closing:
                self._output.append(close_html)
            else:
                self._push(source_tag, close_html)

        if (
            source_tag == "div"
            and "page-callout" in attr_map.get("class", "").split()
        ):
            self._output.append(
                '<div class="pdf-callout-label">NOTE</div>'
            )

    def handle_starttag(self, tag, attrs):
        self._start_tag(tag, attrs)

    def handle_startendtag(self, tag, attrs):
        self._start_tag(tag, attrs, self_closing=True)

    def handle_endtag(self, tag):
        source_tag = str(tag or "").lower()
        if self._blocked_depth:
            self._blocked_depth = max(0, self._blocked_depth - 1)
            return

        match_index = -1
        for index in range(len(self._stack) - 1, -1, -1):
            if self._stack[index].get("source") == source_tag:
                match_index = index
                break
        if match_index < 0:
            return
        while len(self._stack) > match_index:
            entry = self._stack.pop()
            self._output.append(entry.get("close", ""))

    def handle_data(self, data):
        if not self._blocked_depth:
            self._output.append(html.escape(str(data or ""), quote=False))

    def handle_entityref(self, name):
        if not self._blocked_depth:
            self._output.append(f"&{name};")

    def handle_charref(self, name):
        if not self._blocked_depth:
            self._output.append(f"&#{name};")

    def rendered_html(self):
        while self._stack:
            self._output.append(self._stack.pop().get("close", ""))
        return "".join(self._output)


HEIF_IMAGE_EXTENSIONS = {".heic", ".heics", ".heif", ".heifs", ".hif"}
HEIF_SUPPORT_ENABLED = False
HEIF_SUPPORT_ERROR = ""


def _register_heif_support():
    global HEIF_SUPPORT_ENABLED, HEIF_SUPPORT_ERROR
    try:
        pillow_heif = importlib.import_module("pillow_heif")
        register_heif_opener = getattr(pillow_heif, "register_heif_opener", None)
        if not callable(register_heif_opener):
            raise AttributeError("pillow_heif.register_heif_opener is unavailable.")
        register_heif_opener()
        HEIF_SUPPORT_ENABLED = True
        HEIF_SUPPORT_ERROR = ""
        return True
    except Exception as exc:
        HEIF_SUPPORT_ENABLED = False
        HEIF_SUPPORT_ERROR = str(exc)
        return False


def _is_heif_image_path(path):
    return os.path.splitext(str(path or "").strip())[1].lower() in HEIF_IMAGE_EXTENSIONS


def _open_local_pil_image(path):
    try:
        return PILImage.open(path)
    except UnidentifiedImageError as exc:
        if not _is_heif_image_path(path):
            raise
        if not HEIF_SUPPORT_ENABLED:
            raise ValueError(
                "HEIC/HEIF images are not supported in this build. Install pillow-heif and rebuild the app."
            ) from exc
        raise ValueError(
            "Could not read HEIC/HEIF image data. Re-save the image as JPG or PNG and try again."
        ) from exc


PANEL_SCHEDULE_MAX_IMAGE_EDGE = 2048


def _open_panel_schedule_image(path):
    img = _open_local_pil_image(path)
    try:
        img = ImageOps.exif_transpose(img)
    except Exception:
        pass
    resampling = getattr(
        getattr(PILImage, 'Resampling', PILImage),
        'LANCZOS',
        getattr(PILImage, 'LANCZOS', getattr(PILImage, 'BICUBIC', 3))
    )
    img.thumbnail(
        (PANEL_SCHEDULE_MAX_IMAGE_EDGE, PANEL_SCHEDULE_MAX_IMAGE_EDGE),
        resampling,
    )
    return img


_register_heif_support()

# Helper functions for date parsing and status management
STATUS_CANON = ["Waiting", "Working",
                "Pending Review", "Complete", "Delivered"]
STATUS_PRIORITY = ['Delivered', 'Complete',
                   'Pending Review', 'Working', 'Waiting']
LABEL_TO_KEY = {
    "Waiting": "waiting",
    "Working": "working",
    "Pending Review": "pendingReview",
    "Complete": "complete",
    "Delivered": "delivered"
}
KEY_TO_LABEL = {v: k for k, v in LABEL_TO_KEY.items()}

APP_UPDATE_REPO = "jacobhusband/ACIES-Multi-Tool"
APP_INSTALLER_NAME = "acies-scheduler-setup.exe"
GITHUB_API_BASE = "https://api.github.com"
KNOWN_PLUGIN_BUNDLES = [
    "ElectricalCommands.AutoLispCommands.bundle",
    "ElectricalCommands.CleanCADCommands.bundle",
    "ElectricalCommands.GeneralCommands.bundle",
    "ElectricalCommands.GetAttributesCommands.bundle",
    "ElectricalCommands.PlotCommands.bundle",
    "ElectricalCommands.T24Commands.bundle",
    "ElectricalCommands.LFSCommands.bundle",
    "ElectricalCommands.TextCommands.bundle",
]
HIDDEN_PLUGIN_BUNDLES = {
    "ElectricalCommands.GetAttributesCommands.bundle",
}


def load_app_version():
    """Read version from VERSION file; fallback to '0.0.0'."""
    version_file = Path(__file__).resolve().parent / "VERSION"
    try:
        return version_file.read_text(encoding="utf-8").strip() or "0.0.0"
    except Exception as e:
        logging.warning(f"Could not read VERSION file: {e}")
        return "0.0.0"


APP_VERSION = load_app_version()
BUNDLE_RELEASE_TAG = f"v{APP_VERSION}"
GOOGLE_OAUTH_CLIENT_ID_ENV = "GOOGLE_OAUTH_CLIENT_ID"
GOOGLE_OAUTH_CLIENT_SECRET_ENV = "GOOGLE_CLIENT_SECRET"
GOOGLE_OAUTH_AUTH_ENDPOINT = "https://accounts.google.com/o/oauth2/v2/auth"
GOOGLE_OAUTH_TOKEN_ENDPOINT = "https://oauth2.googleapis.com/token"
GOOGLE_OAUTH_REVOKE_ENDPOINT = "https://oauth2.googleapis.com/revoke"
GOOGLE_OAUTH_USERINFO_ENDPOINT = "https://openidconnect.googleapis.com/v1/userinfo"
GOOGLE_OAUTH_SCOPES = ["openid", "email", "profile"]
GOOGLE_OAUTH_CALLBACK_PATH = "/"
GOOGLE_OAUTH_TIMEOUT_SECONDS = 180
OUTLOOK_SCAN_SOURCE_DESKTOP = "desktop-outlook"
OUTLOOK_MAPI_INBOX_FOLDER_ID = 6
OUTLOOK_MAPI_PR_INTERNET_MESSAGE_ID = "http://schemas.microsoft.com/mapi/proptag/0x1035001F"
OUTLOOK_MAPI_PR_SMTP_ADDRESS = "http://schemas.microsoft.com/mapi/proptag/0x39FE001F"
OUTLOOK_SCAN_DEDUPE_SKIP_REASON_THREAD = (
    "Removed from the AI batch because this reply only repeated thread history already seen in newer emails."
)
OUTLOOK_SCAN_DEDUPE_SKIP_REASON_DUPLICATE = (
    "Removed from the AI batch because its content duplicated another email in the selected timeframe."
)
OUTLOOK_SCAN_FETCH_LIMIT = 200
OUTLOOK_SCAN_AI_LIMIT = 40
OUTLOOK_SCAN_PROMPT_MAX_CHARS = 120000
OUTLOOK_SCAN_PROMPT_EMAIL_BUDGET_CHARS = 90000
OUTLOOK_SCAN_PROMPT_DELIVERABLE_BUDGET_CHARS = 25000
OUTLOOK_SCAN_EMAIL_BODY_PROMPT_CHARS = 1200
OUTLOOK_SCAN_RETRY_AI_LIMIT = 20
OUTLOOK_SCAN_RETRY_PROMPT_EMAIL_BUDGET_CHARS = 45000
OUTLOOK_SCAN_RETRY_EMAIL_BODY_PROMPT_CHARS = 600
OUTLOOK_SCAN_PROMPT_SKIP_REASON = "Excluded from the AI batch prompt to keep the scan bounded."
OUTLOOK_SCAN_RETRY_SKIP_REASON = (
    "Excluded from the reduced AI retry batch after the first request timed out."
)
EMAIL_INTAKE_PROJECT_CONTEXT_MAX_PROJECTS = 200
EMAIL_INTAKE_PROJECT_CONTEXT_BUDGET_CHARS = 25000
FIREBASE_API_KEY_ENV = "FIREBASE_API_KEY"
FIREBASE_AUTH_DOMAIN_ENV = "FIREBASE_AUTH_DOMAIN"
FIREBASE_PROJECT_ID_ENV = "FIREBASE_PROJECT_ID"
FIREBASE_APP_ID_ENV = "FIREBASE_APP_ID"
FIREBASE_STORAGE_BUCKET_ENV = "FIREBASE_STORAGE_BUCKET"
FIREBASE_MESSAGING_SENDER_ID_ENV = "FIREBASE_MESSAGING_SENDER_ID"


def build_default_cloud_sync_settings():
    return {
        'enabled': False,
        'firebaseUid': '',
        'lastSyncedAt': '',
        'migrationCompleted': False,
    }


def build_default_workflow_cad_defaults():
    return {
        'manageLayersDwgSource': 'electricalTopLevel',
        'cleanXrefsDwgSource': 'electricalXrefsToNewestArch',
        'cleanXrefsSearchZipArchives': True,
    }


def _normalize_workflow_cad_defaults(value):
    defaults = build_default_workflow_cad_defaults()
    source = value if isinstance(value, dict) else {}
    manage_source = str(
        source.get('manageLayersDwgSource') or defaults['manageLayersDwgSource']
    ).strip()
    clean_source = str(
        source.get('cleanXrefsDwgSource') or defaults['cleanXrefsDwgSource']
    ).strip()
    return {
        'manageLayersDwgSource': (
            manage_source
            if manage_source in {'electricalTopLevel', 'manual'}
            else defaults['manageLayersDwgSource']
        ),
        'cleanXrefsDwgSource': (
            clean_source
            if clean_source in {'electricalXrefsToNewestArch', 'manual'}
            else defaults['cleanXrefsDwgSource']
        ),
        'cleanXrefsSearchZipArchives': (
            source.get('cleanXrefsSearchZipArchives', True) is not False
        ),
    }


# Schema describing each tool that can appear in a user-defined workflow.
# 'invoke' is `(api_instance, launch_context, activity_id, params)` -> dict result.
# 'params' lists the per-step parameter spec the workflow-builder UI introspects via
# Api.get_workflow_tools() so it can render the correct editor controls.
WORKFLOW_TOOL_REGISTRY = {
    'backupDrawings': {
        'displayName': 'Backup DWGs',
        'description': 'Archive configured discipline folders and Xrefs into a timestamped Archive folder.',
        'invoke': (lambda api, ctx, aid, params:
                   api.backup_project_drawings(None, ctx)),
        'params': [],
        'requiredInputs': [
            {'key': 'projectFolder', 'type': 'folder', 'label': 'Project folder',
             'help': 'Discipline folders and Xrefs will be archived from this folder.'},
        ],
    },
    'manageLayers': {
        'displayName': 'Freeze/Thaw Layers',
        'description': 'Freeze or thaw layers across selected DWGs. Patterns support wildcards (* and ?).',
        'invoke': (lambda api, ctx, aid, params:
                   api.run_manage_layers_script(ctx, aid, params)),
        'params': [
            {'key': 'freezePatterns', 'label': 'Freeze patterns',
             'type': 'patternList',
             'help': 'Layer-name wildcards to freeze (e.g. *cloud*, *rev*). Leave empty to skip.'},
            {'key': 'thawPatterns', 'label': 'Thaw patterns',
             'type': 'patternList',
             'help': 'Layer-name wildcards to thaw. Leave empty to skip.'},
        ],
        'requiredInputs': [
            {'key': 'dwgFiles', 'type': 'dwgFiles', 'label': 'DWG files',
             'help': 'These DWGs will be scanned and matching layers frozen or thawed.'},
        ],
    },
    'cleanXrefs': {
        'displayName': 'Prepare CAD for XREF',
        'description': 'Strip XREF paths, set colors by layer, purge, and audit.',
        'invoke': (lambda api, ctx, aid, params:
                   api.run_clean_xrefs_script(ctx, aid, params)),
        'params': [
            {'key': 'stripXrefs', 'label': 'Strip XREF paths', 'type': 'bool', 'default': True},
            {'key': 'setByLayer', 'label': 'Set color by layer', 'type': 'bool', 'default': True},
            {'key': 'purge', 'label': 'Purge', 'type': 'bool', 'default': True},
            {'key': 'audit', 'label': 'Audit', 'type': 'bool', 'default': True},
            {'key': 'hatchColor', 'label': 'Adjust hatch color', 'type': 'bool', 'default': True},
        ],
        'requiredInputs': [
            {'key': 'dwgFiles', 'type': 'dwgFiles', 'label': 'DWG files',
             'help': 'These DWGs will have XREF paths stripped, colors set by layer, purged, and audited.'},
        ],
    },
    'publishDwgs': {
        'displayName': 'Publish DWGs',
        'description': 'Plot DWG layouts and combine outputs.',
        'invoke': (lambda api, ctx, aid, params:
                   api.run_publish_script(ctx, aid, params)),
        'params': [
            {'key': 'autoDetectPaperSize', 'label': 'Auto-detect paper size',
             'type': 'bool', 'default': True},
            {'key': 'shrinkPercent', 'label': 'Shrink %',
             'type': 'int', 'default': 100, 'min': 25, 'max': 200},
            {'key': 'stripPdfLayers', 'label': 'Remove PDF layers',
             'type': 'bool', 'default': True},
        ],
        'requiredInputs': [
            {'key': 'dwgFiles', 'type': 'dwgFiles', 'label': 'DWG files',
             'help': 'These DWGs will be plotted and combined into output PDFs.'},
        ],
    },
}


def get_workflow_tool_descriptors():
    """Return a JSON-safe list describing available workflow step tools (no callables)."""
    return [
        {
            'toolId': tool_id,
            'displayName': entry['displayName'],
            'description': entry.get('description', ''),
            'params': list(entry.get('params') or []),
            'requiredInputs': list(entry.get('requiredInputs') or []),
        }
        for tool_id, entry in WORKFLOW_TOOL_REGISTRY.items()
    ]


def _normalize_layer_pattern_list(value):
    """Normalize layer-name pattern input to a clean list of non-empty strings.

    Accepts list/tuple of strings, a single string (possibly semicolon- or comma-delimited),
    or None. Strips whitespace, drops empties, preserves order, de-duplicates case-insensitively.
    """
    if value is None:
        return []
    if isinstance(value, (list, tuple)):
        raw = value
    else:
        text = str(value)
        for sep in (';', ','):
            if sep in text:
                raw = text.split(sep)
                break
        else:
            raw = [text]
    seen = set()
    result = []
    for entry in raw:
        cleaned = str(entry or '').strip()
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(cleaned)
    return result


DISCIPLINE_OPTIONS = ('Electrical', 'Mechanical', 'Plumbing')


def _normalize_settings_discipline_list(value):
    if isinstance(value, (list, tuple)):
        raw = value
    elif isinstance(value, str):
        raw = re.split(r'[,/;]+', value)
    else:
        raw = []

    lookup = {item.lower(): item for item in DISCIPLINE_OPTIONS}
    seen = set()
    result = []
    for entry in raw:
        normalized = lookup.get(str(entry or '').strip().lower())
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        result.append(normalized)
    return result or ['Electrical']


def _normalize_active_discipline(value, configured_disciplines):
    lookup = {item.lower(): item for item in DISCIPLINE_OPTIONS}
    normalized = lookup.get(str(value or '').strip().lower())
    if normalized:
        return normalized
    configured = _normalize_settings_discipline_list(configured_disciplines)
    return configured[0] if configured else 'Electrical'


def build_default_user_settings():
    return {
        'userName': '',
        'discipline': ['Electrical'],
        'activeDiscipline': 'Electrical',
        'apiKey': '',
        'autocadPath': '',
        'showSetupHelp': True,
        'cleanDwgOptions': {
            'stripXrefs': True,
            'setByLayer': True,
            'purge': True,
            'audit': True,
            'hatchColor': True
        },
        'publishDwgOptions': {
            'autoDetectPaperSize': True,
            'shrinkPercent': 100,
            'stripPdfLayers': True
        },
        'manageLayersOptions': {
            'scanAllLayers': True,
            'freezePatterns': [],
            'thawPatterns': []
        },
        'workflowCadDefaults': build_default_workflow_cad_defaults(),
        'workroomAutoSelectCadFiles': True,
        'enableUnderConstructionTools': False,
        'googleAuth': None,
        'cloudSync': build_default_cloud_sync_settings(),
        'workflows': [],
    }


def _get_legacy_windows_app_data_dir():
    user_profile = os.getenv('USERPROFILE') or os.path.expanduser('~')
    appdata_base = os.getenv('APPDATA') or os.path.join(
        user_profile, 'AppData', 'Roaming')
    return os.path.join(appdata_base, 'ProjectManagementApp')


def _load_legacy_user_settings():
    if sys.platform != "win32":
        return None
    legacy_path = os.path.join(_get_legacy_windows_app_data_dir(), "settings.json")
    try:
        with open(legacy_path, 'r', encoding='utf-8') as f:
            payload = json.load(f)
        return payload if isinstance(payload, dict) else None
    except (FileNotFoundError, json.JSONDecodeError):
        return None


def _merge_missing_user_settings_from_legacy(current_settings, legacy_settings):
    current = current_settings if isinstance(current_settings, dict) else {}
    legacy = legacy_settings if isinstance(legacy_settings, dict) else {}
    if not legacy:
        return current, False

    merged = deepcopy(current)
    changed = False

    for key in (
        "userName",
        "apiKey",
        "autocadPath",
        "theme",
        "defaultPmInitials",
        "autocadVersion",
    ):
        current_value = str(merged.get(key) or "").strip()
        legacy_value = str(legacy.get(key) or "").strip()
        if not current_value and legacy_value:
            merged[key] = legacy_value
            changed = True

    current_discipline = merged.get("discipline")
    legacy_discipline = legacy.get("discipline")
    if (
        (not isinstance(current_discipline, list) or not current_discipline)
        and isinstance(legacy_discipline, list)
        and legacy_discipline
    ):
        merged["discipline"] = deepcopy(legacy_discipline)
        changed = True

    if (
        merged.get("showSetupHelp", True) is not False
        and legacy.get("showSetupHelp") is False
    ):
        merged["showSetupHelp"] = False
        changed = True

    if merged.get("autoPrimary") is not True and legacy.get("autoPrimary") is True:
        merged["autoPrimary"] = True
        changed = True

    current_templates = merged.get("lightingTemplates")
    legacy_templates = legacy.get("lightingTemplates")
    if (
        (not isinstance(current_templates, list) or not current_templates)
        and isinstance(legacy_templates, list)
        and legacy_templates
    ):
        merged["lightingTemplates"] = deepcopy(legacy_templates)
        changed = True

    for auth_key in ("googleAuth",):
        if not isinstance(merged.get(auth_key), dict) and isinstance(legacy.get(auth_key), dict):
            merged[auth_key] = deepcopy(legacy.get(auth_key))
            changed = True

    return merged, changed


def _sanitize_user_settings_payload(settings):
    normalized = deepcopy(settings) if isinstance(settings, dict) else {}
    changed = False

    if "microsoftAuth" in normalized:
        normalized.pop("microsoftAuth", None)
        changed = True

    if (
        ("freezeLayerOptions" in normalized or "thawLayerOptions" in normalized)
        and "manageLayersOptions" not in normalized
    ):
        legacy_freeze = normalized.get("freezeLayerOptions") or {}
        legacy_thaw = normalized.get("thawLayerOptions") or {}
        if isinstance(legacy_freeze, dict) and "scanAllLayers" in legacy_freeze:
            scan_all = bool(legacy_freeze.get("scanAllLayers", True))
        elif isinstance(legacy_thaw, dict) and "scanAllLayers" in legacy_thaw:
            scan_all = bool(legacy_thaw.get("scanAllLayers", True))
        else:
            scan_all = True
        normalized["manageLayersOptions"] = {"scanAllLayers": scan_all}
        changed = True
    if "freezeLayerOptions" in normalized:
        normalized.pop("freezeLayerOptions", None)
        changed = True
    if "thawLayerOptions" in normalized:
        normalized.pop("thawLayerOptions", None)
        changed = True

    defaults = build_default_user_settings()
    for key, value in defaults.items():
        if key not in normalized:
            normalized[key] = deepcopy(value)
            changed = True

    default_publish_options = defaults.get("publishDwgOptions") or {}
    source_publish_options = normalized.get("publishDwgOptions")
    if not isinstance(source_publish_options, dict):
        normalized["publishDwgOptions"] = deepcopy(default_publish_options)
        changed = True
    else:
        merged_publish_options = deepcopy(default_publish_options)
        merged_publish_options.update(source_publish_options)
        if source_publish_options != merged_publish_options:
            normalized["publishDwgOptions"] = merged_publish_options
            changed = True

    workflow_cad_defaults = _normalize_workflow_cad_defaults(
        normalized.get("workflowCadDefaults")
    )
    if normalized.get("workflowCadDefaults") != workflow_cad_defaults:
        normalized["workflowCadDefaults"] = workflow_cad_defaults
        changed = True

    normalized_disciplines = _normalize_settings_discipline_list(
        normalized.get("discipline")
    )
    if normalized.get("discipline") != normalized_disciplines:
        normalized["discipline"] = normalized_disciplines
        changed = True

    normalized_active_discipline = _normalize_active_discipline(
        normalized.get("activeDiscipline"),
        normalized_disciplines,
    )
    if normalized.get("activeDiscipline") != normalized_active_discipline:
        normalized["activeDiscipline"] = normalized_active_discipline
        changed = True

    return normalized, changed


def utc_now_iso():
    return datetime.datetime.now(datetime.timezone.utc).replace(
        microsecond=0
    ).isoformat().replace("+00:00", "Z")


def parse_utc_iso(value):
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        return datetime.datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except Exception:
        return None


def _normalize_version(raw):
    """Turn a version string like 'v1.2.3' into '1.2.3'."""
    return str(raw or '').strip().lstrip('vV')


def _version_tuple(raw):
    parts = []
    for chunk in _normalize_version(raw).split('.'):
        try:
            parts.append(int(chunk))
        except ValueError:
            break
    return tuple(parts or [0])


def _is_remote_newer(remote, current):
    return _version_tuple(remote) > _version_tuple(current)


def _bundle_name_from_asset_name(asset_name, release_tag):
    bundle_name = str(asset_name or "")
    if release_tag:
        bundle_name = bundle_name.replace(f"-{release_tag}.zip", ".bundle")
    if bundle_name.endswith(".zip"):
        bundle_name = bundle_name[:-4] + ".bundle"
    return bundle_name


def parse_due_str(s):
    """Parse due date string similar to JS parseDueStr."""
    if not s:
        return None
    s = s.strip()
    try:
        # Try ISO format first
        return datetime.datetime.fromisoformat(s.replace('Z', '+00:00'))
    except:
        pass
    s = s.replace('.', '/').replace(' ', '')
    parts = s.split('/')
    if len(parts) == 3:
        try:
            mm, dd, yy = parts
            if len(yy) == 2:
                yy = '20' + yy
            iso = f"{yy}-{mm.zfill(2)}-{dd.zfill(2)}T12:00:00"
            return datetime.datetime.fromisoformat(iso)
        except:
            pass
    try:
        return datetime.datetime.strptime(s, "%m/%d/%Y")
    except:
        pass
    return None


def sync_status_arrays(task):
    """Sync status arrays similar to JS syncStatusArrays."""
    if not isinstance(task.get('statuses'), list):
        task['statuses'] = []
    from_tags = task.get('statusTags', [])
    for key in from_tags:
        label = KEY_TO_LABEL.get(key)
        if label and label not in task['statuses']:
            task['statuses'].append(label)
    task['statuses'] = list({s for s in task['statuses'] if s in STATUS_CANON})
    task['statusTags'] = [LABEL_TO_KEY[s]
                          for s in task['statuses'] if LABEL_TO_KEY.get(s)]
    task['status'] = next(
        (label for label in STATUS_PRIORITY if label in task['statuses']), '')


def get_default_documents_dir():
    user_profile = os.getenv('USERPROFILE') or os.path.expanduser('~')
    onedrive_docs = os.path.join(
        user_profile, 'OneDrive - ACIES Engineering', 'Documents')
    if os.path.isdir(onedrive_docs):
        return onedrive_docs
    return os.path.join(user_profile, 'Documents')


_LEGACY_MIGRATION_DONE = False


def _get_windows_documents_dir():
    user_profile = os.getenv('USERPROFILE')
    if user_profile:
        return os.path.join(user_profile, 'Documents')
    return os.path.expanduser('~')


def _migrate_legacy_app_data():
    global _LEGACY_MIGRATION_DONE
    if _LEGACY_MIGRATION_DONE or sys.platform != "win32":
        return
    _LEGACY_MIGRATION_DONE = True

    user_profile = os.getenv('USERPROFILE') or os.path.expanduser('~')
    appdata_base = os.getenv('APPDATA') or os.path.join(
        user_profile, 'AppData', 'Roaming')
    old_dir = os.path.join(appdata_base, 'ProjectManagementApp')
    if not os.path.isdir(old_dir):
        return

    new_dir = os.path.join(_get_windows_documents_dir(), 'ProjectManagementApp')
    os.makedirs(new_dir, exist_ok=True)

    try:
        items = os.listdir(old_dir)
    except OSError as exc:
        logging.warning(f"Could not list legacy data directory: {exc}")
        return

    for name in items:
        src = os.path.join(old_dir, name)
        dst = os.path.join(new_dir, name)
        if os.path.exists(dst):
            logging.info(
                f"Skipping legacy data item because target exists: {dst}")
            continue
        try:
            if os.path.isdir(src):
                shutil.copytree(src, dst)
            else:
                shutil.copy2(src, dst)
            logging.info(f"Migrated legacy data item: {src} -> {dst}")
        except Exception as exc:
            logging.warning(f"Failed to migrate legacy data item {src}: {exc}")


def build_timesheet_filename(user_name, week_key):
    tokens = [t for t in str(user_name or '').strip().split() if t]
    if len(tokens) >= 2:
        initials = (tokens[0][0] + tokens[-1][0]).upper()
    elif tokens:
        initials = (tokens[0][:2]).upper()
        if len(initials) == 1:
            initials = initials + initials
    else:
        initials = "TS"

    week_start = None
    if week_key:
        try:
            week_start = datetime.date.fromisoformat(week_key)
        except ValueError:
            week_start = None
    if not week_start:
        week_start = datetime.date.today()

    monday = week_start + datetime.timedelta(
        days=(7 + 0 - week_start.weekday()) % 7
    )
    friday = monday + datetime.timedelta(days=4)

    return f"{initials}-TS-{monday.strftime('%m%d')}-{friday.strftime('%m%d%Y')}.xlsx"



TIMESHEET_DAY_ORDER = ["mon", "tue", "wed", "thu", "fri", "sat", "sun"]


def _coerce_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _get_timesheet_project_key(project_id="", project_name=""):
    normalized_id = str(project_id or "").strip()
    if normalized_id:
        return f"id:{normalized_id.casefold()}"
    normalized_name = str(project_name or "").strip()
    if normalized_name:
        return f"name:{normalized_name.casefold()}"
    return ""


def _get_timesheet_day_from_expense_date(value):
    if not value:
        return None
    try:
        return TIMESHEET_DAY_ORDER[datetime.date.fromisoformat(str(value)).weekday()]
    except ValueError:
        return None


def _build_export_mileage_by_row(entries, expense_projects):
    if not entries:
        return []

    mileage_by_row = [0.0] * len(entries)
    if not expense_projects:
        return mileage_by_row

    day_rows = {day: [] for day in TIMESHEET_DAY_ORDER}
    project_rows = {}
    project_day_rows = {}
    project_day_totals = {}

    for row_index, entry in enumerate(entries):
        project_key = _get_timesheet_project_key(
            entry.get("projectId", ""),
            entry.get("projectName", ""),
        )
        if project_key:
            project_rows.setdefault(project_key, []).append(row_index)

        hours = entry.get("hours") or {}
        for day in TIMESHEET_DAY_ORDER:
            if _coerce_float(hours.get(day, 0), 0.0) <= 0:
                continue
            day_rows[day].append(row_index)
            if project_key:
                project_day_rows.setdefault((project_key, day), []).append(row_index)

    for project in expense_projects:
        project_key = _get_timesheet_project_key(
            project.get("projectId", ""),
            project.get("projectName", ""),
        )
        for expense_entry in project.get("entries", []):
            day = _get_timesheet_day_from_expense_date(expense_entry.get("date"))
            mileage = _coerce_float(expense_entry.get("mileage", 0), 0.0)
            if not day or mileage <= 0:
                continue
            bucket = (project_key, day)
            project_day_totals[bucket] = project_day_totals.get(bucket, 0.0) + mileage

    for (project_key, day), mileage in project_day_totals.items():
        target_rows = []
        if project_key:
            target_rows = project_day_rows.get((project_key, day), [])
        if not target_rows:
            target_rows = day_rows.get(day, [])
        if not target_rows and project_key:
            target_rows = project_rows.get(project_key, [])
        if not target_rows:
            target_rows = [0]
        mileage_by_row[target_rows[0]] += mileage

    return mileage_by_row


def _format_long_date(date_value=None):
    """Format date like 'January 1, 2026'."""
    if date_value is None:
        date_value = datetime.date.today()
    if isinstance(date_value, datetime.datetime):
        date_value = date_value.date()
    return f"{date_value.strftime('%B')} {date_value.day}, {date_value.strftime('%Y')}"


def _replace_text_in_paragraph(paragraph, replacements):
    if not paragraph.text:
        return
    original = paragraph.text
    updated = original
    for find_text, replace_text in replacements.items():
        if find_text:
            updated = updated.replace(find_text, replace_text)
    if updated == original:
        return

    for run in paragraph.runs:
        for find_text, replace_text in replacements.items():
            if find_text and find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)

    if paragraph.text != updated:
        paragraph.text = updated


def _delete_docx_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def _remove_to_section_docx(doc):
    to_remove = []
    removing = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        lower = text.lower()
        if not removing and lower.startswith("to"):
            removing = True
            to_remove.append(paragraph)
            continue
        if removing:
            if "project" in lower or "*project name*" in lower:
                removing = False
                continue
            to_remove.append(paragraph)
    for paragraph in to_remove:
        _delete_docx_paragraph(paragraph)


def _replace_text_in_docx(doc, replacements):
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, replacements)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)


def _update_docx_file(file_path, replacements, remove_to_section=False):
    from docx import Document

    doc = Document(file_path)
    if remove_to_section:
        _remove_to_section_docx(doc)
    if replacements:
        _replace_text_in_docx(doc, replacements)
    doc.save(file_path)


def _remove_to_section_word(doc):
    removing = False
    idx = 1
    while idx <= doc.Paragraphs.Count:
        paragraph = doc.Paragraphs(idx)
        text = paragraph.Range.Text.replace("\r", "").strip()
        lower = text.lower()
        if not removing and lower.startswith("to"):
            removing = True
        if removing:
            if "project" in lower or "*project name*" in lower:
                removing = False
                idx += 1
                continue
            paragraph.Range.Delete()
            continue
        idx += 1


def _update_word_file_via_com(file_path, replacements, remove_to_section=False):
    import win32com.client

    # Launch an isolated Word instance so closing the template document does not
    # shut down any Word windows the user already has open.
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = None
    try:
        doc = word.Documents.Open(file_path)
        for find_text, replace_text in replacements.items():
            if not find_text:
                continue
            rng = doc.Range()
            rng.Find.ClearFormatting()
            rng.Find.Replacement.ClearFormatting()
            rng.Find.Execute(
                FindText=find_text,
                ReplaceWith=replace_text,
                Replace=2,  # wdReplaceAll
            )
        if remove_to_section:
            _remove_to_section_word(doc)
        doc.Save()
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def _apply_template_context(file_path, context=None, options=None):
    context = context or {}
    options = options or {}
    remove_to_section = bool(options.get("removeToSection"))

    replacements = {
        "*current date*": _format_long_date(),
    }
    deliverable_name = str(context.get("deliverableName") or "").strip()
    project_name = str(context.get("projectName") or "").strip()
    if deliverable_name:
        replacements["*deliverable*"] = deliverable_name
    if project_name:
        replacements["*project name*"] = project_name

    if not replacements and not remove_to_section:
        return

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        _update_docx_file(file_path, replacements, remove_to_section=remove_to_section)
    elif ext == ".doc":
        _update_word_file_via_com(file_path, replacements, remove_to_section=remove_to_section)


# --- Google GenAI (new client) ---
# Uses GOOGLE_API_KEY from environment/.env

# --- Helper function to get the application data path ---


def get_app_data_dir():
    """
    Determines the correct, cross-platform directory for storing user data.

    On Windows, we store data in the user's Documents folder to bypass
    Windows Store Python's file system virtualization, which redirects
    writes to AppData to a package-specific virtualized location.
    """
    if sys.platform == "win32":
        _migrate_legacy_app_data()
        base_dir = _get_windows_documents_dir()
    elif sys.platform == "darwin":
        base_dir = os.path.join(os.path.expanduser(
            '~'), 'Library', 'Application Support')
    else:
        base_dir = os.path.join(os.path.expanduser('~'), '.local', 'share')
    if not base_dir:
        base_dir = os.path.expanduser('~')
    app_data_dir = os.path.join(base_dir, 'ProjectManagementApp')
    os.makedirs(app_data_dir, exist_ok=True)
    return app_data_dir


def get_app_data_path(file_name="tasks.json"):
    """
    Determines the correct, cross-platform path for storing user data
    and returns the full path for the given file_name.
    """
    return os.path.join(get_app_data_dir(), file_name)


# --- Configuration ---
BASE_DIR = Path(__file__).resolve().parent
loaded = load_dotenv(BASE_DIR / '.env', override=True)  # <— force .env to win
logging.info(f".env loaded: {loaded} (override=True)")
TASKS_FILE = get_app_data_path("tasks.json")
NOTES_FILE = get_app_data_path("notes.json")
SETTINGS_FILE = get_app_data_path("settings.json")
TIMESHEETS_FILE = get_app_data_path("timesheets.json")
TEMPLATES_FILE = get_app_data_path("templates.json")
CAD_AUTO_SELECT_TRACE_FILE = get_app_data_path("cad_auto_select_trace.log")
CHECKLISTS_FILE = get_app_data_path("checklists.json")
SYNC_BACKUPS_DIR = os.path.join(get_app_data_dir(), "sync_backups")
SYNC_METADATA_FILE = get_app_data_path("local_project_sync_metadata.json")
LIGHTING_SCHEDULE_SYNC_FILE = "T24LightingFixtureSchedule.sync.json"
LIGHTING_SCHEDULE_DB_FILE = get_app_data_path("lighting_schedules.db")
LIGHTING_PLAN_DB_FILE = get_app_data_path("lighting_plans.db")
LIGHTING_PLAN_SNAPSHOT_FILE = "ACIESLightingPlan.snapshot.json"
LIGHTING_PLAN_INSTRUCTIONS_FILE = "ACIESLightingPlan.instructions.json"
PROJECT_CHECKLIST_DB_FILE = get_app_data_path("project_checklists.db")
EMAIL_CAPTURE_DB_FILE = get_app_data_path("email_capture.db")
EMAIL_CAPTURE_SCHEMA_VERSION = 1
EMAIL_CAPTURE_FIRST_SCAN_DAYS = 7
EMAIL_CAPTURE_MAX_WINDOW_DAYS = 14
EMAIL_CAPTURE_SCAN_OVERLAP_MINUTES = 60
EMAIL_CAPTURE_MAX_MESSAGES_PER_SCAN = 600
SYNC_TRACKED_FILES = {
    "settings": SETTINGS_FILE,
    "tasks": TASKS_FILE,
    "notes": NOTES_FILE,
    "timesheets": TIMESHEETS_FILE,
    "templates": TEMPLATES_FILE,
    "checklists": CHECKLISTS_FILE,
}
LIGHTING_SCHEDULE_FIELDS = [
    "mark",
    "description",
    "manufacturer",
    "modelNumber",
    "mounting",
    "volts",
    "watts",
    "notes",
]
LIGHTING_SCHEDULE_ROW_META_FIELDS = [
    "symbolAssetPath",
    "symbolAlt",
    "starterFixtureKey",
]
LIGHTING_SCHEDULE_CA_STARTER_MODELS = {
    "L1": "IC1JBPF 07LM 30K 90CRI 120 FRPC",
    "L2": "6020EN3-15",
    "L3": "49119EN3-962",
    "L4": "WF4 ADJ SWW5 90CRI MW M6",
    "L5": "8909PEN3-12",
    "L6": "WF4 REG SWW5 90CRI MW M6",
    "L7": "T24M-2C-TUBS-SP-30K-WH",
    "L8": "JSBC 6IN 30K 90CRI WH",
    "L9": "TR24M-2C-WH",
    "L10": "T24M LINEA",
    "L11": "HTLHD-WW-16",
}
LIGHTING_SCHEDULE_DEFAULT_GENERAL_NOTES = "\n".join([
    "A.  VERIFY ALL CEILING TYPES PRIOR TO ORDERING FIXTURES.",
    "B.  VERIFY ALL OPERATING VOLTAGE PRIOR TO ORDERING FIXTURES.",
    "C.  COORDINATE THE HEIGHT OF ALL SUSPENDED FIXTURES WITH THE OWNER, ARCHITECT AND ENGINEER PRIOR TO INSTALLATION.",
    "D.  ALL LAMPS SHALL HAVE A COLOR TEMPERATURE OF 3500 DEG. KELVIN AND A CRI OF 85 UNLESS SPECIFICALLY NOTED.",
    'E.  FIXTURES DESIGNATED WITH A "1/2 SHADE" OR "FULL SHADE" AND ALL EXIT SIGNS SHALL BE CIRCUITED TO THE CENTRAL EMERGENCY BATTERY INVERTER OR PROVIDED WITH INTEGRAL BATTERY BACKUP AS REQUIRED. EM BACKUP POWER SHALL BE SUITABLE TO PROVIDE FULL POWER TO FIXTURES FOR A MINIMUM OF 90 MINUTES.',
])
LIGHTING_SCHEDULE_DEFAULT_NOTES = "\n".join([
    "1.  CONFIRM THE DRIVER TYPE WITH VENDOR. LIGHT DRIVER SHALL BE COMPATIBLE WITH THE DIMMER. REFER TO THE SENSOR SCHEDULE FOR DETAILS.",
    "2.  COORDINATE THE FINAL MANUFACTURER AND MODEL WITH ARCHITECT.",
])
LIGHTING_PROJECT_SEGMENT_REGEX = re.compile(
    r"^(\d{6})(?!\d)(?:\s*(?:[-_]\s*)?(.*))?$"
)


def _normalize_sync_dwg_path(dwg_path):
    raw = str(dwg_path or "").strip().strip('"').strip("'")
    if not raw:
        raise ValueError("DWG path is required.")
    normalized = os.path.normpath(raw)
    if not os.path.isabs(normalized):
        raise ValueError("DWG path must be an absolute path.")
    if os.path.splitext(normalized)[1].lower() != ".dwg":
        raise ValueError("DWG path must end with .dwg.")
    folder = os.path.dirname(normalized)
    if not folder:
        raise ValueError("DWG path must include a parent folder.")
    return normalized


def _resolve_lighting_schedule_sync_path(dwg_path):
    normalized_dwg = _normalize_sync_dwg_path(dwg_path)
    folder = os.path.dirname(normalized_dwg)
    sync_path = os.path.join(folder, LIGHTING_SCHEDULE_SYNC_FILE)
    return normalized_dwg, sync_path


def _normalize_t24_output_json_path(json_path):
    raw = str(json_path or "").strip().strip('"').strip("'")
    if not raw:
        raise ValueError("JSON path is required.")
    normalized = os.path.normpath(raw)
    if not os.path.isabs(normalized):
        raise ValueError("JSON path must be an absolute path.")
    if os.path.splitext(normalized)[1].lower() != ".json":
        raise ValueError("JSON path must end with .json.")
    if os.path.basename(normalized).lower() != "t24output.json":
        raise ValueError("Expected file name: T24Output.json.")
    return normalized


def _read_json_file_strict(path):
    try:
        with open(path, "r", encoding="utf-8-sig") as f:
            raw_text = f.read()
    except UnicodeDecodeError as exc:
        raise ValueError(f"Sync file '{path}' is not valid UTF-8 text: {exc}.")

    if raw_text.startswith("\ufeff"):
        raw_text = raw_text.lstrip("\ufeff")

    try:
        return json.loads(raw_text)
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"Malformed JSON in sync file '{path}': {exc.msg} "
            f"(line {exc.lineno}, column {exc.colno})."
        )


def _atomic_write_json_file(path, payload):
    folder = os.path.dirname(path)
    os.makedirs(folder, exist_ok=True)
    fd, temp_path = tempfile.mkstemp(
        prefix="lfs_sync_",
        suffix=".tmp",
        dir=folder,
        text=True,
    )
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        os.replace(temp_path, path)
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except OSError:
                pass


# --- Local Project Manager sync baseline metadata ---
# Records the last-synced mtime/size of both sides for each file copied by the
# Work Locally tool so later comparisons can tell "one side changed" apart from
# "both sides changed" (conflict) and "deleted locally" apart from "never copied".
_SYNC_METADATA_LOCK = threading.Lock()
SYNC_BASELINE_MTIME_EPSILON_SECONDS = 2.0


def _normalize_sync_metadata_pair_key(local_project_path, server_project_path):
    normalized_local = os.path.normpath(str(local_project_path or '').strip()).lower()
    normalized_server = os.path.normpath(str(server_project_path or '').strip()).lower()
    return f"{normalized_local}|{normalized_server}"


def _normalize_sync_metadata_file_key(relative_path):
    return os.path.normpath(str(relative_path or '').strip()).lower()


def _load_sync_metadata():
    metadata_path = SYNC_METADATA_FILE
    if not metadata_path or not os.path.exists(metadata_path):
        return {'version': 1, 'pairs': {}}
    try:
        payload = _read_json_file_strict(metadata_path)
    except Exception as exc:
        logging.warning(f"Could not read sync metadata file '{metadata_path}': {exc}")
        return {'version': 1, 'pairs': {}}
    if not isinstance(payload, dict) or not isinstance(payload.get('pairs'), dict):
        logging.warning(f"Sync metadata file '{metadata_path}' has an unexpected shape; starting fresh.")
        return {'version': 1, 'pairs': {}}
    payload.setdefault('version', 1)
    return payload


def _save_sync_metadata(metadata):
    _atomic_write_json_file(SYNC_METADATA_FILE, metadata)


def _get_sync_baseline_files(local_project_path, server_project_path):
    pair_key = _normalize_sync_metadata_pair_key(local_project_path, server_project_path)
    with _SYNC_METADATA_LOCK:
        metadata = _load_sync_metadata()
        pair_entry = metadata.get('pairs', {}).get(pair_key)
        if not isinstance(pair_entry, dict):
            return {}
        files = pair_entry.get('files')
        return copy.deepcopy(files) if isinstance(files, dict) else {}


def _sync_baseline_entry_effectively_equal(existing_entry, new_entry):
    if not isinstance(existing_entry, dict):
        return False
    for mtime_key in ('localMtime', 'serverMtime'):
        try:
            existing_value = float(existing_entry.get(mtime_key) or 0.0)
            new_value = float(new_entry.get(mtime_key) or 0.0)
        except (TypeError, ValueError):
            return False
        if abs(existing_value - new_value) > SYNC_BASELINE_MTIME_EPSILON_SECONDS:
            return False
    for size_key in ('localSize', 'serverSize'):
        if existing_entry.get(size_key) != new_entry.get(size_key):
            return False
    return True


def _update_sync_baseline_files(
    local_project_path,
    server_project_path,
    upsert_entries=None,
    remove_relative_paths=None,
):
    """Upsert/remove per-file baseline entries for a local/server pair.

    Failures are logged and swallowed: baseline maintenance must never break a
    compare or copy operation.
    """
    try:
        pair_key = _normalize_sync_metadata_pair_key(local_project_path, server_project_path)
        with _SYNC_METADATA_LOCK:
            metadata = _load_sync_metadata()
            pairs = metadata.setdefault('pairs', {})
            pair_entry = pairs.get(pair_key)
            if not isinstance(pair_entry, dict):
                pair_entry = {}
            files = pair_entry.get('files')
            if not isinstance(files, dict):
                files = {}
            changed = False

            for entry in upsert_entries or []:
                relative_path = str((entry or {}).get('relativePath') or '').strip()
                if not relative_path:
                    continue
                file_key = _normalize_sync_metadata_file_key(relative_path)
                new_entry = {
                    'relativePath': os.path.normpath(relative_path),
                    'localMtime': float(entry.get('localMtime') or 0.0),
                    'localSize': entry.get('localSize'),
                    'serverMtime': float(entry.get('serverMtime') or 0.0),
                    'serverSize': entry.get('serverSize'),
                    'syncedAt': datetime.datetime.now().isoformat(),
                }
                if _sync_baseline_entry_effectively_equal(files.get(file_key), new_entry):
                    continue
                files[file_key] = new_entry
                changed = True

            for relative_path in remove_relative_paths or []:
                file_key = _normalize_sync_metadata_file_key(relative_path)
                if file_key in files:
                    del files[file_key]
                    changed = True

            if not changed:
                return

            pair_entry['localProjectPath'] = os.path.normpath(str(local_project_path or '').strip())
            pair_entry['serverProjectPath'] = os.path.normpath(str(server_project_path or '').strip())
            pair_entry['updatedAt'] = datetime.datetime.now().isoformat()
            pair_entry['files'] = files
            pairs[pair_key] = pair_entry
            _save_sync_metadata(metadata)
    except Exception as exc:
        logging.warning(f"Could not update sync baseline metadata: {exc}")


def _get_file_modified_iso(path):
    if not path or not os.path.exists(path):
        return ""
    try:
        modified = datetime.datetime.fromtimestamp(
            os.path.getmtime(path),
            tz=datetime.timezone.utc,
        )
        return modified.replace(microsecond=0).isoformat().replace("+00:00", "Z")
    except OSError:
        return ""


def _get_local_sync_metadata():
    files = {}
    for key, path in SYNC_TRACKED_FILES.items():
        files[key] = {
            "path": os.path.normpath(path),
            "exists": os.path.exists(path),
            "modified": _get_file_modified_iso(path),
        }
    return files


def _create_cloud_sync_backup(reason="", metadata=None):
    created_at = utc_now_iso()
    safe_reason = re.sub(r"[^a-z0-9]+", "-", str(reason or "").strip().lower()).strip("-")
    suffix = safe_reason[:40] or "sync"
    backup_dir = os.path.join(
        SYNC_BACKUPS_DIR,
        f"{created_at.replace(':', '').replace('-', '')}_{suffix}_{uuid.uuid4().hex[:8]}",
    )
    os.makedirs(backup_dir, exist_ok=True)

    copied = []
    for key, source_path in SYNC_TRACKED_FILES.items():
        if not os.path.exists(source_path):
            continue
        destination_path = os.path.join(backup_dir, os.path.basename(source_path))
        shutil.copy2(source_path, destination_path)
        copied.append(
            {
                "key": key,
                "sourcePath": os.path.normpath(source_path),
                "backupPath": os.path.normpath(destination_path),
                "modified": _get_file_modified_iso(source_path),
            }
        )

    backup_metadata = {
        "createdAt": created_at,
        "reason": str(reason or "").strip(),
        "files": copied,
        "localSyncMetadata": _get_local_sync_metadata(),
    }
    if isinstance(metadata, dict) and metadata:
        backup_metadata["metadata"] = metadata

    _atomic_write_json_file(os.path.join(backup_dir, "metadata.json"), backup_metadata)
    return {
        "path": os.path.normpath(backup_dir),
        "createdAt": created_at,
        "files": copied,
    }


def _normalize_lighting_schedule_text(value):
    return str(value or "").replace("\r\n", "\n").replace("\r", "\n")


def _get_lighting_schedule_json_value(payload, field, default=None):
    if not isinstance(payload, dict):
        return default
    if field in payload:
        return payload[field]
    pascal_field = field[:1].upper() + field[1:]
    return payload.get(pascal_field, default)


def _create_default_lighting_schedule_row(seed=None):
    seed = seed or {}
    row = {}
    for field in LIGHTING_SCHEDULE_FIELDS:
        row[field] = _normalize_lighting_schedule_text(
            _get_lighting_schedule_json_value(seed, field)
        )
    for field in LIGHTING_SCHEDULE_ROW_META_FIELDS:
        row[field] = _normalize_lighting_schedule_text(
            _get_lighting_schedule_json_value(seed, field)
        )
    if not row["symbolAssetPath"] and not row["starterFixtureKey"]:
        expected_model = LIGHTING_SCHEDULE_CA_STARTER_MODELS.get(row["mark"].upper())
        if expected_model and expected_model.casefold() in row["modelNumber"].casefold():
            fixture_number = row["mark"][1:]
            row["starterFixtureKey"] = f"ca-2025-res-l{fixture_number}"
            row["symbolAssetPath"] = (
                f"assets/lighting/ca-residential-l{fixture_number}.png"
            )
            row["symbolAlt"] = row["symbolAlt"] or (
                f"{row['description']} symbol"
                if row["description"]
                else f"Fixture {row['mark']} symbol"
            )
    return row


def _create_default_lighting_schedule():
    return {
        "rows": [_create_default_lighting_schedule_row()],
        "generalNotes": LIGHTING_SCHEDULE_DEFAULT_GENERAL_NOTES,
        "notes": LIGHTING_SCHEDULE_DEFAULT_NOTES,
        "includeSymbolColumn": False,
        "targetDwgPath": "",
    }


def _normalize_lighting_schedule_payload(schedule):
    normalized = (
        dict(schedule)
        if isinstance(schedule, dict)
        else _create_default_lighting_schedule()
    )
    rows = _get_lighting_schedule_json_value(normalized, "rows")
    if not isinstance(rows, list) or not rows:
        normalized["rows"] = [_create_default_lighting_schedule_row()]
    else:
        normalized["rows"] = [
            _create_default_lighting_schedule_row(row)
            for row in rows
            if isinstance(row, dict)
        ] or [_create_default_lighting_schedule_row()]

    normalized["generalNotes"] = _normalize_lighting_schedule_text(
        _get_lighting_schedule_json_value(normalized, "generalNotes")
        if _get_lighting_schedule_json_value(normalized, "generalNotes") is not None
        else LIGHTING_SCHEDULE_DEFAULT_GENERAL_NOTES
    )
    normalized["notes"] = _normalize_lighting_schedule_text(
        _get_lighting_schedule_json_value(normalized, "notes")
        if _get_lighting_schedule_json_value(normalized, "notes") is not None
        else LIGHTING_SCHEDULE_DEFAULT_NOTES
    )
    normalized["includeSymbolColumn"] = (
        _get_lighting_schedule_json_value(normalized, "includeSymbolColumn") is True
        or any(
            row.get("symbolAssetPath") or row.get("starterFixtureKey")
            for row in normalized["rows"]
        )
    )
    normalized["targetDwgPath"] = os.path.normpath(
        str(normalized.get("targetDwgPath") or "").strip()
    ) if str(normalized.get("targetDwgPath") or "").strip() else ""
    return normalized


def _extract_lighting_schedule_project_id_from_path(raw_path):
    normalized = str(raw_path or "").strip().replace("/", "\\")
    if not normalized:
        return ""

    parts = [part.strip() for part in normalized.split("\\") if part.strip()]
    for part in reversed(parts):
        match = LIGHTING_PROJECT_SEGMENT_REGEX.match(part)
        if match:
            return match.group(1)
    return ""


def _resolve_lighting_schedule_project_id(project_like):
    if isinstance(project_like, str):
        value = str(project_like).strip()
        return value

    if not isinstance(project_like, dict):
        return ""

    explicit_id = str(project_like.get("id") or "").strip()
    if explicit_id:
        return explicit_id

    for key in ("path", "localProjectPath", "workroomRootPath"):
        candidate = _extract_lighting_schedule_project_id_from_path(
            project_like.get(key)
        )
        if candidate:
            return candidate

    display_name = str(project_like.get("name") or project_like.get("nick") or "").strip()
    if display_name:
        return f"name:{display_name.lower()}"

    return ""


def _open_lighting_schedule_db():
    conn = sqlite3.connect(LIGHTING_SCHEDULE_DB_FILE, timeout=5)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode = WAL;")
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS lighting_schedule_projects (
            project_id TEXT PRIMARY KEY,
            schedule_json TEXT NOT NULL,
            target_dwg_path TEXT NOT NULL DEFAULT '',
            table_handle TEXT NOT NULL DEFAULT '',
            version INTEGER NOT NULL DEFAULT 1,
            updated_at_utc TEXT NOT NULL,
            updated_by TEXT NOT NULL DEFAULT 'unknown'
        );
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS lighting_schedule_links (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id TEXT NOT NULL,
            dwg_path TEXT NOT NULL UNIQUE,
            table_handle TEXT NOT NULL DEFAULT '',
            last_applied_version INTEGER NOT NULL DEFAULT 0,
            last_seen_at_utc TEXT NOT NULL DEFAULT '',
            FOREIGN KEY(project_id) REFERENCES lighting_schedule_projects(project_id)
                ON DELETE CASCADE
        );
        """
    )
    conn.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_lighting_schedule_links_project_id
        ON lighting_schedule_links(project_id);
        """
    )
    return conn


def _normalize_lighting_schedule_record(row):
    if row is None:
        return None

    try:
        schedule_payload = json.loads(row["schedule_json"])
    except (TypeError, ValueError, KeyError):
        schedule_payload = _create_default_lighting_schedule()

    schedule = _normalize_lighting_schedule_payload(schedule_payload)
    schedule["targetDwgPath"] = os.path.normpath(
        str(row["target_dwg_path"] or "").strip()
    ) if str(row["target_dwg_path"] or "").strip() else ""

    return {
        "projectId": str(row["project_id"] or "").strip(),
        "version": int(row["version"] or 0),
        "updatedAtUtc": str(row["updated_at_utc"] or "").strip(),
        "updatedBy": str(row["updated_by"] or "").strip(),
        "targetDwgPath": schedule["targetDwgPath"],
        "tableHandle": str(row["table_handle"] or "").strip(),
        "schedule": {
            "rows": schedule["rows"],
            "generalNotes": schedule["generalNotes"],
            "notes": schedule["notes"],
        },
    }


def _get_lighting_schedule_record(project_id):
    resolved_id = _resolve_lighting_schedule_project_id(project_id)
    if not resolved_id:
        return None

    with closing(_open_lighting_schedule_db()) as conn:
        row = conn.execute(
            """
            SELECT project_id, schedule_json, target_dwg_path, table_handle, version,
                   updated_at_utc, updated_by
            FROM lighting_schedule_projects
            WHERE project_id = ?
            """,
            (resolved_id,),
        ).fetchone()
        return _normalize_lighting_schedule_record(row)


def _get_lighting_schedule_version(project_id):
    record = _get_lighting_schedule_record(project_id)
    if not record:
        return {
            "exists": False,
            "projectId": _resolve_lighting_schedule_project_id(project_id),
            "version": 0,
            "updatedAtUtc": "",
            "updatedBy": "",
        }

    return {
        "exists": True,
        "projectId": record["projectId"],
        "version": record["version"],
        "updatedAtUtc": record["updatedAtUtc"],
        "updatedBy": record["updatedBy"],
    }


def _get_lighting_schedule_links(project_id):
    resolved_id = _resolve_lighting_schedule_project_id(project_id)
    if not resolved_id:
        return []

    with closing(_open_lighting_schedule_db()) as conn:
        rows = conn.execute(
            """
            SELECT id, project_id, dwg_path, table_handle, last_applied_version, last_seen_at_utc
            FROM lighting_schedule_links
            WHERE project_id = ?
            ORDER BY dwg_path COLLATE NOCASE
            """,
            (resolved_id,),
        ).fetchall()

    links = []
    for row in rows:
        links.append(
            {
                "id": int(row["id"]),
                "projectId": str(row["project_id"] or "").strip(),
                "dwgPath": str(row["dwg_path"] or "").strip(),
                "tableHandle": str(row["table_handle"] or "").strip(),
                "lastAppliedVersion": int(row["last_applied_version"] or 0),
                "lastSeenAtUtc": str(row["last_seen_at_utc"] or "").strip(),
            }
        )
    return links


def _upsert_lighting_schedule_link(
    conn,
    project_id,
    dwg_path,
    table_handle="",
    last_applied_version=0,
):
    normalized_dwg = os.path.normpath(str(dwg_path or "").strip())
    if not normalized_dwg:
        return

    now_iso = datetime.datetime.utcnow().isoformat()
    existing = conn.execute(
        """
        SELECT table_handle, last_applied_version
        FROM lighting_schedule_links
        WHERE dwg_path = ?
        """,
        (normalized_dwg,),
    ).fetchone()

    next_table_handle = str(table_handle or "").strip()
    if not next_table_handle and existing:
        next_table_handle = str(existing["table_handle"] or "").strip()

    next_applied_version = int(last_applied_version or 0)
    if existing and next_applied_version <= 0:
        next_applied_version = int(existing["last_applied_version"] or 0)

    conn.execute(
        """
        INSERT INTO lighting_schedule_links (
            project_id, dwg_path, table_handle, last_applied_version, last_seen_at_utc
        )
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(dwg_path) DO UPDATE SET
            project_id = excluded.project_id,
            table_handle = excluded.table_handle,
            last_applied_version = excluded.last_applied_version,
            last_seen_at_utc = excluded.last_seen_at_utc
        """,
        (
            project_id,
            normalized_dwg,
            next_table_handle,
            next_applied_version,
            now_iso,
        ),
    )


def _save_lighting_schedule_record(
    project_id,
    payload,
    updated_by="desktop",
):
    resolved_id = _resolve_lighting_schedule_project_id(project_id)
    if not resolved_id:
        raise ValueError("Lighting schedule project ID is required.")
    if not isinstance(payload, dict):
        raise ValueError("Lighting schedule payload must be a JSON object.")

    normalized_schedule = _normalize_lighting_schedule_payload(payload.get("schedule"))
    target_dwg_path = payload.get("targetDwgPath")
    if target_dwg_path is None:
        target_dwg_path = normalized_schedule.get("targetDwgPath")
    normalized_target_dwg_path = (
        os.path.normpath(str(target_dwg_path or "").strip())
        if str(target_dwg_path or "").strip()
        else ""
    )
    table_handle = str(payload.get("tableHandle") or "").strip()
    expected_version = payload.get("expectedVersion", None)

    now_iso = datetime.datetime.utcnow().isoformat()
    canonical_schedule_json = json.dumps(
        {
            "rows": normalized_schedule["rows"],
            "generalNotes": normalized_schedule["generalNotes"],
            "notes": normalized_schedule["notes"],
            "includeSymbolColumn": normalized_schedule["includeSymbolColumn"],
        },
        ensure_ascii=False,
        separators=(",", ":"),
    )

    with closing(_open_lighting_schedule_db()) as conn:
        existing = conn.execute(
            """
            SELECT project_id, schedule_json, target_dwg_path, table_handle, version
            FROM lighting_schedule_projects
            WHERE project_id = ?
            """,
            (resolved_id,),
        ).fetchone()

        previous_version = int(existing["version"] or 0) if existing else 0
        next_version = previous_version + 1 if existing else 1
        current_table_handle = table_handle
        if existing and not current_table_handle:
            current_table_handle = str(existing["table_handle"] or "").strip()

        conn.execute(
            """
            INSERT INTO lighting_schedule_projects (
                project_id, schedule_json, target_dwg_path, table_handle,
                version, updated_at_utc, updated_by
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(project_id) DO UPDATE SET
                schedule_json = excluded.schedule_json,
                target_dwg_path = excluded.target_dwg_path,
                table_handle = excluded.table_handle,
                version = excluded.version,
                updated_at_utc = excluded.updated_at_utc,
                updated_by = excluded.updated_by
            """,
            (
                resolved_id,
                canonical_schedule_json,
                normalized_target_dwg_path,
                current_table_handle,
                next_version,
                now_iso,
                str(updated_by or "desktop").strip() or "desktop",
            ),
        )

        if normalized_target_dwg_path:
            _upsert_lighting_schedule_link(
                conn,
                resolved_id,
                normalized_target_dwg_path,
                current_table_handle,
                payload.get("lastAppliedVersion", 0),
            )

        conn.commit()

    record = _get_lighting_schedule_record(resolved_id)
    if record is None:
        raise RuntimeError("Lighting schedule save completed but record could not be reloaded.")

    record["conflict"] = (
        expected_version is not None
        and str(expected_version).strip() != ""
        and int(expected_version) != previous_version
    )
    record["previousVersion"] = previous_version
    return record


def _normalize_lighting_plan_snapshot_path(json_path):
    raw = str(json_path or "").strip().strip('"').strip("'")
    if not raw:
        raise ValueError("Lighting plan snapshot path is required.")
    normalized = os.path.normpath(raw)
    if not os.path.isabs(normalized):
        raise ValueError("Lighting plan snapshot path must be absolute.")
    if os.path.splitext(normalized)[1].lower() != ".json":
        raise ValueError("Lighting plan snapshot path must end with .json.")
    return normalized


def _open_lighting_plan_db():
    conn = sqlite3.connect(LIGHTING_PLAN_DB_FILE, timeout=5)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode = WAL;")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS lighting_plan_projects (
            project_id TEXT PRIMARY KEY,
            snapshot_path TEXT NOT NULL DEFAULT '',
            snapshot_json TEXT NOT NULL,
            analysis_json TEXT NOT NULL,
            instruction_path TEXT NOT NULL DEFAULT '',
            version INTEGER NOT NULL DEFAULT 1,
            updated_at_utc TEXT NOT NULL,
            updated_by TEXT NOT NULL DEFAULT 'unknown'
        );
        """
    )
    return conn


def _normalize_lighting_plan_record(row):
    if row is None:
        return None
    try:
        analysis = json.loads(row["analysis_json"])
    except (TypeError, ValueError, KeyError):
        analysis = {}
    return {
        "projectId": str(row["project_id"] or "").strip(),
        "snapshotPath": str(row["snapshot_path"] or "").strip(),
        "instructionPath": str(row["instruction_path"] or "").strip(),
        "version": int(row["version"] or 0),
        "updatedAtUtc": str(row["updated_at_utc"] or "").strip(),
        "updatedBy": str(row["updated_by"] or "").strip(),
        "analysis": analysis if isinstance(analysis, dict) else {},
    }


def _get_lighting_plan_record(project_id):
    resolved_id = _resolve_lighting_schedule_project_id(project_id)
    if not resolved_id:
        return None
    with closing(_open_lighting_plan_db()) as conn:
        row = conn.execute(
            """
            SELECT project_id, snapshot_path, snapshot_json, analysis_json,
                   instruction_path, version, updated_at_utc, updated_by
            FROM lighting_plan_projects
            WHERE project_id = ?
            """,
            (resolved_id,),
        ).fetchone()
    return _normalize_lighting_plan_record(row)


def _save_lighting_plan_record(
    project_id,
    snapshot,
    snapshot_path="",
    schedule_rows=None,
    updated_by="desktop",
):
    resolved_id = _resolve_lighting_schedule_project_id(project_id)
    if not resolved_id:
        raise ValueError("Lighting plan project ID is required.")
    analysis = analyze_lighting_plan(snapshot, schedule_rows or [])
    normalized_snapshot_path = (
        _normalize_lighting_plan_snapshot_path(snapshot_path)
        if str(snapshot_path or "").strip()
        else ""
    )
    now_iso = datetime.datetime.now(datetime.timezone.utc).isoformat()
    with closing(_open_lighting_plan_db()) as conn:
        existing = conn.execute(
            "SELECT version, instruction_path FROM lighting_plan_projects WHERE project_id = ?",
            (resolved_id,),
        ).fetchone()
        next_version = int(existing["version"] or 0) + 1 if existing else 1
        instruction_path = str(existing["instruction_path"] or "") if existing else ""
        conn.execute(
            """
            INSERT INTO lighting_plan_projects (
                project_id, snapshot_path, snapshot_json, analysis_json,
                instruction_path, version, updated_at_utc, updated_by
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(project_id) DO UPDATE SET
                snapshot_path = excluded.snapshot_path,
                snapshot_json = excluded.snapshot_json,
                analysis_json = excluded.analysis_json,
                instruction_path = excluded.instruction_path,
                version = excluded.version,
                updated_at_utc = excluded.updated_at_utc,
                updated_by = excluded.updated_by
            """,
            (
                resolved_id,
                normalized_snapshot_path,
                json.dumps(snapshot, ensure_ascii=False, separators=(",", ":")),
                json.dumps(analysis, ensure_ascii=False, separators=(",", ":")),
                instruction_path,
                next_version,
                now_iso,
                str(updated_by or "desktop").strip() or "desktop",
            ),
        )
        conn.commit()
    return _get_lighting_plan_record(resolved_id)


def _import_lighting_plan_snapshot(project_id, snapshot_path, schedule_rows=None):
    normalized_path = _normalize_lighting_plan_snapshot_path(snapshot_path)
    if not os.path.isfile(normalized_path):
        raise FileNotFoundError(
            f"Lighting plan snapshot was not found at: {normalized_path}"
        )
    snapshot = _read_json_file_strict(normalized_path)
    return _save_lighting_plan_record(
        project_id,
        snapshot,
        snapshot_path=normalized_path,
        schedule_rows=schedule_rows,
        updated_by="desktop-import",
    )


def _export_lighting_plan_instructions(project_id, output_path=None):
    resolved_id = _resolve_lighting_schedule_project_id(project_id)
    if not resolved_id:
        raise ValueError("Lighting plan project ID is required.")
    with closing(_open_lighting_plan_db()) as conn:
        row = conn.execute(
            """
            SELECT snapshot_path, analysis_json
            FROM lighting_plan_projects
            WHERE project_id = ?
            """,
            (resolved_id,),
        ).fetchone()
        if row is None:
            raise ValueError("Import a lighting plan snapshot before generating CAD instructions.")
        analysis = json.loads(row["analysis_json"])
        if not isinstance(analysis, dict):
            raise ValueError("Stored lighting plan analysis is invalid.")
        summary = analysis.get("summary") if isinstance(analysis.get("summary"), dict) else {}
        if int(summary.get("errorCount") or 0) > 0:
            raise ValueError(
                "Lighting plan has validation errors. Resolve them and rescan before generating CAD instructions."
            )
        instructions = analysis.get("instructions")
        if not isinstance(instructions, dict):
            raise ValueError("Stored lighting plan does not contain CAD instructions.")

        raw_output_path = str(output_path or "").strip()
        if raw_output_path:
            normalized_output_path = _normalize_lighting_plan_snapshot_path(raw_output_path)
        else:
            snapshot_path = str(row["snapshot_path"] or "").strip()
            if not snapshot_path:
                raise ValueError("The imported snapshot has no source folder for CAD instructions.")
            normalized_output_path = os.path.join(
                os.path.dirname(snapshot_path), LIGHTING_PLAN_INSTRUCTIONS_FILE
            )

        _atomic_write_json_file(normalized_output_path, instructions)
        conn.execute(
            """
            UPDATE lighting_plan_projects
            SET instruction_path = ?, updated_at_utc = ?, updated_by = 'desktop-export'
            WHERE project_id = ?
            """,
            (
                normalized_output_path,
                datetime.datetime.now(datetime.timezone.utc).isoformat(),
                resolved_id,
            ),
        )
        conn.commit()
    record = _get_lighting_plan_record(resolved_id)
    return {
        "path": normalized_output_path,
        "tagCount": len(instructions.get("tags") or []),
        "record": record,
    }


def _migrate_project_lighting_schedules(projects):
    if not isinstance(projects, list) or not projects:
        return

    with closing(_open_lighting_schedule_db()) as conn:
        for project in projects:
            if not isinstance(project, dict):
                continue

            project_id = _resolve_lighting_schedule_project_id(project)
            if not project_id:
                continue

            existing = conn.execute(
                "SELECT version FROM lighting_schedule_projects WHERE project_id = ?",
                (project_id,),
            ).fetchone()
            if existing:
                continue

            legacy_schedule = _normalize_lighting_schedule_payload(
                project.get("lightingSchedule")
            )
            target_dwg_path = legacy_schedule.get("targetDwgPath") or ""
            now_iso = datetime.datetime.utcnow().isoformat()
            conn.execute(
                """
                INSERT INTO lighting_schedule_projects (
                    project_id, schedule_json, target_dwg_path, table_handle,
                    version, updated_at_utc, updated_by
                )
                VALUES (?, ?, ?, '', 1, ?, 'migration')
                ON CONFLICT(project_id) DO NOTHING
                """,
                (
                    project_id,
                    json.dumps(
                        {
                            "rows": legacy_schedule["rows"],
                            "generalNotes": legacy_schedule["generalNotes"],
                            "notes": legacy_schedule["notes"],
                        },
                        ensure_ascii=False,
                        separators=(",", ":"),
                    ),
                    target_dwg_path,
                    now_iso,
                ),
            )
            if target_dwg_path:
                _upsert_lighting_schedule_link(
                    conn,
                    project_id,
                    target_dwg_path,
                    "",
                    0,
                )
        conn.commit()


def _overlay_projects_with_lighting_schedule_records(projects):
    if not isinstance(projects, list):
        return projects

    _migrate_project_lighting_schedules(projects)
    for project in projects:
        if not isinstance(project, dict):
            continue

        project_id = _resolve_lighting_schedule_project_id(project)
        if not project_id:
            continue

        record = _get_lighting_schedule_record(project_id)
        if not record:
            continue

        existing_schedule = _normalize_lighting_schedule_payload(
            project.get("lightingSchedule")
        )
        project["lightingSchedule"] = {
            **existing_schedule,
            "rows": record["schedule"]["rows"],
            "generalNotes": record["schedule"]["generalNotes"],
            "notes": record["schedule"]["notes"],
            "targetDwgPath": record["targetDwgPath"],
            "_storeVersion": record["version"],
            "_storeUpdatedAtUtc": record["updatedAtUtc"],
            "_storeUpdatedBy": record["updatedBy"],
            "_tableHandle": record["tableHandle"],
        }
    return projects


def _normalize_project_checklist_project_id(project_like):
    return _resolve_lighting_schedule_project_id(project_like)


def _create_default_project_checklist_state():
    return {"checklists": []}


def _normalize_project_checklist_state(state):
    source = state if isinstance(state, dict) else {}
    raw_checklists = source.get("checklists")
    raw_checklists = raw_checklists if isinstance(raw_checklists, list) else []

    checklists = []
    seen_checklists = set()
    for raw_entry in raw_checklists:
        if not isinstance(raw_entry, dict):
            continue

        checklist_id = str(
            raw_entry.get("checklistId") or raw_entry.get("id") or ""
        ).strip()
        if not checklist_id or checklist_id in seen_checklists:
            continue
        seen_checklists.add(checklist_id)

        completed_items = []
        seen_items = set()
        for raw_item_id in raw_entry.get("completedItems") or []:
            item_id = str(raw_item_id or "").strip()
            if not item_id or item_id in seen_items:
                continue
            seen_items.add(item_id)
            completed_items.append(item_id)

        raw_notes = raw_entry.get("itemNotes")
        raw_notes = raw_notes if isinstance(raw_notes, dict) else {}
        item_notes = {}
        for raw_item_id, raw_note in raw_notes.items():
            item_id = str(raw_item_id or "").strip()
            if not item_id:
                continue
            item_notes[item_id] = str(raw_note if raw_note is not None else "")

        checklists.append(
            {
                "checklistId": checklist_id,
                "completedItems": completed_items,
                "itemNotes": item_notes,
            }
        )

    return {"checklists": checklists}


def _open_project_checklist_db():
    conn = sqlite3.connect(PROJECT_CHECKLIST_DB_FILE, timeout=5)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode = WAL;")
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS project_checklist_records (
            project_id TEXT PRIMARY KEY,
            state_json TEXT NOT NULL,
            version INTEGER NOT NULL DEFAULT 1,
            updated_at_utc TEXT NOT NULL,
            updated_by TEXT NOT NULL DEFAULT 'unknown'
        );
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS project_checklist_links (
            dwg_path TEXT PRIMARY KEY,
            project_id TEXT NOT NULL,
            last_seen_at_utc TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES project_checklist_records(project_id)
                ON DELETE CASCADE
        );
        """
    )
    conn.commit()
    return conn


def _normalize_project_checklist_record(row):
    if row is None:
        return None

    try:
        state_payload = json.loads(row["state_json"])
    except (TypeError, ValueError, KeyError):
        state_payload = _create_default_project_checklist_state()

    return {
        "projectId": str(row["project_id"] or "").strip(),
        "version": int(row["version"] or 0),
        "updatedAtUtc": str(row["updated_at_utc"] or "").strip(),
        "updatedBy": str(row["updated_by"] or "").strip(),
        "state": _normalize_project_checklist_state(state_payload),
    }


def _get_project_checklist_record(project_id):
    resolved_id = _normalize_project_checklist_project_id(project_id)
    if not resolved_id:
        return None

    conn = _open_project_checklist_db()
    try:
        row = conn.execute(
            """
            SELECT project_id, state_json, version, updated_at_utc, updated_by
            FROM project_checklist_records
            WHERE project_id = ?
            """,
            (resolved_id,),
        ).fetchone()
        return _normalize_project_checklist_record(row)
    finally:
        conn.close()


def _get_project_checklist_version(project_id):
    record = _get_project_checklist_record(project_id)
    if not record:
        return {
            "exists": False,
            "projectId": _normalize_project_checklist_project_id(project_id),
            "version": 0,
            "updatedAtUtc": "",
            "updatedBy": "",
        }

    return {
        "exists": True,
        "projectId": record["projectId"],
        "version": record["version"],
        "updatedAtUtc": record["updatedAtUtc"],
        "updatedBy": record["updatedBy"],
    }


def _normalize_project_checklist_dwg_path(dwg_path):
    raw = str(dwg_path or "").strip().strip('"').strip("'")
    if not raw:
        return ""
    return os.path.normpath(raw)


def _get_project_checklist_link(dwg_path):
    normalized_dwg = _normalize_project_checklist_dwg_path(dwg_path)
    if not normalized_dwg:
        return None

    conn = _open_project_checklist_db()
    try:
        row = conn.execute(
            """
            SELECT dwg_path, project_id, last_seen_at_utc
            FROM project_checklist_links
            WHERE dwg_path = ?
            """,
            (normalized_dwg,),
        ).fetchone()
    finally:
        conn.close()

    if row is None:
        return None
    return {
        "dwgPath": str(row["dwg_path"] or "").strip(),
        "projectId": str(row["project_id"] or "").strip(),
        "lastSeenAtUtc": str(row["last_seen_at_utc"] or "").strip(),
    }


def _upsert_project_checklist_link(conn, project_id, dwg_path):
    resolved_id = _normalize_project_checklist_project_id(project_id)
    normalized_dwg = _normalize_project_checklist_dwg_path(dwg_path)
    if not resolved_id or not normalized_dwg:
        return

    now_iso = datetime.datetime.utcnow().isoformat()
    conn.execute(
        """
        INSERT INTO project_checklist_links (
            dwg_path, project_id, last_seen_at_utc
        )
        VALUES (?, ?, ?)
        ON CONFLICT(dwg_path) DO UPDATE SET
            project_id = excluded.project_id,
            last_seen_at_utc = excluded.last_seen_at_utc
        """,
        (normalized_dwg, resolved_id, now_iso),
    )


def _save_project_checklist_record(project_id, payload, updated_by="desktop"):
    resolved_id = _normalize_project_checklist_project_id(project_id)
    if not resolved_id:
        raise ValueError("Project checklist project ID is required.")
    if not isinstance(payload, dict):
        raise ValueError("Project checklist payload must be a JSON object.")

    state = _normalize_project_checklist_state(payload.get("state"))
    expected_version = payload.get("expectedVersion", None)
    canonical_state_json = json.dumps(
        state,
        ensure_ascii=False,
        separators=(",", ":"),
    )
    now_iso = datetime.datetime.utcnow().isoformat()

    conn = _open_project_checklist_db()
    try:
        conn.execute("BEGIN IMMEDIATE")
        existing = conn.execute(
            """
            SELECT project_id, version
            FROM project_checklist_records
            WHERE project_id = ?
            """,
            (resolved_id,),
        ).fetchone()

        previous_version = int(existing["version"] or 0) if existing else 0
        next_version = previous_version + 1 if existing else 1
        conn.execute(
            """
            INSERT INTO project_checklist_records (
                project_id, state_json, version, updated_at_utc, updated_by
            )
            VALUES (?, ?, ?, ?, ?)
            ON CONFLICT(project_id) DO UPDATE SET
                state_json = excluded.state_json,
                version = excluded.version,
                updated_at_utc = excluded.updated_at_utc,
                updated_by = excluded.updated_by
            """,
            (
                resolved_id,
                canonical_state_json,
                next_version,
                now_iso,
                str(updated_by or "desktop").strip() or "desktop",
            ),
        )

        dwg_path = payload.get("dwgPath") or payload.get("dwg_path")
        if dwg_path:
            _upsert_project_checklist_link(conn, resolved_id, dwg_path)
        conn.commit()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        raise
    finally:
        conn.close()

    record = _get_project_checklist_record(resolved_id)
    if record is None:
        raise RuntimeError("Project checklist save completed but record could not be reloaded.")

    try:
        expected_version_int = int(expected_version)
    except (TypeError, ValueError):
        expected_version_int = None
    record["conflict"] = (
        expected_version_int is not None
        and expected_version_int != previous_version
    )
    record["previousVersion"] = previous_version
    return record


def _save_project_checklist_link(project_id, dwg_path):
    resolved_id = _normalize_project_checklist_project_id(project_id)
    normalized_dwg = _normalize_project_checklist_dwg_path(dwg_path)
    if not resolved_id:
        raise ValueError("Project checklist project ID is required.")
    if not normalized_dwg:
        raise ValueError("DWG path is required.")

    existing = _get_project_checklist_record(resolved_id)
    if existing is None:
        _save_project_checklist_record(
            resolved_id,
            {"state": _create_default_project_checklist_state()},
            updated_by="link",
        )

    conn = _open_project_checklist_db()
    try:
        conn.execute("BEGIN IMMEDIATE")
        _upsert_project_checklist_link(conn, resolved_id, normalized_dwg)
        conn.commit()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        raise
    finally:
        conn.close()

    return _get_project_checklist_link(normalized_dwg)


def get_bundled_template_path(template_name: str) -> Path:
    """
    Resolve template files for both dev runs and packaged desktop runs.
    """
    candidates = [BASE_DIR / "templates" / template_name]

    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        candidates.extend([
            exe_dir / "_internal" / "templates" / template_name,
            exe_dir / "templates" / template_name,
        ])

    for candidate in candidates:
        if candidate.exists():
            return candidate

    return candidates[0]

EXPENSE_SHEET_NAME = "Project Expense Sheet"
EXPENSE_SECTION_ROWS = 12
EXPENSE_HEADER_ROWS = 3
EXPENSE_DEFAULT_DATA_ROWS = 5
EXPENSE_FOOTER_ROWS = 4
EXPENSE_INSERTED_DATA_TEMPLATE_OFFSET = EXPENSE_HEADER_ROWS + 1
EXPENSE_IMAGE_ROW_OFFSET = 6


DEFAULT_TEMPLATES = [
    {
        "name": "Narrative of Changes",
        "discipline": "General",
        "fileType": "docx",
        "sourcePath": str(get_bundled_template_path("Narrative of Changes.docx")),
        "description": "Standard narrative of changes template with MEP sections"
    },
    {
        "name": "Plan Check Comments",
        "discipline": "General",
        "fileType": "doc",
        "sourcePath": str(get_bundled_template_path("PCC.doc")),
        "description": "Plan check comments response table template"
    }
]
TEMPLATE_KEY_BY_NAME = {
    "narrative of changes": "narrative",
    "plan check comments": "planCheck",
    "plan check response letter": "planCheck",
}
TEMPLATE_DEFAULT_FILENAME_BY_KEY = {
    "narrative": "Narrative of Changes",
    "planCheck": "Plan Check Comments",
}

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

# --- Circuit Breaker AI helpers ---
CB_TEMPLATE_PATH = BASE_DIR / "CircuitBreakerAI" / "ElectricalPanels" / "Template.xlsx"
CB_RATE_LIMIT_MAX_REQUESTS = 2
CB_RATE_LIMIT_WINDOW = 61
_cb_api_call_timestamps = []

CB_COL_L_NOTE = "B"
CB_COL_L_TYPE = "C"
CB_COL_L_POLE = "D"
CB_COL_L_TRIP = "E"
CB_COL_L_DESC = "F"
CB_COL_L_KVA = "I"
CB_COL_VOLTAGE = "G"
CB_COL_R_KVA = "K"
CB_COL_R_DESC = "L"
CB_COL_R_POLE = "O"
CB_COL_R_TRIP = "P"
CB_COL_R_TYPE = "Q"
CB_COL_R_NOTE = "R"


class CircuitItem(BaseModel):
    circuit_number: int = Field(
        ..., description="The first circuit number of the breaker")
    description: str = Field(..., description="Load description")
    breaker_amps: str = Field(..., description="Amperage (e.g., '20')")
    poles: int = Field(1, description="Number of poles (1, 2, or 3)")
    load_type: str = Field(..., description="Code: 'C', 'D', 'G', 'K', 'M'")
    kva: str = Field(
        "",
        description="Visible load in kVA from an existing directory schedule, or blank if not shown",
    )
    phase_kva: List[str] = Field(
        [],
        description="Visible kVA values for each occupied circuit row, ordered top-to-bottom",
    )


class PanelData(BaseModel):
    panel_name: str = Field(..., description="Panel Name")
    voltage: str = Field(..., description="Voltage")
    bus_rating: str = Field(..., description="Bus Rating")
    aic_rating: str = Field(
        "",
        description="Panel AIC rating when visibly shown, or blank if not shown",
    )
    phase: str = Field(..., description="Phase")
    wire: str = Field(..., description="Wire")
    mounting: str = Field(..., description="Mounting")
    enclosure: str = Field(..., description="Enclosure")
    circuits: List[CircuitItem] = Field(
        ..., description="List of detected breakers")


def cb_enforce_rate_limit():
    global _cb_api_call_timestamps
    now = time.time()
    _cb_api_call_timestamps = [
        t for t in _cb_api_call_timestamps if now - t < CB_RATE_LIMIT_WINDOW
    ]
    if len(_cb_api_call_timestamps) >= CB_RATE_LIMIT_MAX_REQUESTS:
        earliest_call = _cb_api_call_timestamps[0]
        wait_time = (earliest_call + CB_RATE_LIMIT_WINDOW) - now
        if wait_time > 0:
            time.sleep(wait_time)
    _cb_api_call_timestamps.append(time.time())


def cb_clean_text(text):
    if text is None:
        return ""
    return str(text).upper().strip()


def cb_calculate_estimated_load(amps_str, description, load_type):
    desc = cb_clean_text(description)

    if any(x in desc for x in ["SPARE", "SPACE", "UNUSED"]):
        return 0.00

    try:
        numeric_amps = float("".join(filter(str.isdigit, str(amps_str))))
        breaker_kva = (numeric_amps * 120) / 1000
    except Exception:
        numeric_amps = 0
        breaker_kva = 0

    if any(x in desc for x in ["A/C", "AC", "AIR COND", "CONDENSER", "HVAC", "COOLING", "COMPRESSOR"]):
        return round(breaker_kva * 0.70, 2)

    if any(x in desc for x in ["WATER HEATER", "WH", "HWH", "HOT WATER"]):
        return round(breaker_kva * 0.60, 2)

    if any(x in desc for x in ["FAN", "EXHAUST", "VENTILAT", "VENT FAN"]):
        return round(breaker_kva * 0.50, 2)

    if "ATM" in desc:
        return 0.60

    if any(x in desc for x in ["POLE LIGHT", "POLE LT", "PARKING LOT", "LOT LIGHT"]):
        return random.choice([1.0, 1.1, 1.2])

    if any(x in desc for x in ["EXT LIGHT", "EXTERIOR LIGHT", "OUTSIDE LIGHT", "OUTDOOR LIGHT",
                                "EXT LT", "EXTERIOR LT", "SIGN", "FACADE", "BUILDING LIGHT"]):
        return random.choice([0.6, 0.7, 0.8, 0.9, 1.0])

    if any(x in desc for x in ["LIGHT", "LT", "LTG", "LIGHTING", "LAMP"]):
        large_room_keywords = ["LOBBY", "HALL", "CONFERENCE", "MEETING", "AUDITORIUM", "GYM",
                               "WAREHOUSE", "SHOWROOM", "DINING", "MAIN", "OPEN", "COMMON"]
        medium_room_keywords = ["OFFICE", "BREAK", "STORAGE", "KITCHEN", "LUNCH"]

        if any(kw in desc for kw in large_room_keywords):
            return random.choice([0.7, 0.8, 0.9])
        if any(kw in desc for kw in medium_room_keywords):
            return random.choice([0.5, 0.6, 0.7])
        return random.choice([0.3, 0.4, 0.5])

    if any(x in desc for x in ["PLUG", "RECEP", "DUPLEX", "OUTLET", "REC", "RCPT"]):
        high_recep_keywords = ["KITCHEN", "BREAK", "LAB", "WORKSTATION", "COMPUTER", "SERVER",
                               "OFFICE", "CONF", "MEETING", "NURSE", "MEDICAL"]
        medium_recep_keywords = ["STORAGE", "UTILITY", "MECH", "CLOSET", "REST", "BATH"]

        if any(kw in desc for kw in high_recep_keywords):
            return random.choice([0.72, 0.9])
        if any(kw in desc for kw in medium_recep_keywords):
            return random.choice([0.36, 0.54])
        return random.choice([0.36, 0.54, 0.72, 0.9])

    l_type = cb_clean_text(load_type)
    factor = 0.50 if l_type in ['C', 'G'] else 0.80
    return round((numeric_amps * factor * 120) / 1000, 2)


def cb_normalize_extracted_kva(raw_kva):
    if raw_kva is None or isinstance(raw_kva, bool):
        return ""
    if isinstance(raw_kva, (int, float)):
        value = float(raw_kva)
    else:
        text = str(raw_kva).strip()
        if not text:
            return ""
        upper = text.upper()
        if upper in {"N/A", "NA", "NONE", "UNKNOWN", "UNREADABLE", "BLANK", "-", "--", "---"}:
            return ""
        match = re.search(r"[-+]?\d+(?:\.\d+)?", upper.replace(",", ""))
        if not match:
            return ""
        try:
            value = float(match.group(0))
        except Exception:
            return ""
        if "VA" in upper and "KVA" not in upper and value >= 10:
            value = value / 1000
    if not math.isfinite(value) or value < 0:
        return ""
    return round(value, 3)


def cb_normalize_extracted_phase_kva(circuit):
    raw_values = getattr(circuit, "phase_kva", None)
    if isinstance(raw_values, (list, tuple)):
        values = [cb_normalize_extracted_kva(value) for value in raw_values]
        if any(value != "" for value in values):
            return values
    return [cb_normalize_extracted_kva(getattr(circuit, "kva", ""))]


def cb_fix_nema_type(raw_type):
    t = str(raw_type).upper()
    return "NEMA 3R" if any(x in t for x in ["3R", "OUT", "EXT", "WEATHER"]) else "NEMA 1"


def cb_resolve_ditto_marks(circuits: List[CircuitItem]) -> List[CircuitItem]:
    ditto_pattern = re.compile(r'^[\""\u2018\u2019\u201c\u201d\.]+$|^(SAME|DO)$', re.IGNORECASE)
    odds = sorted([c for c in circuits if c.circuit_number % 2 != 0], key=lambda x: x.circuit_number)
    evens = sorted([c for c in circuits if c.circuit_number % 2 == 0], key=lambda x: x.circuit_number)

    def process_column(col_list):
        last_desc = ""
        for ckt in col_list:
            if ditto_pattern.match(ckt.description.strip()):
                if last_desc:
                    ckt.description = last_desc
            elif ckt.description.strip() and ckt.description.strip() != "---":
                last_desc = ckt.description
        return col_list

    return process_column(odds) + process_column(evens)


def cb_make_unique_sheet_name(raw_name, existing_names):
    base = cb_clean_text(raw_name) or "PANEL"
    base = base.replace("/", "-").replace("\\", "-")
    base = base[:31] if base else "PANEL"
    if base not in existing_names:
        return base
    suffix = 2
    while True:
        suffix_str = f"-{suffix}"
        trimmed = base[: max(1, 31 - len(suffix_str))]
        candidate = f"{trimmed}{suffix_str}"
        if candidate not in existing_names:
            return candidate
        suffix += 1


def cb_ensure_template_sheet(wb):
    if "TEMPLATE" in wb.sheetnames:
        return
    if not CB_TEMPLATE_PATH.exists():
        raise ValueError("Panel schedule template not found.")
    template_wb = openpyxl.load_workbook(CB_TEMPLATE_PATH)
    try:
        source = template_wb["TEMPLATE"] if "TEMPLATE" in template_wb.sheetnames else template_wb.active
        target = wb.create_sheet("TEMPLATE")
        WorksheetCopy(source, target).copy_worksheet()
    finally:
        template_wb.close()


def cb_update_excel_workbook(panel_data: PanelData, workbook_path: str, use_extracted_kva=False) -> str:
    wb = openpyxl.load_workbook(workbook_path)
    cb_ensure_template_sheet(wb)

    source = wb["TEMPLATE"]
    target = wb.copy_worksheet(source)
    target.sheet_view.showGridLines = False

    safe_name = cb_make_unique_sheet_name(panel_data.panel_name, wb.sheetnames)
    target.title = safe_name
    target["A3"] = f"(E) PANEL '{safe_name}'"

    target[f"{CB_COL_VOLTAGE}2"] = cb_clean_text(panel_data.voltage)
    target["G3"] = cb_clean_text(panel_data.bus_rating)
    target["K2"] = cb_clean_text(panel_data.wire)
    target["K3"] = cb_clean_text(panel_data.phase)
    target["K4"] = cb_fix_nema_type(panel_data.enclosure)
    target["N2"] = cb_clean_text(panel_data.mounting)
    aic_rating = cb_clean_text(getattr(panel_data, "aic_rating", ""))
    if aic_rating:
        target["N3"] = aic_rating

    start_row = 8
    max_row = 28
    occupied_slots = set()

    circuits = cb_resolve_ditto_marks(panel_data.circuits or [])
    sorted_ckts = sorted(circuits, key=lambda x: x.circuit_number)

    for ckt in sorted_ckts:
        try:
            c_num = int(ckt.circuit_number)
        except Exception:
            continue
        if c_num <= 0 or c_num > 42:
            continue

        row_idx = start_row + math.ceil(c_num / 2) - 1
        if row_idx > max_row:
            continue

        is_odd = (c_num % 2 != 0)
        side = "L" if is_odd else "R"
        if (side, row_idx) in occupied_slots:
            continue

        if is_odd:
            c_note, c_type, c_pole, c_trip, c_desc, c_kva = (
                CB_COL_L_NOTE, CB_COL_L_TYPE, CB_COL_L_POLE, CB_COL_L_TRIP, CB_COL_L_DESC, CB_COL_L_KVA
            )
        else:
            c_note, c_type, c_pole, c_trip, c_desc, c_kva = (
                CB_COL_R_NOTE, CB_COL_R_TYPE, CB_COL_R_POLE, CB_COL_R_TRIP, CB_COL_R_DESC, CB_COL_R_KVA
            )

        desc_text = cb_clean_text(ckt.description)
        is_spare = any(x in desc_text for x in ["SPARE", "SPACE", "UNUSED"])

        if is_spare:
            kva_val = ""
            phase_kva_values = []
        elif use_extracted_kva:
            phase_kva_values = cb_normalize_extracted_phase_kva(ckt)
            kva_val = phase_kva_values[0] if phase_kva_values else ""
        else:
            kva_val = cb_calculate_estimated_load(
                ckt.breaker_amps, ckt.description, ckt.load_type
            )
            phase_kva_values = []
        note_val = "" if is_spare else "1"
        type_val = "" if is_spare else cb_clean_text(ckt.load_type)

        target[f"{c_desc}{row_idx}"] = desc_text
        target[f"{c_pole}{row_idx}"] = ckt.poles
        target[f"{c_trip}{row_idx}"] = cb_clean_text(ckt.breaker_amps)
        target[f"{c_kva}{row_idx}"] = kva_val
        target[f"{c_note}{row_idx}"] = note_val
        target[f"{c_type}{row_idx}"] = type_val

        occupied_slots.add((side, row_idx))

        if ckt.poles and ckt.poles > 1:
            for i in range(1, ckt.poles):
                ext_row = row_idx + i
                if ext_row <= max_row:
                    target[f"{c_desc}{ext_row}"] = "---"
                    target[f"{c_pole}{ext_row}"] = "-"
                    target[f"{c_trip}{ext_row}"] = "-"
                    if use_extracted_kva:
                        ext_kva_val = phase_kva_values[i] if i < len(phase_kva_values) else ""
                    else:
                        ext_kva_val = kva_val
                    target[f"{c_kva}{ext_row}"] = ext_kva_val
                    target[f"{c_type}{ext_row}"] = type_val
                    target[f"{c_note}{ext_row}"] = note_val
                    occupied_slots.add((side, ext_row))

    wb.save(workbook_path)
    return safe_name

# --- API Class ---


def _open_email_capture_db():
    conn = sqlite3.connect(EMAIL_CAPTURE_DB_FILE, timeout=5)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode = WAL;")
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS email_capture_meta (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL DEFAULT ''
        );
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS captured_emails (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entry_id TEXT NOT NULL DEFAULT '',
            internet_message_id TEXT NOT NULL DEFAULT '',
            conversation_id TEXT NOT NULL DEFAULT '',
            subject TEXT NOT NULL DEFAULT '',
            sender_name TEXT NOT NULL DEFAULT '',
            sender_address TEXT NOT NULL DEFAULT '',
            received_at_utc TEXT NOT NULL DEFAULT '',
            body_preview TEXT NOT NULL DEFAULT '',
            to_recipients_json TEXT NOT NULL DEFAULT '[]',
            cc_recipients_json TEXT NOT NULL DEFAULT '[]',
            directedness TEXT NOT NULL DEFAULT 'unknown',
            ai_status TEXT NOT NULL DEFAULT 'pending',
            ai_project_json TEXT NOT NULL DEFAULT '',
            ai_suggestion_json TEXT NOT NULL DEFAULT '',
            ai_skip_reason TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'new',
            status_changed_at_utc TEXT NOT NULL DEFAULT '',
            accepted_project_id TEXT NOT NULL DEFAULT '',
            accepted_deliverable_id TEXT NOT NULL DEFAULT '',
            first_seen_at_utc TEXT NOT NULL,
            last_scan_at_utc TEXT NOT NULL DEFAULT ''
        );
        """
    )
    conn.execute(
        """
        CREATE UNIQUE INDEX IF NOT EXISTS idx_captured_emails_imid
            ON captured_emails(internet_message_id) WHERE internet_message_id <> '';
        """
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_captured_emails_status ON captured_emails(status);"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_captured_emails_received ON captured_emails(received_at_utc);"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_captured_emails_conversation ON captured_emails(conversation_id);"
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS follow_ups (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            kind TEXT NOT NULL DEFAULT 'email',
            captured_email_id INTEGER,
            project_id TEXT NOT NULL DEFAULT '',
            deliverable_id TEXT NOT NULL DEFAULT '',
            title TEXT NOT NULL DEFAULT '',
            note TEXT NOT NULL DEFAULT '',
            waiting_on TEXT NOT NULL DEFAULT '',
            due_date TEXT NOT NULL DEFAULT '',
            snoozed_until TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'open',
            created_at_utc TEXT NOT NULL,
            resolved_at_utc TEXT NOT NULL DEFAULT '',
            FOREIGN KEY(captured_email_id) REFERENCES captured_emails(id) ON DELETE SET NULL
        );
        """
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_follow_ups_status ON follow_ups(status);"
    )
    row = conn.execute(
        "SELECT value FROM email_capture_meta WHERE key = 'schema_version'"
    ).fetchone()
    if row is None:
        conn.execute(
            "INSERT INTO email_capture_meta (key, value) VALUES ('schema_version', ?)",
            (str(EMAIL_CAPTURE_SCHEMA_VERSION),),
        )
    conn.commit()
    return conn


def _get_email_capture_meta(conn, key, default=""):
    row = conn.execute(
        "SELECT value FROM email_capture_meta WHERE key = ?",
        (str(key or ""),),
    ).fetchone()
    if row is None:
        return default
    return str(row["value"] or "")


def _set_email_capture_meta(conn, key, value):
    conn.execute(
        """
        INSERT INTO email_capture_meta (key, value) VALUES (?, ?)
        ON CONFLICT(key) DO UPDATE SET value = excluded.value
        """,
        (str(key or ""), str(value or "")),
    )


def _parse_email_capture_json(raw, fallback):
    try:
        parsed = json.loads(raw)
    except (TypeError, ValueError):
        return fallback
    if fallback is None:
        return parsed
    return parsed if isinstance(parsed, type(fallback)) else fallback


def _normalize_captured_email_record(row):
    if row is None:
        return None
    return {
        "id": int(row["id"]),
        "entryId": str(row["entry_id"] or "").strip(),
        "internetMessageId": str(row["internet_message_id"] or "").strip(),
        "conversationId": str(row["conversation_id"] or "").strip(),
        "subject": str(row["subject"] or "").strip(),
        "senderName": str(row["sender_name"] or "").strip(),
        "senderAddress": str(row["sender_address"] or "").strip(),
        "receivedAtUtc": str(row["received_at_utc"] or "").strip(),
        "bodyPreview": str(row["body_preview"] or ""),
        "toRecipients": _parse_email_capture_json(row["to_recipients_json"], []),
        "ccRecipients": _parse_email_capture_json(row["cc_recipients_json"], []),
        "directedness": str(row["directedness"] or "unknown"),
        "aiStatus": str(row["ai_status"] or "pending"),
        "aiProject": _parse_email_capture_json(row["ai_project_json"], None),
        "aiSuggestion": _parse_email_capture_json(row["ai_suggestion_json"], None),
        "aiSkipReason": str(row["ai_skip_reason"] or ""),
        "status": str(row["status"] or "new"),
        "statusChangedAtUtc": str(row["status_changed_at_utc"] or ""),
        "acceptedProjectId": str(row["accepted_project_id"] or ""),
        "acceptedDeliverableId": str(row["accepted_deliverable_id"] or ""),
        "firstSeenAtUtc": str(row["first_seen_at_utc"] or ""),
        "lastScanAtUtc": str(row["last_scan_at_utc"] or ""),
    }


def _normalize_follow_up_record(row):
    if row is None:
        return None
    record = {
        "id": int(row["id"]),
        "kind": str(row["kind"] or "email"),
        "capturedEmailId": row["captured_email_id"],
        "projectId": str(row["project_id"] or ""),
        "deliverableId": str(row["deliverable_id"] or ""),
        "title": str(row["title"] or "").strip(),
        "note": str(row["note"] or ""),
        "waitingOn": str(row["waiting_on"] or "").strip(),
        "dueDate": str(row["due_date"] or "").strip(),
        "snoozedUntil": str(row["snoozed_until"] or "").strip(),
        "status": str(row["status"] or "open"),
        "createdAtUtc": str(row["created_at_utc"] or ""),
        "resolvedAtUtc": str(row["resolved_at_utc"] or ""),
    }
    keys = row.keys() if hasattr(row, "keys") else []
    if "email_subject" in keys:
        record["emailSubject"] = str(row["email_subject"] or "").strip()
    if "email_entry_id" in keys:
        record["emailEntryId"] = str(row["email_entry_id"] or "").strip()
    if "email_sender_name" in keys:
        record["emailSenderName"] = str(row["email_sender_name"] or "").strip()
    return record


def _normalize_email_capture_recipient_list(raw_list):
    recipients = []
    if not isinstance(raw_list, list):
        return recipients
    for entry in raw_list:
        if isinstance(entry, dict):
            name = str(entry.get("name") or "").strip()
            address = str(entry.get("address") or "").strip()
        else:
            name = ""
            address = str(entry or "").strip()
        if name or address:
            recipients.append({"name": name, "address": address})
    return recipients


def _classify_email_directedness(summary, owner_identity):
    """Pure classification: 'to' | 'named' | 'cc' | 'unknown'.

    owner_identity = {"addresses": [...], "names": [...]}
    """
    summary = summary if isinstance(summary, dict) else {}
    identity = owner_identity if isinstance(owner_identity, dict) else {}
    addresses = {
        str(address or "").strip().lower()
        for address in identity.get("addresses") or []
        if str(address or "").strip()
    }
    names = [
        str(name or "").strip()
        for name in identity.get("names") or []
        if str(name or "").strip()
    ]
    names_lower = {name.lower() for name in names}

    def _recipient_matches(recipient):
        address = str(recipient.get("address") or "").strip().lower()
        name = str(recipient.get("name") or "").strip().lower()
        if address and address in addresses:
            return True
        if name and name in names_lower:
            return True
        return False

    to_recipients = _normalize_email_capture_recipient_list(summary.get("toRecipients"))
    cc_recipients = _normalize_email_capture_recipient_list(summary.get("ccRecipients"))

    if any(_recipient_matches(recipient) for recipient in to_recipients):
        return "to"

    body_head = str(summary.get("bodyPreview") or "")[:1200]
    if body_head:
        name_patterns = set()
        for name in names:
            if len(name) >= 3:
                name_patterns.add(name)
            first_token = name.split()[0] if name.split() else ""
            if len(first_token) >= 3:
                name_patterns.add(first_token)
        for pattern in name_patterns:
            if re.search(r"\b" + re.escape(pattern) + r"\b", body_head, re.IGNORECASE):
                return "named"

    if any(_recipient_matches(recipient) for recipient in cc_recipients):
        return "cc"
    return "unknown"


def _upsert_captured_email(conn, summary, directedness, now_iso):
    """Insert or refresh a captured email row. Returns 'inserted' or 'updated'.

    Dedupe: internet_message_id first, entry_id fallback. Never regresses
    status or ai_status on re-see; refreshes entry_id (EntryIDs change when
    messages move folders) and last_scan_at_utc.
    """
    summary = summary if isinstance(summary, dict) else {}
    entry_id = str(summary.get("id") or "").strip()
    internet_message_id = str(summary.get("internetMessageId") or "").strip()
    sender = summary.get("from") if isinstance(summary.get("from"), dict) else {}
    to_recipients = _normalize_email_capture_recipient_list(summary.get("toRecipients"))
    cc_recipients = _normalize_email_capture_recipient_list(summary.get("ccRecipients"))

    existing = None
    if internet_message_id:
        existing = conn.execute(
            "SELECT id FROM captured_emails WHERE internet_message_id = ?",
            (internet_message_id,),
        ).fetchone()
    if existing is None and entry_id:
        existing = conn.execute(
            "SELECT id FROM captured_emails WHERE entry_id = ? AND entry_id <> ''",
            (entry_id,),
        ).fetchone()

    if existing is not None:
        conn.execute(
            """
            UPDATE captured_emails
            SET entry_id = ?, last_scan_at_utc = ?
            WHERE id = ?
            """,
            (entry_id, str(now_iso or ""), int(existing["id"])),
        )
        return "updated"

    conn.execute(
        """
        INSERT INTO captured_emails (
            entry_id, internet_message_id, conversation_id, subject,
            sender_name, sender_address, received_at_utc, body_preview,
            to_recipients_json, cc_recipients_json, directedness,
            first_seen_at_utc, last_scan_at_utc
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            entry_id,
            internet_message_id,
            str(summary.get("conversationId") or "").strip(),
            str(summary.get("subject") or "").strip(),
            str(sender.get("name") or "").strip(),
            str(sender.get("address") or "").strip(),
            str(summary.get("receivedDateTime") or "").strip(),
            str(summary.get("bodyPreview") or ""),
            json.dumps(to_recipients, ensure_ascii=True),
            json.dumps(cc_recipients, ensure_ascii=True),
            str(directedness or "unknown"),
            str(now_iso or ""),
            str(now_iso or ""),
        ),
    )
    return "inserted"


def _captured_email_row_to_summary(record):
    """Rebuild an Outlook-scan message summary dict from a normalized record."""
    record = record if isinstance(record, dict) else {}
    return {
        "id": str(record.get("entryId") or ""),
        "subject": str(record.get("subject") or ""),
        "bodyPreview": str(record.get("bodyPreview") or ""),
        "receivedDateTime": str(record.get("receivedAtUtc") or ""),
        "webLink": "",
        "internetMessageId": str(record.get("internetMessageId") or ""),
        "conversationId": str(record.get("conversationId") or ""),
        "hasAttachments": False,
        "source": OUTLOOK_SCAN_SOURCE_DESKTOP,
        "from": {
            "name": str(record.get("senderName") or ""),
            "address": str(record.get("senderAddress") or ""),
        },
        "toRecipients": record.get("toRecipients") or [],
        "ccRecipients": record.get("ccRecipients") or [],
    }


def _follow_up_effective_due(record):
    record = record if isinstance(record, dict) else {}
    return str(record.get("snoozedUntil") or "").strip() or str(record.get("dueDate") or "").strip()


class Api:
    # --- FIX: Removed self.window from __init__ to prevent circular reference ---
    def __init__(self):
        # --- Configuration for Bundle Management ---
        self.github_repo = "jacobhusband/ElectricalCommands"
        # AutoCAD plugin releases are tracked in a separate repo; resolved dynamically.
        self.release_tag = None
        self.app_version = APP_VERSION
        self.app_update_repo = APP_UPDATE_REPO
        self.app_installer_name = APP_INSTALLER_NAME

        appdata_path = os.getenv('APPDATA')
        if not appdata_path:
            raise EnvironmentError(
                "CRITICAL: Could not determine the APPDATA directory.")
        self.app_plugins_folder = os.path.join(
            appdata_path, 'Autodesk', 'ApplicationPlugins')
        os.makedirs(self.app_plugins_folder, exist_ok=True)

        test_mode_raw = str(os.getenv('ACIES_TEST_MODE', '')).strip().lower()
        self.test_mode = test_mode_raw in ('1', 'true', 'yes', 'on')
        self._test_mode_records = []
        self._workroom_cad_file_cache = {}
        self._panel_schedule_running = False
        self._panel_schedule_thread = None
        self._panel_schedule_job_record = None
        self._panel_schedule_job_lock = threading.Lock()
        self._email_capture_scan_running = False
        self._email_capture_scan_lock = threading.Lock()
        if self.test_mode:
            logging.info(
                "Api initialized in ACIES_TEST_MODE. CAD tools run in deterministic validation mode.")

    # --- Application update helpers ---
    def _fetch_latest_release(self):
        """Fetch latest release metadata for this application.

        GitHub returns 404 for /releases/latest when no published release exists,
        so we fall back to the first item from /releases.
        """
        endpoints = [
            f"{GITHUB_API_BASE}/repos/{self.app_update_repo}/releases/latest",
            f"{GITHUB_API_BASE}/repos/{self.app_update_repo}/releases?per_page=1"
        ]

        data = None
        last_error = None

        for url in endpoints:
            try:
                response = requests.get(url, timeout=10)
                if response.status_code == 404:
                    continue
                response.raise_for_status()
                payload = response.json()
                data = payload[0] if isinstance(
                    payload, list) and payload else payload
                if data:
                    break
            except Exception as e:
                last_error = e

        if not data:
            raise Exception(
                "No published releases found for the updater repo. "
                "Publish a release on GitHub so the updater can find it."
            ) from last_error

        tag_name = data.get('tag_name') or ''
        latest_version = _normalize_version(tag_name)

        asset = None
        for a in data.get('assets', []):
            if a.get('name', '').lower() == self.app_installer_name.lower():
                asset = a
                break

        download_url = asset.get('browser_download_url') if asset else None

        return {
            'tag': tag_name,
            'latest_version': latest_version,
            'download_url': download_url,
            'release_notes': data.get('body') or '',
            'html_url': data.get('html_url') or ''
        }

    def get_app_update_status(self):
        """Check GitHub for a newer installer."""
        try:
            release = self._fetch_latest_release()
            latest_version = release['latest_version']
            download_url = release['download_url']
            update_available = bool(download_url) and _is_remote_newer(
                latest_version, self.app_version)
            return {
                'status': 'success',
                'current_version': self.app_version,
                'latest_version': latest_version,
                'update_available': update_available,
                'download_url': download_url,
                'release_notes': release['release_notes'],
                'release_page': release['html_url'],
            }
        except Exception as e:
            logging.error(f"Update check failed: {e}")
            return {
                'status': 'error',
                'message': str(e),
                'current_version': self.app_version
            }

    def _fetch_latest_bundle_release(self):
        """Fetch latest AutoCAD plugin release (or latest tag as fallback)."""
        endpoints = [
            f"{GITHUB_API_BASE}/repos/{self.github_repo}/releases/latest",
            f"{GITHUB_API_BASE}/repos/{self.github_repo}/releases?per_page=1"
        ]

        data = None
        last_error = None

        for url in endpoints:
            try:
                response = requests.get(url, timeout=10)
                if response.status_code == 404:
                    continue
                response.raise_for_status()
                payload = response.json()
                data = payload[0] if isinstance(
                    payload, list) and payload else payload
                if data:
                    break
            except Exception as e:
                last_error = e

        if data:
            return {
                'tag': data.get('tag_name') or '',
                'assets': data.get('assets', []),
                'release_notes': data.get('body') or '',
                'html_url': data.get('html_url') or ''
            }

        # Fallback: look at tags if no releases are published yet.
        try:
            tag_resp = requests.get(
                f"{GITHUB_API_BASE}/repos/{self.github_repo}/tags?per_page=1",
                timeout=10
            )
            tag_resp.raise_for_status()
            tags = tag_resp.json()
            if isinstance(tags, list) and tags:
                tag_name = tags[0].get('name') or ''
                release_data = {}
                try:
                    rel_resp = requests.get(
                        f"{GITHUB_API_BASE}/repos/{self.github_repo}/releases/tags/{tag_name}",
                        timeout=10
                    )
                    if rel_resp.status_code != 404:
                        rel_resp.raise_for_status()
                        release_data = rel_resp.json() or {}
                except Exception as inner:
                    logging.warning(
                        f"Could not fetch release data for tag {tag_name}: {inner}")

                return {
                    'tag': tag_name,
                    'assets': release_data.get('assets', []),
                    'release_notes': release_data.get('body', ''),
                    'html_url': release_data.get('html_url', '')
                }
        except Exception as e:
            last_error = e

        raise Exception(
            "No published AutoCAD plugin releases or tags found in the plugin repository."
        ) from last_error

    def download_and_install_app_update(self, download_url=None):
        """Download the newest installer and launch it silently."""
        if not download_url:
            return {'status': 'error', 'message': 'No download URL provided (publish a release asset named acies-scheduler-setup.exe).'}

        target = Path(tempfile.gettempdir()) / self.app_installer_name
        app_path = sys.executable if getattr(sys, "frozen", False) else None

        try:
            with requests.get(download_url, stream=True, timeout=60) as r:
                r.raise_for_status()
                with open(target, 'wb') as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)

            # Prepare the restart command separately to avoid f-string backslash syntax error
            restart_cmd = f'Start-Process -FilePath "{app_path}"' if app_path else ""

            # Use a small helper process to show a message, run installer (visible), and relaunch app after completion
            ps_script = (
                '$msg = "Updating ACIES Scheduler...' + "`n" +
                'Installer will run and the app will reopen automatically." ; '
                'Add-Type -AssemblyName PresentationFramework; '
                '[System.Windows.MessageBox]::Show($msg, "Updating", "OK", "Information") | Out-Null; '
                f'Start-Process -FilePath "{target}" '
                f'-ArgumentList \'/SILENT /NORESTART /CLOSEAPPLICATIONS /RESTARTAPPLICATIONS\' -Wait; '
                'Start-Sleep -Seconds 2; '
                f'{restart_cmd}'
            )
            subprocess.Popen(
                ["powershell", "-NoProfile", "-WindowStyle",
                    "Hidden", "-Command", ps_script],
                creationflags=subprocess.CREATE_NO_WINDOW
            )

            # Exit current app shortly to free files for installer
            if app_path:
                threading.Thread(target=lambda: (
                    time.sleep(1), os._exit(0)), daemon=True).start()

            return {
                'status': 'success',
                'message': 'Installer launched. Follow prompts to finish.',
                'installer_path': str(target)
            }
        except Exception as e:
            logging.error(f"Failed to download/install update: {e}")
            return {'status': 'error', 'message': str(e)}

    def check_bundle_installed(self, bundle_name):
        """
        Checks if a specific bundle directory exists in the ApplicationPlugins folder.
        Returns a simple boolean status.
        """
        if not bundle_name or not bundle_name.endswith('.bundle'):
            return {'status': 'error', 'message': 'Invalid bundle name provided.'}

        bundle_path = os.path.join(self.app_plugins_folder, bundle_name)
        is_installed = os.path.isdir(bundle_path)

        logging.info(
            f"Checking for bundle '{bundle_name}': {'Installed' if is_installed else 'Not Installed'}")

        return {'status': 'success', 'is_installed': is_installed}

    def _is_autocad_running(self):
        """Checks if the AutoCAD process (acad.exe) is currently running."""
        try:
            # Use tasklist to check for the process name.
            # This is standard on Windows and doesn't require installing psutil.
            output = subprocess.check_output(
                'tasklist /FI "IMAGENAME eq acad.exe" /FO CSV',
                shell=True, creationflags=subprocess.CREATE_NO_WINDOW
            ).decode()
            # If "acad.exe" is in the output, it's running.
            return "acad.exe" in output.lower()
        except Exception as e:
            logging.error(f"Error checking for AutoCAD process: {e}")
            return False  # Assume not running if check fails to avoid blocking user

    def get_bundle_statuses(self):
        """
        Fetches remote assets and compares against local bundles + version.txt.
        """
        logging.info("Fetching bundle statuses...")
        try:
            local_bundles = {
                name for name in os.listdir(self.app_plugins_folder)
                if str(name).endswith('.bundle')
            }

            release_tag = BUNDLE_RELEASE_TAG
            assets = []
            try:
                release_info = self._fetch_latest_bundle_release()
                self.release_tag = release_info.get('tag') or BUNDLE_RELEASE_TAG
                release_tag = self.release_tag
                assets = release_info.get('assets', []) or []
                if not assets and release_tag:
                    api_url = f"{GITHUB_API_BASE}/repos/{self.github_repo}/releases/tags/{release_tag}"
                    tag_response = requests.get(api_url, timeout=10)
                    if tag_response.status_code != 404:
                        tag_response.raise_for_status()
                        assets = tag_response.json().get('assets', [])
            except Exception as e:
                self.release_tag = release_tag
                logging.warning(
                    f"Could not refresh plugin release assets; falling back to known bundle catalog: {e}"
                )

            asset_by_bundle = {}
            for asset in assets:
                asset_name = asset.get('name')
                if not asset_name or 'Source code' in asset_name or not asset_name.endswith('.zip'):
                    continue
                bundle_name = _bundle_name_from_asset_name(asset_name, release_tag)
                if bundle_name:
                    asset_by_bundle[bundle_name] = asset

            bundle_names = list(KNOWN_PLUGIN_BUNDLES)
            for bundle_name in sorted(asset_by_bundle):
                if bundle_name not in bundle_names:
                    bundle_names.append(bundle_name)
            for bundle_name in sorted(local_bundles):
                if bundle_name not in bundle_names:
                    bundle_names.append(bundle_name)
            bundle_names = [
                bundle_name for bundle_name in bundle_names
                if bundle_name not in HIDDEN_PLUGIN_BUNDLES
            ]

            statuses = []
            for bundle_name in bundle_names:
                bundle_path = os.path.join(self.app_plugins_folder, bundle_name)
                is_installed = bundle_name in local_bundles
                local_version = None
                asset = asset_by_bundle.get(bundle_name)

                if is_installed:
                    version_file = os.path.join(bundle_path, 'version.txt')
                    if os.path.exists(version_file):
                        with open(version_file, 'r') as f:
                            local_version = f.read().strip()

                if asset:
                    if not is_installed:
                        state = 'not_installed'
                    elif local_version != release_tag:
                        state = 'update_available'
                    else:
                        state = 'installed'
                else:
                    state = 'installed' if is_installed else 'not_published'

                status = {
                    'name': bundle_name.replace('.bundle', ''),
                    'bundle_name': bundle_name,
                    'state': state,
                    'local_version': local_version or 'unknown',
                    'remote_version': release_tag,
                    'asset': asset,
                }
                statuses.append(status)

            return {'status': 'success', 'data': statuses}

        except Exception as e:
            logging.error(f"Error getting bundle statuses: {e}")
            return {'status': 'error', 'message': str(e)}

    def install_single_bundle(self, asset):
        """Downloads, extracts bundle, and writes version.txt."""

        # 1. Safety Check: AutoCAD locks files (DLLs). It must be closed to Update/Install.
        if self._is_autocad_running():
            return {
                'status': 'error',
                'message': 'AutoCAD is currently running. Please close AutoCAD and try again to prevent file locking errors.'
            }

        asset_name = asset.get('name')
        download_url = asset.get('browser_download_url')
        release_tag = self.release_tag or BUNDLE_RELEASE_TAG
        bundle_name = asset_name.replace(f"-{release_tag}.zip", ".bundle")
        if bundle_name == asset_name:
            bundle_name = asset_name.replace(".zip", ".bundle")
        extract_path = os.path.join(self.app_plugins_folder, bundle_name)

        logging.info(f"Installing bundle: {bundle_name}")
        try:
            response = requests.get(download_url, stream=True, timeout=30)
            response.raise_for_status()

            os.makedirs(extract_path, exist_ok=True)

            with zipfile.ZipFile(io.BytesIO(response.content)) as z:
                z.extractall(extract_path)

            # --- NEW: Write version.txt ---
            with open(os.path.join(extract_path, 'version.txt'), 'w') as f:
                f.write(release_tag)

            logging.info(
                f"Successfully installed {bundle_name} version {release_tag}")
            return {
                'status': 'success',
                'bundle_name': bundle_name,
                'bundlePath': extract_path,
                'pluginsFolderPath': self.app_plugins_folder,
            }
        except Exception as e:
            logging.error(f"Failed to install {bundle_name}: {e}")
            return {'status': 'error', 'message': f"Failed to install {bundle_name}: {e}"}

    def uninstall_bundle(self, bundle_name):
        """Removes a bundle directory."""

        # 1. Safety Check: AutoCAD locks files.
        if self._is_autocad_running():
            return {
                'status': 'error',
                'message': 'AutoCAD is currently running. Please close AutoCAD before uninstalling to ensure all files can be removed.'
            }

        if not bundle_name or not bundle_name.endswith('.bundle'):
            return {'status': 'error', 'message': 'Invalid bundle name.'}

        bundle_path = os.path.join(self.app_plugins_folder, bundle_name)

        try:
            if os.path.isdir(bundle_path):
                shutil.rmtree(bundle_path)
                return {
                    'status': 'success',
                    'bundle_name': bundle_name,
                    'pluginsFolderPath': self.app_plugins_folder,
                }
            else:
                return {'status': 'error', 'message': 'Bundle not found.'}
        except Exception as e:
            logging.error(f"Failed to uninstall {bundle_name}: {e}")
            return {'status': 'error', 'message': f"Failed to uninstall {bundle_name}: {e}"}

    def _run_script_with_progress(self, command, tool_id, activity_id=None, wait=False):
        """
        Runs a script in a separate thread, captures its stdout, and sends
        progress updates to the frontend.
        """
        result_holder = {'status': 'started'}

        def script_runner():
            try:
                print(f"DEBUG THREAD: Starting script for {tool_id}")
                startupinfo = None
                if sys.platform == "win32":
                    startupinfo = subprocess.STARTUPINFO()
                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                    startupinfo.wShowWindow = subprocess.SW_HIDE

                # FIX: Handle both string and list commands
                if isinstance(command, list):
                    process = subprocess.Popen(
                        command,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.STDOUT,
                        text=True,
                        encoding='utf-8',
                        errors='replace',
                        startupinfo=startupinfo
                    )
                    self._trace_cad_auto_select(
                        'script_subprocess_spawn',
                        tool_id=tool_id,
                        command_type='argv',
                        command=command,
                    )
                else:
                    process = subprocess.Popen(
                        command,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.STDOUT,
                        text=True,
                        encoding='utf-8',
                        errors='replace',
                        shell=True,
                        startupinfo=startupinfo
                    )
                    self._trace_cad_auto_select(
                        'script_subprocess_spawn',
                        tool_id=tool_id,
                        command_type='shell_string',
                        command=command,
                    )

                print(f"DEBUG THREAD: Process started, reading output...")
                last_progress_error = ""
                for line in iter(process.stdout.readline, ''):
                    line = line.strip()
                    print(f"DEBUG THREAD OUTPUT: {line}")
                    if line.startswith("PROGRESS:"):
                        message = line[len("PROGRESS:"):].strip()
                        if message.startswith("TRACE"):
                            self._trace_cad_auto_select(
                                'script_trace',
                                tool_id=tool_id,
                                message=message,
                            )
                            continue
                        if message.startswith("ERROR:"):
                            last_progress_error = message[len(
                                "ERROR:"):].strip() or message
                        self._notify_tool_status(
                            tool_id,
                            message,
                            activity_id=activity_id,
                        )

                process.stdout.close()
                return_code = process.wait()
                print(f"DEBUG THREAD: Process finished with return code {return_code}")
                self._trace_cad_auto_select(
                    'script_subprocess_exit',
                    tool_id=tool_id,
                    return_code=return_code,
                )

                if return_code == 0:
                    self._notify_tool_status(
                        tool_id,
                        "DONE",
                        activity_id=activity_id,
                    )
                    result_holder.update({
                        'status': 'success',
                        'returnCode': return_code,
                    })
                else:
                    if last_progress_error:
                        error_message = (
                            f"{last_progress_error} (exit code {return_code})."
                        )
                    else:
                        error_message = f"Script finished with error code {return_code}."
                    self._notify_tool_status(
                        tool_id,
                        f"ERROR: {error_message}",
                        activity_id=activity_id,
                    )
                    result_holder.update({
                        'status': 'error',
                        'message': error_message,
                        'returnCode': return_code,
                    })

            except Exception as e:
                print(f"DEBUG THREAD ERROR: {e}")
                import traceback
                traceback.print_exc()
                logging.error(f"Failed to execute script for {tool_id}: {e}")
                self._notify_tool_status(
                    tool_id,
                    f"ERROR: {str(e)}",
                    activity_id=activity_id,
                )
                result_holder.update({
                    'status': 'error',
                    'message': str(e),
                })

        thread = threading.Thread(target=script_runner)
        thread.start()
        if wait:
            thread.join()
            return dict(result_holder)
        return {'status': 'started'}

    def _build_powershell_script_command(self, script_path, *args):
        command = [
            'powershell.exe',
            '-ExecutionPolicy',
            'Bypass',
            '-File',
            str(script_path),
        ]
        for arg in args:
            if arg is None:
                continue
            command.append(str(arg))
        return command

    def get_user_settings(self):
        """Reads and returns user settings from settings.json."""
        settings = None
        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                settings = json.load(f)
        except FileNotFoundError:
            settings = None
        except json.JSONDecodeError:
            bak_path = SETTINGS_FILE + '.bak'
            try:
                with open(bak_path, 'r', encoding='utf-8') as f:
                    settings = json.load(f)
                logging.info(f"Recovered user settings from backup file: {bak_path}")
            except (FileNotFoundError, json.JSONDecodeError):
                settings = None
        if settings is None:
            settings = build_default_user_settings()

        settings, sanitized = _sanitize_user_settings_payload(settings)
        repaired_settings, repaired = _merge_missing_user_settings_from_legacy(
            settings,
            _load_legacy_user_settings(),
        )
        repaired_settings, repaired_sanitized = _sanitize_user_settings_payload(
            repaired_settings
        )
        if sanitized or repaired or repaired_sanitized:
            if repaired:
                logging.info("Recovered missing user settings from the legacy AppData settings file.")
            self.save_user_settings(repaired_settings)
            return repaired_settings
        return repaired_settings

    def save_user_settings(self, data):
        """Saves user settings to settings.json."""
        try:
            normalized_data, _ = _sanitize_user_settings_payload(data)
            if os.path.exists(SETTINGS_FILE):
                shutil.copy2(SETTINGS_FILE, SETTINGS_FILE + '.bak')
            with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
                json.dump(normalized_data, f, ensure_ascii=False, indent=2)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error saving user settings: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_cloud_sync_config(self):
        config = {
            "apiKey": str(os.getenv(FIREBASE_API_KEY_ENV) or "").strip(),
            "authDomain": str(os.getenv(FIREBASE_AUTH_DOMAIN_ENV) or "").strip(),
            "projectId": str(os.getenv(FIREBASE_PROJECT_ID_ENV) or "").strip(),
            "appId": str(os.getenv(FIREBASE_APP_ID_ENV) or "").strip(),
            "storageBucket": str(os.getenv(FIREBASE_STORAGE_BUCKET_ENV) or "").strip(),
            "messagingSenderId": str(os.getenv(FIREBASE_MESSAGING_SENDER_ID_ENV) or "").strip(),
        }
        enabled = all(
            config.get(key)
            for key in ("apiKey", "authDomain", "projectId", "appId")
        )
        return {
            "status": "success",
            "enabled": enabled,
            "config": config,
        }

    def get_local_sync_metadata(self):
        try:
            return {
                "status": "success",
                "files": _get_local_sync_metadata(),
            }
        except Exception as e:
            logging.error(f"Error loading local sync metadata: {e}")
            return {"status": "error", "message": str(e), "files": {}}

    def create_cloud_sync_backup(self, reason="", metadata=None):
        try:
            result = _create_cloud_sync_backup(reason=reason, metadata=metadata)
            return {
                "status": "success",
                **result,
            }
        except Exception as e:
            logging.error(f"Error creating cloud sync backup: {e}")
            return {"status": "error", "message": str(e)}

    def _get_google_oauth_client_id(self):
        return (os.getenv(GOOGLE_OAUTH_CLIENT_ID_ENV) or "").strip()

    def _get_google_oauth_client_secret(self):
        return (os.getenv(GOOGLE_OAUTH_CLIENT_SECRET_ENV) or "").strip()

    def _base64url_encode(self, raw_bytes):
        return base64.urlsafe_b64encode(raw_bytes).rstrip(b"=").decode("ascii")

    def _base64url_decode(self, raw_text):
        text = str(raw_text or "").strip()
        if not text:
            return b""
        padding = "=" * ((4 - (len(text) % 4)) % 4)
        return base64.urlsafe_b64decode(f"{text}{padding}")

    def _generate_google_pkce_pair(self):
        verifier = secrets.token_urlsafe(64)
        challenge = self._base64url_encode(
            hashlib.sha256(verifier.encode("ascii")).digest()
        )
        return verifier, challenge

    def _extract_google_error_message(self, response):
        try:
            payload = response.json() or {}
        except Exception:
            payload = {}
        message = str(
            payload.get("error_description")
            or payload.get("error")
            or response.text
            or "Google sign-in failed."
        ).strip()
        if "client_secret is missing" in message.lower():
            if not self._get_google_oauth_client_secret():
                return (
                    "Google sign-in may require a client secret for this OAuth client. "
                    f"Add {GOOGLE_OAUTH_CLIENT_SECRET_ENV} to .env and restart the app, "
                    "or verify the Google OAuth client settings."
                )
            return (
                "Google sign-in was rejected by Google OAuth. Verify that "
                f"{GOOGLE_OAUTH_CLIENT_ID_ENV} and {GOOGLE_OAUTH_CLIENT_SECRET_ENV} "
                "match the same Google OAuth client."
            )
        return message

    def _build_google_auth_record(self, token_payload, profile_payload, existing_auth=None):
        existing_auth = existing_auth if isinstance(existing_auth, dict) else {}
        issued_at = datetime.datetime.now(datetime.timezone.utc)
        expires_in = 0
        try:
            expires_in = int(token_payload.get("expires_in") or 0)
        except Exception:
            expires_in = 0
        expires_at = (
            issued_at + datetime.timedelta(seconds=max(expires_in, 0))
        ).replace(microsecond=0).isoformat().replace("+00:00", "Z")

        refresh_token = str(token_payload.get("refresh_token") or "").strip()
        if not refresh_token:
            refresh_token = str(existing_auth.get("refreshToken") or "").strip()
        id_token = str(token_payload.get("id_token") or "").strip()
        if not id_token:
            id_token = str(existing_auth.get("idToken") or "").strip()

        return {
            "provider": "google",
            "subject": str(profile_payload.get("sub") or existing_auth.get("subject") or "").strip(),
            "email": str(profile_payload.get("email") or existing_auth.get("email") or "").strip(),
            "displayName": str(profile_payload.get("name") or existing_auth.get("displayName") or "").strip(),
            "avatarUrl": str(profile_payload.get("picture") or existing_auth.get("avatarUrl") or "").strip(),
            "signedInAt": str(existing_auth.get("signedInAt") or utc_now_iso()),
            "tokenIssuedAt": utc_now_iso(),
            "expiresAt": expires_at,
            "idToken": id_token,
            "accessToken": str(token_payload.get("access_token") or "").strip(),
            "refreshToken": refresh_token,
            "tokenType": str(token_payload.get("token_type") or "Bearer").strip() or "Bearer",
            "scope": str(token_payload.get("scope") or " ".join(GOOGLE_OAUTH_SCOPES)).strip(),
        }

    def _sanitize_google_auth_record(self, auth_record):
        if not isinstance(auth_record, dict):
            return {
                "signedIn": False,
                "provider": "google",
                "email": "",
                "displayName": "",
                "avatarUrl": "",
                "signedInAt": "",
                "expiresAt": "",
                "hasRefreshToken": False,
            }
        return {
            "signedIn": True,
            "provider": str(auth_record.get("provider") or "google").strip() or "google",
            "email": str(auth_record.get("email") or "").strip(),
            "displayName": str(auth_record.get("displayName") or "").strip(),
            "avatarUrl": str(auth_record.get("avatarUrl") or "").strip(),
            "signedInAt": str(auth_record.get("signedInAt") or "").strip(),
            "expiresAt": str(auth_record.get("expiresAt") or "").strip(),
            "hasRefreshToken": bool(str(auth_record.get("refreshToken") or "").strip()),
        }

    def _build_google_sync_session(self, auth_record):
        if not isinstance(auth_record, dict):
            return {
                "signedIn": False,
                "idToken": "",
                "accessToken": "",
                "firebaseReady": False,
                "auth": self._sanitize_google_auth_record(None),
            }
        id_token = str(auth_record.get("idToken") or "").strip()
        access_token = str(auth_record.get("accessToken") or "").strip()
        return {
            "signedIn": True,
            "idToken": id_token,
            "accessToken": access_token,
            "firebaseReady": bool(id_token or access_token),
            "auth": self._sanitize_google_auth_record(auth_record),
        }

    def _load_google_auth_record(self):
        settings = self.get_user_settings()
        auth_record = settings.get("googleAuth")
        return auth_record if isinstance(auth_record, dict) else None

    def _persist_google_auth_record(self, auth_record):
        settings = self.get_user_settings()
        settings["googleAuth"] = auth_record if isinstance(auth_record, dict) and auth_record else None
        result = self.save_user_settings(settings)
        if result.get("status") != "success":
            raise RuntimeError(result.get("message") or "Failed to save Google sign-in state.")
        return settings

    def _build_google_oauth_status_html(self, title, body, accent="#2563eb"):
        safe_title = html.escape(str(title or "").strip() or "Google Sign-In")
        safe_body = html.escape(str(body or "").strip() or "Return to the app.")
        return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{safe_title}</title>
  <style>
    body {{
      margin: 0;
      min-height: 100vh;
      display: grid;
      place-items: center;
      padding: 24px;
      background: #0f172a;
      color: #e2e8f0;
      font-family: Inter, Segoe UI, sans-serif;
    }}
    .card {{
      width: min(440px, 100%);
      border-radius: 20px;
      padding: 28px;
      background: rgba(15, 23, 42, 0.92);
      border: 1px solid rgba(148, 163, 184, 0.22);
      box-shadow: 0 24px 60px rgba(2, 6, 23, 0.45);
    }}
    h1 {{
      margin: 0 0 12px;
      font-size: 1.45rem;
      color: {accent};
    }}
    p {{
      margin: 0;
      line-height: 1.6;
      color: #cbd5e1;
    }}
  </style>
</head>
<body>
  <div class="card">
    <h1>{safe_title}</h1>
    <p>{safe_body}</p>
  </div>
</body>
</html>"""

    def _start_google_oauth_callback_server(self):
        class GoogleOAuthCallbackServer(http.server.ThreadingHTTPServer):
            allow_reuse_address = True

        class GoogleOAuthCallbackHandler(http.server.BaseHTTPRequestHandler):
            def log_message(self, fmt, *args):
                return

            def do_GET(self):
                parsed = urlparse(self.path)
                if parsed.path != GOOGLE_OAUTH_CALLBACK_PATH:
                    self.send_response(404)
                    self.send_header("Content-Type", "text/plain; charset=utf-8")
                    self.end_headers()
                    self.wfile.write(b"Not found.")
                    return

                params = parse_qs(parsed.query or "", keep_blank_values=True)
                self.server.oauth_result = {
                    "code": (params.get("code") or [""])[0],
                    "state": (params.get("state") or [""])[0],
                    "error": (params.get("error") or [""])[0],
                    "error_description": (params.get("error_description") or [""])[0],
                }

                error_code = str(self.server.oauth_result.get("error") or "").strip().lower()
                if error_code:
                    title = "Google Sign-In Cancelled" if error_code == "access_denied" else "Google Sign-In Failed"
                    body = (
                        "You can close this window and return to the app."
                        if error_code == "access_denied"
                        else "There was a problem completing Google sign-in. Return to the app for details."
                    )
                    page = self.server.owner._build_google_oauth_status_html(
                        title,
                        body,
                        "#f59e0b" if error_code == "access_denied" else "#ef4444",
                    )
                else:
                    page = self.server.owner._build_google_oauth_status_html(
                        "Google Sign-In Complete",
                        "Return to the app. Your sign-in will finish automatically.",
                    )

                self.send_response(200)
                self.send_header("Content-Type", "text/html; charset=utf-8")
                self.send_header("Cache-Control", "no-store")
                self.end_headers()
                self.wfile.write(page.encode("utf-8"))
                self.server.oauth_event.set()
                threading.Thread(target=self.server.shutdown, daemon=True).start()

        server = GoogleOAuthCallbackServer(("127.0.0.1", 0), GoogleOAuthCallbackHandler)
        server.owner = self
        server.oauth_event = threading.Event()
        server.oauth_result = {}
        server_thread = threading.Thread(target=server.serve_forever, daemon=True)
        server_thread.start()
        return server

    def _exchange_google_auth_code(self, client_id, code, code_verifier, redirect_uri):
        data = {
            "client_id": client_id,
            "code": code,
            "code_verifier": code_verifier,
            "grant_type": "authorization_code",
            "redirect_uri": redirect_uri,
        }
        client_secret = self._get_google_oauth_client_secret()
        if client_secret:
            data["client_secret"] = client_secret
        response = requests.post(
            GOOGLE_OAUTH_TOKEN_ENDPOINT,
            data=data,
            timeout=20,
        )
        if response.status_code != 200:
            raise RuntimeError(self._extract_google_error_message(response))
        payload = response.json() or {}
        if not str(payload.get("access_token") or "").strip():
            raise RuntimeError("Google sign-in completed without an access token.")
        return payload

    def _refresh_google_auth_record_if_needed(self, auth_record):
        if not isinstance(auth_record, dict):
            return None

        refresh_token = str(auth_record.get("refreshToken") or "").strip()
        client_id = self._get_google_oauth_client_id()
        expires_at = parse_utc_iso(auth_record.get("expiresAt"))
        if not refresh_token or not client_id or not expires_at:
            return auth_record

        now = datetime.datetime.now(datetime.timezone.utc)
        if expires_at > now + datetime.timedelta(minutes=5):
            return auth_record

        try:
            data = {
                "client_id": client_id,
                "grant_type": "refresh_token",
                "refresh_token": refresh_token,
            }
            client_secret = self._get_google_oauth_client_secret()
            if client_secret:
                data["client_secret"] = client_secret
            response = requests.post(
                GOOGLE_OAUTH_TOKEN_ENDPOINT,
                data=data,
                timeout=20,
            )
            if response.status_code != 200:
                message = self._extract_google_error_message(response)
                try:
                    payload = response.json() or {}
                except Exception:
                    payload = {}
                if str(payload.get("error") or "").strip().lower() == "invalid_grant":
                    self._persist_google_auth_record(None)
                    return None
                logging.warning(f"Could not refresh Google token: {message}")
                return auth_record

            token_payload = response.json() or {}
            refreshed_auth = self._build_google_auth_record(
                token_payload,
                {
                    "sub": auth_record.get("subject"),
                    "email": auth_record.get("email"),
                    "name": auth_record.get("displayName"),
                    "picture": auth_record.get("avatarUrl"),
                },
                existing_auth=auth_record,
            )
            refreshed_auth["signedInAt"] = str(auth_record.get("signedInAt") or utc_now_iso())
            self._persist_google_auth_record(refreshed_auth)
            return refreshed_auth
        except Exception as exc:
            logging.warning(f"Could not refresh Google auth state: {exc}")
            return auth_record

    def _fetch_google_user_profile(self, access_token):
        response = requests.get(
            GOOGLE_OAUTH_USERINFO_ENDPOINT,
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=20,
        )
        if response.status_code != 200:
            raise RuntimeError(self._extract_google_error_message(response))
        payload = response.json() or {}
        if not str(payload.get("email") or "").strip():
            raise RuntimeError("Google sign-in completed without an email address.")
        return payload

    def get_google_auth_state(self):
        try:
            auth_record = self._refresh_google_auth_record_if_needed(
                self._load_google_auth_record()
            )
            return {
                "status": "success",
                "auth": self._sanitize_google_auth_record(auth_record),
            }
        except Exception as e:
            logging.error(f"Error loading Google auth state: {e}")
            return {
                "status": "error",
                "message": str(e),
                "auth": self._sanitize_google_auth_record(None),
            }

    def get_google_sync_session(self):
        try:
            auth_record = self._refresh_google_auth_record_if_needed(
                self._load_google_auth_record()
            )
            return {
                "status": "success",
                **self._build_google_sync_session(auth_record),
            }
        except Exception as e:
            logging.error(f"Error loading Google sync session: {e}")
            return {
                "status": "error",
                "message": str(e),
                **self._build_google_sync_session(None),
            }

    def sign_in_with_google(self):
        client_id = self._get_google_oauth_client_id()
        if not client_id:
            return {
                "status": "error",
                "message": (
                    "Google sign-in is not configured. Add "
                    f"{GOOGLE_OAUTH_CLIENT_ID_ENV} to .env and restart the app."
                ),
            }

        existing_auth = self._load_google_auth_record()
        callback_server = self._start_google_oauth_callback_server()
        callback_port = callback_server.server_address[1]
        redirect_uri = f"http://127.0.0.1:{callback_port}{GOOGLE_OAUTH_CALLBACK_PATH}"
        code_verifier, code_challenge = self._generate_google_pkce_pair()
        state = secrets.token_urlsafe(32)
        auth_url = (
            f"{GOOGLE_OAUTH_AUTH_ENDPOINT}?"
            + urlencode(
                {
                    "client_id": client_id,
                    "redirect_uri": redirect_uri,
                    "response_type": "code",
                    "scope": " ".join(GOOGLE_OAUTH_SCOPES),
                    "state": state,
                    "access_type": "offline",
                    "prompt": "consent select_account",
                    "code_challenge": code_challenge,
                    "code_challenge_method": "S256",
                }
            )
        )

        try:
            open_result = self.open_url(auth_url)
            if open_result.get("status") != "success":
                raise RuntimeError(open_result.get("message") or "Could not open the browser for Google sign-in.")

            if not callback_server.oauth_event.wait(GOOGLE_OAUTH_TIMEOUT_SECONDS):
                return {
                    "status": "error",
                    "message": "Google sign-in timed out. Please try again.",
                }

            result = callback_server.oauth_result or {}
            error_code = str(result.get("error") or "").strip().lower()
            if error_code:
                if error_code == "access_denied":
                    return {
                        "status": "cancelled",
                        "message": "Google sign-in was cancelled.",
                    }
                error_detail = str(result.get("error_description") or result.get("error") or "").strip()
                return {
                    "status": "error",
                    "message": error_detail or "Google sign-in did not complete.",
                }

            if str(result.get("state") or "").strip() != state:
                return {
                    "status": "error",
                    "message": "Google sign-in returned an invalid state token.",
                }

            code = str(result.get("code") or "").strip()
            if not code:
                return {
                    "status": "error",
                    "message": "Google sign-in finished without an authorization code.",
                }

            token_payload = self._exchange_google_auth_code(
                client_id,
                code,
                code_verifier,
                redirect_uri,
            )
            profile_payload = self._fetch_google_user_profile(
                str(token_payload.get("access_token") or "").strip()
            )
            auth_record = self._build_google_auth_record(
                token_payload,
                profile_payload,
                existing_auth=existing_auth,
            )
            self._persist_google_auth_record(auth_record)
            return {
                "status": "success",
                "auth": self._sanitize_google_auth_record(auth_record),
                "syncSession": self._build_google_sync_session(auth_record),
            }
        except Exception as e:
            logging.error(f"Error signing in with Google: {e}")
            return {'status': 'error', 'message': str(e)}
        finally:
            try:
                callback_server.shutdown()
            except Exception:
                pass
            try:
                callback_server.server_close()
            except Exception:
                pass

    def sign_out_google(self):
        try:
            auth_record = self._load_google_auth_record()
            revoke_token = ""
            if isinstance(auth_record, dict):
                revoke_token = str(
                    auth_record.get("refreshToken") or auth_record.get("accessToken") or ""
                ).strip()

            if revoke_token:
                try:
                    requests.post(
                        GOOGLE_OAUTH_REVOKE_ENDPOINT,
                        data={"token": revoke_token},
                        headers={"Content-Type": "application/x-www-form-urlencoded"},
                        timeout=10,
                    )
                except Exception as revoke_exc:
                    logging.warning(f"Google token revoke failed during sign-out: {revoke_exc}")

            self._persist_google_auth_record(None)
            return {
                "status": "success",
                "auth": self._sanitize_google_auth_record(None),
            }
        except Exception as e:
            logging.error(f"Error signing out of Google: {e}")
            return {'status': 'error', 'message': str(e)}

    def _normalize_outlook_scan_date(self, scan_date):
        raw = str(scan_date or "").strip()
        if not raw:
            return ""
        try:
            selected = datetime.date.fromisoformat(raw)
        except ValueError:
            return ""
        today = datetime.datetime.now().astimezone().date()
        if selected > today:
            selected = today
        return selected.isoformat()

    def _build_outlook_scan_range(self, timeframe="week", scan_date=""):
        now = datetime.datetime.now().astimezone()
        tzinfo = now.tzinfo
        normalized_scan_date = self._normalize_outlook_scan_date(scan_date)
        if normalized_scan_date:
            selected = datetime.date.fromisoformat(normalized_scan_date)
            start_local = datetime.datetime.combine(
                selected,
                datetime.time.min,
                tzinfo=tzinfo,
            )
            end_local = datetime.datetime.combine(
                selected,
                datetime.time.max,
                tzinfo=tzinfo,
            )
            return {
                "scanDate": normalized_scan_date,
                "timeframe": "",
                "startUtc": start_local.astimezone(datetime.timezone.utc),
                "endUtc": end_local.astimezone(datetime.timezone.utc),
            }

        today = now.date()
        normalized_timeframe = str(timeframe or "week").strip().lower()
        if normalized_timeframe not in {"week", "month"}:
            normalized_timeframe = "week"
        if normalized_timeframe == "month":
            start_local = datetime.datetime.combine(
                today.replace(day=1),
                datetime.time.min,
                tzinfo=tzinfo,
            )
        else:
            weekday = today.weekday()
            days_since_sunday = (weekday + 1) % 7
            start_date = today - datetime.timedelta(days=days_since_sunday)
            start_local = datetime.datetime.combine(
                start_date,
                datetime.time.min,
                tzinfo=tzinfo,
            )
        end_local = datetime.datetime.combine(
            today,
            datetime.time.max,
            tzinfo=tzinfo,
        )
        return {
            "scanDate": "",
            "timeframe": normalized_timeframe,
            "startUtc": start_local.astimezone(datetime.timezone.utc),
            "endUtc": end_local.astimezone(datetime.timezone.utc),
        }

    def _format_outlook_scan_period_label(self, timeframe="week", scan_date=""):
        normalized_scan_date = self._normalize_outlook_scan_date(scan_date)
        if normalized_scan_date:
            try:
                selected = datetime.date.fromisoformat(normalized_scan_date)
                return selected.strftime("%m/%d/%Y")
            except ValueError:
                pass
        return "this month" if str(timeframe or "").strip().lower() == "month" else "this week"

    def _build_outlook_scan_error_result(self, message, timeframe="week", source="", scan_date=""):
        return {
            "status": "error",
            "message": str(message or "").strip() or "Outlook scan failed.",
            "candidateCount": 0,
            "scannedCount": 0,
            "emailsIncludedCount": 0,
            "deliverablesIncludedCount": 0,
            "threadsDetected": 0,
            "dedupedEmailCount": 0,
            "dedupeSkippedEmailCount": 0,
            "analysisMode": "batch",
            "suggestions": [],
            "skippedMessages": [],
            "promptTruncated": False,
            "truncated": False,
            "timeframe": timeframe,
            "scanDate": self._normalize_outlook_scan_date(scan_date),
            "source": str(source or "").strip(),
        }

    def _count_outlook_scan_project_context_deliverables(self, project_context):
        total = 0
        if not isinstance(project_context, list):
            return total
        for project in project_context:
            if not isinstance(project, dict):
                continue
            deliverables = project.get("deliverables")
            if isinstance(deliverables, list):
                total += len(deliverables)
        return total

    def _notify_outlook_scan_progress(self, payload):
        try:
            if not webview.windows:
                return
            payload_json = json.dumps(payload or {}, ensure_ascii=True)
            webview.windows[0].evaluate_js(
                f"window.updateOutlookScanProgress({payload_json})"
            )
        except Exception as exc:
            logging.debug(f"_notify_outlook_scan_progress failed: {exc}")

    def _send_outlook_scan_progress(self, stage, message="", **kwargs):
        payload = {
            "stage": str(stage or "").strip(),
            "message": str(message or "").strip(),
        }
        allowed_keys = {
            "active",
            "source",
            "timeframe",
            "scanDate",
            "totalEmails",
            "processedEmails",
            "includedEmails",
            "skippedEmails",
            "deliverablesInPeriod",
            "relevantEmails",
            "threadsDetected",
            "dedupedEmailCount",
            "dedupeSkippedEmailCount",
        }
        for key, value in (kwargs or {}).items():
            if key not in allowed_keys or value is None:
                continue
            if key in {
                "totalEmails",
                "processedEmails",
                "includedEmails",
                "skippedEmails",
                "deliverablesInPeriod",
                "relevantEmails",
                "threadsDetected",
                "dedupedEmailCount",
                "dedupeSkippedEmailCount",
            }:
                try:
                    payload[key] = max(int(value), 0)
                except Exception:
                    continue
            else:
                payload[key] = value
        if "active" not in payload:
            payload["active"] = payload["stage"] not in {"done", "error"}
        self._notify_outlook_scan_progress(payload)

    def _get_outlook_scan_capability(self):
        desktop_available, desktop_reason = self._get_desktop_outlook_availability()
        return {
            "desktopAvailable": desktop_available,
            "desktopReason": desktop_reason,
        }

    def get_outlook_scan_capability(self):
        try:
            capability = self._get_outlook_scan_capability()
            return {
                "status": "success",
                "desktopAvailable": capability.get("desktopAvailable", False),
                "desktopReason": capability.get("desktopReason", ""),
            }
        except Exception as e:
            logging.error(f"Error loading Outlook scan capability: {e}")
            return {
                "status": "error",
                "message": str(e),
                "desktopAvailable": False,
                "desktopReason": str(e),
            }

    def _get_desktop_outlook_availability(self):
        if sys.platform != "win32":
            return False, "Desktop Outlook inbox scan is only available on Windows."
        install_path = self._find_desktop_outlook_install_path()
        if not install_path:
            return False, "Classic Outlook is not installed on this machine."
        try:
            import win32com.client  # noqa: F401
        except Exception:
            return False, "Desktop Outlook support requires pywin32."
        return True, ""

    def _find_desktop_outlook_install_path(self):
        if sys.platform != "win32":
            return ""

        candidate_paths = []
        try:
            import winreg

            registry_locations = [
                (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\OUTLOOK.EXE"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\OUTLOOK.EXE"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\App Paths\OUTLOOK.EXE"),
            ]
            for hive, subkey in registry_locations:
                try:
                    with winreg.OpenKey(hive, subkey) as key:
                        value = ""
                        try:
                            value = winreg.QueryValue(key, None)
                        except Exception:
                            try:
                                value = winreg.QueryValueEx(key, "")[0]
                            except Exception:
                                value = ""
                        value = os.path.abspath(str(value or "").strip()) if value else ""
                        if value:
                            candidate_paths.append(value)
                except Exception:
                    continue
        except Exception:
            pass

        program_files_roots = [
            os.environ.get("ProgramFiles", r"C:\Program Files"),
            os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)"),
        ]
        office_variants = [
            ("Microsoft Office", "root", "Office16"),
            ("Microsoft Office", "root", "Office15"),
            ("Microsoft Office", "Office16"),
            ("Microsoft Office", "Office15"),
            ("Microsoft Office", "Office14"),
        ]
        for root in program_files_roots:
            root = str(root or "").strip()
            if not root:
                continue
            for variant in office_variants:
                candidate_paths.append(os.path.join(root, *variant, "OUTLOOK.EXE"))

        for candidate in candidate_paths:
            full_path = os.path.abspath(str(candidate or "").strip()) if candidate else ""
            if full_path and os.path.exists(full_path):
                return full_path
        return ""

    def _run_with_outlook_com(self, func):
        pythoncom = None
        try:
            import pythoncom as _pythoncom

            pythoncom = _pythoncom
            pythoncom.CoInitialize()
        except Exception:
            pythoncom = None

        try:
            return func()
        finally:
            if pythoncom is not None:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    def _get_desktop_outlook_namespace(self):
        available, reason = self._get_desktop_outlook_availability()
        if not available:
            raise RuntimeError(reason or "Desktop Outlook is unavailable.")
        try:
            import win32com.client
        except Exception as exc:
            raise RuntimeError("Desktop Outlook support requires pywin32.") from exc

        try:
            application = win32com.client.Dispatch("Outlook.Application")
            namespace = application.GetNamespace("MAPI")
            return application, namespace
        except Exception as exc:
            raise RuntimeError(f"Could not access Desktop Outlook: {exc}") from exc

    def _coerce_outlook_datetime(self, raw_value):
        if isinstance(raw_value, datetime.datetime):
            dt_value = raw_value
        else:
            try:
                dt_value = datetime.datetime.fromtimestamp(float(raw_value))
            except Exception:
                return None
        if dt_value.tzinfo is None:
            dt_value = dt_value.replace(tzinfo=datetime.datetime.now().astimezone().tzinfo)
        try:
            return dt_value.astimezone(datetime.timezone.utc)
        except Exception:
            return None

    def _format_outlook_datetime(self, raw_value):
        dt_value = self._coerce_outlook_datetime(raw_value)
        if dt_value is None:
            return ""
        return dt_value.isoformat().replace("+00:00", "Z")

    def _get_desktop_outlook_internet_message_id(self, mail_item):
        accessor = getattr(mail_item, "PropertyAccessor", None)
        if accessor is None:
            return ""
        try:
            return str(accessor.GetProperty(OUTLOOK_MAPI_PR_INTERNET_MESSAGE_ID) or "").strip()
        except Exception:
            return ""

    def _extract_desktop_outlook_recipients(self, mail_item):
        to_recipients = []
        cc_recipients = []
        try:
            recipients = getattr(mail_item, "Recipients", None)
            count = int(getattr(recipients, "Count", 0) or 0)
        except Exception:
            return to_recipients, cc_recipients

        for index in range(1, count + 1):
            try:
                recipient = recipients.Item(index)
            except Exception:
                continue
            try:
                recipient_type = int(getattr(recipient, "Type", 1) or 1)
            except Exception:
                recipient_type = 1
            name = str(getattr(recipient, "Name", "") or "").strip()
            address = ""
            try:
                accessor = getattr(recipient, "PropertyAccessor", None)
                if accessor is not None:
                    address = str(
                        accessor.GetProperty(OUTLOOK_MAPI_PR_SMTP_ADDRESS) or ""
                    ).strip()
            except Exception:
                address = ""
            if not address:
                try:
                    address = str(getattr(recipient, "Address", "") or "").strip()
                except Exception:
                    address = ""
            if not name and not address:
                continue
            entry = {"name": name, "address": address}
            if recipient_type == 2:
                cc_recipients.append(entry)
            else:
                to_recipients.append(entry)
        return to_recipients, cc_recipients

    def _build_desktop_outlook_message_summary(self, mail_item):
        body_text = str(getattr(mail_item, "Body", "") or "").strip()
        preview = re.sub(r"\s+", " ", body_text).strip()
        if len(preview) > 500:
            preview = preview[:497].rstrip() + "..."
        to_recipients, cc_recipients = self._extract_desktop_outlook_recipients(mail_item)
        return {
            "id": str(getattr(mail_item, "EntryID", "") or "").strip(),
            "subject": str(getattr(mail_item, "Subject", "") or "").strip(),
            "bodyPreview": preview,
            "receivedDateTime": self._format_outlook_datetime(getattr(mail_item, "ReceivedTime", None)),
            "webLink": "",
            "internetMessageId": self._get_desktop_outlook_internet_message_id(mail_item),
            "conversationId": str(getattr(mail_item, "ConversationID", "") or "").strip(),
            "hasAttachments": bool(int(getattr(getattr(mail_item, "Attachments", None), "Count", 0) or 0)),
            "source": OUTLOOK_SCAN_SOURCE_DESKTOP,
            "from": {
                "name": str(getattr(mail_item, "SenderName", "") or "").strip(),
                "address": str(getattr(mail_item, "SenderEmailAddress", "") or "").strip(),
            },
            "toRecipients": to_recipients,
            "ccRecipients": cc_recipients,
        }

    def _list_desktop_outlook_inbox_messages(
        self,
        timeframe="week",
        limit=OUTLOOK_SCAN_FETCH_LIMIT,
        scan_date="",
        start_utc=None,
        end_utc=None,
    ):
        if start_utc is None or end_utc is None:
            scan_range = self._build_outlook_scan_range(timeframe=timeframe, scan_date=scan_date)
            start_utc = scan_range["startUtc"]
            end_utc = scan_range["endUtc"]
        max_items = max(int(limit or 0), 1)

        def _read_messages():
            _, namespace = self._get_desktop_outlook_namespace()
            inbox = namespace.GetDefaultFolder(OUTLOOK_MAPI_INBOX_FOLDER_ID)
            items = inbox.Items
            try:
                items.Sort("[ReceivedTime]", True)
            except Exception:
                pass

            messages = []
            has_more = False
            current = None
            try:
                current = items.GetFirst()
            except Exception:
                current = None

            while current is not None:
                message_class = str(getattr(current, "MessageClass", "") or "").strip().lower()
                if message_class and not message_class.startswith("ipm.note"):
                    try:
                        current = items.GetNext()
                    except Exception:
                        current = None
                    continue

                received_dt = self._coerce_outlook_datetime(getattr(current, "ReceivedTime", None))
                if received_dt is not None and end_utc is not None and received_dt > end_utc:
                    try:
                        current = items.GetNext()
                    except Exception:
                        current = None
                    continue
                if received_dt is not None and received_dt < start_utc:
                    break

                summary = self._build_desktop_outlook_message_summary(current)
                if summary.get("id"):
                    messages.append(summary)
                    if len(messages) >= max_items:
                        has_more = True
                        break

                try:
                    current = items.GetNext()
                except Exception:
                    current = None

            return messages, has_more

        return self._run_with_outlook_com(_read_messages)

    def _get_desktop_outlook_message_body_text(self, message_id):
        target_id = str(message_id or "").strip()
        if not target_id:
            raise RuntimeError("Desktop Outlook message is missing an identifier.")

        def _read_message():
            _, namespace = self._get_desktop_outlook_namespace()
            mail_item = namespace.GetItemFromID(target_id)
            summary = self._build_desktop_outlook_message_summary(mail_item)
            body_text = str(getattr(mail_item, "Body", "") or "").strip()
            if not body_text:
                html_body = str(getattr(mail_item, "HTMLBody", "") or "").strip()
                if html_body:
                    body_text = self._html_to_plain_text(html_body)
            if not body_text:
                body_text = str(summary.get("bodyPreview") or "").strip()
            return summary, body_text

        return self._run_with_outlook_com(_read_message)

    def open_outlook_desktop_message(self, message_id):
        payload = message_id
        if isinstance(payload, str):
            payload = {"messageId": payload}
        if not isinstance(payload, dict):
            payload = {}
        target_id = str(payload.get("messageId") or payload.get("id") or "").strip()
        if not target_id:
            return {"status": "error", "message": "Missing Outlook desktop message identifier."}

        try:
            def _open_message():
                _, namespace = self._get_desktop_outlook_namespace()
                mail_item = namespace.GetItemFromID(target_id)
                mail_item.Display()

            self._run_with_outlook_com(_open_message)
            return {"status": "success"}
        except Exception as exc:
            logging.error(f"Error opening Desktop Outlook message: {exc}")
            return {"status": "error", "message": str(exc)}

    def _outlook_message_candidate_score(self, message_summary):
        summary = message_summary if isinstance(message_summary, dict) else {}
        subject = str(summary.get("subject") or "")
        preview = str(summary.get("bodyPreview") or "")
        sender = str((summary.get("from") or {}).get("address") or "")
        combined = f"{subject}\n{preview}".lower()
        if not combined.strip():
            return 0

        score = 1
        if re.search(r"\b\d{5,}\b", combined):
            score += 2
        if re.search(r"[A-Z]:\\\\", f"{subject}\n{preview}"):
            score += 2
        if re.search(
            r"\b(rfi|submittal|record drawings|record set|coordination|meeting|bulletin|revision|dd\d+|cd\d+|ifp|ifc|survey|due|deliverable)\b",
            combined,
        ):
            score += 2
        if re.search(r"\b(review|comments|deadline|schedule|submit|issuance|issue)\b", combined):
            score += 1
        if summary.get("hasAttachments"):
            score += 1
        if "noreply" in sender.lower() and score < 3:
            return 0
        return score

    def _html_to_plain_text(self, value):
        raw = str(value or "")
        if not raw:
            return ""
        stripped = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
        stripped = re.sub(r"(?i)<br\s*/?>", "\n", stripped)
        stripped = re.sub(r"(?i)</(p|div|li|tr|h[1-6])>", "\n", stripped)
        stripped = re.sub(r"(?s)<[^>]+>", " ", stripped)
        text = html.unescape(stripped)
        text = re.sub(r"[ \t\f\v]+", " ", text)
        text = re.sub(r"\r?\n\s*", "\n", text)
        return text.strip()

    def _trim_outlook_scan_prompt_text(self, value, limit=OUTLOOK_SCAN_EMAIL_BODY_PROMPT_CHARS):
        text = str(value or "").strip()
        if not text:
            return ""
        text = re.sub(r"\r\n?", "\n", text)
        text = re.sub(r"[ \t\f\v]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        max_chars = max(int(limit or 0), 0)
        if not max_chars or len(text) <= max_chars:
            return text
        if max_chars <= 3:
            return text[:max_chars]
        return text[: max_chars - 3].rstrip() + "..."

    def _normalize_outlook_scan_dedupe_text(self, value):
        text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\xa0", " ")
        text = re.sub(r"[ \t\f\v]+", " ", text)
        text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _looks_like_outlook_reply_header_block(self, lines, start_idx):
        labels = []
        for idx in range(start_idx, min(start_idx + 6, len(lines))):
            raw_line = str(lines[idx] or "").strip()
            if not raw_line:
                if labels:
                    break
                continue
            match = re.match(r"^(from|sent|to|cc|subject):", raw_line, re.IGNORECASE)
            if not match:
                break
            label = match.group(1).lower()
            if label not in labels:
                labels.append(label)
        return "subject" in labels and ("from" in labels or "sent" in labels)

    def _strip_outlook_reply_history(self, value):
        text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
        if not text.strip():
            return ""
        lines = text.split("\n")
        kept_lines = []
        for idx, raw_line in enumerate(lines):
            line = str(raw_line or "")
            stripped = line.strip()
            if re.match(
                r"^-{2,}\s*(original message|forwarded message)\s*-{0,}\s*$",
                stripped,
                re.IGNORECASE,
            ):
                break
            if re.match(r"^begin forwarded message:?\s*$", stripped, re.IGNORECASE):
                break
            if re.match(r"^on .+ wrote:\s*$", stripped, re.IGNORECASE):
                break
            if self._looks_like_outlook_reply_header_block(lines, idx):
                break
            if stripped.startswith(">"):
                continue
            kept_lines.append(line)
        cleaned = "\n".join(kept_lines)
        return self._normalize_outlook_scan_dedupe_text(cleaned)

    def _extract_outlook_scan_unique_body(self, value):
        text = str(value or "")
        stripped = self._strip_outlook_reply_history(text)
        if stripped:
            text = stripped
        return self._normalize_outlook_scan_dedupe_text(text)

    def _looks_like_outlook_banner_block(self, value):
        block = self._normalize_outlook_scan_dedupe_text(value)
        if not block:
            return False
        lower = block.lower()
        if (
            "external email" in lower
            or "this email originated from outside" in lower
            or "use caution with links" in lower
        ):
            return len(block.split("\n")) <= 4
        return False

    def _looks_like_outlook_device_footer_line(self, value):
        line = str(value or "").strip().lower()
        if not line:
            return False
        return bool(
            re.match(r"^sent from my (iphone|ipad|android|mobile device).*$", line)
            or re.match(r"^get outlook for (ios|android)$", line)
            or line == "sent from outlook for ios"
        )

    def _looks_like_outlook_signature_intro_line(self, value):
        line = str(value or "").strip()
        lower = line.lower()
        if not lower:
            return False
        if line in {"--", "__", "___"}:
            return True
        if self._looks_like_outlook_device_footer_line(lower):
            return True
        return bool(
            re.match(
                r"^(thanks|thanks again|many thanks|thank you|best|best regards|kind regards|warm regards|regards|sincerely|respectfully|cheers|all the best)[,!.]?$",
                lower,
            )
        )

    def _looks_like_outlook_disclaimer_block(self, value):
        block = self._normalize_outlook_scan_dedupe_text(value)
        if not block:
            return False
        if self._looks_like_outlook_banner_block(block):
            return True
        lines = [str(line or "").strip() for line in block.split("\n") if str(line or "").strip()]
        if lines and all(self._looks_like_outlook_device_footer_line(line) for line in lines):
            return True
        lower = block.lower()
        phrases = (
            "this email and any attachments",
            "this message and any attachments",
            "intended only for the named recipient",
            "intended only for the addressee",
            "intended only for the person",
            "if you are not the intended recipient",
            "please delete this email",
            "please notify the sender",
            "confidential",
            "privileged",
            "disclosure",
            "do not disseminate",
            "do not copy",
            "do not distribute",
            "views or opinions expressed",
            "environment before printing",
        )
        matches = sum(1 for phrase in phrases if phrase in lower)
        return matches >= 2 or (
            matches >= 1
            and len(block) >= 80
            and ("recipient" in lower or "confidential" in lower or "privileged" in lower)
        )

    def _looks_like_outlook_signature_block(self, value):
        block = self._normalize_outlook_scan_dedupe_text(value)
        if not block:
            return False
        lines = [str(line or "").strip() for line in block.split("\n") if str(line or "").strip()]
        if not lines or len(lines) > 8:
            return False
        if self._looks_like_outlook_signature_intro_line(lines[0]):
            return True
        lower = block.lower()
        short_lines = sum(1 for line in lines if len(line) <= 60)
        has_contact = bool(
            re.search(r"@|https?://|www\.|\b(?:cell|mobile|office|direct|phone|tel|fax)\b", lower)
        )
        title_hits = sum(
            1
            for keyword_pattern in (
                r"\bproject manager\b",
                r"\bmanager\b",
                r"\bengineer\b",
                r"\bdesigner\b",
                r"\bprincipal\b",
                r"\bassociate\b",
                r"\bcoordinator\b",
                r"\bdirector\b",
                r"\bconsultant\b",
                r"\bengineering\b",
                r"\bsolutions\b",
                r"\binc\.?\b",
                r"\bllc\b",
            )
            if re.search(keyword_pattern, lower)
        )
        if has_contact and short_lines >= max(2, len(lines) - 1):
            return True
        return title_hits >= 2 and short_lines >= max(2, len(lines) - 1)

    def _strip_outlook_noise_blocks(self, value):
        blocks = [
            self._normalize_outlook_scan_dedupe_text(raw_block)
            for raw_block in re.split(r"\n\s*\n+", str(value or ""))
        ]
        blocks = [block for block in blocks if block]
        while blocks and self._looks_like_outlook_banner_block(blocks[0]):
            blocks.pop(0)
        while len(blocks) > 1 and (
            self._looks_like_outlook_disclaimer_block(blocks[-1])
            or self._looks_like_outlook_signature_block(blocks[-1])
        ):
            blocks.pop()
        return "\n\n".join(blocks).strip()

    def _strip_outlook_signature_lines(self, value):
        normalized = self._normalize_outlook_scan_dedupe_text(value)
        if not normalized:
            return ""
        lines = normalized.split("\n")
        non_empty_indexes = [idx for idx, line in enumerate(lines) if str(line or "").strip()]
        if len(non_empty_indexes) < 2:
            return normalized
        search_start = max(0, non_empty_indexes[max(len(non_empty_indexes) - 8, 0)])
        for idx in range(search_start, len(lines)):
            if not self._looks_like_outlook_signature_intro_line(lines[idx]):
                continue
            candidate = self._normalize_outlook_scan_dedupe_text("\n".join(lines[:idx]))
            if candidate:
                return candidate
        return normalized

    def _build_outlook_scan_analysis_body(self, value):
        unique_body = self._extract_outlook_scan_unique_body(value)
        if not unique_body:
            return ""
        cleaned = self._strip_outlook_noise_blocks(unique_body)
        cleaned = self._strip_outlook_signature_lines(cleaned)
        cleaned = self._strip_outlook_noise_blocks(cleaned)
        cleaned = self._normalize_outlook_scan_dedupe_text(cleaned)
        return cleaned or unique_body

    def _build_outlook_scan_prompt_payload(
        self,
        entry,
        body_limit=OUTLOOK_SCAN_EMAIL_BODY_PROMPT_CHARS,
    ):
        entry = entry if isinstance(entry, dict) else {}
        message = entry.get("message") if isinstance(entry.get("message"), dict) else {}
        existing_payload = (
            entry.get("promptPayload") if isinstance(entry.get("promptPayload"), dict) else {}
        )
        prompt_ref = str(entry.get("promptRef") or existing_payload.get("messageRef") or "").strip()
        body_source = str(
            entry.get("analysisBody") or existing_payload.get("body") or ""
        ).strip()
        message_from = message.get("from") if isinstance(message.get("from"), dict) else {}
        payload_from = (
            existing_payload.get("from") if isinstance(existing_payload.get("from"), dict) else {}
        )
        return {
            "messageRef": prompt_ref,
            "subject": str(message.get("subject") or existing_payload.get("subject") or "").strip(),
            "from": {
                "name": str(message_from.get("name") or payload_from.get("name") or "").strip(),
            },
            "receivedDateTime": str(
                message.get("receivedDateTime") or existing_payload.get("receivedDateTime") or ""
            ).strip(),
            "body": self._trim_outlook_scan_prompt_text(body_source, body_limit),
        }

    def _build_outlook_scan_prompt_config(self, reduced=False):
        if reduced:
            return {
                "bodyLimit": OUTLOOK_SCAN_RETRY_EMAIL_BODY_PROMPT_CHARS,
                "emailBudgetChars": OUTLOOK_SCAN_RETRY_PROMPT_EMAIL_BUDGET_CHARS,
                "maxEmails": OUTLOOK_SCAN_RETRY_AI_LIMIT,
                "skipReason": OUTLOOK_SCAN_RETRY_SKIP_REASON,
            }
        return {
            "bodyLimit": OUTLOOK_SCAN_EMAIL_BODY_PROMPT_CHARS,
            "emailBudgetChars": OUTLOOK_SCAN_PROMPT_EMAIL_BUDGET_CHARS,
            "maxEmails": OUTLOOK_SCAN_AI_LIMIT,
            "skipReason": OUTLOOK_SCAN_PROMPT_SKIP_REASON,
        }

    def _is_outlook_scan_timeout_error(self, error):
        message = str(error or "").strip().lower()
        return (
            "deadline_exceeded" in message
            or "deadline exceeded" in message
            or "deadline expired" in message
            or "timed out" in message
            or "timeout" in message
            or ("504" in message and "deadline" in message)
        )

    def _dedupe_outlook_scan_skipped_messages(self, skipped_messages):
        if not isinstance(skipped_messages, list):
            return []
        deduped = []
        seen = set()
        for skipped in skipped_messages:
            if not isinstance(skipped, dict):
                continue
            message = skipped.get("message") if isinstance(skipped.get("message"), dict) else {}
            reason = str(skipped.get("reason") or "").strip()
            if reason == OUTLOOK_SCAN_RETRY_SKIP_REASON:
                reason = OUTLOOK_SCAN_PROMPT_SKIP_REASON
            message_id = str(message.get("id") or "").strip().lower()
            subject = str(message.get("subject") or "").strip().lower()
            message_key = message_id or subject or json.dumps(message, ensure_ascii=True, sort_keys=True)
            key = (message_key, reason)
            if key in seen:
                continue
            seen.add(key)
            deduped.append({
                "message": message,
                "reason": reason,
            })
        return deduped

    def _split_outlook_scan_dedupe_blocks(self, value):
        text = self._normalize_outlook_scan_dedupe_text(value)
        if not text:
            return []
        blocks = []
        seen_keys = set()
        for raw_block in re.split(r"\n\s*\n+", text):
            block_text = self._normalize_outlook_scan_dedupe_text(raw_block)
            if not block_text:
                continue
            block_key = re.sub(r"\s+", " ", block_text).strip().lower()
            if not block_key or block_key in seen_keys:
                continue
            seen_keys.add(block_key)
            blocks.append({
                "text": block_text,
                "key": block_key,
            })
        return blocks

    def _dedupe_outlook_scan_entries(self, hydrated_entries):
        if not isinstance(hydrated_entries, list):
            return [], [], {
                "threadsDetected": 0,
                "dedupedEmailCount": 0,
                "dedupeSkippedEmailCount": 0,
            }

        entries_by_thread = {}
        threadless_entries = []
        for entry in hydrated_entries:
            if not isinstance(entry, dict):
                continue
            message = entry.get("message") if isinstance(entry.get("message"), dict) else {}
            prompt_payload = (
                entry.get("promptPayload")
                if isinstance(entry.get("promptPayload"), dict)
                else {}
            )
            conversation_id = str(
                message.get("conversationId") or prompt_payload.get("conversationId") or ""
            ).strip()
            if conversation_id:
                entries_by_thread.setdefault(conversation_id.lower(), []).append(entry)
            else:
                threadless_entries.append(entry)

        deduped_entries = {}
        skipped_messages = []
        deduped_email_count = 0
        dedupe_skipped_email_count = 0

        for thread_entries in entries_by_thread.values():
            seen_blocks = set()
            ordered_entries = sorted(
                thread_entries,
                key=lambda entry: (
                    -(entry.get("receivedSort") or 0),
                    str(entry.get("promptRef") or ""),
                ),
            )
            for entry in ordered_entries:
                prompt_payload = (
                    entry.get("promptPayload")
                    if isinstance(entry.get("promptPayload"), dict)
                    else {}
                )
                raw_body = str(entry.get("analysisBody") or prompt_payload.get("body") or "").strip()
                normalized_original = self._normalize_outlook_scan_dedupe_text(raw_body)
                unique_body = self._extract_outlook_scan_unique_body(raw_body)
                deduped_blocks = []
                for block in self._split_outlook_scan_dedupe_blocks(unique_body):
                    if block["key"] in seen_blocks:
                        continue
                    seen_blocks.add(block["key"])
                    deduped_blocks.append(block["text"])
                deduped_body = "\n\n".join(deduped_blocks).strip()
                if not deduped_body:
                    dedupe_skipped_email_count += 1
                    skipped_messages.append({
                        "message": entry.get("message") or {},
                        "reason": OUTLOOK_SCAN_DEDUPE_SKIP_REASON_THREAD,
                    })
                    continue
                normalized_deduped = self._normalize_outlook_scan_dedupe_text(deduped_body)
                if normalized_deduped and normalized_deduped != normalized_original:
                    deduped_email_count += 1
                deduped_entries[str(entry.get("promptRef") or "")] = {
                    **entry,
                    "analysisBody": deduped_body,
                }

        seen_threadless_bodies = set()
        for entry in threadless_entries:
            prompt_payload = (
                entry.get("promptPayload")
                if isinstance(entry.get("promptPayload"), dict)
                else {}
            )
            raw_body = str(entry.get("analysisBody") or prompt_payload.get("body") or "").strip()
            normalized_original = self._normalize_outlook_scan_dedupe_text(raw_body)
            deduped_body = self._extract_outlook_scan_unique_body(raw_body)
            normalized_deduped = self._normalize_outlook_scan_dedupe_text(deduped_body)
            if not normalized_deduped:
                dedupe_skipped_email_count += 1
                skipped_messages.append({
                    "message": entry.get("message") or {},
                    "reason": OUTLOOK_SCAN_DEDUPE_SKIP_REASON_THREAD,
                })
                continue
            content_key = re.sub(r"\s+", " ", normalized_deduped).strip().lower()
            if content_key in seen_threadless_bodies:
                dedupe_skipped_email_count += 1
                skipped_messages.append({
                    "message": entry.get("message") or {},
                    "reason": OUTLOOK_SCAN_DEDUPE_SKIP_REASON_DUPLICATE,
                })
                continue
            seen_threadless_bodies.add(content_key)
            if normalized_deduped != normalized_original:
                deduped_email_count += 1
            deduped_entries[str(entry.get("promptRef") or "")] = {
                **entry,
                "analysisBody": deduped_body,
            }

        ordered_entries = []
        for entry in hydrated_entries:
            prompt_ref = str(entry.get("promptRef") or "")
            if prompt_ref in deduped_entries:
                ordered_entries.append(deduped_entries[prompt_ref])

        return ordered_entries, skipped_messages, {
            "threadsDetected": len(entries_by_thread),
            "dedupedEmailCount": deduped_email_count,
            "dedupeSkippedEmailCount": dedupe_skipped_email_count,
        }

    def _normalize_outlook_scan_task_list(self, raw_tasks):
        normalized_tasks = []
        if not isinstance(raw_tasks, list):
            return normalized_tasks
        for raw_task in raw_tasks:
            if isinstance(raw_task, dict):
                text = str(raw_task.get("text") or "").strip()
                done = bool(raw_task.get("done"))
                links = raw_task.get("links") if isinstance(raw_task.get("links"), list) else []
            else:
                text = str(raw_task or "").strip()
                done = False
                links = []
            if not text:
                continue
            normalized_tasks.append({
                "text": text,
                "done": done,
                "links": links,
            })
        return normalized_tasks

    def _normalize_outlook_scan_project_context(self, raw_context):
        if not isinstance(raw_context, list):
            return []

        normalized = []
        for raw_project in raw_context:
            if not isinstance(raw_project, dict):
                continue
            deliverables = []
            raw_deliverables = raw_project.get("deliverables")
            raw_deliverables = raw_deliverables if isinstance(raw_deliverables, list) else []
            for raw_deliverable in raw_deliverables:
                if not isinstance(raw_deliverable, dict):
                    continue
                name = str(raw_deliverable.get("name") or "").strip()
                due = str(raw_deliverable.get("due") or "").strip()
                status = str(raw_deliverable.get("status") or "").strip()
                if not (name or due or status):
                    continue
                deliverables.append({
                    "name": name,
                    "due": due,
                    "status": status,
                })
            deliverables.sort(
                key=lambda deliverable: (
                    1 if parse_due_str(deliverable.get("due")) is None else 0,
                    parse_due_str(deliverable.get("due")) or datetime.datetime.max,
                    str(deliverable.get("name") or "").lower(),
                )
            )
            project = {
                "id": str(raw_project.get("id") or "").strip(),
                "name": str(raw_project.get("name") or "").strip(),
                "nick": str(raw_project.get("nick") or "").strip(),
                "path": str(raw_project.get("path") or "").strip(),
                "deliverables": deliverables,
            }
            if any(project.get(key) for key in ("id", "name", "nick", "path")) or deliverables:
                normalized.append(project)

        normalized.sort(
            key=lambda project: (
                min(
                    (
                        parse_due_str(deliverable.get("due"))
                        for deliverable in project.get("deliverables") or []
                        if parse_due_str(deliverable.get("due")) is not None
                    ),
                    default=datetime.datetime.max,
                ),
                str(project.get("id") or "").lower(),
                str(project.get("name") or "").lower(),
                str(project.get("path") or "").lower(),
            )
        )
        return normalized

    def _select_outlook_scan_prompt_context(
        self,
        hydrated_entries,
        project_context,
        timeframe="week",
        scan_date="",
        body_limit=OUTLOOK_SCAN_EMAIL_BODY_PROMPT_CHARS,
        email_budget_chars=OUTLOOK_SCAN_PROMPT_EMAIL_BUDGET_CHARS,
        max_emails=OUTLOOK_SCAN_AI_LIMIT,
        skip_reason=OUTLOOK_SCAN_PROMPT_SKIP_REASON,
    ):
        prioritized_emails = sorted(
            hydrated_entries,
            key=lambda entry: (
                -int(entry.get("score") or 0),
                -(entry.get("receivedSort") or 0),
                str(entry.get("promptRef") or ""),
            ),
        )

        flat_deliverables = []
        for project in project_context:
            for deliverable in project.get("deliverables") or []:
                flat_deliverables.append({
                    "projectId": str(project.get("id") or "").strip(),
                    "projectName": str(project.get("name") or "").strip(),
                    "projectNick": str(project.get("nick") or "").strip(),
                    "projectPath": str(project.get("path") or "").strip(),
                    "deliverable": str(deliverable.get("name") or "").strip(),
                    "due": str(deliverable.get("due") or "").strip(),
                    "status": str(deliverable.get("status") or "").strip(),
                })
        flat_deliverables.sort(
            key=lambda deliverable: (
                1 if parse_due_str(deliverable.get("due")) is None else 0,
                parse_due_str(deliverable.get("due")) or datetime.datetime.max,
                str(deliverable.get("projectId") or "").lower(),
                str(deliverable.get("projectName") or "").lower(),
                str(deliverable.get("deliverable") or "").lower(),
            )
        )

        included_emails = []
        email_chars = 0
        skipped_messages = []
        for entry in prioritized_emails:
            max_email_count = max(int(max_emails or 0), 0)
            if max_email_count and len(included_emails) >= max_email_count:
                skipped_messages.append({
                    "message": entry.get("message") or {},
                    "reason": skip_reason,
                })
                continue
            prompt_payload = self._build_outlook_scan_prompt_payload(
                entry,
                body_limit=body_limit,
            )
            serialized = json.dumps(prompt_payload, ensure_ascii=True, separators=(",", ":"))
            if included_emails and email_chars + len(serialized) > email_budget_chars:
                skipped_messages.append({
                    "message": entry.get("message") or {},
                    "reason": skip_reason,
                })
                continue
            included_emails.append({
                **entry,
                "promptPayload": prompt_payload,
            })
            email_chars += len(serialized)

        included_deliverables = []
        deliverable_chars = 0
        for deliverable in flat_deliverables:
            serialized = json.dumps(
                deliverable,
                ensure_ascii=True,
                separators=(",", ":"),
            )
            if included_deliverables and deliverable_chars + len(serialized) > OUTLOOK_SCAN_PROMPT_DELIVERABLE_BUDGET_CHARS:
                continue
            included_deliverables.append(deliverable)
            deliverable_chars += len(serialized)

        prompt_truncated = (
            len(included_emails) < len(prioritized_emails)
            or len(included_deliverables) < len(flat_deliverables)
        )

        while True:
            prompt = self._build_outlook_scan_batch_prompt(
                included_emails,
                included_deliverables,
                "",
                [],
                timeframe,
                scan_date=scan_date,
                prompt_truncated=prompt_truncated,
            )
            if len(prompt) <= OUTLOOK_SCAN_PROMPT_MAX_CHARS:
                break
            if len(included_emails) > 1:
                removed = included_emails.pop()
                skipped_messages.append({
                    "message": removed.get("message") or {},
                    "reason": skip_reason,
                })
                prompt_truncated = True
                continue
            if included_deliverables:
                included_deliverables.pop()
                prompt_truncated = True
                continue
            break

        included_email_refs = {str(entry.get("promptRef") or "") for entry in included_emails}
        trimmed_messages = [
            skipped
            for skipped in skipped_messages
            if str(((skipped.get("message") or {}).get("id")) or "").strip()
        ]

        seen_message_ids = {
            str(((skipped.get("message") or {}).get("id")) or "").strip().lower()
            for skipped in trimmed_messages
        }
        for entry in prioritized_emails:
            prompt_ref = str(entry.get("promptRef") or "")
            message_id = str(((entry.get("message") or {}).get("id")) or "").strip().lower()
            if prompt_ref in included_email_refs or (message_id and message_id in seen_message_ids):
                continue
            trimmed_messages.append({
                "message": entry.get("message") or {},
                "reason": skip_reason,
            })
            if message_id:
                seen_message_ids.add(message_id)

        return included_emails, included_deliverables, prompt_truncated, trimmed_messages

    def _build_outlook_scan_batch_prompt(
        self,
        included_emails,
        included_deliverables,
        user_name,
        discipline,
        timeframe,
        scan_date="",
        prompt_truncated=False,
    ):
        current_date = datetime.date.today().strftime("%m/%d/%Y")
        disciplines_str = ", ".join(discipline) if isinstance(discipline, list) else (discipline or "Engineering")
        task_notes_contract = self._build_deliverable_task_notes_prompt_contract(
            disciplines_str,
        )
        period_label = self._format_outlook_scan_period_label(
            timeframe=timeframe,
            scan_date=scan_date,
        )
        if self._normalize_outlook_scan_date(scan_date):
            period_instruction = f"received on {period_label}"
            deliverable_instruction = "due on that same day"
        else:
            period_instruction = f"for {period_label}"
            deliverable_instruction = "due in that same period"
        email_payload = [
            entry.get("promptPayload") or {}
            for entry in included_emails
            if isinstance(entry, dict)
        ]
        deliverable_payload = [
            deliverable
            for deliverable in included_deliverables
            if isinstance(deliverable, dict)
        ]
        email_json = json.dumps(email_payload, ensure_ascii=True, separators=(",", ":"))
        deliverable_json = json.dumps(deliverable_payload, ensure_ascii=True, separators=(",", ":"))
        truncation_note = (
            "The provided emails or deliverables were trimmed to fit one AI request. "
            "Only reason over the included data."
            if prompt_truncated
            else "The included emails and deliverables are the full context provided for this scan."
        )
        return f"""
You are an intelligent assistant for {user_name}, a(n) {disciplines_str} engineering project manager.
Review the batched Outlook emails {period_instruction} and compare them against the current deliverables {deliverable_instruction}.
Suggest only NEW deliverables that should be added, if any. Do not suggest a deliverable if the same or an equivalent deliverable already appears in CURRENT_DELIVERABLES_IN_PERIOD.

Use only evidence from the included emails. Every suggestion must cite one or more supportingMessageRefs from INCLUDED_EMAILS.
Each included email body is a reduced plain-text thread view. Quoted history, signatures, device footers, and disclaimer boilerplate may already be removed. Treat the body as the unique useful text for that reply when provided.
Prefer concise, standardized deliverable names when possible (for example: DD60, DD90, CD60, CD90, CD100, CDF, RFI, RFI #2, Submittal, Lighting Submittal, Controls Submittal, Record Set, Record Drawings, IFP, Site Survey, Survey Report, ASR, ASR #2, PCC, PCC #3, Bulletin #2, Coordination, Meeting, Revision).
Use the provided project identifiers when available. If a project must be inferred from the emails, return the best available id/name/path fields.
Focus only on the primary {disciplines_str} engineering tasks mentioned. Ignore work for other disciplines.
Use the current date {current_date} when interpreting due dates. If a due date is not clear, leave it empty.
{task_notes_contract}
{truncation_note}

Return ONLY valid JSON with this exact shape:
{{"suggestions":[{{"projectId":"","projectName":"","projectNick":"","projectPath":"","deliverable":"","due":"","tasks":[""],"notes":"","supportingMessageRefs":["E001"]}}]}}

If no additions are warranted, return:
{{"suggestions":[]}}

INCLUDED_EMAILS:
{email_json}

CURRENT_DELIVERABLES_IN_PERIOD:
{deliverable_json}
""".strip()

    def _normalize_outlook_scan_batch_suggestions(self, raw_payload):
        if isinstance(raw_payload, list):
            raw_suggestions = raw_payload
        elif isinstance(raw_payload, dict):
            if isinstance(raw_payload.get("suggestions"), list):
                raw_suggestions = raw_payload.get("suggestions") or []
            elif any(
                key in raw_payload
                for key in (
                    "deliverable",
                    "deliverableName",
                    "project",
                    "projectId",
                    "projectName",
                )
            ):
                raw_suggestions = [raw_payload]
            else:
                raise RuntimeError("AI returned invalid Outlook scan suggestions.")
        else:
            raise RuntimeError("AI returned invalid Outlook scan suggestions.")

        if not isinstance(raw_suggestions, list):
            raise RuntimeError("AI returned invalid Outlook scan suggestions.")
        if not raw_suggestions:
            return []

        suggestions = []
        for raw_suggestion in raw_suggestions:
            if not isinstance(raw_suggestion, dict):
                continue
            project = raw_suggestion.get("project")
            project = project if isinstance(project, dict) else {}
            deliverable_payload = raw_suggestion.get("deliverable")
            deliverable_payload = deliverable_payload if isinstance(deliverable_payload, dict) else {}

            project_info = {
                "id": str(project.get("id") or raw_suggestion.get("projectId") or "").strip(),
                "name": str(project.get("name") or raw_suggestion.get("projectName") or "").strip(),
                "nick": str(project.get("nick") or raw_suggestion.get("projectNick") or "").strip(),
                "path": str(project.get("path") or raw_suggestion.get("projectPath") or "").strip(),
            }
            deliverable_name = str(
                deliverable_payload.get("name")
                or raw_suggestion.get("deliverable")
                or raw_suggestion.get("deliverableName")
                or raw_suggestion.get("name")
                or ""
            ).strip()
            supporting_refs = (
                raw_suggestion.get("supportingMessageRefs")
                or raw_suggestion.get("supportingMessageIds")
                or raw_suggestion.get("messageRefs")
                or []
            )
            if isinstance(supporting_refs, str):
                supporting_refs = [supporting_refs]
            if not isinstance(supporting_refs, list):
                supporting_refs = []
            normalized_refs = []
            seen_refs = set()
            for raw_ref in supporting_refs:
                ref = str(raw_ref or "").strip()
                if not ref or ref in seen_refs:
                    continue
                seen_refs.add(ref)
                normalized_refs.append(ref)

            if not deliverable_name:
                continue
            if not any(project_info.values()):
                continue

            suggestions.append({
                "project": project_info,
                "deliverable": {
                    "name": deliverable_name,
                    "due": str(deliverable_payload.get("due") or raw_suggestion.get("due") or "").strip(),
                    "notes": str(deliverable_payload.get("notes") or raw_suggestion.get("notes") or "").strip(),
                    "tasks": self._normalize_outlook_scan_task_list(
                        deliverable_payload.get("tasks")
                        if isinstance(deliverable_payload.get("tasks"), list)
                        else raw_suggestion.get("tasks")
                    ),
                },
                "supportingMessageRefs": normalized_refs,
            })

        if suggestions:
            return suggestions
        raise RuntimeError("AI returned invalid Outlook scan suggestions.")

    def _request_outlook_scan_batch_suggestions(self, prompt, api_key):
        final_api_key = self._resolve_google_ai_api_key(api_key)
        try:
            self._ensure_aiohttp()
            client = genai.Client(
                api_key=final_api_key,
                http_options=types.HttpOptions(timeout=120000),
            )
            response = client.models.generate_content(
                model="gemini-3-flash-preview",
                contents=[
                    types.Content(
                        role="user",
                        parts=[types.Part.from_text(text=prompt)],
                    ),
                ],
                config=types.GenerateContentConfig(
                    temperature=0,
                    response_mime_type="application/json",
                ),
            )

            raw_text = response.text
            if raw_text is None:
                if hasattr(response, "parts") and response.parts:
                    raw_text = "".join(
                        part.text for part in response.parts
                        if hasattr(part, "text") and part.text
                    )
                if not raw_text:
                    logging.error(f"AI returned empty Outlook scan response. Full response: {response}")
                    raise RuntimeError("AI returned an empty Outlook scan response. Please try again.")

            cleaned = raw_text.strip()
            if not cleaned:
                raise RuntimeError("AI returned an empty Outlook scan response. Please try again.")

            logging.debug(f"Outlook batch AI response text: {cleaned[:500]}...")
            return self._normalize_outlook_scan_batch_suggestions(json.loads(cleaned))
        except json.JSONDecodeError as exc:
            logging.error(f"Failed to parse Outlook scan AI response as JSON: {exc}")
            raise RuntimeError("AI returned invalid Outlook scan suggestions JSON. Please try again.")
        except Exception as exc:
            msg = str(exc)
            lower = msg.lower()
            logging.error(f"Error processing Outlook scan batch with AI: {type(exc).__name__}: {exc}")
            if (
                "api key expired" in lower
                or "api_key_invalid" in lower
                or "invalid api key" in lower
            ):
                raise RuntimeError(
                    "Your Google API key is expired/invalid. "
                    "Create a new key in Google AI Studio, update your settings, then try again."
                )
            if "model" in lower and ("not found" in lower or "does not exist" in lower):
                raise RuntimeError(
                    "AI model not available. The Gemini 3 Flash model may not be accessible with your API key."
                )
            if "quota" in lower or "rate limit" in lower:
                raise RuntimeError("API rate limit exceeded. Please wait a moment and try again.")
            raise RuntimeError(f"AI error: {msg}")

    def _analyze_outlook_scan_batch(
        self,
        message_summaries,
        body_loader,
        project_context,
        api_key,
        user_name,
        discipline,
        timeframe,
        source,
        has_more=False,
        scan_date="",
        progress_sender=None,
    ):
        if progress_sender is None:
            progress_sender = self._send_outlook_scan_progress
        hydrated_entries = []
        skipped_messages = []
        relevant_email_count = 0
        total_emails = len(message_summaries) if isinstance(message_summaries, list) else 0
        deliverables_in_period = self._count_outlook_scan_project_context_deliverables(
            project_context
        )

        progress_sender(
            "hydrating",
            "Reading Outlook emails...",
            source=source,
            timeframe=timeframe,
            scanDate=scan_date,
            totalEmails=total_emails,
            processedEmails=0,
            skippedEmails=0,
            deliverablesInPeriod=deliverables_in_period,
        )

        for idx, summary in enumerate(message_summaries):
            score = self._outlook_message_candidate_score(summary)
            if score > 0:
                relevant_email_count += 1
            try:
                hydrated_summary, body_text = body_loader(summary)
                body_content = self._build_outlook_scan_analysis_body(
                    body_text or (hydrated_summary or {}).get("bodyPreview") or ""
                )
                if not body_content:
                    skipped_messages.append({
                        "message": hydrated_summary or summary,
                        "reason": "Message body was empty.",
                    })
                    continue
                received_dt = parse_utc_iso((hydrated_summary or {}).get("receivedDateTime"))
                received_sort = received_dt.timestamp() if received_dt else 0
                prompt_ref = f"E{idx + 1:03d}"
                hydrated_entries.append({
                    "promptRef": prompt_ref,
                    "score": score,
                    "receivedSort": received_sort,
                    "message": hydrated_summary or summary,
                    "analysisBody": body_content,
                })
            except Exception as exc:
                skipped_messages.append({
                    "message": summary,
                    "reason": str(exc),
                })
            processed_emails = idx + 1
            if processed_emails == 1 or processed_emails % 10 == 0 or processed_emails == total_emails:
                progress_sender(
                    "hydrating",
                    f"Processed {processed_emails} of {total_emails} emails.",
                    source=source,
                    timeframe=timeframe,
                    scanDate=scan_date,
                    totalEmails=total_emails,
                    processedEmails=processed_emails,
                    skippedEmails=len(skipped_messages),
                    deliverablesInPeriod=deliverables_in_period,
                    relevantEmails=relevant_email_count,
                )

        hydrated_entries, dedupe_skipped_messages, dedupe_stats = self._dedupe_outlook_scan_entries(
            hydrated_entries
        )
        skipped_messages.extend(dedupe_skipped_messages)

        prompt_config = self._build_outlook_scan_prompt_config()
        included_emails, included_deliverables, prompt_truncated, trimmed_messages = self._select_outlook_scan_prompt_context(
            hydrated_entries,
            project_context,
            timeframe=timeframe,
            scan_date=scan_date,
            body_limit=prompt_config["bodyLimit"],
            email_budget_chars=prompt_config["emailBudgetChars"],
            max_emails=prompt_config["maxEmails"],
            skip_reason=prompt_config["skipReason"],
        )
        skipped_messages.extend(trimmed_messages)
        skipped_messages = self._dedupe_outlook_scan_skipped_messages(skipped_messages)

        progress_sender(
            "preparing_ai",
            "Preparing the batched AI review...",
            source=source,
            timeframe=timeframe,
            scanDate=scan_date,
            totalEmails=total_emails,
            processedEmails=total_emails,
            includedEmails=len(included_emails),
            skippedEmails=len(skipped_messages),
            deliverablesInPeriod=deliverables_in_period,
            relevantEmails=relevant_email_count,
            threadsDetected=dedupe_stats["threadsDetected"],
            dedupedEmailCount=dedupe_stats["dedupedEmailCount"],
            dedupeSkippedEmailCount=dedupe_stats["dedupeSkippedEmailCount"],
        )

        batch_suggestions = []
        if included_emails:
            progress_sender(
                "reviewing_ai",
                "Reviewing the included emails with AI...",
                source=source,
                timeframe=timeframe,
                scanDate=scan_date,
                totalEmails=total_emails,
                processedEmails=total_emails,
                includedEmails=len(included_emails),
                skippedEmails=len(skipped_messages),
                deliverablesInPeriod=deliverables_in_period,
                relevantEmails=relevant_email_count,
                threadsDetected=dedupe_stats["threadsDetected"],
                dedupedEmailCount=dedupe_stats["dedupedEmailCount"],
                dedupeSkippedEmailCount=dedupe_stats["dedupeSkippedEmailCount"],
            )
            prompt = self._build_outlook_scan_batch_prompt(
                included_emails,
                included_deliverables,
                user_name,
                discipline,
                timeframe,
                scan_date=scan_date,
                prompt_truncated=prompt_truncated,
            )
            try:
                batch_suggestions = self._request_outlook_scan_batch_suggestions(prompt, api_key)
            except RuntimeError as exc:
                if not self._is_outlook_scan_timeout_error(exc):
                    raise
                retry_prompt_config = self._build_outlook_scan_prompt_config(reduced=True)
                included_emails, included_deliverables, retry_prompt_truncated, retry_trimmed_messages = self._select_outlook_scan_prompt_context(
                    hydrated_entries,
                    project_context,
                    timeframe=timeframe,
                    scan_date=scan_date,
                    body_limit=retry_prompt_config["bodyLimit"],
                    email_budget_chars=retry_prompt_config["emailBudgetChars"],
                    max_emails=retry_prompt_config["maxEmails"],
                    skip_reason=retry_prompt_config["skipReason"],
                )
                skipped_messages.extend(retry_trimmed_messages)
                skipped_messages = self._dedupe_outlook_scan_skipped_messages(skipped_messages)
                prompt_truncated = prompt_truncated or retry_prompt_truncated
                progress_sender(
                    "reviewing_ai",
                    "AI request timed out. Retrying with a reduced text-only batch...",
                    source=source,
                    timeframe=timeframe,
                    scanDate=scan_date,
                    totalEmails=total_emails,
                    processedEmails=total_emails,
                    includedEmails=len(included_emails),
                    skippedEmails=len(skipped_messages),
                    deliverablesInPeriod=deliverables_in_period,
                    relevantEmails=relevant_email_count,
                    threadsDetected=dedupe_stats["threadsDetected"],
                    dedupedEmailCount=dedupe_stats["dedupedEmailCount"],
                    dedupeSkippedEmailCount=dedupe_stats["dedupeSkippedEmailCount"],
                )
                retry_prompt = self._build_outlook_scan_batch_prompt(
                    included_emails,
                    included_deliverables,
                    user_name,
                    discipline,
                    timeframe,
                    scan_date=scan_date,
                    prompt_truncated=prompt_truncated,
                )
                try:
                    batch_suggestions = self._request_outlook_scan_batch_suggestions(
                        retry_prompt,
                        api_key,
                    )
                except RuntimeError as retry_exc:
                    if self._is_outlook_scan_timeout_error(retry_exc):
                        raise RuntimeError(
                            "AI timed out while reviewing Outlook emails, even after retrying with a reduced text-only batch. Try scanning a quieter day or narrowing the email set."
                        ) from retry_exc
                    raise

        progress_sender(
            "matching",
            "Matching scan results to your current projects...",
            source=source,
            timeframe=timeframe,
            scanDate=scan_date,
            totalEmails=total_emails,
            processedEmails=total_emails,
            includedEmails=len(included_emails),
            skippedEmails=len(skipped_messages),
            deliverablesInPeriod=deliverables_in_period,
            relevantEmails=relevant_email_count,
            threadsDetected=dedupe_stats["threadsDetected"],
            dedupedEmailCount=dedupe_stats["dedupedEmailCount"],
            dedupeSkippedEmailCount=dedupe_stats["dedupeSkippedEmailCount"],
        )

        messages_by_ref = {
            str(entry.get("promptRef") or ""): entry.get("message") or {}
            for entry in included_emails
        }
        enriched_suggestions = []
        for suggestion in batch_suggestions:
            related_messages = []
            supporting_message_ids = []
            seen_ids = set()
            for ref in suggestion.get("supportingMessageRefs") or []:
                message = messages_by_ref.get(str(ref or "").strip())
                if not isinstance(message, dict) or not message:
                    continue
                message_id = str(message.get("id") or "").strip()
                dedupe_key = message_id.lower() if message_id else str(ref or "").strip().lower()
                if dedupe_key in seen_ids:
                    continue
                seen_ids.add(dedupe_key)
                related_messages.append(message)
                if message_id:
                    supporting_message_ids.append(message_id)
            enriched_suggestions.append({
                **suggestion,
                "supportingMessageIds": supporting_message_ids,
                "relatedMessages": related_messages,
            })

        result = {
            "status": "success",
            "analysisMode": "batch",
            "candidateCount": len(included_emails),
            "relevantEmailCount": relevant_email_count,
            "scannedCount": total_emails,
            "emailsIncludedCount": len(included_emails),
            "deliverablesIncludedCount": len(included_deliverables),
            "suggestions": enriched_suggestions,
            "skippedMessages": skipped_messages,
            "promptTruncated": prompt_truncated,
            "truncated": bool(has_more) or prompt_truncated,
            "threadsDetected": dedupe_stats["threadsDetected"],
            "dedupedEmailCount": dedupe_stats["dedupedEmailCount"],
            "dedupeSkippedEmailCount": dedupe_stats["dedupeSkippedEmailCount"],
            "timeframe": timeframe,
            "scanDate": self._normalize_outlook_scan_date(scan_date),
            "source": source,
        }
        selected_label = self._format_outlook_scan_period_label(
            timeframe=timeframe,
            scan_date=scan_date,
        )
        if total_emails <= 0:
            completion_message = (
                "Scan complete. No emails were found for the selected day."
                if self._normalize_outlook_scan_date(scan_date)
                else "Scan complete. No emails were found for the selected timeframe."
            )
        elif enriched_suggestions:
            completion_message = (
                f"Scan complete. Found {len(enriched_suggestions)} deliverable "
                f"suggestion{'s' if len(enriched_suggestions) != 1 else ''}."
            )
        else:
            completion_message = "Scan complete. No new deliverables were suggested."
        if self._normalize_outlook_scan_date(scan_date):
            completion_message += f" Day: {selected_label}."
        if prompt_truncated:
            completion_message += " Prompt trimmed to keep the scan bounded."
        progress_sender(
            "done",
            completion_message,
            active=False,
            source=source,
            timeframe=timeframe,
            scanDate=scan_date,
            totalEmails=total_emails,
            processedEmails=total_emails,
            includedEmails=len(included_emails),
            skippedEmails=len(skipped_messages),
            deliverablesInPeriod=deliverables_in_period,
            relevantEmails=relevant_email_count,
            threadsDetected=dedupe_stats["threadsDetected"],
            dedupedEmailCount=dedupe_stats["dedupedEmailCount"],
            dedupeSkippedEmailCount=dedupe_stats["dedupeSkippedEmailCount"],
        )
        return result

    def scan_outlook_inbox(self, payload, api_key, user_name, discipline):
        timeframe = "week"
        scan_date = ""
        source = ""
        try:
            request_payload = payload or {}
            if isinstance(request_payload, str):
                request_payload = json.loads(request_payload)
            if not isinstance(request_payload, dict):
                request_payload = {}
            raw_scan_date = request_payload.get("scanDate")
            scan_date = self._normalize_outlook_scan_date(raw_scan_date)
            timeframe = str(request_payload.get("timeframe") or "week").strip().lower()
            if timeframe not in ("week", "month"):
                timeframe = "week"
            if not scan_date and (
                raw_scan_date is not None or "timeframe" not in request_payload
            ):
                scan_date = self._normalize_outlook_scan_date(
                    datetime.datetime.now().astimezone().date().isoformat()
                )
            project_context = self._normalize_outlook_scan_project_context(
                request_payload.get("projectContext")
            )
            deliverables_in_period = self._count_outlook_scan_project_context_deliverables(
                project_context
            )
            self._send_outlook_scan_progress(
                "starting",
                "Starting Outlook inbox scan...",
                timeframe=timeframe,
                scanDate=scan_date,
                deliverablesInPeriod=deliverables_in_period,
            )

            desktop_available, desktop_reason = self._get_desktop_outlook_availability()
            source = OUTLOOK_SCAN_SOURCE_DESKTOP
            if not desktop_available:
                message = desktop_reason or "Desktop Outlook is unavailable on this machine."
                self._send_outlook_scan_progress(
                    "error",
                    message,
                    active=False,
                    source=source,
                    timeframe=timeframe,
                    scanDate=scan_date,
                    deliverablesInPeriod=deliverables_in_period,
                )
                return self._build_outlook_scan_error_result(
                    message,
                    timeframe=timeframe,
                    source=source,
                    scan_date=scan_date,
                )

            self._send_outlook_scan_progress(
                "listing",
                "Loading emails from Desktop Outlook...",
                source=source,
                timeframe=timeframe,
                scanDate=scan_date,
                deliverablesInPeriod=deliverables_in_period,
            )
            try:
                message_summaries, has_more = self._list_desktop_outlook_inbox_messages(
                    timeframe=timeframe,
                    limit=OUTLOOK_SCAN_FETCH_LIMIT,
                    scan_date=scan_date,
                )
            except Exception as exc:
                logging.error(f"Desktop Outlook inbox scan failed: {exc}")
                message = f"Desktop Outlook could not be read: {exc}"
                self._send_outlook_scan_progress(
                    "error",
                    message,
                    active=False,
                    source=source,
                    timeframe=timeframe,
                    scanDate=scan_date,
                    deliverablesInPeriod=deliverables_in_period,
                )
                return self._build_outlook_scan_error_result(
                    message,
                    timeframe=timeframe,
                    source=source,
                    scan_date=scan_date,
                )

            return self._analyze_outlook_scan_batch(
                message_summaries,
                lambda summary: self._get_desktop_outlook_message_body_text(summary.get("id")),
                project_context,
                api_key,
                user_name,
                discipline,
                timeframe,
                source,
                has_more=has_more,
                scan_date=scan_date,
            )
        except json.JSONDecodeError:
            message = "Invalid Outlook scan payload."
            self._send_outlook_scan_progress(
                "error",
                message,
                active=False,
                source=source,
                timeframe=timeframe,
                scanDate=scan_date,
            )
            return self._build_outlook_scan_error_result(
                message,
                timeframe=timeframe,
                scan_date=scan_date,
            )
        except Exception as e:
            logging.error(f"Error scanning Outlook inbox: {e}")
            message = str(e)
            self._send_outlook_scan_progress(
                "error",
                message,
                active=False,
                source=source,
                timeframe=timeframe,
                scanDate=scan_date,
            )
            return self._build_outlook_scan_error_result(
                message,
                timeframe=timeframe,
                source=source,
                scan_date=scan_date,
            )

    # --- Email capture inbox & follow-ups ---

    def _notify_email_capture_progress(self, payload):
        try:
            if not webview.windows:
                return
            payload_json = json.dumps(payload or {}, ensure_ascii=True)
            webview.windows[0].evaluate_js(
                f"window.updateEmailCaptureProgress({payload_json})"
            )
        except Exception as exc:
            logging.debug(f"_notify_email_capture_progress failed: {exc}")

    def _send_email_capture_progress(self, stage, message="", **kwargs):
        payload = {
            "stage": str(stage or "").strip(),
            "message": str(message or "").strip(),
        }
        for key, value in (kwargs or {}).items():
            if value is None:
                continue
            payload[key] = value
        if "active" not in payload:
            payload["active"] = payload["stage"] not in {"done", "error", "skipped"}
        self._notify_email_capture_progress(payload)

    def _resolve_outlook_current_user_smtp(self):
        def _read_identity():
            _, namespace = self._get_desktop_outlook_namespace()
            current_user = getattr(namespace, "CurrentUser", None)
            address_entry = getattr(current_user, "AddressEntry", None)
            if address_entry is None:
                return ""
            try:
                exchange_user = address_entry.GetExchangeUser()
            except Exception:
                exchange_user = None
            if exchange_user is not None:
                address = str(getattr(exchange_user, "PrimarySmtpAddress", "") or "").strip()
                if address:
                    return address
            return str(getattr(address_entry, "Address", "") or "").strip()

        try:
            return str(self._run_with_outlook_com(_read_identity) or "").strip()
        except Exception:
            return ""

    def _get_email_capture_owner_identity(self, cached_smtp=""):
        addresses = []
        names = []
        try:
            settings = self.get_user_settings()
        except Exception:
            settings = {}
        if not isinstance(settings, dict):
            settings = {}
        user_name = str(settings.get("userName") or "").strip()
        if user_name:
            names.append(user_name)
        google_auth = settings.get("googleAuth")
        if isinstance(google_auth, dict):
            google_email = str(google_auth.get("email") or "").strip()
            if google_email:
                addresses.append(google_email)
        smtp = str(cached_smtp or "").strip()
        if smtp:
            addresses.append(smtp)
        return {"addresses": addresses, "names": names}

    def _resolve_email_capture_scan_window(self, conn):
        now_utc = datetime.datetime.now(datetime.timezone.utc)
        watermark_raw = _get_email_capture_meta(conn, "last_scan_completed_at_utc")
        watermark = parse_utc_iso(watermark_raw)
        if watermark is None:
            start_utc = now_utc - datetime.timedelta(days=EMAIL_CAPTURE_FIRST_SCAN_DAYS)
        else:
            start_utc = watermark - datetime.timedelta(
                minutes=EMAIL_CAPTURE_SCAN_OVERLAP_MINUTES
            )
        earliest_allowed = now_utc - datetime.timedelta(days=EMAIL_CAPTURE_MAX_WINDOW_DAYS)
        if start_utc < earliest_allowed:
            start_utc = earliest_allowed
        return start_utc, now_utc

    def run_email_capture_scan(self, payload, api_key="", user_name="", discipline=None):
        try:
            request_payload = payload or {}
            if isinstance(request_payload, str):
                request_payload = json.loads(request_payload)
            if not isinstance(request_payload, dict):
                request_payload = {}
        except json.JSONDecodeError:
            return {"status": "error", "message": "Invalid email capture scan payload."}

        mode = str(request_payload.get("mode") or "manual").strip().lower()
        if mode not in {"auto", "manual"}:
            mode = "manual"
        project_context = self._normalize_outlook_scan_project_context(
            request_payload.get("projectContext")
        )

        with self._email_capture_scan_lock:
            if self._email_capture_scan_running:
                return {"status": "skipped", "reason": "scan-in-progress"}
            self._email_capture_scan_running = True

        try:
            return self._run_email_capture_scan_locked(
                mode, project_context, api_key, user_name, discipline
            )
        except Exception as exc:
            logging.error(f"Email capture scan failed: {exc}")
            self._send_email_capture_progress("error", str(exc), mode=mode)
            return {"status": "error", "message": str(exc), "mode": mode}
        finally:
            with self._email_capture_scan_lock:
                self._email_capture_scan_running = False

    def _run_email_capture_scan_locked(self, mode, project_context, api_key, user_name, discipline):
        desktop_available, desktop_reason = self._get_desktop_outlook_availability()
        if not desktop_available:
            message = desktop_reason or "Desktop Outlook is unavailable on this machine."
            conn = _open_email_capture_db()
            try:
                _set_email_capture_meta(conn, "last_scan_status", "unavailable")
                _set_email_capture_meta(
                    conn,
                    "last_scan_summary_json",
                    json.dumps({"status": "error", "message": message}, ensure_ascii=True),
                )
                conn.commit()
            finally:
                conn.close()
            self._send_email_capture_progress("error", message, mode=mode)
            return {"status": "error", "message": message, "mode": mode}

        self._send_email_capture_progress(
            "listing", "Checking Outlook for new emails...", mode=mode
        )

        conn = _open_email_capture_db()
        try:
            start_utc, end_utc = self._resolve_email_capture_scan_window(conn)

            cached_smtp = _get_email_capture_meta(conn, "owner_smtp_address")
            if not cached_smtp:
                cached_smtp = self._resolve_outlook_current_user_smtp()
                if cached_smtp:
                    _set_email_capture_meta(conn, "owner_smtp_address", cached_smtp)
                    conn.commit()
            owner_identity = self._get_email_capture_owner_identity(cached_smtp)

            message_summaries, has_more = self._list_desktop_outlook_inbox_messages(
                limit=EMAIL_CAPTURE_MAX_MESSAGES_PER_SCAN,
                start_utc=start_utc,
                end_utc=end_utc,
            )

            now_iso = utc_now_iso()
            new_count = 0
            cc_count = 0
            for summary in message_summaries:
                directedness = _classify_email_directedness(summary, owner_identity)
                outcome = _upsert_captured_email(conn, summary, directedness, now_iso)
                if outcome == "inserted":
                    new_count += 1
                    if directedness in ("cc", "unknown"):
                        cc_count += 1

            _set_email_capture_meta(conn, "last_scan_completed_at_utc", now_iso)
            _set_email_capture_meta(conn, "last_scan_status", "success")
            conn.commit()

            self._send_email_capture_progress(
                "capturing",
                f"Captured {new_count} new email{'s' if new_count != 1 else ''}.",
                mode=mode,
                newCount=new_count,
            )

            pending_rows = conn.execute(
                """
                SELECT * FROM captured_emails
                WHERE ai_status = 'pending' AND status = 'new'
                ORDER BY received_at_utc DESC
                LIMIT ?
                """,
                (OUTLOOK_SCAN_AI_LIMIT,),
            ).fetchall()
            pending_records = [_normalize_captured_email_record(row) for row in pending_rows]
        finally:
            conn.close()

        matched_count = 0
        ai_error = ""
        try:
            resolved_api_key = self._resolve_google_ai_api_key(api_key)
        except Exception:
            resolved_api_key = ""

        if pending_records and resolved_api_key:
            try:
                matched_count = self._run_email_capture_ai_stage(
                    pending_records,
                    project_context,
                    resolved_api_key,
                    user_name,
                    discipline,
                    mode,
                )
            except Exception as exc:
                logging.error(f"Email capture AI stage failed: {exc}")
                ai_error = str(exc)

        conn = _open_email_capture_db()
        try:
            pending_ai_count = conn.execute(
                "SELECT COUNT(*) FROM captured_emails WHERE ai_status = 'pending' AND status = 'new'"
            ).fetchone()[0]
            result = {
                "status": "success",
                "mode": mode,
                "newCount": new_count,
                "ccCount": cc_count,
                "matchedCount": matched_count,
                "pendingAiCount": int(pending_ai_count or 0),
                "truncated": bool(has_more),
                "windowStartUtc": start_utc.isoformat().replace("+00:00", "Z"),
                "windowEndUtc": end_utc.isoformat().replace("+00:00", "Z"),
            }
            if ai_error:
                result["aiError"] = ai_error
            _set_email_capture_meta(
                conn, "last_scan_summary_json", json.dumps(result, ensure_ascii=True)
            )
            conn.commit()
        finally:
            conn.close()

        completion_message = (
            f"Email scan complete. {new_count} new email{'s' if new_count != 1 else ''} captured."
        )
        self._send_email_capture_progress(
            "done",
            completion_message,
            mode=mode,
            newCount=new_count,
            matchedCount=matched_count,
        )
        return result

    def _run_email_capture_ai_stage(
        self, pending_records, project_context, api_key, user_name, discipline, mode
    ):
        summaries = [_captured_email_row_to_summary(record) for record in pending_records]
        record_ids_by_entry = {
            str(record.get("entryId") or ""): int(record.get("id"))
            for record in pending_records
            if str(record.get("entryId") or "")
        }

        analysis = self._analyze_outlook_scan_batch(
            summaries,
            lambda summary: self._get_desktop_outlook_message_body_text(summary.get("id")),
            project_context,
            api_key,
            user_name,
            discipline,
            "week",
            OUTLOOK_SCAN_SOURCE_DESKTOP,
            has_more=False,
            progress_sender=self._send_email_capture_progress,
        )

        retryable_skip_reasons = {
            OUTLOOK_SCAN_PROMPT_SKIP_REASON,
            OUTLOOK_SCAN_RETRY_SKIP_REASON,
        }
        matched_count = 0
        conn = _open_email_capture_db()
        try:
            for suggestion in analysis.get("suggestions") or []:
                project_json = json.dumps(suggestion.get("project") or {}, ensure_ascii=True)
                suggestion_json = json.dumps(
                    suggestion.get("deliverable") or {}, ensure_ascii=True
                )
                for entry_id in suggestion.get("supportingMessageIds") or []:
                    record_id = record_ids_by_entry.get(str(entry_id or ""))
                    if record_id is None:
                        continue
                    cursor = conn.execute(
                        """
                        UPDATE captured_emails
                        SET ai_status = 'matched', ai_project_json = ?, ai_suggestion_json = ?
                        WHERE id = ? AND ai_status = 'pending'
                        """,
                        (project_json, suggestion_json, record_id),
                    )
                    if cursor.rowcount:
                        matched_count += 1

            for skipped in analysis.get("skippedMessages") or []:
                message = skipped.get("message") if isinstance(skipped, dict) else {}
                reason = str(skipped.get("reason") or "") if isinstance(skipped, dict) else ""
                entry_id = str((message or {}).get("id") or "")
                record_id = record_ids_by_entry.get(entry_id)
                if record_id is None or reason in retryable_skip_reasons:
                    continue
                conn.execute(
                    """
                    UPDATE captured_emails
                    SET ai_status = 'skipped', ai_skip_reason = ?
                    WHERE id = ? AND ai_status = 'pending'
                    """,
                    (reason, record_id),
                )

            batch_ids = list(record_ids_by_entry.values())
            if batch_ids:
                skipped_entry_ids = {
                    str(((skipped or {}).get("message") or {}).get("id") or "")
                    for skipped in analysis.get("skippedMessages") or []
                    if isinstance(skipped, dict)
                    and str(skipped.get("reason") or "") in retryable_skip_reasons
                }
                retryable_ids = {
                    record_ids_by_entry[entry_id]
                    for entry_id in skipped_entry_ids
                    if entry_id in record_ids_by_entry
                }
                no_match_ids = [
                    record_id for record_id in batch_ids if record_id not in retryable_ids
                ]
                if no_match_ids:
                    placeholders = ",".join("?" for _ in no_match_ids)
                    conn.execute(
                        f"""
                        UPDATE captured_emails
                        SET ai_status = 'no_match'
                        WHERE id IN ({placeholders}) AND ai_status = 'pending'
                        """,
                        no_match_ids,
                    )
            conn.commit()
        finally:
            conn.close()
        return matched_count

    def list_captured_emails(self, payload=None):
        request = payload if isinstance(payload, dict) else {}
        status_filter = str(request.get("status") or "new").strip().lower()
        if status_filter not in {"new", "accepted", "dismissed", "all"}:
            status_filter = "new"
        directedness_filter = str(request.get("directedness") or "all").strip().lower()
        try:
            limit = max(1, min(int(request.get("limit") or 200), 500))
        except (TypeError, ValueError):
            limit = 200
        try:
            offset = max(0, int(request.get("offset") or 0))
        except (TypeError, ValueError):
            offset = 0

        conditions = []
        params = []
        if status_filter != "all":
            conditions.append("status = ?")
            params.append(status_filter)
        if directedness_filter == "directed":
            conditions.append("directedness IN ('to', 'named')")
        elif directedness_filter == "cc":
            conditions.append("directedness IN ('cc', 'unknown')")
        where_clause = f"WHERE {' AND '.join(conditions)}" if conditions else ""

        try:
            conn = _open_email_capture_db()
            try:
                rows = conn.execute(
                    f"""
                    SELECT * FROM captured_emails
                    {where_clause}
                    ORDER BY received_at_utc DESC
                    LIMIT ? OFFSET ?
                    """,
                    (*params, limit, offset),
                ).fetchall()
                counts_row = conn.execute(
                    """
                    SELECT
                        SUM(CASE WHEN status = 'new' THEN 1 ELSE 0 END) AS new_count,
                        SUM(CASE WHEN status = 'new' AND directedness IN ('to', 'named')
                            THEN 1 ELSE 0 END) AS new_directed,
                        SUM(CASE WHEN status = 'new' AND directedness IN ('cc', 'unknown')
                            THEN 1 ELSE 0 END) AS new_cc,
                        SUM(CASE WHEN status = 'new' AND ai_status = 'pending'
                            THEN 1 ELSE 0 END) AS pending_ai
                    FROM captured_emails
                    """
                ).fetchone()
                return {
                    "status": "success",
                    "items": [_normalize_captured_email_record(row) for row in rows],
                    "counts": {
                        "new": int(counts_row["new_count"] or 0),
                        "newDirected": int(counts_row["new_directed"] or 0),
                        "newCc": int(counts_row["new_cc"] or 0),
                        "pendingAi": int(counts_row["pending_ai"] or 0),
                    },
                }
            finally:
                conn.close()
        except Exception as exc:
            logging.error(f"Error listing captured emails: {exc}")
            return {"status": "error", "message": str(exc)}

    def update_captured_email_status(self, payload):
        request = payload if isinstance(payload, dict) else {}
        try:
            record_id = int(request.get("id"))
        except (TypeError, ValueError):
            return {"status": "error", "message": "Missing captured email id."}
        new_status = str(request.get("status") or "").strip().lower()
        allowed_transitions = {
            ("new", "accepted"),
            ("new", "dismissed"),
            ("dismissed", "new"),
        }

        try:
            conn = _open_email_capture_db()
            try:
                row = conn.execute(
                    "SELECT * FROM captured_emails WHERE id = ?",
                    (record_id,),
                ).fetchone()
                if row is None:
                    return {"status": "error", "message": "Captured email not found."}
                current_status = str(row["status"] or "new")
                if (current_status, new_status) not in allowed_transitions:
                    return {
                        "status": "error",
                        "message": f"Cannot change email status from '{current_status}' to '{new_status}'.",
                    }
                conn.execute(
                    """
                    UPDATE captured_emails
                    SET status = ?, status_changed_at_utc = ?,
                        accepted_project_id = ?, accepted_deliverable_id = ?
                    WHERE id = ?
                    """,
                    (
                        new_status,
                        utc_now_iso(),
                        str(request.get("acceptedProjectId") or "").strip(),
                        str(request.get("acceptedDeliverableId") or "").strip(),
                        record_id,
                    ),
                )
                conn.commit()
                updated = conn.execute(
                    "SELECT * FROM captured_emails WHERE id = ?",
                    (record_id,),
                ).fetchone()
                return {
                    "status": "success",
                    "item": _normalize_captured_email_record(updated),
                }
            finally:
                conn.close()
        except Exception as exc:
            logging.error(f"Error updating captured email status: {exc}")
            return {"status": "error", "message": str(exc)}

    def _normalize_follow_up_date(self, value):
        raw = str(value or "").strip()
        if not raw:
            return ""
        try:
            return datetime.date.fromisoformat(raw).isoformat()
        except ValueError:
            return ""

    def create_follow_up(self, payload):
        request = payload if isinstance(payload, dict) else {}
        kind = str(request.get("kind") or "email").strip().lower()
        if kind not in {"email", "deliverable", "custom"}:
            kind = "email"
        captured_email_id = request.get("capturedEmailId")
        try:
            captured_email_id = int(captured_email_id) if captured_email_id is not None else None
        except (TypeError, ValueError):
            captured_email_id = None
        title = str(request.get("title") or "").strip()

        try:
            conn = _open_email_capture_db()
            try:
                if captured_email_id is not None:
                    email_row = conn.execute(
                        "SELECT subject FROM captured_emails WHERE id = ?",
                        (captured_email_id,),
                    ).fetchone()
                    if email_row is None:
                        return {"status": "error", "message": "Captured email not found."}
                    if not title:
                        title = str(email_row["subject"] or "").strip()
                if not title:
                    return {"status": "error", "message": "Follow-up needs a title."}
                cursor = conn.execute(
                    """
                    INSERT INTO follow_ups (
                        kind, captured_email_id, project_id, deliverable_id,
                        title, note, waiting_on, due_date, created_at_utc
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        kind,
                        captured_email_id,
                        str(request.get("projectId") or "").strip(),
                        str(request.get("deliverableId") or "").strip(),
                        title,
                        str(request.get("note") or "").strip(),
                        str(request.get("waitingOn") or "").strip(),
                        self._normalize_follow_up_date(request.get("dueDate")),
                        utc_now_iso(),
                    ),
                )
                conn.commit()
                row = conn.execute(
                    "SELECT * FROM follow_ups WHERE id = ?",
                    (cursor.lastrowid,),
                ).fetchone()
                return {"status": "success", "item": _normalize_follow_up_record(row)}
            finally:
                conn.close()
        except Exception as exc:
            logging.error(f"Error creating follow-up: {exc}")
            return {"status": "error", "message": str(exc)}

    def list_follow_ups(self, payload=None):
        request = payload if isinstance(payload, dict) else {}
        status_filter = str(request.get("status") or "open").strip().lower()
        if status_filter not in {"open", "resolved", "all"}:
            status_filter = "open"

        conditions = ""
        params = []
        if status_filter != "all":
            conditions = "WHERE f.status = ?"
            params.append(status_filter)

        try:
            conn = _open_email_capture_db()
            try:
                rows = conn.execute(
                    f"""
                    SELECT f.*,
                        e.subject AS email_subject,
                        e.entry_id AS email_entry_id,
                        e.sender_name AS email_sender_name
                    FROM follow_ups f
                    LEFT JOIN captured_emails e ON e.id = f.captured_email_id
                    {conditions}
                    ORDER BY f.created_at_utc DESC
                    """,
                    params,
                ).fetchall()
                items = [_normalize_follow_up_record(row) for row in rows]
                items.sort(
                    key=lambda item: (
                        item["status"] != "open",
                        _follow_up_effective_due(item) == "",
                        _follow_up_effective_due(item),
                        item["createdAtUtc"],
                    )
                )
                return {"status": "success", "items": items}
            finally:
                conn.close()
        except Exception as exc:
            logging.error(f"Error listing follow-ups: {exc}")
            return {"status": "error", "message": str(exc)}

    def update_follow_up(self, payload):
        request = payload if isinstance(payload, dict) else {}
        try:
            follow_up_id = int(request.get("id"))
        except (TypeError, ValueError):
            return {"status": "error", "message": "Missing follow-up id."}
        action = str(request.get("action") or "").strip().lower()

        try:
            conn = _open_email_capture_db()
            try:
                row = conn.execute(
                    "SELECT * FROM follow_ups WHERE id = ?",
                    (follow_up_id,),
                ).fetchone()
                if row is None:
                    return {"status": "error", "message": "Follow-up not found."}

                if action == "resolve":
                    conn.execute(
                        "UPDATE follow_ups SET status = 'resolved', resolved_at_utc = ? WHERE id = ?",
                        (utc_now_iso(), follow_up_id),
                    )
                elif action == "reopen":
                    conn.execute(
                        "UPDATE follow_ups SET status = 'open', resolved_at_utc = '' WHERE id = ?",
                        (follow_up_id,),
                    )
                elif action == "snooze":
                    try:
                        days = max(1, int(request.get("days") or 3))
                    except (TypeError, ValueError):
                        days = 3
                    snoozed_until = (
                        datetime.date.today() + datetime.timedelta(days=days)
                    ).isoformat()
                    conn.execute(
                        "UPDATE follow_ups SET snoozed_until = ? WHERE id = ?",
                        (snoozed_until, follow_up_id),
                    )
                else:
                    updates = []
                    params = []
                    if "title" in request:
                        title = str(request.get("title") or "").strip()
                        if not title:
                            return {"status": "error", "message": "Follow-up needs a title."}
                        updates.append("title = ?")
                        params.append(title)
                    if "note" in request:
                        updates.append("note = ?")
                        params.append(str(request.get("note") or "").strip())
                    if "waitingOn" in request:
                        updates.append("waiting_on = ?")
                        params.append(str(request.get("waitingOn") or "").strip())
                    if "dueDate" in request:
                        updates.append("due_date = ?")
                        params.append(self._normalize_follow_up_date(request.get("dueDate")))
                        updates.append("snoozed_until = ''")
                    if not updates:
                        return {"status": "error", "message": "No follow-up changes provided."}
                    params.append(follow_up_id)
                    conn.execute(
                        f"UPDATE follow_ups SET {', '.join(updates)} WHERE id = ?",
                        params,
                    )
                conn.commit()
                updated = conn.execute(
                    "SELECT * FROM follow_ups WHERE id = ?",
                    (follow_up_id,),
                ).fetchone()
                return {"status": "success", "item": _normalize_follow_up_record(updated)}
            finally:
                conn.close()
        except Exception as exc:
            logging.error(f"Error updating follow-up: {exc}")
            return {"status": "error", "message": str(exc)}

    def get_email_capture_badge_counts(self):
        try:
            conn = _open_email_capture_db()
            try:
                email_counts = conn.execute(
                    """
                    SELECT
                        SUM(CASE WHEN status = 'new' AND directedness IN ('to', 'named')
                            THEN 1 ELSE 0 END) AS new_directed,
                        SUM(CASE WHEN status = 'new' AND directedness IN ('cc', 'unknown')
                            THEN 1 ELSE 0 END) AS new_cc
                    FROM captured_emails
                    """
                ).fetchone()
                follow_up_rows = conn.execute(
                    "SELECT * FROM follow_ups WHERE status = 'open'"
                ).fetchall()
            finally:
                conn.close()
            today = datetime.date.today().isoformat()
            open_follow_ups = len(follow_up_rows)
            overdue = 0
            for row in follow_up_rows:
                effective_due = _follow_up_effective_due(_normalize_follow_up_record(row))
                if effective_due and effective_due < today:
                    overdue += 1
            return {
                "status": "success",
                "newDirected": int(email_counts["new_directed"] or 0),
                "newCc": int(email_counts["new_cc"] or 0),
                "openFollowUps": open_follow_ups,
                "overdueFollowUps": overdue,
            }
        except Exception as exc:
            logging.error(f"Error loading email capture badge counts: {exc}")
            return {"status": "error", "message": str(exc)}

    def get_email_capture_scan_state(self):
        try:
            conn = _open_email_capture_db()
            try:
                last_scan = _get_email_capture_meta(conn, "last_scan_completed_at_utc")
                last_status = _get_email_capture_meta(conn, "last_scan_status")
                summary = _parse_email_capture_json(
                    _get_email_capture_meta(conn, "last_scan_summary_json"), None
                )
            finally:
                conn.close()
            with self._email_capture_scan_lock:
                running = self._email_capture_scan_running
            return {
                "status": "success",
                "running": running,
                "lastScanCompletedAtUtc": last_scan,
                "lastScanStatus": last_status,
                "lastScanSummary": summary if isinstance(summary, dict) else None,
            }
        except Exception as exc:
            logging.error(f"Error loading email capture scan state: {exc}")
            return {"status": "error", "message": str(exc)}

    def get_installed_autocad_versions(self):
        """Scans for installed AutoCAD versions in the typical directory."""
        versions = []
        base_dir = r"C:\Program Files\Autodesk"
        if not os.path.exists(base_dir):
            return {'status': 'success', 'versions': []}

        for year in range(2022, 2027):  # 2022 to 2026
            folder_name = f"AutoCAD {year}"
            folder_path = os.path.join(base_dir, folder_name)
            exe_path = os.path.join(folder_path, "accoreconsole.exe")
            if os.path.exists(exe_path):
                versions.append({'year': year, 'path': exe_path})

        return {'status': 'success', 'versions': versions}

    def _ensure_aiohttp(self):
        """
        Validate that aiohttp is available with ClientSession. This prevents
        cryptic AttributeErrors during GenAI calls when the dependency is
        missing or only partially installed.
        """
        try:
            import aiohttp  # type: ignore
        except Exception as exc:
            raise RuntimeError(
                "Missing dependency: aiohttp. Install/upgrade it (e.g. "
                "pip install \"aiohttp>=3.13\") and rebuild the app."
            ) from exc

        if not hasattr(aiohttp, "ClientSession"):
            raise RuntimeError(
                "aiohttp is installed but ClientSession is unavailable. "
                "Reinstall it with `pip install --force-reinstall aiohttp>=3.13`."
            )

    def get_chat_response(self, chat_history):
        settings = self.get_user_settings()
        api_key = settings.get('apiKey', '').strip()

        if not api_key:
            api_key = (os.getenv("GOOGLE_API_KEY") or "").strip()

        if not api_key:
            return {
                'status': 'error',
                'response': 'API key not found. Please configure it in the AI settings.'
            }

        try:
            # Validate input
            if not isinstance(chat_history, list):
                return {
                    'status': 'error',
                    'response': 'Invalid chat history format.'
                }

            if len(chat_history) == 0:
                return {
                    'status': 'error',
                'response': 'Chat history is empty.'
            }

            self._ensure_aiohttp()
            client = genai.Client(api_key=api_key)
            model = "gemini-3-flash-preview"

            # Convert JavaScript chat history to Gemini API format
            contents = []
            for msg in chat_history:
                role = msg.get('role', 'user')
                # Map 'model' role to 'assistant' if needed
                if role == 'model':
                    role = 'model'  # Gemini uses 'model' for assistant responses

                parts = msg.get('parts', [])
                if not parts:
                    continue

                # Extract text from parts
                text_parts = []
                for part in parts:
                    if isinstance(part, dict) and 'text' in part:
                        text_parts.append(
                            types.Part.from_text(text=part['text']))
                    elif isinstance(part, str):
                        text_parts.append(types.Part.from_text(text=part))

                if text_parts:
                    contents.append(types.Content(role=role, parts=text_parts))

            if not contents:
                return {
                    'status': 'error',
                    'response': 'No valid messages found in chat history.'
                }

            tools = [types.Tool(google_search=types.GoogleSearch())]

            generate_content_config = types.GenerateContentConfig(
                temperature=0,
                top_p=0.7,
                thinking_config=types.ThinkingConfig(
                    thinking_budget=-1,
                ),
                tools=tools,
            )

            response = client.models.generate_content(
                model=model,
                contents=contents,
                config=generate_content_config,
            )

            return {'status': 'success', 'response': response.text}

        except Exception as e:
            msg = str(e)
            lower = msg.lower()
            if ("api key expired" in lower or
                "api_key_invalid" in lower or
                    "invalid api key" in lower):
                return {'status': 'error',
                        'response': ('Your Google API key is expired/invalid. '
                                     'Create a new key in Google AI Studio, '
                                     'update your settings, then try again.')}
            logging.error(f"Error getting chat response from AI: {e}")
            return {'status': 'error', 'response': f"An AI error occurred: {msg}"}

    def _resolve_google_ai_api_key(self, api_key):
        final_api_key = (api_key or "").strip()
        if not final_api_key:
            final_api_key = (os.getenv("GOOGLE_API_KEY") or "").strip()
        if not final_api_key:
            raise RuntimeError(
                "AI API key is not configured. Please provide it in the app settings or set GOOGLE_API_KEY in your .env file."
            )
        return final_api_key

    def _build_email_analysis_prompt(self, email_text, user_name, discipline, project_context=None):
        current_date = datetime.date.today().strftime("%m/%d/%Y")
        disciplines_str = ', '.join(discipline) if isinstance(
            discipline, list) else (discipline or 'Engineering')
        task_notes_contract = self._build_deliverable_task_notes_prompt_contract(
            disciplines_str,
            include_json_field_names=False,
        )
        known_projects = self._normalize_email_intake_project_context(project_context)
        known_projects_section = ""
        if known_projects:
            known_projects_json = json.dumps(
                known_projects,
                ensure_ascii=True,
                separators=(",", ":"),
            )
            known_projects_section = f"""
Known existing projects are provided below as JSON.
- Use KNOWN_PROJECTS only when the email evidence clearly supports a match.
- If the email mentions a known project id/job number, exact or close project name, nickname, project path, path leaf, client, address, building, or location that clearly matches one entry, return that known project's exact "id", "name", and "path".
- Prefer exact known-project identity over composing a new project name.
- Use "nick" and "pathLeaf" only as matching evidence; do not add extra JSON keys.
- If multiple known projects could match and the evidence is ambiguous, do not force a known project; extract the best supported fields from the email instead.

KNOWN_PROJECTS:
{known_projects_json}
""".strip()
        known_projects_text = f"\n{known_projects_section}\n" if known_projects_section else ""
        return f"""
You are an intelligent assistant for {user_name}, a(n) {disciplines_str} engineering project manager. Your task is to analyze an email and extract specific project details. Focus ONLY on the primary {disciplines_str} engineering tasks mentioned. Ignore tasks for other disciplines.
Analyze the following email text and extract the information into a valid JSON object with the following keys: "id", "name", "due", "path", "deliverable", "tasks", "notes".
{known_projects_text}
- "id": Find a project number or project ID (e.g., "250597", "P-12345", "Job #1042"). Look in the subject line, headers, and body. This could be called a job number, project number, project ID, or similar. If none, leave it empty.
- "name": Determine the project name, typically including the client and address or building name (e.g., "BofA, 22004 Sherman Way, Canoga Park, CA"). Include enough detail to uniquely identify the project. If no formal name is found and no known project is clearly supported, compose one from the client name and location mentioned.
- "due": Find the due date and format it as "MM/DD/YY". The current date is {current_date}. If the year is not specified in the email, assume the current year or the next year if the date would be in the past. Ensure the due date is on or after today. If multiple dates, choose the most relevant upcoming one.
- "path": Find the main project file path (e.g., "M:\\\\Gensler\\\\...").
- "deliverable": Infer the deliverable name from the email if possible. Prefer concise, standardized names when present (for example: DD60, DD90, CD60, CD90, CD100, CDF, RFI, RFI #2, Submittal, Lighting Submittal, Controls Submittal, Record Set, Record Drawings, IFP, Site Survey, Survey Report, ASR, ASR #2, PCC, PCC #3, Bulletin #2, Coordination, Meeting, Revision). If no deliverable is clear, leave it empty.
{task_notes_contract}
If a piece of information is not found, the value should be an empty string "" for strings, or an empty array [] for tasks.
Here is the email:
---
{email_text}
---
Return ONLY the JSON object.
""".strip()

    def _normalize_email_project_data(self, project_data):
        if isinstance(project_data, list):
            project_data = project_data[0] if project_data else {}
        if not isinstance(project_data, dict):
            project_data = {}
        project_data.setdefault("id", "")
        project_data.setdefault("name", "")
        project_data.setdefault("due", "")
        project_data.setdefault("path", "")
        project_data.setdefault("deliverable", "")
        project_data.setdefault("tasks", [])
        project_data.setdefault("notes", "")
        normalized_tasks = []
        if isinstance(project_data.get("tasks"), list):
            for task in project_data["tasks"]:
                if isinstance(task, dict):
                    text = str(task.get("text") or "").strip()
                    done = bool(task.get("done"))
                    links = task.get("links") if isinstance(task.get("links"), list) else []
                else:
                    text = str(task).strip()
                    done = False
                    links = []
                if not text:
                    continue
                normalized_tasks.append({
                    "text": text,
                    "done": done,
                    "links": links,
                })
        project_data["tasks"] = normalized_tasks
        return project_data

    def _normalize_email_intake_project_context(self, raw_context):
        if not isinstance(raw_context, list):
            return []

        normalized = []
        for raw_project in raw_context:
            if not isinstance(raw_project, dict):
                continue
            path = str(raw_project.get("path") or "").strip()
            path_leaf = str(
                raw_project.get("pathLeaf")
                or raw_project.get("path_leaf")
                or ""
            ).strip()
            if not path_leaf and path:
                path_leaf = re.split(r"[\\/]+", path.rstrip("\\/"))[-1].strip()
            project = {
                "id": str(raw_project.get("id") or "").strip(),
                "name": str(raw_project.get("name") or "").strip(),
                "nick": str(raw_project.get("nick") or "").strip(),
                "path": path,
                "pathLeaf": path_leaf,
            }
            if any(project.values()):
                normalized.append(project)

        normalized.sort(
            key=lambda project: (
                str(project.get("id") or "").lower(),
                str(project.get("name") or "").lower(),
                str(project.get("nick") or "").lower(),
                str(project.get("path") or "").lower(),
                str(project.get("pathLeaf") or "").lower(),
            )
        )

        capped = []
        serialized_chars = 2
        for project in normalized:
            if len(capped) >= EMAIL_INTAKE_PROJECT_CONTEXT_MAX_PROJECTS:
                break
            serialized = json.dumps(project, ensure_ascii=True, separators=(",", ":"))
            next_chars = serialized_chars + len(serialized) + (1 if capped else 0)
            if capped and next_chars > EMAIL_INTAKE_PROJECT_CONTEXT_BUDGET_CHARS:
                break
            if not capped and next_chars > EMAIL_INTAKE_PROJECT_CONTEXT_BUDGET_CHARS:
                continue
            capped.append(project)
            serialized_chars = next_chars
        return capped

    def _build_deliverable_task_notes_prompt_contract(
        self,
        disciplines_str,
        include_json_field_names=True,
    ):
        tasks_label = '"tasks"' if include_json_field_names else 'tasks'
        notes_label = '"notes"' if include_json_field_names else 'notes'
        return "\n".join([
            f'- {tasks_label}: Create a JSON array of strings listing every actionable {disciplines_str} engineering step needed to complete the deliverable. Include revisions, submissions, coordination, and follow-up work. Be concise. Example: ["Revise lighting plans per architect comments", "Update panel schedules", "Issue updated PDF set for resubmittal"].',
            f'- {notes_label}: Provide only non-actionable context the user should note, such as constraints, approvals, dependencies, reminders, delivery method, or special instructions. Do NOT put action items, next steps, or completion work in {notes_label}. Example: "Awaiting architect approval before final issue." If there is no notable non-actionable context, return "".',
        ])

    def _extract_project_data_from_email_text(
        self,
        email_text,
        api_key,
        user_name,
        discipline,
        project_context=None,
    ):
        final_api_key = self._resolve_google_ai_api_key(api_key)
        prompt = self._build_email_analysis_prompt(
            email_text,
            user_name,
            discipline,
            project_context,
        )
        try:
            self._ensure_aiohttp()
            client = genai.Client(
                api_key=final_api_key,
                http_options=types.HttpOptions(timeout=120000),
            )
            response = client.models.generate_content(
                model="gemini-3-flash-preview",
                contents=[
                    types.Content(
                        role="user",
                        parts=[types.Part.from_text(text=prompt)],
                    ),
                ],
                config=types.GenerateContentConfig(
                    temperature=0,
                    response_mime_type="application/json",
                ),
            )

            logging.debug(f"AI response object: {response}")
            raw_text = response.text
            if raw_text is None:
                if hasattr(response, 'parts') and response.parts:
                    raw_text = ''.join(
                        part.text for part in response.parts
                        if hasattr(part, 'text') and part.text
                    )
                if not raw_text:
                    logging.error(f"AI returned empty response. Full response: {response}")
                    raise RuntimeError('AI returned an empty response. Please try again.')

            cleaned = raw_text.strip()
            if not cleaned:
                raise RuntimeError('AI returned an empty response. Please try again.')

            logging.debug(f"AI response text: {cleaned[:500]}...")
            return self._normalize_email_project_data(json.loads(cleaned))
        except json.JSONDecodeError as e:
            logging.error(f"Failed to parse AI response as JSON: {e}")
            raise RuntimeError('AI returned invalid JSON. Please try again.')
        except Exception as e:
            msg = str(e)
            lower = msg.lower()
            logging.error(f"Error processing email with AI: {type(e).__name__}: {e}")
            if ("api key expired" in lower or
                "api_key_invalid" in lower or
                    "invalid api key" in lower):
                raise RuntimeError(
                    'Your Google API key is expired/invalid. '
                    'Create a new key in Google AI Studio, '
                    'update your settings, then try again.'
                )
            if "model" in lower and ("not found" in lower or "does not exist" in lower):
                raise RuntimeError(
                    'AI model not available. The Gemini 3 Flash model may not be accessible with your API key.'
                )
            if "quota" in lower or "rate limit" in lower:
                raise RuntimeError('API rate limit exceeded. Please wait a moment and try again.')
            raise RuntimeError(f"AI error: {msg}")

    def process_email_with_ai(self, email_text, api_key, user_name, discipline, project_context=None):
        """
        Processes email text using Google GenAI to extract project details.
        """
        try:
            project_data = self._extract_project_data_from_email_text(
                email_text,
                api_key,
                user_name,
                discipline,
                project_context,
            )
            return {'status': 'success', 'data': project_data}
        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    def get_tasks(self):
        """Reads and returns the content of tasks.json."""
        try:
            payload = _read_json_file_strict(TASKS_FILE)
            return _overlay_projects_with_lighting_schedule_records(payload)
        except FileNotFoundError:
            return []
        except ValueError as e:
            logging.error(f"Error loading tasks from {TASKS_FILE}: {e}")
            return []

    def save_tasks(self, data):
        """Saves data to tasks.json and creates a backup."""
        try:
            if os.path.exists(TASKS_FILE):
                shutil.copy2(TASKS_FILE, TASKS_FILE + '.bak')
            with open(TASKS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error saving tasks: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_lighting_schedule_sync(self, dwg_path):
        """Read project-specific light fixture schedule sync JSON near the DWG."""
        try:
            _, sync_path = _resolve_lighting_schedule_sync_path(dwg_path)
            exists = os.path.exists(sync_path)
            modified = None
            payload = None
            if exists:
                payload = _read_json_file_strict(sync_path)
                mtime = os.path.getmtime(sync_path)
                modified = datetime.datetime.fromtimestamp(mtime).isoformat()

            result = {
                'status': 'success',
                'exists': exists,
                'path': sync_path,
                'modified': modified,
            }
            if payload is not None:
                result['data'] = payload
            return result
        except ValueError as e:
            logging.warning(f"Lighting schedule sync read validation failed: {e}")
            return {'status': 'error', 'message': str(e)}
        except FileNotFoundError:
            return {'status': 'success', 'exists': False, 'path': '', 'modified': None}
        except Exception as e:
            logging.error(f"Error reading lighting schedule sync: {e}")
            return {'status': 'error', 'message': str(e)}

    def save_lighting_schedule_sync(self, dwg_path, payload):
        """Write project-specific light fixture schedule sync JSON near the DWG."""
        try:
            if not isinstance(payload, dict):
                raise ValueError("Sync payload must be a JSON object.")

            _, sync_path = _resolve_lighting_schedule_sync_path(dwg_path)
            sync_dir = os.path.dirname(sync_path)
            if not os.path.isdir(sync_dir):
                raise ValueError(f"DWG folder does not exist: {sync_dir}")

            try:
                json.dumps(payload)
            except (TypeError, ValueError) as exc:
                raise ValueError(f"Sync payload is not valid JSON: {exc}")

            _atomic_write_json_file(sync_path, payload)
            mtime = os.path.getmtime(sync_path)
            modified = datetime.datetime.fromtimestamp(mtime).isoformat()
            return {
                'status': 'success',
                'path': sync_path,
                'modified': modified,
            }
        except ValueError as e:
            logging.warning(f"Lighting schedule sync save validation failed: {e}")
            return {'status': 'error', 'message': str(e)}
        except Exception as e:
            logging.error(f"Error saving lighting schedule sync: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_lighting_schedule_record(self, project_id):
        """Load the canonical lighting schedule record from SQLite."""
        try:
            record = _get_lighting_schedule_record(project_id)
            return {
                'status': 'success',
                'exists': bool(record),
                'projectId': _resolve_lighting_schedule_project_id(project_id),
                'data': record,
            }
        except Exception as e:
            logging.error(f"Error loading lighting schedule record: {e}")
            return {'status': 'error', 'message': str(e)}

    def save_lighting_schedule_record(self, project_id, payload):
        """Persist the canonical lighting schedule record to SQLite."""
        try:
            safe_payload = payload if isinstance(payload, dict) else {}
            record = _save_lighting_schedule_record(
                project_id,
                safe_payload,
                updated_by=safe_payload.get('updatedBy', 'desktop'),
            )
            return {
                'status': 'success',
                'projectId': record['projectId'],
                'data': record,
            }
        except Exception as e:
            logging.error(f"Error saving lighting schedule record: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_lighting_schedule_version(self, project_id):
        """Return only the current version metadata for polling."""
        try:
            return {
                'status': 'success',
                'data': _get_lighting_schedule_version(project_id),
            }
        except Exception as e:
            logging.error(f"Error loading lighting schedule version: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_lighting_schedule_links(self, project_id):
        """Return linked DWG/table records for a project."""
        try:
            return {
                'status': 'success',
                'projectId': _resolve_lighting_schedule_project_id(project_id),
                'data': _get_lighting_schedule_links(project_id),
            }
        except Exception as e:
            logging.error(f"Error loading lighting schedule links: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_lighting_plan_record(self, project_id):
        """Load the latest Phase 1 lighting-plan scan and analysis."""
        try:
            record = _get_lighting_plan_record(project_id)
            return {
                'status': 'success',
                'exists': bool(record),
                'projectId': _resolve_lighting_schedule_project_id(project_id),
                'data': record,
            }
        except Exception as e:
            logging.error(f"Error loading lighting plan record: {e}")
            return {'status': 'error', 'message': str(e)}

    def import_lighting_plan_snapshot(self, project_id, snapshot_path):
        """Import an AutoCAD lighting scan and run deterministic Phase 1 analysis."""
        try:
            schedule_record = _get_lighting_schedule_record(project_id)
            schedule_rows = (
                schedule_record.get('schedule', {}).get('rows', [])
                if isinstance(schedule_record, dict)
                else []
            )
            record = _import_lighting_plan_snapshot(
                project_id,
                snapshot_path,
                schedule_rows=schedule_rows,
            )
            return {
                'status': 'success',
                'projectId': record['projectId'],
                'data': record,
            }
        except (FileNotFoundError, ValueError, LightingPlanValidationError) as e:
            logging.warning(f"Lighting plan import validation failed: {e}")
            return {'status': 'error', 'message': str(e)}
        except Exception as e:
            logging.error(f"Error importing lighting plan snapshot: {e}")
            return {'status': 'error', 'message': str(e)}

    def export_lighting_plan_instructions(self, project_id):
        """Write reviewed tag instructions beside the imported AutoCAD snapshot."""
        try:
            result = _export_lighting_plan_instructions(project_id)
            return {
                'status': 'success',
                'projectId': _resolve_lighting_schedule_project_id(project_id),
                'path': result['path'],
                'tagCount': result['tagCount'],
                'data': result['record'],
            }
        except (ValueError, LightingPlanValidationError) as e:
            logging.warning(f"Lighting plan instruction export blocked: {e}")
            return {'status': 'error', 'message': str(e)}
        except Exception as e:
            logging.error(f"Error exporting lighting plan instructions: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_project_checklist_record(self, project_id):
        """Load the project-scoped checklist completion record from SQLite."""
        try:
            record = _get_project_checklist_record(project_id)
            return {
                'status': 'success',
                'exists': bool(record),
                'projectId': _normalize_project_checklist_project_id(project_id),
                'data': record,
            }
        except Exception as e:
            logging.error(f"Error loading project checklist record: {e}")
            return {'status': 'error', 'message': str(e)}

    def save_project_checklist_record(self, project_id, payload):
        """Persist project-scoped checklist completion state to SQLite."""
        try:
            safe_payload = payload if isinstance(payload, dict) else {}
            record = _save_project_checklist_record(
                project_id,
                safe_payload,
                updated_by=safe_payload.get('updatedBy', 'desktop'),
            )
            return {
                'status': 'success',
                'projectId': record['projectId'],
                'data': record,
            }
        except Exception as e:
            logging.error(f"Error saving project checklist record: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_project_checklist_version(self, project_id):
        """Return only project checklist version metadata for polling."""
        try:
            return {
                'status': 'success',
                'data': _get_project_checklist_version(project_id),
            }
        except Exception as e:
            logging.error(f"Error loading project checklist version: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_project_checklist_link(self, dwg_path):
        """Return the project checklist DWG-to-project link, if one exists."""
        try:
            link = _get_project_checklist_link(dwg_path)
            return {
                'status': 'success',
                'exists': bool(link),
                'data': link,
            }
        except Exception as e:
            logging.error(f"Error loading project checklist link: {e}")
            return {'status': 'error', 'message': str(e)}

    def save_project_checklist_link(self, project_id, dwg_path):
        """Persist a DWG-to-project link for project checklist lookups."""
        try:
            link = _save_project_checklist_link(project_id, dwg_path)
            return {
                'status': 'success',
                'data': link,
            }
        except Exception as e:
            logging.error(f"Error saving project checklist link: {e}")
            return {'status': 'error', 'message': str(e)}

    def read_t24_output_json(self, json_path):
        """Read and validate TXTSUMEXPORT output (T24Output.json)."""
        try:
            normalized_path = _normalize_t24_output_json_path(json_path)
            if not os.path.isfile(normalized_path):
                raise FileNotFoundError(
                    f"T24Output.json was not found at: {normalized_path}")

            payload = _read_json_file_strict(normalized_path)
            if not isinstance(payload, list):
                raise ValueError(
                    "T24Output.json must contain a JSON array of room entries.")

            rows = []
            for idx, item in enumerate(payload, start=1):
                if not isinstance(item, dict):
                    raise ValueError(
                        f"Entry #{idx} must be a JSON object with RoomType and SquareFeet.")

                room_type = str(item.get("RoomType", "")).strip()
                if not room_type:
                    raise ValueError(f"Entry #{idx} is missing RoomType.")

                square_feet_raw = item.get("SquareFeet", None)
                try:
                    square_feet = float(square_feet_raw)
                except (TypeError, ValueError):
                    raise ValueError(
                        f"Entry #{idx} has invalid SquareFeet value: {square_feet_raw!r}.")

                if not math.isfinite(square_feet) or square_feet < 0:
                    raise ValueError(
                        f"Entry #{idx} has non-finite or negative SquareFeet: {square_feet_raw!r}.")

                rows.append({
                    "roomType": room_type,
                    "squareFeet": round(square_feet, 4)
                })

            total_square_feet = round(
                sum(row["squareFeet"] for row in rows), 4)
            mtime = os.path.getmtime(normalized_path)
            modified = datetime.datetime.fromtimestamp(mtime).isoformat()

            return {
                "status": "success",
                "path": normalized_path,
                "modified": modified,
                "data": {
                    "rows": rows,
                    "totalSquareFeet": total_square_feet
                }
            }
        except FileNotFoundError as e:
            return {"status": "error", "message": str(e)}
        except ValueError as e:
            logging.warning(f"T24Output.json validation failed: {e}")
            return {"status": "error", "message": str(e)}
        except Exception as e:
            logging.error(f"Error reading T24Output.json: {e}")
            return {"status": "error", "message": str(e)}

    def get_notes(self):
        """Reads and returns the content of notes.json."""
        try:
            with open(NOTES_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            return {}

    def save_notes(self, data):
        """Saves notes data to notes.json and creates a backup."""
        try:
            if os.path.exists(NOTES_FILE):
                shutil.copy2(NOTES_FILE, NOTES_FILE + '.bak')
            with open(NOTES_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error saving notes: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_timesheets(self):
        """Reads and returns the content of timesheets.json."""
        try:
            with open(TIMESHEETS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict) and not isinstance(data.get('expenses'), dict):
                    data['expenses'] = {}
                return data
        except (FileNotFoundError, json.JSONDecodeError):
            return {'weeks': {}, 'expenses': {}, 'lastModified': None}

    def save_timesheets(self, data):
        """Saves timesheets data to timesheets.json and creates a backup."""
        try:
            if os.path.exists(TIMESHEETS_FILE):
                shutil.copy2(TIMESHEETS_FILE, TIMESHEETS_FILE + '.bak')
            with open(TIMESHEETS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error saving timesheets: {e}")
            return {'status': 'error', 'message': str(e)}

    # ===================== TEMPLATES API =====================

    def get_templates(self):
        """Reads and returns templates data, installing defaults on first run."""
        try:
            with open(TEMPLATES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Check if defaults need installation
            if not data.get('defaultTemplatesInstalled'):
                data = self._install_default_templates(data)

            data = self._ensure_default_templates(data)
            return data
        except (FileNotFoundError, json.JSONDecodeError):
            # First run - create initial structure with defaults
            data = {'templates': [], 'defaultTemplatesInstalled': False, 'lastModified': None}
            data = self._install_default_templates(data)
            data = self._ensure_default_templates(data)
            return data

    def _install_default_templates(self, data):
        """Install default templates on first run."""
        for tpl_def in DEFAULT_TEMPLATES:
            if os.path.exists(tpl_def['sourcePath']):
                template = {
                    'id': f"tpl_{os.urandom(6).hex()}",
                    'name': tpl_def['name'],
                    'discipline': tpl_def['discipline'],
                    'fileType': tpl_def['fileType'],
                    'sourcePath': tpl_def['sourcePath'],
                    'isDefault': True,
                    'dateAdded': datetime.datetime.now().isoformat(),
                    'description': tpl_def.get('description', '')
                }
                data['templates'].append(template)

        data['defaultTemplatesInstalled'] = True
        data['lastModified'] = datetime.datetime.now().isoformat()
        self.save_templates(data)
        return data

    def _ensure_default_templates(self, data):
        """Ensure bundled default templates exist and update paths if needed."""
        templates = data.get('templates') or []
        updated = False
        for tpl_def in DEFAULT_TEMPLATES:
            name_key = str(tpl_def.get('name', '')).strip().lower()
            expected_path = tpl_def.get('sourcePath')
            existing = next(
                (
                    t for t in templates
                    if t.get('isDefault') and str(t.get('name', '')).strip().lower() == name_key
                ),
                None
            )

            if existing:
                existing_path = str(existing.get('sourcePath') or '')
                expected_abs = os.path.abspath(expected_path) if expected_path else ''
                existing_abs = os.path.abspath(existing_path) if existing_path else ''
                if expected_abs and (not os.path.exists(existing_path) or existing_abs != expected_abs):
                    existing['sourcePath'] = expected_path
                    existing['fileType'] = tpl_def.get('fileType')
                    existing['discipline'] = tpl_def.get('discipline')
                    existing['description'] = tpl_def.get('description', '')
                    updated = True
                continue

            if expected_path and os.path.exists(expected_path):
                template = {
                    'id': f"tpl_{os.urandom(6).hex()}",
                    'name': tpl_def['name'],
                    'discipline': tpl_def['discipline'],
                    'fileType': tpl_def['fileType'],
                    'sourcePath': expected_path,
                    'isDefault': True,
                    'dateAdded': datetime.datetime.now().isoformat(),
                    'description': tpl_def.get('description', '')
                }
                templates.append(template)
                updated = True

        if updated:
            data['templates'] = templates
            data['defaultTemplatesInstalled'] = True
            data['lastModified'] = datetime.datetime.now().isoformat()
            self.save_templates(data)
        return data

    def save_templates(self, data):
        """Saves templates data with backup."""
        try:
            if os.path.exists(TEMPLATES_FILE):
                shutil.copy2(TEMPLATES_FILE, TEMPLATES_FILE + '.bak')
            with open(TEMPLATES_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error saving templates: {e}")
            return {'status': 'error', 'message': str(e)}

    def _normalize_template_key(self, template_key):
        raw = str(template_key or '').strip().lower()
        compact = re.sub(r'[^a-z0-9]+', '', raw)
        alias_map = {
            'narrative': 'narrative',
            'narrativeofchanges': 'narrative',
            'plancheck': 'planCheck',
            'plancheckcomments': 'planCheck',
            'plancheckresponseletter': 'planCheck',
            'pcc': 'planCheck',
        }
        if raw in TEMPLATE_KEY_BY_NAME:
            return TEMPLATE_KEY_BY_NAME[raw]
        if raw in alias_map:
            return alias_map[raw]
        return alias_map.get(compact, '')

    def _infer_template_key(self, template):
        if not isinstance(template, dict):
            return ''
        name_key = str(template.get('name') or '').strip().lower()
        if name_key in TEMPLATE_KEY_BY_NAME:
            return TEMPLATE_KEY_BY_NAME[name_key]
        source_name = os.path.basename(str(template.get('sourcePath') or '')).lower()
        if 'narrative of changes' in source_name:
            return 'narrative'
        if 'plan check' in source_name or 'pcc' in source_name:
            return 'planCheck'
        return ''

    def _find_template_by_key(self, template_key):
        normalized_key = self._normalize_template_key(template_key)
        if not normalized_key:
            return None, ''
        templates_data = self.get_templates()
        templates = templates_data.get('templates', []) if isinstance(
            templates_data, dict) else []
        for template in templates:
            if self._infer_template_key(template) == normalized_key:
                return template, normalized_key
        return None, normalized_key

    def _sanitize_template_filename_stem(self, value, fallback):
        text = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", str(value or '').strip())
        text = re.sub(r'\s+', ' ', text).strip().strip('.')
        text = text[:120]
        return text or fallback

    def _get_template_output_name_parts(self, template, template_key='', context=None, new_name=None):
        normalized_key = self._normalize_template_key(
            template_key or self._infer_template_key(template)
        )
        source_path = str(template.get('sourcePath') or '').strip()
        extension = os.path.splitext(source_path)[1]
        if not extension:
            ext_hint = str(template.get('fileType') or '').strip().lstrip('.')
            extension = f'.{ext_hint}' if ext_hint else ''

        if new_name is not None and str(new_name).strip():
            base_name = str(new_name).strip()
            if extension and base_name.lower().endswith(extension.lower()):
                base_name = base_name[:-len(extension)]
            return base_name, extension, normalized_key

        if normalized_key in TEMPLATE_DEFAULT_FILENAME_BY_KEY:
            return TEMPLATE_DEFAULT_FILENAME_BY_KEY[normalized_key], extension, normalized_key

        fallback_name = (
            str(template.get('name') or '').strip()
            or Path(source_path).stem
            or 'Template'
        )
        return fallback_name, extension, normalized_key

    def _resolve_template_destination_path(self, folder_path, base_name, extension, conflict_policy='timestamp'):
        policy = str(conflict_policy or 'timestamp').strip().lower()
        safe_base = self._sanitize_template_filename_stem(base_name, 'Template')
        ext = str(extension or '').strip()
        if ext and not ext.startswith('.'):
            ext = f'.{ext}'
        candidate_path = os.path.join(folder_path, f'{safe_base}{ext}')
        if not os.path.exists(candidate_path):
            return candidate_path, os.path.basename(candidate_path), False
        if policy == 'overwrite':
            return candidate_path, os.path.basename(candidate_path), False
        if policy != 'timestamp':
            return candidate_path, os.path.basename(candidate_path), True

        timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H%M')
        suffix_index = 0
        while True:
            suffix = f' - {timestamp}' if suffix_index == 0 else f' - {timestamp} ({suffix_index + 1})'
            candidate_name = self._sanitize_template_filename_stem(
                f'{safe_base}{suffix}', safe_base)
            candidate_path = os.path.join(folder_path, f'{candidate_name}{ext}')
            if not os.path.exists(candidate_path):
                return candidate_path, os.path.basename(candidate_path), False
            suffix_index += 1

    def add_template(self, name, discipline, source_path, description=''):
        """Adds a new template to the collection."""
        try:
            if not os.path.exists(source_path):
                return {'status': 'error', 'message': 'Source file does not exist'}

            # Detect file type
            ext = os.path.splitext(source_path)[1].lower().lstrip('.')
            if ext not in ['doc', 'docx', 'dwg', 'xlsx', 'xls']:
                return {'status': 'error', 'message': f'Unsupported file type: .{ext}'}

            data = self.get_templates()
            template = {
                'id': f"tpl_{os.urandom(6).hex()}",
                'name': name,
                'discipline': discipline,
                'fileType': ext,
                'sourcePath': source_path,
                'isDefault': False,
                'dateAdded': datetime.datetime.now().isoformat(),
                'description': description
            }
            data['templates'].append(template)
            data['lastModified'] = datetime.datetime.now().isoformat()
            self.save_templates(data)
            return {'status': 'success', 'template': template}
        except Exception as e:
            logging.error(f"Error adding template: {e}")
            return {'status': 'error', 'message': str(e)}

    def remove_template(self, template_id):
        """Removes a template by ID."""
        try:
            data = self.get_templates()
            original_count = len(data['templates'])
            data['templates'] = [t for t in data['templates'] if t['id'] != template_id]

            if len(data['templates']) == original_count:
                return {'status': 'error', 'message': 'Template not found'}

            data['lastModified'] = datetime.datetime.now().isoformat()
            self.save_templates(data)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error removing template: {e}")
            return {'status': 'error', 'message': str(e)}

    def copy_template_to_folder(
        self,
        template_id,
        destination_folder,
        new_name=None,
        context=None,
        options=None,
        conflict_policy='timestamp',
    ):
        """Copies a template file to a folder, with optional open-after-save behavior."""
        try:
            data = self.get_templates()
            template = next((t for t in data['templates'] if t['id'] == template_id), None)

            if not template:
                return {'status': 'error', 'message': 'Template not found'}

            source = template['sourcePath']
            if not os.path.exists(source):
                return {'status': 'error', 'message': f'Template file not found: {source}'}

            if isinstance(destination_folder, (list, tuple)):
                destination_folder = destination_folder[0] if destination_folder else ''

            destination_folder = os.path.normpath(str(destination_folder or '').strip())
            if not destination_folder:
                return {'status': 'error', 'message': 'Destination folder is required'}

            options_payload = dict(options or {}) if isinstance(options, dict) else {}
            context_payload = dict(context or {}) if isinstance(context, dict) else {}

            if not os.path.isdir(destination_folder):
                create_destination = False
                if options_payload:
                    create_destination = bool(options_payload.get('createDestination'))
                if create_destination:
                    os.makedirs(destination_folder, exist_ok=True)
                else:
                    return {'status': 'error', 'message': 'Destination folder does not exist'}

            base_name, extension, normalized_key = self._get_template_output_name_parts(
                template,
                template_key=options_payload.get('templateKey'),
                context=context_payload,
                new_name=new_name,
            )
            if not extension:
                return {'status': 'error', 'message': 'Template file extension could not be determined.'}

            dest_path, filename, has_conflict = self._resolve_template_destination_path(
                destination_folder,
                base_name,
                extension,
                conflict_policy=conflict_policy,
            )
            if has_conflict:
                return {'status': 'error', 'message': 'File already exists at destination'}

            shutil.copy2(source, dest_path)
            if context_payload or options_payload:
                template_options = dict(options_payload)
                if normalized_key and 'templateKey' not in template_options:
                    template_options['templateKey'] = normalized_key
                _apply_template_context(
                    dest_path,
                    context=context_payload,
                    options=template_options,
                )

            result = {
                'status': 'success',
                'path': dest_path,
                'filename': filename,
                'outputFilePath': dest_path,
                'outputFolderPath': self._find_project_root_by_id(destination_folder) or destination_folder,
            }

            return result
        except Exception as e:
            logging.error(f"Error copying template: {e}")
            return {'status': 'error', 'message': str(e)}

    def copy_template_to_path(self, template_id, destination_path, context=None, options=None):
        """Copies a template file to the specified file path."""
        try:
            data = self.get_templates()
            template = next((t for t in data['templates'] if t['id'] == template_id), None)

            if not template:
                return {'status': 'error', 'message': 'Template not found'}

            source = template['sourcePath']
            if not os.path.exists(source):
                return {'status': 'error', 'message': f'Template file not found: {source}'}

            if isinstance(destination_path, (list, tuple)):
                destination_path = destination_path[0] if destination_path else ''

            dest_path = str(destination_path or '').strip()
            if not dest_path:
                return {'status': 'error', 'message': 'Destination path is required'}

            ext = os.path.splitext(source)[1]
            if ext and not dest_path.lower().endswith(ext.lower()):
                dest_path = dest_path + ext

            dest_dir = os.path.dirname(dest_path)
            if not dest_dir:
                return {'status': 'error', 'message': 'Destination path is invalid'}
            if not os.path.isdir(dest_dir):
                os.makedirs(dest_dir, exist_ok=True)

            if os.path.exists(dest_path):
                return {'status': 'error', 'message': 'File already exists at destination'}

            shutil.copy2(source, dest_path)
            if context or options:
                _apply_template_context(dest_path, context=context, options=options)
            return {
                'status': 'success',
                'path': dest_path,
                'outputFilePath': dest_path,
                'outputFolderPath': dest_dir,
            }
        except Exception as e:
            logging.error(f"Error copying template to path: {e}")
            return {'status': 'error', 'message': str(e)}

    def create_template_for_workroom(self, template_key, launch_context=None, context=None, conflict_policy='timestamp'):
        """Auto-create a template in the active Workroom project's saved path."""
        try:
            settings = self.get_user_settings()
            workroom_context = self._resolve_workroom_context(settings, launch_context)
            source = str(workroom_context.get('source') or '').strip().lower()
            project_path = str(workroom_context.get('project_path') or '').strip()
            if source != 'workroom':
                logging.info(
                    "create_template_for_workroom: fallback to manual save dialog "
                    f"(source={source or 'none'})."
                )
                return {
                    'status': 'fallback',
                    'autoCreated': False,
                    'fallbackUsed': True,
                    'message': 'Template auto-create is only enabled for Workroom launches.',
                }
            if not project_path:
                logging.info(
                    "create_template_for_workroom: fallback to manual save dialog (missing project path)."
                )
                return {
                    'status': 'fallback',
                    'autoCreated': False,
                    'fallbackUsed': True,
                    'message': 'Saved project path is missing.',
                }
            if not os.path.isdir(project_path):
                logging.info(
                    "create_template_for_workroom: fallback to manual save dialog "
                    f"(project path is not a directory: {project_path})."
                )
                return {
                    'status': 'fallback',
                    'autoCreated': False,
                    'fallbackUsed': True,
                    'message': 'Saved project path is invalid.',
                }

            template, normalized_key = self._find_template_by_key(template_key)
            if not normalized_key:
                return {
                    'status': 'error',
                    'autoCreated': False,
                    'fallbackUsed': False,
                    'message': f'Unsupported template key: {template_key}',
                }
            if not template:
                return {
                    'status': 'error',
                    'autoCreated': False,
                    'fallbackUsed': False,
                    'message': f'Template not found for key: {normalized_key}',
                }

            source_path = str(template.get('sourcePath') or '').strip()
            if not source_path or not os.path.exists(source_path):
                return {
                    'status': 'error',
                    'autoCreated': False,
                    'fallbackUsed': False,
                    'message': f'Template file not found: {source_path or "<empty>"}',
                }

            extension = os.path.splitext(source_path)[1]
            if not extension:
                ext_hint = str(template.get('fileType') or '').strip().lstrip('.')
                extension = f'.{ext_hint}' if ext_hint else ''
            if not extension:
                return {
                    'status': 'error',
                    'autoCreated': False,
                    'fallbackUsed': False,
                    'message': 'Template file extension could not be determined.',
                }

            default_name = TEMPLATE_DEFAULT_FILENAME_BY_KEY.get(
                normalized_key,
                str(template.get('name') or 'Template').strip() or 'Template',
            )
            destination_path, filename, has_conflict = self._resolve_template_destination_path(
                project_path,
                default_name,
                extension,
                conflict_policy=conflict_policy,
            )
            if has_conflict:
                return {
                    'status': 'error',
                    'autoCreated': False,
                    'fallbackUsed': False,
                    'message': 'File already exists at destination.',
                }

            context_payload = dict(context or {}) if isinstance(context, dict) else {}
            if 'projectName' not in context_payload:
                context_payload['projectName'] = os.path.basename(
                    project_path.rstrip("\\/"))
            options_payload = {'templateKey': normalized_key}

            shutil.copy2(source_path, destination_path)
            _apply_template_context(
                destination_path, context=context_payload, options=options_payload)

            logging.info(
                "create_template_for_workroom: auto-created template "
                f"(key={normalized_key}, path={destination_path})."
            )
            return {
                'status': 'success',
                'path': destination_path,
                'filename': filename,
                'autoCreated': True,
                'fallbackUsed': False,
                'message': 'Template created.',
                'outputFilePath': destination_path,
                'outputFolderPath': project_path,
            }
        except Exception as e:
            logging.error(f"create_template_for_workroom failed: {e}")
            return {
                'status': 'error',
                'autoCreated': False,
                'fallbackUsed': False,
                'message': str(e),
            }

    def _resolve_dialog_directory(self, default_dir=None):
        directory = str(default_dir or '').strip()
        if directory:
            normalized = os.path.normpath(directory)
            normalized_copy_path = self._to_windows_extended_path(normalized)
            if os.path.isfile(normalized_copy_path):
                normalized = os.path.dirname(normalized)
            elif not os.path.isdir(normalized_copy_path):
                parent_dir = os.path.dirname(normalized)
                if parent_dir and os.path.isdir(self._to_windows_extended_path(parent_dir)):
                    normalized = parent_dir
                else:
                    normalized = ''
            directory = os.path.normpath(normalized) if normalized else ''
        return directory or get_default_documents_dir()

    def select_folder(self, default_dir=None):
        """Shows a folder selection dialog."""
        try:
            window = webview.windows[0]
            directory = self._resolve_dialog_directory(default_dir)
            folder_path = window.create_file_dialog(
                webview.FOLDER_DIALOG,
                directory=directory,
            )
            if not folder_path:
                return {'status': 'cancelled', 'path': None}
            # folder_path is a tuple, get first element
            return {'status': 'success', 'path': folder_path[0] if folder_path else None}
        except TypeError:
            try:
                window = webview.windows[0]
                folder_path = window.create_file_dialog(webview.FOLDER_DIALOG)
                if not folder_path:
                    return {'status': 'cancelled', 'path': None}
                return {'status': 'success', 'path': folder_path[0] if folder_path else None}
            except Exception as e:
                logging.error(f"Error in folder dialog fallback: {e}")
                return {'status': 'error', 'message': str(e)}
        except Exception as e:
            logging.error(f"Error in folder dialog: {e}")
            return {'status': 'error', 'message': str(e)}

    def select_template_output_folder(self, default_dir=None):
        """Shows a folder dialog for template-tool output."""
        try:
            window = webview.windows[0]
            directory = str(default_dir or '').strip() or get_default_documents_dir()
            if os.path.isfile(directory):
                directory = os.path.dirname(directory)
            if not os.path.isdir(directory):
                directory = get_default_documents_dir()

            folder_path = window.create_file_dialog(
                webview.FOLDER_DIALOG,
                directory=directory,
            )
            if not folder_path:
                return {'status': 'cancelled', 'path': None}
            if isinstance(folder_path, (list, tuple)):
                folder_path = folder_path[0] if folder_path else ''
            return {'status': 'success', 'path': folder_path}
        except TypeError:
            return self.select_folder(default_dir)
        except Exception as e:
            logging.error(f"Error in template output folder dialog: {e}")
            return {'status': 'error', 'message': str(e)}

    def select_template_save_location(self, default_dir=None, default_name=None, file_type=None):
        """Shows a save dialog for template output."""
        try:
            window = webview.windows[0]
            directory = str(default_dir or '').strip() or get_default_documents_dir()
            save_filename = str(default_name or 'Template').strip() or 'Template'

            ext = str(file_type or '').lower().lstrip('.')
            if ext and not save_filename.lower().endswith(f".{ext}"):
                save_filename = f"{save_filename}.{ext}"

            file_type_map = {
                'doc': 'Word Document (*.doc)',
                'docx': 'Word Document (*.docx)',
                'dwg': 'AutoCAD Drawing (*.dwg)',
                'xlsx': 'Excel Spreadsheet (*.xlsx)',
                'xls': 'Excel Spreadsheet (*.xls)',
            }
            file_types = None
            if ext and ext in file_type_map:
                file_types = (file_type_map[ext],)

            file_path = window.create_file_dialog(
                webview.FileDialog.SAVE,
                directory=directory,
                save_filename=save_filename,
                file_types=file_types
            )
            if not file_path:
                return {'status': 'cancelled', 'path': None}
            if isinstance(file_path, (list, tuple)):
                file_path = file_path[0] if file_path else ''
            return {'status': 'success', 'path': file_path}
        except Exception as e:
            logging.error(f"Error in save dialog: {e}")
            return {'status': 'error', 'message': str(e)}

    def select_template_file(self):
        """Shows file dialog for selecting template files."""
        try:
            window = webview.windows[0]
            file_types = (
                'Document Files (*.doc;*.docx)',
                'AutoCAD Files (*.dwg)',
                'Excel Files (*.xlsx;*.xls)',
                'All Supported (*.doc;*.docx;*.dwg;*.xlsx;*.xls)'
            )
            file_paths = window.create_file_dialog(
                webview.OPEN_DIALOG,
                allow_multiple=False,
                file_types=file_types
            )
            if not file_paths:
                return {'status': 'cancelled', 'path': None}
            return {'status': 'success', 'path': file_paths[0]}
        except Exception as e:
            logging.error(f"Error in file dialog: {e}")
            return {'status': 'error', 'message': str(e)}

    def verify_template_exists(self, template_id):
        """Verifies if a template's source file still exists."""
        try:
            data = self.get_templates()
            template = next((t for t in data['templates'] if t['id'] == template_id), None)
            if not template:
                return {'status': 'error', 'message': 'Template not found'}
            exists = os.path.exists(template['sourcePath'])
            return {'status': 'success', 'exists': exists, 'path': template['sourcePath']}
        except Exception as e:
            return {'status': 'error', 'message': str(e)}

    # ===================== CHECKLISTS API =====================

    def get_checklists(self):
        """Reads and returns checklists data."""
        try:
            with open(CHECKLISTS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            # First run - return empty structure
            return {'checklists': [], 'lastModified': None}

    def save_checklists(self, data):
        """Saves checklists data with backup."""
        try:
            if os.path.exists(CHECKLISTS_FILE):
                shutil.copy2(CHECKLISTS_FILE, CHECKLISTS_FILE + '.bak')
            with open(CHECKLISTS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error saving checklists: {e}")
            return {'status': 'error', 'message': str(e)}

    def _find_expense_signature_row(self, worksheet):
        for row_idx in range(1, worksheet.max_row + 1):
            value = worksheet.cell(row=row_idx, column=1).value
            if isinstance(value, str) and value.strip().upper() == "EMPLOYEE:":
                return row_idx
        raise ValueError("Could not locate the expense-sheet signature block.")

    def _get_expense_section_starts(self, worksheet, signature_row):
        section_starts = []
        for row_idx in range(1, signature_row):
            value = worksheet.cell(row=row_idx, column=1).value
            if isinstance(value, str) and value.strip().upper().startswith("PROJECT:"):
                section_starts.append(row_idx)
        return section_starts

    def _capture_expense_section_template(self, worksheet, start_row):
        section_end_row = start_row + EXPENSE_SECTION_ROWS - 1
        row_merge_map = {}

        for merged_range in worksheet.merged_cells.ranges:
            if merged_range.min_row < start_row or merged_range.max_row > section_end_row:
                continue
            row_offset = merged_range.min_row - start_row
            row_merge_map.setdefault(row_offset, []).append(
                (merged_range.min_col, merged_range.max_col)
            )

        rows = []
        for row_offset in range(EXPENSE_SECTION_ROWS):
            source_row = start_row + row_offset
            row_dimension = worksheet.row_dimensions[source_row]
            rows.append({
                "height": row_dimension.height,
                "hidden": row_dimension.hidden,
                "outlineLevel": row_dimension.outlineLevel,
                "collapsed": row_dimension.collapsed,
                "merges": list(row_merge_map.get(row_offset, [])),
                "cells": [
                    {
                        "value": worksheet.cell(row=source_row, column=col_idx).value,
                        "style": copy(worksheet.cell(row=source_row, column=col_idx)._style),
                    }
                    for col_idx in range(1, worksheet.max_column + 1)
                ],
            })

        return {
            "max_column": worksheet.max_column,
            "rows": rows,
        }

    def _apply_expense_template_row(self, worksheet, template, row_offset, target_row, copy_values=True):
        row_template = template["rows"][row_offset]

        for col_idx, cell_template in enumerate(row_template["cells"], start=1):
            target_cell = worksheet.cell(row=target_row, column=col_idx)
            target_cell._style = copy(cell_template["style"])
            target_cell.value = cell_template["value"] if copy_values else None

        target_dimension = worksheet.row_dimensions[target_row]
        target_dimension.height = row_template["height"]
        target_dimension.hidden = row_template["hidden"]
        target_dimension.outlineLevel = row_template["outlineLevel"]
        target_dimension.collapsed = row_template["collapsed"]

        for min_col, max_col in row_template["merges"]:
            worksheet.merge_cells(
                start_row=target_row,
                start_column=min_col,
                end_row=target_row,
                end_column=max_col,
            )

    def _apply_expense_section_template(self, worksheet, template, target_start_row):
        for row_offset in range(EXPENSE_SECTION_ROWS):
            self._apply_expense_template_row(
                worksheet,
                template,
                row_offset,
                target_start_row + row_offset,
                copy_values=True,
            )

    def _reset_expense_section_merges(self, worksheet, section_start, actual_data_rows):
        section_end = section_start + EXPENSE_HEADER_ROWS + actual_data_rows + EXPENSE_FOOTER_ROWS - 1
        merges_to_remove = []
        for merged_range in worksheet.merged_cells.ranges:
            if merged_range.min_row < section_start or merged_range.max_row > section_end:
                continue
            if merged_range.max_col > 5:
                continue
            merges_to_remove.append(str(merged_range))

        for merged_range in merges_to_remove:
            worksheet.unmerge_cells(merged_range)

        header_row = section_start + 2
        worksheet.merge_cells(
            start_row=header_row,
            start_column=2,
            end_row=header_row,
            end_column=3,
        )

        data_start_row = section_start + EXPENSE_HEADER_ROWS
        subtotal_row = data_start_row + actual_data_rows
        for row_idx in range(data_start_row, subtotal_row):
            worksheet.merge_cells(
                start_row=row_idx,
                start_column=2,
                end_row=row_idx,
                end_column=3,
            )

        total_row = subtotal_row + 2
        worksheet.merge_cells(
            start_row=total_row,
            start_column=4,
            end_row=total_row,
            end_column=5,
        )

    def _capture_worksheet_row_template(self, worksheet, source_row):
        row_dimension = worksheet.row_dimensions[source_row]
        return {
            "height": row_dimension.height,
            "hidden": row_dimension.hidden,
            "outlineLevel": row_dimension.outlineLevel,
            "collapsed": row_dimension.collapsed,
            "cells": [
                {
                    "value": worksheet.cell(row=source_row, column=col_idx).value,
                    "style": copy(worksheet.cell(row=source_row, column=col_idx)._style),
                }
                for col_idx in range(1, worksheet.max_column + 1)
            ],
        }

    def _apply_worksheet_row_template(self, worksheet, template, target_row, copy_values=True):
        for col_idx, cell_template in enumerate(template["cells"], start=1):
            target_cell = worksheet.cell(row=target_row, column=col_idx)
            target_cell._style = copy(cell_template["style"])
            target_cell.value = cell_template["value"] if copy_values else None

        target_dimension = worksheet.row_dimensions[target_row]
        target_dimension.height = template["height"]
        target_dimension.hidden = template["hidden"]
        target_dimension.outlineLevel = template["outlineLevel"]
        target_dimension.collapsed = template["collapsed"]

    def _render_project_expense_sheet(self, worksheet, projects, mileage_rate):
        if not projects:
            return {"image_paths": [], "image_start_row": None}

        signature_row = self._find_expense_signature_row(worksheet)
        section_starts = self._get_expense_section_starts(worksheet, signature_row)
        if not section_starts:
            raise ValueError("No formatted expense sections were found in the template.")

        section_template = self._capture_expense_section_template(worksheet, section_starts[0])
        first_section_start = section_starts[0]
        expense_detail_number_format = worksheet.cell(
            row=first_section_start + EXPENSE_HEADER_ROWS,
            column=5,
        ).number_format
        expense_subtotal_number_format = worksheet.cell(
            row=first_section_start + EXPENSE_HEADER_ROWS + EXPENSE_DEFAULT_DATA_ROWS,
            column=5,
        ).number_format
        expense_rate_number_format = worksheet.cell(
            row=first_section_start + EXPENSE_HEADER_ROWS + EXPENSE_DEFAULT_DATA_ROWS + 1,
            column=4,
        ).number_format
        expense_total_number_format = worksheet.cell(
            row=first_section_start + EXPENSE_HEADER_ROWS + EXPENSE_DEFAULT_DATA_ROWS + 2,
            column=4,
        ).number_format

        while len(section_starts) < len(projects):
            worksheet.insert_rows(signature_row, EXPENSE_SECTION_ROWS)
            self._apply_expense_section_template(worksheet, section_template, signature_row)
            section_starts.append(signature_row)
            signature_row += EXPENSE_SECTION_ROWS

        section_layouts = []

        for project_index, project in enumerate(projects):
            section_start = section_starts[project_index]
            entries = project.get('entries', []) or []
            num_entries = max(len(entries), 1)
            extra_rows = max(0, num_entries - EXPENSE_DEFAULT_DATA_ROWS)

            if extra_rows:
                insert_at = section_start + EXPENSE_HEADER_ROWS + EXPENSE_DEFAULT_DATA_ROWS
                worksheet.insert_rows(insert_at, extra_rows)
                for row_offset in range(extra_rows):
                    self._apply_expense_template_row(
                        worksheet,
                        section_template,
                        EXPENSE_INSERTED_DATA_TEMPLATE_OFFSET,
                        insert_at + row_offset,
                        copy_values=False,
                    )
                for later_index in range(project_index + 1, len(section_starts)):
                    section_starts[later_index] += extra_rows
                signature_row += extra_rows

            data_start_row = section_start + EXPENSE_HEADER_ROWS
            actual_data_rows = max(num_entries, EXPENSE_DEFAULT_DATA_ROWS)
            subtotal_row = data_start_row + actual_data_rows
            rate_row = subtotal_row + 1
            total_row = rate_row + 1

            section_layouts.append({
                "project": project,
                "section_start": section_start,
                "data_start_row": data_start_row,
                "actual_data_rows": actual_data_rows,
                "subtotal_row": subtotal_row,
                "rate_row": rate_row,
                "total_row": total_row,
            })

        for section_layout in section_layouts:
            self._reset_expense_section_merges(
                worksheet,
                section_layout["section_start"],
                section_layout["actual_data_rows"],
            )

        image_paths = []
        for section_layout in section_layouts:
            project = section_layout["project"]
            entries = project.get('entries', []) or []
            section_start = section_layout["section_start"]
            data_start_row = section_layout["data_start_row"]
            subtotal_row = section_layout["subtotal_row"]
            rate_row = section_layout["rate_row"]
            total_row = section_layout["total_row"]

            worksheet.cell(
                row=section_start,
                column=1,
                value=f"PROJECT: {project.get('projectName', '')}",
            )
            worksheet.cell(
                row=section_start + 1,
                column=1,
                value=f"JOB #: {project.get('projectId', '')}",
            )

            for row_idx in range(data_start_row, subtotal_row):
                worksheet.cell(row=row_idx, column=1, value=None)
                worksheet.cell(row=row_idx, column=2, value=None)
                worksheet.cell(row=row_idx, column=4, value=None)
                worksheet.cell(row=row_idx, column=5, value=None)
                worksheet.cell(row=row_idx, column=5).number_format = expense_detail_number_format

            for entry_index, entry in enumerate(entries):
                row_idx = data_start_row + entry_index
                worksheet.cell(row=row_idx, column=1, value=entry.get('date', ''))
                worksheet.cell(row=row_idx, column=2, value=entry.get('description', ''))

                mileage = entry.get('mileage', 0) or 0
                expense = entry.get('expense', 0) or 0

                worksheet.cell(row=row_idx, column=4, value=mileage if mileage else '')
                expense_cell = worksheet.cell(
                    row=row_idx,
                    column=5,
                    value=expense if expense else '',
                )
                expense_cell.number_format = expense_detail_number_format

            worksheet.cell(row=subtotal_row, column=3, value="SUBTOTAL")
            worksheet.cell(
                row=subtotal_row,
                column=4,
                value=f'=SUM(D{data_start_row}:D{subtotal_row - 1})',
            )
            worksheet.cell(
                row=subtotal_row,
                column=5,
                value=f'=SUM(E{data_start_row}:E{subtotal_row - 1})',
            )
            worksheet.cell(row=subtotal_row, column=5).number_format = expense_subtotal_number_format

            worksheet.cell(
                row=rate_row,
                column=3,
                value=f"{mileage_rate:.2f} CENTS PER MILE",
            )
            worksheet.cell(
                row=rate_row,
                column=4,
                value=f'=D{subtotal_row}*{mileage_rate}',
            )
            worksheet.cell(row=rate_row, column=4).number_format = expense_rate_number_format
            worksheet.cell(row=rate_row, column=5, value=None)

            worksheet.cell(row=total_row, column=3, value="TOTAL")
            worksheet.cell(
                row=total_row,
                column=4,
                value=f'=D{rate_row}+E{subtotal_row}',
            )
            worksheet.cell(row=total_row, column=4).number_format = expense_total_number_format

            for image in project.get('images', []):
                image_path = image.get('path', '')
                if image_path:
                    resolved_image_path = self._resolve_expense_image_path(image_path)
                    if resolved_image_path:
                        image_paths.append(resolved_image_path)

        return {
            "image_paths": image_paths,
            "image_start_row": signature_row + EXPENSE_IMAGE_ROW_OFFSET,
        }

    def _add_expense_sheet_images(self, worksheet, image_paths, image_start_row):
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.utils.units import points_to_pixels

        temp_files_to_cleanup = []
        if not image_paths or image_start_row is None:
            return temp_files_to_cleanup

        target_width = self._get_expense_image_target_width_pixels(worksheet)

        for image_path in image_paths:
            if not os.path.isfile(image_path):
                continue

            try:
                ext = os.path.splitext(image_path)[1].lower()
                if ext == '.pdf':
                    continue

                with _open_local_pil_image(image_path) as source_img:
                    if hasattr(source_img, 'n_frames') and source_img.n_frames > 1:
                        source_img.seek(0)
                    pil_img = ImageOps.exif_transpose(source_img).copy()
                    if pil_img.mode in ('RGBA', 'P', 'LA'):
                        pil_img = pil_img.convert('RGB')
                    elif pil_img.mode != 'RGB':
                        pil_img = pil_img.convert('RGB')

                source_width, source_height = pil_img.size
                if source_width <= 0 or source_height <= 0:
                    continue

                temp_img_path = os.path.join(
                    tempfile.gettempdir(),
                    f"expense_img_{os.urandom(4).hex()}.png",
                )
                pil_img.save(temp_img_path, 'PNG')
                temp_files_to_cleanup.append(temp_img_path)

                image = XLImage(temp_img_path)
                image.width = max(int(round(target_width)), 1)
                image.height = max(
                    int(round(target_width * (source_height / source_width))),
                    1,
                )
                worksheet.add_image(image, f"A{image_start_row}")

                rows_for_image = 0
                covered_height = 0
                next_row = image_start_row
                while covered_height < image.height:
                    row_height_points = worksheet.row_dimensions[next_row].height
                    if row_height_points is None:
                        row_height_points = worksheet.sheet_format.defaultRowHeight or 15
                    covered_height += max(points_to_pixels(row_height_points), 1)
                    rows_for_image += 1
                    next_row += 1
                image_start_row += rows_for_image + 1
            except UnidentifiedImageError:
                continue
            except Exception as img_err:
                logging.warning(f"Could not add image {image_path}: {img_err}")

        return temp_files_to_cleanup

    def _get_expense_image_target_width_pixels(self, worksheet, start_column=1, end_column=5):
        total_width = 0
        default_width = worksheet.sheet_format.defaultColWidth or 8.43
        for column_index in range(start_column, end_column + 1):
            column_letter = openpyxl.utils.get_column_letter(column_index)
            column_width = worksheet.column_dimensions[column_letter].width
            total_width += self._excel_column_width_to_pixels(
                default_width if column_width is None else column_width
            )
        return max(total_width, 1)

    def _excel_column_width_to_pixels(self, column_width):
        width = float(column_width or 8.43)
        return int(math.floor(((256 * width + math.floor(128 / 7)) / 256) * 7))

    def _find_nearest_existing_directory(self, path):
        current = os.path.normpath(str(path or "").strip())
        while current:
            if os.path.isdir(current):
                return current
            parent = os.path.dirname(current)
            if parent == current:
                break
            current = parent
        return ""

    def _resolve_expense_image_path(self, raw_path):
        normalized_path = self._coerce_local_file_path(raw_path)
        if not normalized_path:
            return ""
        if os.path.isfile(normalized_path):
            return normalized_path

        filename = os.path.basename(normalized_path)
        if not filename:
            return ""

        search_root = self._find_nearest_existing_directory(os.path.dirname(normalized_path))
        if not search_root:
            return ""

        original_parent = os.path.basename(os.path.dirname(normalized_path).rstrip("\\/")).strip().lower()
        exact_parent_matches = []
        any_matches = []

        try:
            for candidate in Path(search_root).rglob(filename):
                candidate_path = os.path.normpath(str(candidate))
                if not os.path.isfile(candidate_path):
                    continue
                any_matches.append(candidate_path)
                candidate_parent = os.path.basename(
                    os.path.dirname(candidate_path).rstrip("\\/")
                ).strip().lower()
                if original_parent and candidate_parent == original_parent:
                    exact_parent_matches.append(candidate_path)
        except Exception:
            return ""

        if len(exact_parent_matches) == 1:
            return exact_parent_matches[0]
        if len(any_matches) == 1:
            return any_matches[0]
        if exact_parent_matches:
            return sorted(exact_parent_matches)[0]
        return ""

    def export_timesheet_excel(self, data):
        """Exports timesheet data to an Excel file using a template."""
        try:
            import openpyxl

            template_path = get_bundled_template_path("Template_Timesheet.xlsx")

            if not template_path.exists():
                return {'status': 'error', 'message': f'Template file not found: {template_path}'}

            # Determine output path
            week_key = data.get('weekKey', '')
            user_name = data.get('userName', '')
            file_path = data.get('filePath')
            if not file_path:
                default_dir = get_default_documents_dir()
                default_name = build_timesheet_filename(user_name, week_key)
                window = webview.windows[0]
                file_path = window.create_file_dialog(
                    webview.FileDialog.SAVE,
                    directory=default_dir,
                    save_filename=default_name,
                    file_types=('Excel Files (*.xlsx)',)
                )
                if not file_path:
                    return {'status': 'cancelled'}
            if isinstance(file_path, (list, tuple)):
                file_path = file_path[0] if file_path else ''
            if not file_path:
                return {'status': 'cancelled'}
            if not str(file_path).lower().endswith('.xlsx'):
                file_path = f"{file_path}.xlsx"

            # Copy template to output location
            shutil.copy2(str(template_path), file_path)

            # Open the copied file and modify the "time log" sheet
            wb = openpyxl.load_workbook(file_path)

            # Get the "time log" sheet
            if "time log" not in wb.sheetnames:
                return {'status': 'error', 'message': 'Sheet "time log" not found in template'}

            ws = wb["time log"]
            expense_data = data.get('expenses', {})
            expense_projects = expense_data.get('projects', [])

            # Row 1: Employee name (column M = 13)
            ws.cell(row=1, column=13, value=data.get('userName', 'Employee'))

            # Row 2: Week of (column M = 13)
            ws.cell(row=2, column=13, value=data.get('weekDisplay', ''))

            # Data rows start at row 5
            entries = data.get('entries', [])
            data_start_row = 5
            insert_after_row = 22
            insert_at_row = insert_after_row + 1
            base_totals_row = 26
            max_template_project_rows_before_insert = insert_after_row - data_start_row + 1
            overflow_rows = max(0, len(entries) - max_template_project_rows_before_insert)

            if overflow_rows:
                inserted_row_template = self._capture_worksheet_row_template(ws, insert_after_row)
                ws.insert_rows(insert_at_row, overflow_rows)
                for row_offset in range(overflow_rows):
                    self._apply_worksheet_row_template(
                        ws,
                        inserted_row_template,
                        insert_at_row + row_offset,
                        copy_values=False,
                    )

            day_order = TIMESHEET_DAY_ORDER
            export_mileage_by_row = _build_export_mileage_by_row(entries, expense_projects)

            for row_idx, entry in enumerate(entries, data_start_row):
                # Column A (1): PROJECT #
                ws.cell(row=row_idx, column=1, value=entry.get('projectId', ''))
                # Column B (2): Task #
                ws.cell(row=row_idx, column=2, value=entry.get('taskNumber', ''))
                # Column C (3): PROJECT NAME
                ws.cell(row=row_idx, column=3, value=entry.get('projectName', ''))
                # Column D (4): FUNCTION (M E P F)
                ws.cell(row=row_idx, column=4, value=entry.get('function', ''))
                # Column E (5): PM
                ws.cell(row=row_idx, column=5, value=entry.get('pmInitials', ''))
                # Columns F-I (6-9): DESIGN/ENG, DRAFTING, TRAVEL, OTHERS - leave empty
                # Column J (10): SERVICE & WORK DESCRIPTION
                ws.cell(row=row_idx, column=10, value=entry.get('serviceDescription', ''))

                # Columns K-Q (11-17): MON through SUN
                hours = entry.get('hours', {})
                for day_idx, day in enumerate(day_order):
                    h = hours.get(day, 0)
                    ws.cell(row=row_idx, column=11 + day_idx, value=h if h else '')

                # Column R (18): MILEAGE derived from project expense dates.
                mileage = export_mileage_by_row[row_idx - data_start_row] if row_idx - data_start_row < len(export_mileage_by_row) else 0
                ws.cell(row=row_idx, column=18, value=mileage if mileage else '')

            # Totals row (row 26 in template based on CSV)
            totals_row = base_totals_row + overflow_rows
            data_end_row = totals_row - 1

            for day_idx, _day in enumerate(day_order):
                col_idx = 11 + day_idx
                col_letter = openpyxl.utils.get_column_letter(col_idx)
                ws.cell(
                    row=totals_row,
                    column=col_idx,
                    value=f'=SUM({col_letter}{data_start_row}:{col_letter}{data_end_row})',
                )

            mileage_col_letter = openpyxl.utils.get_column_letter(18)
            ws.cell(
                row=totals_row,
                column=18,
                value=f'=SUM({mileage_col_letter}{data_start_row}:{mileage_col_letter}{data_end_row})',
            )

            # =============== EXPENSE SHEET EXPORT ===============
            # If expense data is provided, also populate the expense sheet
            temp_files_to_cleanup = []

            if expense_projects:
                if EXPENSE_SHEET_NAME in wb.sheetnames:
                    ws_exp = wb[EXPENSE_SHEET_NAME]
                    render_result = self._render_project_expense_sheet(
                        ws_exp,
                        expense_projects,
                        expense_data.get('mileageRate', 0.70),
                    )
                    temp_files_to_cleanup.extend(
                        self._add_expense_sheet_images(
                            ws_exp,
                            render_result["image_paths"],
                            render_result["image_start_row"],
                        )
                    )

            # Save the modified file
            wb.save(file_path)

            # Clean up temp files after save
            for temp_file in temp_files_to_cleanup:
                try:
                    os.remove(temp_file)
                except:
                    pass

            # Open file on Windows
            if sys.platform == "win32":
                os.startfile(file_path)

            return {'status': 'success', 'path': file_path}

        except ImportError:
            return {'status': 'error', 'message': 'openpyxl not installed. Run: pip install openpyxl'}
        except Exception as e:
            logging.error(f"Error exporting timesheet: {e}")
            return {'status': 'error', 'message': str(e)}

    def select_expense_images(self):
        """Shows file dialog for selecting expense receipt images."""
        try:
            window = webview.windows[0]
            file_types = (
                'Image Files (*.jpg;*.jpeg;*.png;*.gif;*.bmp;*.heic;*.heif)',
                'PDF Files (*.pdf)',
                'All Files (*.*)'
            )
            file_paths = window.create_file_dialog(
                webview.OPEN_DIALOG,
                allow_multiple=True,
                file_types=file_types
            )
            if not file_paths:
                return {'status': 'cancelled', 'paths': []}
            return {'status': 'success', 'paths': list(file_paths)}
        except Exception as e:
            logging.error(f"Error in expense image selection dialog: {e}")
            return {'status': 'error', 'message': str(e)}

    def resolve_expense_attachment_path(self, path):
        """Resolves a stored expense attachment path to an existing local file."""
        try:
            resolved_path = self._resolve_expense_image_path(path)
            if not resolved_path or not os.path.isfile(resolved_path):
                return {'status': 'error', 'message': 'Attachment file not found.'}
            return {
                'status': 'success',
                'path': resolved_path,
                'filename': os.path.basename(resolved_path),
            }
        except Exception as e:
            logging.error(f"Error resolving expense attachment path: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_expense_image_preview(self, path, max_size=1600):
        """Returns a browser-safe preview data URL for a local expense image."""
        try:
            resolved_path = self._resolve_expense_image_path(path)
            if not resolved_path or not os.path.isfile(resolved_path):
                return {'status': 'error', 'message': 'Image file not found.'}

            try:
                preview_max_size = max(1, int(max_size or 1600))
            except Exception:
                preview_max_size = 1600

            with _open_local_pil_image(resolved_path) as source_img:
                if hasattr(source_img, 'n_frames') and source_img.n_frames > 1:
                    source_img.seek(0)
                preview = ImageOps.exif_transpose(source_img).copy()
                if preview.mode in ('P', 'LA'):
                    preview = preview.convert('RGBA')
                elif preview.mode not in ('RGB', 'RGBA'):
                    preview = preview.convert('RGB')

            resampling = getattr(
                getattr(PILImage, 'Resampling', PILImage),
                'LANCZOS',
                getattr(PILImage, 'LANCZOS', getattr(PILImage, 'BICUBIC', 3))
            )
            preview.thumbnail((preview_max_size, preview_max_size), resampling)

            buffer = io.BytesIO()
            preview.save(buffer, format='PNG')
            encoded = base64.b64encode(buffer.getvalue()).decode('ascii')

            return {
                'status': 'success',
                'dataUrl': f'data:image/png;base64,{encoded}',
                'width': preview.width,
                'height': preview.height,
                'path': resolved_path,
                'filename': os.path.basename(resolved_path),
            }
        except UnidentifiedImageError:
            resolved_path = self._resolve_expense_image_path(path)
            return {
                'status': 'unsupported',
                'message': 'Preview not available for this file type.',
                'path': resolved_path or '',
                'filename': os.path.basename(resolved_path) if resolved_path else '',
            }
        except ValueError as e:
            resolved_path = self._resolve_expense_image_path(path)
            return {
                'status': 'unsupported',
                'message': str(e),
                'path': resolved_path or '',
                'filename': os.path.basename(resolved_path) if resolved_path else '',
            }
        except Exception as e:
            logging.error(f"Error generating expense image preview: {e}")
            return {'status': 'error', 'message': str(e)}

    def _page_assets_root(self):
        """Absolute path to the page_assets media folder (created on demand)."""
        root = get_app_data_path("page_assets")
        os.makedirs(root, exist_ok=True)
        return os.path.realpath(root)

    def _safe_page_owner_key(self, owner_key):
        """Sanitizes an owner key (project/deliverable scope) into a safe folder name."""
        cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", str(owner_key or "").strip()).strip("_")
        return cleaned or "shared"

    def _resolve_page_asset_path(self, asset_path):
        """Resolves a stored page asset reference to an absolute path strictly inside page_assets."""
        raw = str(asset_path or "").strip().replace("\\", "/")
        if not raw:
            return ""
        parts = [p for p in raw.split("/") if p not in ("", ".", "..")]
        if parts and parts[0] == "page_assets":
            parts = parts[1:]
        if not parts:
            return ""
        root = self._page_assets_root()
        candidate = os.path.realpath(os.path.join(root, *parts))
        try:
            if os.path.commonpath([root, candidate]) != root:
                return ""
        except ValueError:
            return ""
        return candidate

    def save_page_asset(self, owner_key, data_url, filename=""):
        """Persists a page image (base64 data URL) to page_assets and returns a relative reference."""
        try:
            raw = str(data_url or "")
            header, _, payload = raw.partition(",")
            if not payload or "base64" not in header.lower():
                return {'status': 'error', 'message': 'Unsupported image data.'}
            mime_match = re.match(r"^data:([^;,]+)", header)
            mime = (mime_match.group(1).lower() if mime_match else 'image/png')
            try:
                blob = base64.b64decode(payload)
            except Exception:
                return {'status': 'error', 'message': 'Could not decode image data.'}
            if not blob:
                return {'status': 'error', 'message': 'Empty image data.'}

            ext_map = {
                'image/png': '.png', 'image/jpeg': '.jpg', 'image/jpg': '.jpg',
                'image/gif': '.gif', 'image/webp': '.webp', 'image/bmp': '.bmp',
                'image/svg+xml': '.svg', 'image/heic': '.heic', 'image/heif': '.heif',
            }
            ext = ext_map.get(mime, '')
            if not ext:
                _, guessed = os.path.splitext(str(filename or ''))
                ext = guessed.lower() if guessed else '.png'

            owner = self._safe_page_owner_key(owner_key)
            owner_dir = os.path.join(self._page_assets_root(), owner)
            os.makedirs(owner_dir, exist_ok=True)
            asset_name = f"{uuid.uuid4().hex}{ext}"
            with open(os.path.join(owner_dir, asset_name), 'wb') as fh:
                fh.write(blob)

            return {'status': 'success', 'assetPath': f"page_assets/{owner}/{asset_name}"}
        except Exception as e:
            logging.error(f"Error saving page asset: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_page_asset(self, asset_path, max_size=1600):
        """Returns a browser-safe preview data URL for a stored page asset."""
        try:
            resolved_path = self._resolve_page_asset_path(asset_path)
            if not resolved_path or not os.path.isfile(resolved_path):
                return {'status': 'error', 'message': 'Image file not found.'}

            if resolved_path.lower().endswith('.svg'):
                with open(resolved_path, 'rb') as fh:
                    encoded = base64.b64encode(fh.read()).decode('ascii')
                return {'status': 'success', 'dataUrl': f'data:image/svg+xml;base64,{encoded}'}

            try:
                preview_max_size = max(1, int(max_size or 1600))
            except Exception:
                preview_max_size = 1600

            with _open_local_pil_image(resolved_path) as source_img:
                if hasattr(source_img, 'n_frames') and source_img.n_frames > 1:
                    source_img.seek(0)
                preview = ImageOps.exif_transpose(source_img).copy()
                if preview.mode in ('P', 'LA'):
                    preview = preview.convert('RGBA')
                elif preview.mode not in ('RGB', 'RGBA'):
                    preview = preview.convert('RGB')

            resampling = getattr(
                getattr(PILImage, 'Resampling', PILImage),
                'LANCZOS',
                getattr(PILImage, 'LANCZOS', getattr(PILImage, 'BICUBIC', 3))
            )
            preview.thumbnail((preview_max_size, preview_max_size), resampling)

            buffer = io.BytesIO()
            preview.save(buffer, format='PNG')
            encoded = base64.b64encode(buffer.getvalue()).decode('ascii')
            return {
                'status': 'success',
                'dataUrl': f'data:image/png;base64,{encoded}',
                'width': preview.width,
                'height': preview.height,
            }
        except ValueError as e:
            return {'status': 'unsupported', 'message': str(e)}
        except Exception as e:
            logging.error(f"Error generating page asset preview: {e}")
            return {'status': 'error', 'message': str(e)}

    def _page_files_root(self):
        """Absolute path to locally managed page file storage."""
        root = get_app_data_path("page_files")
        os.makedirs(root, exist_ok=True)
        return os.path.realpath(root)

    @staticmethod
    def _page_file_links_path():
        return get_app_data_path("page_file_links.json")

    def _read_page_file_links(self):
        path = self._page_file_links_path()
        if not os.path.isfile(path):
            return {}
        payload = _read_json_file_strict(path)
        if not isinstance(payload, dict) or not isinstance(payload.get("links", {}), dict):
            raise ValueError("Managed page file link data is invalid.")
        return payload.get("links", {})

    def _write_page_file_links(self, links):
        _atomic_write_json_file(self._page_file_links_path(), {
            "version": 1,
            "links": links if isinstance(links, dict) else {},
        })

    @staticmethod
    def _page_file_link_id(file_ref):
        match = re.fullmatch(
            r"page_file_links/([A-Fa-f0-9]{32})",
            str(file_ref or "").strip().replace("\\", "/"),
        )
        return match.group(1).lower() if match else ""

    @staticmethod
    def _normalize_external_workbook_path(path, require_exists=True):
        raw = str(path or "").strip().strip('"').strip("'")
        if not raw:
            raise ValueError("Workbook path is required.")
        expanded = os.path.expandvars(os.path.expanduser(raw))
        normalized = os.path.normpath(expanded)
        if not os.path.isabs(normalized):
            raise ValueError("Enter an absolute workbook path.")
        extension = os.path.splitext(normalized)[1].lower()
        if extension not in {".xlsx", ".xlsm", ".xls", ".xlsb", ".csv"}:
            raise ValueError("Choose an Excel workbook or CSV file.")
        if require_exists and not os.path.isfile(normalized):
            raise ValueError("Workbook file was not found.")
        return os.path.abspath(normalized)

    @staticmethod
    def _safe_page_workbook_filename(filename):
        """Return a safe .xlsx filename while retaining the user's readable name."""
        raw = str(filename or "").strip()
        if raw.lower().endswith(".xlsx"):
            raw = raw[:-5]
        raw = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", raw).strip(" .")
        if not raw:
            raise ValueError("Workbook name is required.")
        reserved = {
            "CON", "PRN", "AUX", "NUL",
            *(f"COM{index}" for index in range(1, 10)),
            *(f"LPT{index}" for index in range(1, 10)),
        }
        if raw.split(".", 1)[0].upper() in reserved:
            raw = f"_{raw}"
        raw = raw[:120].rstrip(" .")
        if not raw:
            raise ValueError("Workbook name is required.")
        return f"{raw}.xlsx"

    def _resolve_page_file_path(self, file_ref):
        """Resolve a managed page file reference strictly inside page_files."""
        raw = str(file_ref or "").strip().replace("\\", "/")
        if not raw or raw.startswith("/") or re.match(r"^[A-Za-z]:", raw):
            return ""
        parts = raw.split("/")
        if any(part in ("", ".", "..") for part in parts):
            return ""
        if len(parts) != 4 or parts[0] != "page_files":
            return ""
        if not re.fullmatch(r"[A-Za-z0-9_-]+", parts[1]):
            return ""
        if not re.fullmatch(r"[A-Fa-f0-9]{32}", parts[2]):
            return ""
        if not parts[-1].lower().endswith(".xlsx"):
            return ""
        try:
            if self._safe_page_workbook_filename(parts[-1]) != parts[-1]:
                return ""
        except ValueError:
            return ""
        root = self._page_files_root()
        candidate = os.path.realpath(os.path.join(root, *parts[1:]))
        try:
            if os.path.commonpath([root, candidate]) != root:
                return ""
        except ValueError:
            return ""
        return candidate

    def create_page_workbook(self, owner_key, name):
        """Create a blank workbook in managed local page storage."""
        target_dir = ""
        try:
            file_name = self._safe_page_workbook_filename(name)
            owner = self._safe_page_owner_key(owner_key)
            workbook_id = uuid.uuid4().hex
            target_dir = os.path.join(self._page_files_root(), owner, workbook_id)
            os.makedirs(target_dir, exist_ok=False)
            target_path = os.path.join(target_dir, file_name)

            workbook = openpyxl.Workbook()
            workbook.active.title = "Sheet1"
            workbook.save(target_path)
            workbook.close()

            return {
                'status': 'success',
                'fileRef': f"page_files/{owner}/{workbook_id}/{file_name}",
                'fileName': file_name,
            }
        except ValueError as e:
            return {'status': 'error', 'message': str(e)}
        except Exception as e:
            if target_dir:
                shutil.rmtree(target_dir, ignore_errors=True)
            logging.error(f"Error creating page workbook: {e}")
            return {'status': 'error', 'message': str(e)}

    def link_page_workbook(self, path):
        """Register an existing local workbook without copying or exposing its path in page HTML."""
        try:
            normalized_path = self._normalize_external_workbook_path(path, require_exists=True)
            link_id = uuid.uuid4().hex
            links = self._read_page_file_links()
            links[link_id] = {
                'path': normalized_path,
                'fileName': os.path.basename(normalized_path),
                'createdAt': datetime.datetime.now().isoformat(),
            }
            self._write_page_file_links(links)
            return {
                'status': 'success',
                'fileRef': f"page_file_links/{link_id}",
                'fileName': os.path.basename(normalized_path),
                'storageType': 'external',
            }
        except ValueError as e:
            return {'status': 'error', 'message': str(e)}
        except Exception as e:
            logging.error(f"Error linking page workbook: {e}")
            return {'status': 'error', 'message': str(e)}

    def _get_page_file_reference(self, file_ref):
        resolved_path = self._resolve_page_file_path(file_ref)
        if resolved_path:
            return {
                'storageType': 'managed',
                'path': resolved_path,
                'fileName': os.path.basename(resolved_path),
            }

        link_id = self._page_file_link_id(file_ref)
        if not link_id:
            return None
        links = self._read_page_file_links()
        entry = links.get(link_id)
        if not isinstance(entry, dict):
            return {
                'storageType': 'external',
                'path': '',
                'fileName': '',
                'linkId': link_id,
            }
        try:
            external_path = self._normalize_external_workbook_path(
                entry.get('path'),
                require_exists=False,
            )
        except ValueError:
            external_path = ''
        return {
            'storageType': 'external',
            'path': external_path,
            'fileName': str(entry.get('fileName') or os.path.basename(external_path)).strip(),
            'linkId': link_id,
        }

    def get_page_file_info(self, file_ref):
        """Return local availability and metadata for a managed page file."""
        try:
            reference = self._get_page_file_reference(file_ref)
            if not reference:
                return {'status': 'error', 'message': 'Invalid managed page file reference.'}
            resolved_path = reference.get('path') or ''
            exists = os.path.isfile(resolved_path)
            result = {
                'status': 'success',
                'exists': exists,
                'fileName': reference.get('fileName') or os.path.basename(resolved_path),
                'storageType': reference.get('storageType') or 'managed',
            }
            if exists:
                stat = os.stat(resolved_path)
                result.update({
                    'sizeBytes': stat.st_size,
                    'modified': datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
                })
            return result
        except Exception as e:
            logging.error(f"Error reading page file info: {e}")
            return {'status': 'error', 'message': str(e)}

    def open_page_file(self, file_ref):
        """Open a managed page file with the operating system's registered app."""
        try:
            reference = self._get_page_file_reference(file_ref)
            if not reference:
                return {'status': 'error', 'message': 'Invalid managed page file reference.'}
            resolved_path = reference.get('path') or ''
            if not os.path.isfile(resolved_path):
                return {'status': 'error', 'message': 'Workbook is unavailable on this device.'}
            return self.open_path(resolved_path)
        except Exception as e:
            logging.error(f"Error opening page file: {e}")
            return {'status': 'error', 'message': str(e)}

    def delete_page_file(self, file_ref):
        """Permanently delete one managed page file and its empty container folders."""
        try:
            link_id = self._page_file_link_id(file_ref)
            if link_id:
                links = self._read_page_file_links()
                removed = links.pop(link_id, None)
                if removed is not None:
                    self._write_page_file_links(links)
                return {
                    'status': 'success',
                    'deleted': removed is not None,
                    'storageType': 'external',
                    'externalFileDeleted': False,
                }

            resolved_path = self._resolve_page_file_path(file_ref)
            if not resolved_path:
                return {'status': 'error', 'message': 'Invalid managed page file reference.'}
            deleted = False
            if os.path.isfile(resolved_path):
                os.remove(resolved_path)
                deleted = True

            root = self._page_files_root()
            current = os.path.dirname(resolved_path)
            while current != root:
                try:
                    if os.listdir(current):
                        break
                    os.rmdir(current)
                except FileNotFoundError:
                    pass
                current = os.path.dirname(current)
                if not current or os.path.commonpath([root, current]) != root:
                    break
            return {'status': 'success', 'deleted': deleted, 'storageType': 'managed'}
        except Exception as e:
            logging.error(f"Error deleting page file: {e}")
            return {'status': 'error', 'message': str(e)}

    def delete_page_files(self, file_refs):
        """Best-effort batch deletion used when pages or projects are removed."""
        refs = file_refs if isinstance(file_refs, list) else []
        deleted = 0
        missing = 0
        failed = []
        for file_ref in dict.fromkeys(str(ref or "") for ref in refs if str(ref or "").strip()):
            result = self.delete_page_file(file_ref)
            if result.get('status') == 'success':
                if result.get('deleted'):
                    deleted += 1
                else:
                    missing += 1
            else:
                failed.append({
                    'fileRef': file_ref,
                    'message': result.get('message') or 'Could not delete workbook.',
                })
        return {
            'status': 'success' if not failed else 'partial',
            'deleted': deleted,
            'missing': missing,
            'failed': failed,
        }

    def _resolve_project_page_pdf_image(self, attrs):
        """Returns a browser-safe image source and printable width percentage."""
        raw_width = str(attrs.get("data-width-percent") or "80").strip().rstrip("%")
        try:
            width_percent = int(round(float(raw_width)))
        except (TypeError, ValueError):
            width_percent = 80
        width_percent = max(10, min(100, width_percent))

        asset_path = str(attrs.get("data-asset") or "").strip()
        if not asset_path:
            source = str(attrs.get("src") or "").strip()
            if source.startswith("data:image/"):
                return source, width_percent
            if source.startswith("page_assets/"):
                asset_path = source
            else:
                return None

        preview = self.get_page_asset(asset_path, max_size=2400)
        if preview.get("status") != "success" or not preview.get("dataUrl"):
            return None

        try:
            image_width = float(preview.get("width") or 0)
            image_height = float(preview.get("height") or 0)
            if image_width > 0 and image_height > image_width:
                aspect_ratio = image_height / image_width
                portrait_width_cap = max(15, min(100, int(100 / aspect_ratio)))
                width_percent = min(width_percent, portrait_width_cap)
        except (TypeError, ValueError, ZeroDivisionError):
            pass

        return str(preview["dataUrl"]), width_percent

    def _prepare_project_page_pdf_html(self, raw_html):
        renderer = _ProjectPagePdfHtmlRenderer(
            self._resolve_project_page_pdf_image
        )
        renderer.feed(str(raw_html or ""))
        renderer.close()
        return renderer.rendered_html()

    @staticmethod
    def _project_page_pdf_css():
        return """
            body {
                color: #172033;
                font-family: "Segoe UI", Arial, sans-serif;
                font-size: 10.5pt;
                font-variant-ligatures: none;
                line-height: 1.48;
            }
            .pdf-kicker {
                color: #64748b;
                font-size: 7.5pt;
                font-weight: 700;
                letter-spacing: 1.15px;
                margin: 0 0 3px;
            }
            .pdf-title {
                color: #111827;
                font-size: 27pt;
                line-height: 1.12;
                margin: 0 0 5px;
            }
            .pdf-context {
                color: #64748b;
                font-size: 9pt;
                margin: 0 0 18px;
            }
            .pdf-empty {
                color: #64748b;
                font-style: italic;
            }
            h1 {
                color: #243b66;
                font-size: 18pt;
                line-height: 1.2;
                margin: 18px 0 6px;
            }
            h2 {
                color: #243b66;
                font-size: 15pt;
                line-height: 1.25;
                margin: 15px 0 5px;
            }
            h3 {
                color: #334155;
                font-size: 12pt;
                line-height: 1.3;
                margin: 12px 0 4px;
            }
            p {
                margin: 5px 0;
            }
            ul, ol {
                margin: 6px 0;
                padding-left: 22px;
            }
            a, .page-wiki-link {
                color: #2563eb;
                text-decoration: underline;
            }
            blockquote {
                border-left: 3px solid #94a3b8;
                color: #475569;
                margin: 9px 0;
                padding: 2px 0 2px 11px;
            }
            hr {
                border: 0;
                border-top: 1px solid #cbd5e1;
                margin: 14px 0;
            }
            table {
                border-collapse: collapse;
                margin: 9px 0;
                width: 100%;
            }
            th, td {
                border: 1px solid #cbd5e1;
                padding: 5px 7px;
                text-align: left;
                vertical-align: top;
            }
            th {
                background: #f1f5f9;
                color: #1e293b;
                font-weight: 700;
            }
            pre {
                background: #f1f5f9;
                border: 1px solid #dbe3ee;
                color: #172033;
                font-family: Consolas, "Courier New", monospace;
                font-size: 9pt;
                line-height: 1.42;
                margin: 9px 0;
                padding: 9px 11px;
                white-space: pre-wrap;
            }
            code {
                background: #f1f5f9;
                font-family: Consolas, "Courier New", monospace;
                font-size: 9pt;
            }
            mark {
                background: #fef08a;
                color: inherit;
            }
            .pdf-task-list {
                margin: 6px 0;
            }
            .pdf-task-item {
                border: 0;
                margin: 2px 0;
                width: 100%;
            }
            .pdf-task-item td {
                border: 0;
                padding: 1px 0;
                vertical-align: top;
            }
            .pdf-task-state {
                color: #2563eb;
                font-family: Consolas, "Courier New", monospace;
                width: 27px;
            }
            .pdf-task-content p {
                margin: 0;
            }
            .pdf-task-item.is-complete .pdf-task-content {
                color: #64748b;
                text-decoration: line-through;
            }
            .page-callout {
                background: #f1f5f9;
                border-left: 4px solid #64748b;
                margin: 9px 0;
                padding: 8px 11px;
            }
            .page-callout[data-callout-color="blue"] {
                background: #eff6ff;
                border-left-color: #3b82f6;
            }
            .page-callout[data-callout-color="green"] {
                background: #f0fdf4;
                border-left-color: #22c55e;
            }
            .page-callout[data-callout-color="yellow"] {
                background: #fefce8;
                border-left-color: #eab308;
            }
            .page-callout[data-callout-color="red"] {
                background: #fef2f2;
                border-left-color: #ef4444;
            }
            .page-callout[data-callout-color="purple"] {
                background: #faf5ff;
                border-left-color: #a855f7;
            }
            .pdf-callout-label {
                color: #475569;
                font-size: 7pt;
                font-weight: 700;
                letter-spacing: 0.7px;
                margin-bottom: 2px;
            }
            .page-toggle {
                border-left: 2px solid #cbd5e1;
                margin: 9px 0;
                padding-left: 11px;
            }
            .page-toggle-summary {
                color: #1e293b;
                font-weight: 700;
                margin-bottom: 3px;
            }
            .pdf-image-wrap {
                margin: 11px 0;
            }
            .pdf-image-wrap img {
                height: auto;
                max-width: 100%;
            }
            .pdf-image-missing {
                border: 1px dashed #94a3b8;
                color: #64748b;
                margin: 9px 0;
                padding: 10px;
            }
            .pdf-workbook {
                background: #eef8f1;
                border: 1px solid #9fc9aa;
                border-radius: 5px;
                color: #1f5132;
                margin: 9px 0;
                padding: 8px 10px;
            }
        """

    def publish_project_page_pdf(self, payload):
        """Publishes the current project/global page as a polished PDF."""
        pdf_document = None
        content_pdf_path = ""
        temp_pdf_path = ""
        try:
            if not isinstance(payload, dict):
                return {'status': 'error', 'message': 'Invalid page export payload.'}

            import fitz

            title = str(payload.get("title") or "Untitled").strip() or "Untitled"
            project_name = str(payload.get("projectName") or "").strip()
            kind = str(payload.get("kind") or "project").strip().lower()
            kind_labels = {
                "project": "Project page",
                "subpage": "Project subpage",
                "global": "Page",
            }
            kind_label = kind_labels.get(kind, "Project page")

            output_path = payload.get("outputPath") or payload.get("filePath")
            if not output_path:
                default_directory = self._resolve_dialog_directory(
                    payload.get("defaultDirectory")
                )
                safe_title = self._sanitize_template_filename_stem(
                    title, "Project Page"
                )
                window = webview.windows[0]
                output_path = window.create_file_dialog(
                    webview.FileDialog.SAVE,
                    directory=default_directory,
                    save_filename=f"{safe_title}.pdf",
                    file_types=("PDF Files (*.pdf)",),
                )
                if not output_path:
                    return {'status': 'cancelled'}

            if isinstance(output_path, (list, tuple)):
                output_path = output_path[0] if output_path else ""
            output_path = os.path.abspath(os.path.normpath(str(output_path or "")))
            if not output_path:
                return {'status': 'cancelled'}
            if not output_path.lower().endswith(".pdf"):
                output_path = f"{output_path}.pdf"

            output_directory = os.path.dirname(output_path)
            if not os.path.isdir(output_directory):
                return {
                    'status': 'error',
                    'message': 'The selected output folder does not exist.',
                }

            body_html = self._prepare_project_page_pdf_html(
                payload.get("html") or ""
            ).strip()
            if not body_html:
                body_html = '<p class="pdf-empty">This page has no content.</p>'

            context_html = ""
            if project_name and project_name.casefold() != title.casefold():
                context_html = (
                    f'<div class="pdf-context">'
                    f'{html.escape(project_name)}</div>'
                )
            document_html = (
                f'<div class="pdf-kicker">{html.escape(kind_label.upper())}</div>'
                f'<h1 class="pdf-title">{html.escape(title)}</h1>'
                f'{context_html}{body_html}'
            )

            page_rect = fitz.paper_rect("letter")
            content_rect = fitz.Rect(
                54,
                58,
                page_rect.width - 54,
                page_rect.height - 58,
            )

            def rect_function(_rect_number, _filled):
                return page_rect, content_rect, fitz.Identity

            story = fitz.Story(
                html=document_html,
                user_css=self._project_page_pdf_css(),
            )
            pdf_document = story.write_with_links(rect_function)
            page_count = len(pdf_document)

            content_pdf_path = os.path.join(
                output_directory,
                f".acies-project-page-content-{uuid.uuid4().hex}.tmp.pdf",
            )
            pdf_document.save(
                content_pdf_path,
                garbage=4,
                deflate=True,
            )
            pdf_document.close()
            pdf_document = fitz.open(content_pdf_path)

            running_font_path = os.path.join(
                os.environ.get("WINDIR", r"C:\Windows"),
                "Fonts",
                "segoeui.ttf",
            )
            has_running_font_file = os.path.isfile(running_font_path)

            for page_number, page in enumerate(pdf_document, start=1):
                running_font_kwargs = (
                    {
                        "fontname": f"aciespdfsans{page_number}",
                        "fontfile": running_font_path,
                    }
                    if has_running_font_file
                    else {"fontname": "helv"}
                )
                page.draw_line(
                    fitz.Point(54, 42),
                    fitz.Point(page_rect.width - 54, 42),
                    color=(0.80, 0.84, 0.89),
                    width=0.6,
                )
                page.insert_text(
                    fitz.Point(54, 33),
                    "ACIES",
                    fontsize=7,
                    color=(0.30, 0.36, 0.45),
                    **running_font_kwargs,
                )
                page.insert_textbox(
                    fitz.Rect(page_rect.width - 240, 22, page_rect.width - 54, 35),
                    kind_label.upper(),
                    fontsize=7,
                    color=(0.39, 0.45, 0.55),
                    align=fitz.TEXT_ALIGN_RIGHT,
                    **running_font_kwargs,
                )
                page.draw_line(
                    fitz.Point(54, page_rect.height - 42),
                    fitz.Point(page_rect.width - 54, page_rect.height - 42),
                    color=(0.80, 0.84, 0.89),
                    width=0.6,
                )
                page.insert_textbox(
                    fitz.Rect(
                        54,
                        page_rect.height - 34,
                        page_rect.width - 54,
                        page_rect.height - 20,
                    ),
                    f"Page {page_number} of {page_count}",
                    fontsize=7.5,
                    color=(0.39, 0.45, 0.55),
                    align=fitz.TEXT_ALIGN_RIGHT,
                    **running_font_kwargs,
                )

            metadata = dict(pdf_document.metadata or {})
            metadata.update({
                "title": title,
                "author": "ACIES",
                "subject": kind_label,
                "creator": "ACIES Desktop Application",
            })
            pdf_document.set_metadata(metadata)

            temp_pdf_path = os.path.join(
                output_directory,
                f".acies-project-page-{uuid.uuid4().hex}.tmp.pdf",
            )
            pdf_document.save(
                temp_pdf_path,
                garbage=4,
                deflate=True,
            )
            pdf_document.close()
            pdf_document = None
            os.replace(temp_pdf_path, output_path)
            temp_pdf_path = ""
            os.remove(content_pdf_path)
            content_pdf_path = ""

            return {
                'status': 'success',
                'path': output_path,
                'fileName': os.path.basename(output_path),
                'pageCount': page_count,
                'sizeBytes': os.path.getsize(output_path),
            }
        except Exception as e:
            logging.error(f"Error publishing project page PDF: {e}")
            return {'status': 'error', 'message': str(e)}
        finally:
            if pdf_document is not None:
                try:
                    pdf_document.close()
                except Exception:
                    pass
            if temp_pdf_path and os.path.isfile(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                except Exception:
                    pass
            if content_pdf_path and os.path.isfile(content_pdf_path):
                try:
                    os.remove(content_pdf_path)
                except Exception:
                    pass

    def export_expense_sheet_excel(self, data):
        """Exports expense sheet data to an Excel file with images."""
        try:
            import openpyxl

            template_path = get_bundled_template_path("Template_Timesheet.xlsx")

            if not template_path.exists():
                return {'status': 'error', 'message': f'Template file not found: {template_path}'}

            # Determine output path
            week_key = data.get('weekKey', 'expense')
            user_docs = os.path.join(os.path.expanduser('~'), 'Documents')
            file_path = os.path.join(user_docs, f"Expense_Sheet_{week_key}.xlsx")

            # Copy template to output location
            shutil.copy2(str(template_path), file_path)

            # Open the copied file
            wb = openpyxl.load_workbook(file_path)

            # Get the expense sheet
            if EXPENSE_SHEET_NAME not in wb.sheetnames:
                return {'status': 'error', 'message': f'Sheet "{EXPENSE_SHEET_NAME}" not found in template'}

            ws = wb[EXPENSE_SHEET_NAME]
            render_result = self._render_project_expense_sheet(
                ws,
                data.get('projects', []),
                data.get('mileageRate', 0.70),
            )
            temp_files_to_cleanup = self._add_expense_sheet_images(
                ws,
                render_result["image_paths"],
                render_result["image_start_row"],
            )

            # Save the modified file
            wb.save(file_path)

            # Clean up temp files after save
            for temp_file in temp_files_to_cleanup:
                try:
                    os.remove(temp_file)
                except:
                    pass

            # Open file on Windows
            if sys.platform == "win32":
                os.startfile(file_path)

            return {'status': 'success', 'path': file_path}

        except ImportError:
            return {'status': 'error', 'message': 'openpyxl not installed. Run: pip install openpyxl'}
        except Exception as e:
            logging.error(f"Error exporting expense sheet: {e}")
            return {'status': 'error', 'message': str(e)}

    def mark_overdue_projects_complete(self):
        """Marks all deliverables with due dates before today as complete."""
        try:
            tasks = self.get_tasks()
            today = datetime.date.today()
            count = 0
            for task in tasks:
                deliverables = task.get('deliverables')
                if isinstance(deliverables, list):
                    for deliverable in deliverables:
                        due_str = deliverable.get('due', '')
                        if due_str:
                            due_date = parse_due_str(due_str)
                            if due_date and due_date.date() < today:
                                deliverable['statuses'] = ['Complete']
                                deliverable['status'] = 'Complete'
                                sync_status_arrays(deliverable)
                                if isinstance(deliverable.get('tasks'), list):
                                    for t in deliverable['tasks']:
                                        t['done'] = True
                                count += 1
                else:
                    due_str = task.get('due', '')
                    if due_str:
                        due_date = parse_due_str(due_str)
                        if due_date and due_date.date() < today:
                            task['statuses'] = ['Complete']
                            task['status'] = 'Complete'
                            sync_status_arrays(task)
                            if isinstance(task.get('tasks'), list):
                                for t in task['tasks']:
                                    t['done'] = True
                            count += 1
            if count > 0:
                self.save_tasks(tasks)
            return {'status': 'success', 'count': count}
        except Exception as e:
            logging.error(f"Error marking overdue projects: {e}")
            return {'status': 'error', 'message': str(e)}

    def mark_overdue_projects_delivered(self):
        """Marks all deliverables with due dates before today as delivered."""
        try:
            tasks = self.get_tasks()
            today = datetime.date.today()
            count = 0
            for task in tasks:
                deliverables = task.get('deliverables')
                if isinstance(deliverables, list):
                    for deliverable in deliverables:
                        due_str = deliverable.get('due', '')
                        if due_str:
                            due_date = parse_due_str(due_str)
                            if due_date and due_date.date() < today:
                                deliverable['statuses'] = ['Delivered']
                                deliverable['status'] = 'Delivered'
                                sync_status_arrays(deliverable)
                                if isinstance(deliverable.get('tasks'), list):
                                    for t in deliverable['tasks']:
                                        t['done'] = True
                                count += 1
                else:
                    due_str = task.get('due', '')
                    if due_str:
                        due_date = parse_due_str(due_str)
                        if due_date and due_date.date() < today:
                            task['statuses'] = ['Delivered']
                            task['status'] = 'Delivered'
                            sync_status_arrays(task)
                            if isinstance(task.get('tasks'), list):
                                for t in task['tasks']:
                                    t['done'] = True
                            count += 1
            if count > 0:
                self.save_tasks(tasks)
            return {'status': 'success', 'count': count}
        except Exception as e:
            logging.error(f"Error marking overdue projects delivered: {e}")
            return {'status': 'error', 'message': str(e)}

    def delete_all_notes(self):
        """Deletes all notes data."""
        try:
            if os.path.exists(NOTES_FILE):
                os.remove(NOTES_FILE)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error deleting notes: {e}")
            return {'status': 'error', 'message': str(e)}

    def check_autocad_running(self):
        """Checks if AutoCAD is running."""
        is_running = self._is_autocad_running()
        return {'is_running': is_running}

    def uninstall_all_plugins(self):
        """Uninstalls all installed plugins."""
        try:
            # Check AutoCAD again just in case
            if self._is_autocad_running():
                return {'status': 'error', 'message': 'AutoCAD is currently running. Please close AutoCAD before uninstalling.'}

            bundles = os.listdir(self.app_plugins_folder)
            count = 0
            for bundle in bundles:
                if bundle.endswith('.bundle'):
                    bundle_path = os.path.join(self.app_plugins_folder, bundle)
                    try:
                        shutil.rmtree(bundle_path)
                        count += 1
                        logging.info(f"Uninstalled bundle: {bundle}")
                    except Exception as e:
                        logging.error(f"Failed to uninstall {bundle}: {e}")
            return {'status': 'success', 'count': count}
        except Exception as e:
            logging.error(f"Error uninstalling plugins: {e}")
            return {'status': 'error', 'message': str(e)}

    def open_url(self, url, browser=None):
        """Opens a URL or protocol link using the OS handler or a preferred browser."""
        try:
            target = str(url or "").strip()
            preferred_browser = str(browser or "").strip().lower()
            if not target:
                return {'status': 'error', 'message': 'URL is required.'}

            if target.lower().startswith("file://"):
                local_path = self._coerce_local_email_path(target)
                if local_path:
                    return self.open_path(local_path)

            if sys.platform == "win32":
                if preferred_browser == "edge" and target.lower().startswith(("http://", "https://")):
                    try:
                        os.startfile(f"microsoft-edge:{target}")
                    except OSError:
                        logging.warning("Could not launch Microsoft Edge directly; falling back to the default browser.")
                        os.startfile(target)
                else:
                    os.startfile(target)
            else:
                subprocess.run(
                    ['open', target] if sys.platform == "darwin" else ['xdg-open', target],
                    check=False
                )
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error opening URL: {e}")
            return {'status': 'error', 'message': str(e)}

    def open_notepad_with_text(self, text):
        """Writes text to a temp file and opens it in Notepad on Windows."""
        try:
            content = str(text or "")
            if not content.strip():
                return {'status': 'error', 'message': 'Text is required.'}

            temp_dir = os.path.join(get_app_data_dir(), 'temp')
            os.makedirs(temp_dir, exist_ok=True)

            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            unique = uuid.uuid4().hex[:8]
            output_path = os.path.join(
                temp_dir, f"deliverables_export_{timestamp}_{unique}.txt")
            with open(output_path, 'w', encoding='utf-8', newline='') as f:
                f.write(content)

            if sys.platform == "win32":
                subprocess.Popen(["notepad.exe", output_path])
            else:
                self.open_path(output_path)

            return {'status': 'success', 'path': output_path}
        except Exception as e:
            logging.error(f"Error opening Notepad text export: {e}")
            return {'status': 'error', 'message': str(e)}

    def export_deliverables_excel(self, data):
        """Exports selected deliverables to a basic Excel workbook."""
        try:
            payload = data or {}
            if isinstance(payload, str):
                payload = json.loads(payload)
            if not isinstance(payload, dict):
                return {'status': 'error', 'message': 'Invalid export payload.'}

            raw_entries = payload.get('entries', [])
            if not isinstance(raw_entries, list):
                return {'status': 'error', 'message': 'Deliverable entries are required.'}

            cleaned_entries = []
            for raw_entry in raw_entries:
                if not isinstance(raw_entry, dict):
                    continue

                due_raw = str(raw_entry.get('due') or '').strip()
                due_value = parse_due_str(due_raw)
                deliverable_name = str(raw_entry.get('deliverableName') or '').strip()
                cleaned_entries.append({
                    'projectId': str(raw_entry.get('projectId') or '').strip(),
                    'projectName': str(raw_entry.get('projectName') or '').strip(),
                    'deliverableName': deliverable_name or 'Untitled Deliverable',
                    'due': due_raw,
                    'dueValue': due_value.date() if due_value else None,
                    'statusText': str(raw_entry.get('statusText') or '').strip() or 'None',
                })

            if not cleaned_entries:
                return {'status': 'error', 'message': 'Select at least one deliverable to export.'}

            cleaned_entries.sort(key=lambda entry: (
                0 if entry['dueValue'] else 1,
                -(entry['dueValue'].toordinal()) if entry['dueValue'] else 0,
                entry['projectId'].lower(),
                entry['projectName'].lower(),
                entry['deliverableName'].lower(),
            ))

            file_path = payload.get('filePath')
            if isinstance(file_path, (list, tuple)):
                file_path = file_path[0] if file_path else ''
            file_path = str(file_path or '').strip()

            if not file_path:
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                selection = self.select_template_save_location(
                    default_dir=get_default_documents_dir(),
                    default_name=f"Deliverables_{timestamp}",
                    file_type='xlsx',
                )
                if selection.get('status') != 'success':
                    return selection
                file_path = str(selection.get('path') or '').strip()

            if not file_path:
                return {'status': 'cancelled'}
            if not file_path.lower().endswith('.xlsx'):
                file_path = f"{file_path}.xlsx"

            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            worksheet.title = "Deliverables"

            headers = ("Project ID", "Project Name", "Deliverable", "Due Date", "Status")
            worksheet.append(headers)
            for header_cell in worksheet[1]:
                header_cell.font = openpyxl.styles.Font(bold=True)

            for row_index, entry in enumerate(cleaned_entries, start=2):
                worksheet.cell(row=row_index, column=1, value=entry['projectId'])
                worksheet.cell(row=row_index, column=2, value=entry['projectName'])
                worksheet.cell(row=row_index, column=3, value=entry['deliverableName'])
                due_cell = worksheet.cell(
                    row=row_index,
                    column=4,
                    value=entry['dueValue'] or entry['due'],
                )
                if entry['dueValue']:
                    due_cell.number_format = "mm/dd/yyyy"
                worksheet.cell(row=row_index, column=5, value=entry['statusText'])

            worksheet.freeze_panes = "A2"
            worksheet.auto_filter.ref = worksheet.dimensions

            for column_index, width in {
                1: 14,
                2: 28,
                3: 32,
                4: 14,
                5: 18,
            }.items():
                column_letter = openpyxl.utils.get_column_letter(column_index)
                worksheet.column_dimensions[column_letter].width = width

            workbook.save(file_path)
            workbook.close()

            if sys.platform == "win32":
                os.startfile(file_path)
            else:
                self.open_path(file_path)

            return {'status': 'success', 'path': file_path, 'count': len(cleaned_entries)}
        except ImportError:
            return {'status': 'error', 'message': 'openpyxl not installed. Run: pip install openpyxl'}
        except Exception as e:
            logging.error(f"Error exporting deliverables to Excel: {e}")
            return {'status': 'error', 'message': str(e)}

    def save_dropped_email(self, upload, context=None):
        """Persists a dropped .msg/.eml payload and returns a normalized email reference."""
        try:
            payload = upload or {}
            if isinstance(payload, str):
                payload = json.loads(payload)
            if not isinstance(payload, dict):
                return {'status': 'error', 'message': 'Invalid upload payload.'}

            data_url = payload.get('dataUrl') or payload.get('data_url')
            if not data_url:
                return {'status': 'error', 'message': 'Missing email payload data.'}

            file_name = str(payload.get('name') or 'email.msg').strip() or 'email.msg'
            data, mime = self._decode_data_url(data_url)
            max_size_bytes = 20 * 1024 * 1024
            if len(data) > max_size_bytes:
                return {'status': 'error', 'message': 'Email attachment exceeds the 20 MB limit.'}

            ext = os.path.splitext(file_name)[1].lower()
            if ext not in ('.msg', '.eml'):
                mime_lower = (mime or '').lower()
                if 'rfc822' in mime_lower:
                    ext = '.eml'
                elif 'ms-outlook' in mime_lower:
                    ext = '.msg'
                else:
                    return {'status': 'error', 'message': 'Only .msg or .eml files are supported.'}

            context_data = context or {}
            if isinstance(context_data, str):
                try:
                    context_data = json.loads(context_data)
                except Exception:
                    context_data = {}

            project_hint = self._sanitize_email_path_component(
                context_data.get('projectId') or context_data.get('projectName'),
                'project'
            )
            deliverable_hint = self._sanitize_email_path_component(
                context_data.get('deliverableId') or context_data.get('deliverableName'),
                'deliverable'
            )
            file_stem = self._sanitize_email_path_component(
                os.path.splitext(os.path.basename(file_name))[0],
                'email'
            )

            email_root = self._get_email_links_root()
            target_dir = os.path.abspath(os.path.join(email_root, project_hint, deliverable_hint))
            if not self._is_within_directory(email_root, target_dir):
                return {'status': 'error', 'message': 'Invalid destination path.'}
            os.makedirs(target_dir, exist_ok=True)

            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            unique = uuid.uuid4().hex[:8]
            final_name = f"{timestamp}_{unique}_{file_stem}{ext}"
            output_path = os.path.abspath(os.path.join(target_dir, final_name))
            if not self._is_within_directory(email_root, output_path):
                return {'status': 'error', 'message': 'Invalid output path.'}

            with open(output_path, 'wb') as f:
                f.write(data)

            label = f"{file_stem}{ext}"
            saved_at = datetime.datetime.now().isoformat()
            return {
                'status': 'success',
                'emailRef': {
                    'raw': output_path,
                    'url': Path(output_path).as_uri(),
                    'label': label,
                    'source': 'saved-file',
                    'savedAt': saved_at,
                }
            }
        except Exception as e:
            logging.error(f"Error saving dropped email: {e}")
            return {'status': 'error', 'message': str(e)}

    def delete_saved_email(self, path):
        """Deletes a previously saved dropped email if it lives inside the managed email-links directory."""
        try:
            raw = str(path or '').strip()
            if not raw:
                return {'status': 'error', 'message': 'Path is required.'}

            target_path = os.path.abspath(self._coerce_local_email_path(raw))
            email_root = self._get_email_links_root()
            if not self._is_within_directory(email_root, target_path):
                return {'status': 'error', 'message': 'Path is outside managed email storage.'}

            if os.path.isdir(target_path):
                return {'status': 'error', 'message': 'Expected a file path, not a directory.'}

            deleted = False
            if os.path.exists(target_path):
                os.remove(target_path)
                deleted = True

                current = os.path.dirname(target_path)
                while current and current != email_root and self._is_within_directory(email_root, current):
                    if os.listdir(current):
                        break
                    os.rmdir(current)
                    current = os.path.dirname(current)

            return {'status': 'success', 'deleted': deleted}
        except Exception as e:
            logging.error(f"Error deleting saved email: {e}")
            return {'status': 'error', 'message': str(e)}

    def open_path(self, path):
        """Opens a path in the file explorer."""
        try:
            p = os.path.normpath(path)
            if sys.platform == "win32":
                if os.path.exists(p):
                    os.startfile(p)
                else:
                    parent = os.path.dirname(p)
                    if os.path.exists(parent):
                        os.startfile(parent)
                    else:
                        return {'status': 'error', 'message': 'Path and parent do not exist.'}
            else:
                subprocess.run(['open', p] if sys.platform ==
                               "darwin" else ['xdg-open', p])
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error opening path: {e}")
            return {'status': 'error', 'message': str(e)}

    def copy_file_to_clipboard(self, path):
        """Copies a file object to the Windows clipboard for pasting elsewhere."""
        try:
            raw_path = str(path or '').strip()
            if not raw_path:
                return {'status': 'error', 'message': 'File path is required.'}

            normalized_path = os.path.normpath(raw_path)
            if not os.path.exists(normalized_path):
                return {'status': 'error', 'message': 'File does not exist.'}
            if not os.path.isfile(normalized_path):
                return {'status': 'error', 'message': 'Expected a file path.'}
            if sys.platform != "win32":
                return {'status': 'error', 'message': 'File clipboard copy is only available on Windows.'}

            try:
                import win32clipboard
                import win32con
                import struct
            except Exception:
                return {'status': 'error', 'message': 'Windows clipboard support is unavailable.'}

            hdrop_header = struct.pack("<IiiII", 20, 0, 0, 0, 1)
            hdrop_file_list = f"{normalized_path}\0\0".encode("utf-16le")
            hdrop_data = hdrop_header + hdrop_file_list
            clipboard_open = False
            try:
                win32clipboard.OpenClipboard()
                clipboard_open = True
                win32clipboard.EmptyClipboard()
                win32clipboard.SetClipboardData(
                    win32con.CF_HDROP,
                    hdrop_data,
                )
            finally:
                if clipboard_open:
                    win32clipboard.CloseClipboard()

            return {'status': 'success', 'path': normalized_path}
        except Exception as e:
            logging.error(f"Error copying file to clipboard: {e}")
            return {'status': 'error', 'message': str(e)}

    def reveal_path(self, path):
        """Reveals a file or folder in the system file explorer."""
        try:
            raw_path = str(path or '').strip()
            if not raw_path:
                return {'status': 'error', 'message': 'Path is required.'}

            normalized_path = os.path.normpath(raw_path)
            target_exists = os.path.exists(normalized_path)
            parent_path = os.path.dirname(normalized_path)
            parent_exists = bool(parent_path) and os.path.exists(parent_path)

            if sys.platform == "win32":
                if target_exists and os.path.isfile(normalized_path):
                    subprocess.run(
                        ["explorer.exe", "/select,", normalized_path],
                        check=False,
                    )
                    return {
                        'status': 'success',
                        'mode': 'select-file',
                        'path': normalized_path,
                    }
                if target_exists and os.path.isdir(normalized_path):
                    os.startfile(normalized_path)
                    return {
                        'status': 'success',
                        'mode': 'open-directory',
                        'path': normalized_path,
                    }
                if parent_exists:
                    os.startfile(parent_path)
                    return {
                        'status': 'success',
                        'mode': 'open-parent',
                        'path': parent_path,
                    }
                return {
                    'status': 'error',
                    'message': 'Path and parent do not exist.',
                }

            fallback_path = normalized_path if target_exists else parent_path
            if not fallback_path or not os.path.exists(fallback_path):
                return {'status': 'error', 'message': 'Path and parent do not exist.'}

            subprocess.run(
                ['open', fallback_path] if sys.platform == "darwin" else ['xdg-open', fallback_path],
                check=False,
            )
            return {'status': 'success', 'mode': 'open-path', 'path': fallback_path}
        except Exception as e:
            logging.error(f"Error revealing path: {e}")
            return {'status': 'error', 'message': str(e)}

    def open_directory_strict(self, path):
        """Opens an existing directory without falling back to parent paths."""
        try:
            raw_path = str(path or '').strip()
            if not raw_path:
                return {'status': 'error', 'message': 'Path is required.'}

            directory_path = os.path.normpath(raw_path)
            if not os.path.exists(directory_path):
                return {'status': 'error', 'message': 'Directory does not exist.'}
            if not os.path.isdir(directory_path):
                return {'status': 'error', 'message': 'Expected a directory path.'}

            if sys.platform == "win32":
                os.startfile(directory_path)
            else:
                subprocess.run(
                    ['open', directory_path] if sys.platform == "darwin" else ['xdg-open', directory_path]
                )
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error opening directory: {e}")
            return {'status': 'error', 'message': str(e)}

    def _resolve_local_project_directory(self, project_info):
        """Resolve the local "Local Projects" folder for a project payload.

        Prefers an explicit, existing localProjectPath; otherwise derives the
        folder name from the server path basename (matching the Work Locally copy
        target), falling back to "<id> <name>" when no server path is saved.
        """
        info = project_info if isinstance(project_info, dict) else {}

        explicit_local = os.path.normpath(str(info.get('localProjectPath') or '').strip())
        if (
            explicit_local
            and explicit_local != '.'
            and os.path.isdir(self._to_windows_extended_path(explicit_local))
        ):
            return explicit_local

        server_path = str(info.get('path') or '').strip()
        if server_path:
            target_info = self._get_copy_project_target_info(server_path)
            if target_info.get('status') == 'success' and target_info.get('localProjectPath'):
                return target_info['localProjectPath']

        project_name = ' '.join(
            part for part in (
                str(info.get('id') or '').strip(),
                str(info.get('name') or '').strip(),
            ) if part
        )
        if not project_name:
            return ''

        local_root = os.path.join(_get_windows_documents_dir(), 'Local Projects')
        return os.path.normpath(os.path.join(local_root, project_name))

    def open_project_directory(self, kind, project=None):
        """Opens the server or local directory for a project.

        Backs the project-details right-click menu ("Open Server Directory" /
        "Open Local Directory"). Opens the resolved folder directly without
        falling back to a parent directory.
        """
        mode = 'local' if str(kind or '').strip().lower() == 'local' else 'server'
        try:
            info = project if isinstance(project, dict) else {}

            if mode == 'local':
                directory_path = self._resolve_local_project_directory(info)
                missing_message = 'Local project directory does not exist.'
            else:
                directory_path = os.path.normpath(str(info.get('path') or '').strip())
                if directory_path == '.':
                    directory_path = ''
                missing_message = 'Directory does not exist.'

            if not directory_path:
                return {
                    'status': 'error',
                    'mode': mode,
                    'message': f'No {mode} project path is available for this project.',
                }

            if not os.path.isdir(self._to_windows_extended_path(directory_path)):
                return {'status': 'error', 'mode': mode, 'message': missing_message}

            if sys.platform == "win32":
                os.startfile(directory_path)
            else:
                subprocess.run(
                    ['open', directory_path] if sys.platform == "darwin" else ['xdg-open', directory_path],
                    check=False,
                )
            return {'status': 'success', 'mode': mode, 'path': directory_path}
        except Exception as e:
            logging.error(f"Error opening {mode} project directory: {e}")
            return {'status': 'error', 'mode': mode, 'message': str(e)}

    def get_local_project_copy_info(self, server_project_path):
        """Returns the expected local copy path for a server project and whether it exists."""
        try:
            target_info = self._get_copy_project_target_info(server_project_path)
            if target_info.get('status') != 'success':
                return target_info
            return {
                'status': 'success',
                'serverProjectPath': target_info.get('serverProjectPath') or '',
                'projectName': target_info.get('projectName') or '',
                'path': target_info.get('localProjectPath') or '',
                'exists': bool(target_info.get('localProjectExists')),
            }
        except Exception as e:
            logging.error(f"Error reading local project copy info: {e}")
            return {'status': 'error', 'message': str(e)}

    def open_timesheets_folder(self):
        """Opens the folder where timesheets.json is stored."""
        try:
            folder = os.path.dirname(TIMESHEETS_FILE)
            return self.open_path(folder)
        except Exception as e:
            logging.error(f"Error opening timesheets folder: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_timesheets_info(self):
        """Returns details about the timesheets file path and state."""
        try:
            path = os.path.normpath(TIMESHEETS_FILE)
            exists = os.path.exists(path)
            size = os.path.getsize(path) if exists else 0
            mtime = os.path.getmtime(path) if exists else None
            mtime_iso = None
            if mtime is not None:
                mtime_iso = datetime.datetime.fromtimestamp(mtime).isoformat()
            return {
                'status': 'success',
                'path': path,
                'exists': exists,
                'size': size,
                'modified': mtime_iso
            }
        except Exception as e:
            logging.error(f"Error reading timesheets info: {e}")
            return {'status': 'error', 'message': str(e)}

    def write_timesheets_test_file(self):
        """Writes a small test file next to timesheets.json to verify writes."""
        try:
            folder = os.path.dirname(TIMESHEETS_FILE)
            os.makedirs(folder, exist_ok=True)
            test_path = os.path.join(folder, "timesheets_write_test.txt")
            stamp = datetime.datetime.now().isoformat()
            with open(test_path, "w", encoding="utf-8") as f:
                f.write(f"write_test {stamp}\n")
            exists = os.path.exists(test_path)
            if not exists:
                return {
                    'status': 'error',
                    'message': 'Test file write did not persist.',
                    'path': test_path
                }
            return {'status': 'success', 'path': test_path, 'exists': True}
        except Exception as e:
            logging.error(f"Error writing timesheets test file: {e}")
            return {'status': 'error', 'message': str(e)}

    def create_folder(self, path):
        """Creates a directory."""
        if not path:
            return {'status': 'error', 'message': 'Path cannot be empty.'}
        try:
            p = os.path.normpath(path)
            os.makedirs(p, exist_ok=True)
            return {'status': 'success'}
        except Exception as e:
            logging.error(f"Error creating folder: {e}")
            return {'status': 'error', 'message': str(e)}

    def _get_copy_project_disciplines(self, settings):
        """Normalize settings discipline value into a non-empty list."""
        if not isinstance(settings, dict):
            return ['Electrical']

        raw_value = settings.get('discipline', ['Electrical'])
        if isinstance(raw_value, str):
            candidates = [
                part.strip() for part in re.split(r'[,/;]+', raw_value) if part and part.strip()
            ]
        elif isinstance(raw_value, (list, tuple, set)):
            candidates = [str(item or '').strip() for item in raw_value if str(item or '').strip()]
        else:
            candidates = []

        if not candidates:
            return ['Electrical']

        canonical_map = {
            'electrical': 'Electrical',
            'mechanical': 'Mechanical',
            'plumbing': 'Plumbing',
            'arch': 'Arch',
            'architecture': 'Arch',
            'general': 'Electrical',
        }
        result = []
        seen = set()
        for item in candidates:
            canonical = canonical_map.get(item.lower(), item)
            key = canonical.lower()
            if key in seen:
                continue
            seen.add(key)
            result.append(canonical)

        return result or ['Electrical']

    def _normalize_copy_project_folder_names(self, folder_names):
        if isinstance(folder_names, str):
            folder_names = [folder_names]
        normalized_names = []
        seen_names = set()
        for folder_name in folder_names or []:
            clean_name = str(folder_name or '').strip()
            if not clean_name:
                continue
            if clean_name in {'.', '..'}:
                continue
            if '/' in clean_name or '\\' in clean_name:
                continue
            key = clean_name.lower()
            if key in seen_names:
                continue
            seen_names.add(key)
            normalized_names.append(clean_name)
        return normalized_names

    def _get_copy_project_default_folder_names(self, settings):
        return self._normalize_copy_project_folder_names(
            ['Arch', *self._get_copy_project_disciplines(settings), 'Xrefs', 'Documents', 'RFI']
        )

    def _normalize_copy_project_folder_requests(self, selected_folder_requests):
        if isinstance(selected_folder_requests, dict):
            selected_folder_requests = [selected_folder_requests]

        normalized_requests = []
        seen_names = set()
        for raw_request in selected_folder_requests or []:
            if not isinstance(raw_request, dict):
                continue

            folder_names = self._normalize_copy_project_folder_names([raw_request.get('name')])
            if not folder_names:
                continue

            folder_name = folder_names[0]
            key = folder_name.lower()
            if key in seen_names:
                continue
            seen_names.add(key)

            mode = str(raw_request.get('mode') or '').strip().lower()
            if mode != 'subset':
                mode = 'all'

            selected_child_names = self._normalize_copy_project_folder_names(
                raw_request.get('selectedChildNames')
            )
            include_parent_root_files = bool(raw_request.get('includeParentRootFiles'))

            if mode == 'subset':
                if not selected_child_names and not include_parent_root_files:
                    continue
            else:
                selected_child_names = []
                include_parent_root_files = False

            normalized_requests.append({
                'name': folder_name,
                'mode': mode,
                'selectedChildNames': selected_child_names,
                'includeParentRootFiles': include_parent_root_files,
            })

        return normalized_requests

    def _get_copy_project_target_info(self, server_project_path):
        normalized_server_path = os.path.normpath(str(server_project_path or '').strip())
        if not normalized_server_path:
            return {'status': 'error', 'message': 'Server project path is required.'}

        project_name = os.path.basename(normalized_server_path.rstrip('\\/'))
        if not project_name:
            return {'status': 'error', 'message': 'Invalid server project path.'}

        local_root = os.path.join(_get_windows_documents_dir(), 'Local Projects')
        local_project_path = os.path.normpath(os.path.join(local_root, project_name))
        local_project_exists = os.path.isdir(self._to_windows_extended_path(local_project_path))

        return {
            'status': 'success',
            'serverProjectPath': normalized_server_path,
            'projectName': project_name,
            'localRootPath': local_root,
            'localProjectPath': local_project_path,
            'localProjectExists': local_project_exists,
        }

    def _to_windows_extended_path(self, path):
        """Prefix Windows absolute paths so copy operations can exceed MAX_PATH."""
        normalized = os.path.normpath(str(path or ''))
        if sys.platform != "win32" or not normalized:
            return normalized
        if normalized.startswith('\\\\?\\'):
            return normalized
        if normalized.startswith('\\\\'):
            unc_path = normalized.lstrip('\\')
            return f"\\\\?\\UNC\\{unc_path}"
        if re.match(r'^[A-Za-z]:\\', normalized):
            return f"\\\\?\\{normalized}"
        return normalized

    def _is_copy_project_reparse_point(self, stat_result):
        reparse_flag = getattr(stat_module, 'FILE_ATTRIBUTE_REPARSE_POINT', 0)
        if not reparse_flag:
            return False
        file_attributes = getattr(stat_result, 'st_file_attributes', 0)
        return bool(file_attributes & reparse_flag)

    def _is_copy_project_entry_directory(self, entry):
        try:
            entry_stat = entry.stat(follow_symlinks=False)
        except Exception:
            return False

        try:
            if entry.is_symlink() or self._is_copy_project_reparse_point(entry_stat):
                return False
        except Exception:
            return False

        try:
            return entry.is_dir(follow_symlinks=False)
        except Exception:
            return False

    def _is_copy_project_entry_file(self, entry, entry_stat=None):
        try:
            resolved_stat = entry_stat or entry.stat(follow_symlinks=False)
        except Exception:
            return False

        try:
            if entry.is_symlink() or self._is_copy_project_reparse_point(resolved_stat):
                return False
        except Exception:
            return False

        try:
            return entry.is_file(follow_symlinks=False)
        except Exception:
            return False

    def _format_copy_project_size_label(self, size_bytes):
        size_value = max(0, int(size_bytes or 0))
        gigabyte = 1024 ** 3
        megabyte = 1024 ** 2
        if size_value >= gigabyte:
            return f"{size_value / gigabyte:.1f} GB"
        return f"{size_value / megabyte:.1f} MB"

    def _get_copy_project_folder_size_info(self, folder_path):
        unavailable_result = {
            'sizeBytes': None,
            'sizeLabel': 'Size unavailable',
            'sizeStatus': 'unavailable',
        }

        normalized_folder_path = os.path.normpath(str(folder_path or '').strip())
        if not normalized_folder_path:
            return unavailable_result

        total_size = 0
        pending_paths = [normalized_folder_path]
        while pending_paths:
            current_path = pending_paths.pop()
            current_copy_path = self._to_windows_extended_path(current_path)
            try:
                with os.scandir(current_copy_path) as entries:
                    for entry in entries:
                        try:
                            entry_stat = entry.stat(follow_symlinks=False)
                        except Exception:
                            return unavailable_result

                        try:
                            if entry.is_symlink() or self._is_copy_project_reparse_point(entry_stat):
                                continue
                        except Exception:
                            return unavailable_result

                        try:
                            if entry.is_dir(follow_symlinks=False):
                                pending_paths.append(os.path.join(current_path, entry.name))
                                continue
                            if entry.is_file(follow_symlinks=False):
                                total_size += max(0, int(getattr(entry_stat, 'st_size', 0) or 0))
                        except Exception:
                            return unavailable_result
            except Exception:
                return unavailable_result

        return {
            'sizeBytes': total_size,
            'sizeLabel': self._format_copy_project_size_label(total_size),
            'sizeStatus': 'available',
        }

    def _get_copy_project_has_child_folders(self, folder_path):
        normalized_folder_path = os.path.normpath(str(folder_path or '').strip())
        if not normalized_folder_path:
            return False

        try:
            with os.scandir(self._to_windows_extended_path(normalized_folder_path)) as entries:
                for entry in entries:
                    if self._is_copy_project_entry_directory(entry):
                        return True
        except Exception:
            return False

        return False

    def _get_copy_project_direct_files_size_info(self, folder_path):
        unavailable_result = {
            'sizeBytes': None,
            'sizeLabel': 'Size unavailable',
            'sizeStatus': 'unavailable',
        }

        normalized_folder_path = os.path.normpath(str(folder_path or '').strip())
        if not normalized_folder_path:
            return unavailable_result

        total_size = 0
        try:
            with os.scandir(self._to_windows_extended_path(normalized_folder_path)) as entries:
                for entry in entries:
                    try:
                        entry_stat = entry.stat(follow_symlinks=False)
                    except Exception:
                        return unavailable_result

                    if not self._is_copy_project_entry_file(entry, entry_stat):
                        continue

                    total_size += max(0, int(getattr(entry_stat, 'st_size', 0) or 0))
        except Exception:
            return unavailable_result

        return {
            'sizeBytes': total_size,
            'sizeLabel': self._format_copy_project_size_label(total_size),
            'sizeStatus': 'available',
        }

    def _get_copy_project_path_stat(self, path):
        normalized_path = os.path.normpath(str(path or '').strip())
        if not normalized_path:
            return None
        try:
            return os.stat(self._to_windows_extended_path(normalized_path), follow_symlinks=False)
        except Exception:
            return None

    def _is_copy_project_path_directory(self, path):
        path_stat = self._get_copy_project_path_stat(path)
        if not path_stat or self._is_copy_project_reparse_point(path_stat):
            return False
        return stat_module.S_ISDIR(path_stat.st_mode)

    def _is_copy_project_path_file(self, path):
        path_stat = self._get_copy_project_path_stat(path)
        if not path_stat or self._is_copy_project_reparse_point(path_stat):
            return False
        return stat_module.S_ISREG(path_stat.st_mode)

    def _format_local_project_manager_modified_label(self, timestamp_value):
        try:
            if timestamp_value is None:
                return ''
            return datetime.datetime.fromtimestamp(float(timestamp_value)).strftime("%Y-%m-%d %I:%M %p")
        except Exception:
            return ''

    def _get_copy_project_latest_modified_info(self, folder_path):
        normalized_folder_path = os.path.normpath(str(folder_path or '').strip())
        if not normalized_folder_path:
            return {
                'modifiedAt': '',
                'modifiedLabel': '',
                'modifiedTimestamp': None,
            }

        max_timestamp = None
        root_stat = self._get_copy_project_path_stat(normalized_folder_path)
        if root_stat and not self._is_copy_project_reparse_point(root_stat):
            max_timestamp = float(getattr(root_stat, 'st_mtime', 0.0) or 0.0)

        pending_paths = [normalized_folder_path]
        while pending_paths:
            current_path = pending_paths.pop()
            try:
                with os.scandir(self._to_windows_extended_path(current_path)) as entries:
                    for entry in entries:
                        try:
                            entry_stat = entry.stat(follow_symlinks=False)
                        except Exception:
                            continue

                        try:
                            if entry.is_symlink() or self._is_copy_project_reparse_point(entry_stat):
                                continue
                        except Exception:
                            continue

                        entry_timestamp = float(getattr(entry_stat, 'st_mtime', 0.0) or 0.0)
                        if max_timestamp is None or entry_timestamp > max_timestamp:
                            max_timestamp = entry_timestamp

                        try:
                            if entry.is_dir(follow_symlinks=False):
                                pending_paths.append(os.path.join(current_path, entry.name))
                        except Exception:
                            continue
            except Exception:
                continue

        if max_timestamp is None:
            return {
                'modifiedAt': '',
                'modifiedLabel': '',
                'modifiedTimestamp': None,
            }

        modified_iso = datetime.datetime.fromtimestamp(
            max_timestamp, tz=datetime.timezone.utc
        ).replace(microsecond=0).isoformat().replace("+00:00", "Z")
        return {
            'modifiedAt': modified_iso,
            'modifiedLabel': self._format_local_project_manager_modified_label(max_timestamp),
            'modifiedTimestamp': max_timestamp,
        }

    def _normalize_local_project_manager_relative_paths(self, relative_paths):
        if isinstance(relative_paths, str):
            relative_paths = [relative_paths]

        normalized_paths = []
        seen_paths = set()
        for raw_path in relative_paths or []:
            clean_path = str(raw_path or '').strip()
            if not clean_path:
                continue

            clean_path = clean_path.replace('/', os.sep).replace('\\', os.sep)
            normalized_path = os.path.normpath(clean_path)
            if normalized_path in ('', '.', os.curdir):
                continue
            if os.path.isabs(normalized_path):
                continue

            path_parts = [
                part for part in normalized_path.split(os.sep)
                if part and part not in ('.', os.curdir)
            ]
            if not path_parts or any(part == '..' for part in path_parts):
                continue

            safe_relative_path = os.path.normpath(os.path.join(*path_parts))
            safe_key = safe_relative_path.lower()
            if safe_key in seen_paths:
                continue
            seen_paths.add(safe_key)
            normalized_paths.append(safe_relative_path)

        return normalized_paths

    def _scan_copy_project_files(self, root_path, excluded_root_names=None, excluded_file_extensions=None):
        normalized_root_path = os.path.normpath(str(root_path or '').strip())
        results = []
        scan_errors = []
        excluded_root_name_keys = {
            str(name or '').strip().lower()
            for name in (excluded_root_names or [])
            if str(name or '').strip()
        }
        excluded_ext_keys = {
            str(ext or '').strip().lower()
            for ext in (excluded_file_extensions or [])
            if str(ext or '').strip()
        }

        if not normalized_root_path:
            return {
                'files': results,
                'scanErrors': scan_errors,
            }

        pending_paths = [('', normalized_root_path)]
        while pending_paths:
            relative_root, current_path = pending_paths.pop()
            current_copy_path = self._to_windows_extended_path(current_path)
            try:
                with os.scandir(current_copy_path) as entries:
                    for entry in entries:
                        entry_name = str(entry.name or '').strip()
                        if not entry_name:
                            continue

                        relative_path = os.path.normpath(
                            os.path.join(relative_root, entry_name)
                        ) if relative_root else entry_name

                        try:
                            entry_stat = entry.stat(follow_symlinks=False)
                        except Exception as e:
                            scan_errors.append({
                                'relativePath': relative_path,
                                'path': os.path.join(current_path, entry_name),
                                'error': str(e),
                                'reason': 'scan_error',
                            })
                            continue

                        try:
                            if entry.is_symlink() or self._is_copy_project_reparse_point(entry_stat):
                                continue
                        except Exception as e:
                            scan_errors.append({
                                'relativePath': relative_path,
                                'path': os.path.join(current_path, entry_name),
                                'error': str(e),
                                'reason': 'scan_error',
                            })
                            continue

                        try:
                            if entry.is_dir(follow_symlinks=False):
                                if (
                                    not relative_root
                                    and entry_name.lower() in excluded_root_name_keys
                                ):
                                    continue
                                pending_paths.append(
                                    (relative_path, os.path.join(current_path, entry_name))
                                )
                                continue
                            if entry.is_file(follow_symlinks=False):
                                file_ext = os.path.splitext(entry_name)[1].lower()
                                if file_ext in excluded_ext_keys:
                                    continue
                                file_size = max(0, int(getattr(entry_stat, 'st_size', 0) or 0))
                                results.append({
                                    'relativePath': relative_path,
                                    'path': os.path.join(current_path, entry_name),
                                    'sizeBytes': file_size,
                                    'sizeLabel': self._format_copy_project_size_label(file_size),
                                    'modifiedAt': _get_file_modified_iso(os.path.join(current_path, entry_name)),
                                    'modifiedTimestamp': float(
                                        getattr(entry_stat, 'st_mtime', 0.0) or 0.0
                                    ),
                                })
                        except Exception as e:
                            scan_errors.append({
                                'relativePath': relative_path,
                                'path': os.path.join(current_path, entry_name),
                                'error': str(e),
                                'reason': 'scan_error',
                            })
            except Exception as e:
                scan_errors.append({
                    'relativePath': relative_root,
                    'path': current_path,
                    'error': str(e),
                    'reason': 'scan_error',
                })

        results.sort(key=lambda item: str(item.get('relativePath') or '').lower())
        return {
            'files': results,
            'scanErrors': scan_errors,
        }

    def _get_local_project_manager_path_collision(self, root_path, relative_path):
        normalized_root_path = os.path.normpath(str(root_path or '').strip())
        normalized_relative_path = os.path.normpath(str(relative_path or '').strip())
        if not normalized_root_path or not normalized_relative_path:
            return ''

        relative_parts = [
            part for part in normalized_relative_path.split(os.sep)
            if part and part not in ('.', os.curdir)
        ]
        if not relative_parts:
            return ''

        current_path = normalized_root_path
        for part in relative_parts[:-1]:
            current_path = os.path.normpath(os.path.join(current_path, part))
            if not os.path.exists(self._to_windows_extended_path(current_path)):
                continue
            if not self._is_copy_project_path_directory(current_path):
                return current_path

        return ''

    def _build_local_project_manager_blocked_entry(
        self,
        local_project_path,
        server_project_path,
        relative_path,
        reason,
        local_path='',
        server_path='',
        error='',
    ):
        normalized_relative_path = os.path.normpath(str(relative_path or '').strip())
        normalized_local_project_path = os.path.normpath(str(local_project_path or '').strip())
        normalized_server_project_path = os.path.normpath(str(server_project_path or '').strip())
        local_display_path = str(local_path or '').strip()
        server_display_path = str(server_path or '').strip()

        if not local_display_path and normalized_relative_path:
            local_display_path = os.path.join(normalized_local_project_path, normalized_relative_path)
        if not server_display_path and normalized_relative_path:
            server_display_path = os.path.join(normalized_server_project_path, normalized_relative_path)

        return {
            'relativePath': normalized_relative_path,
            'localPath': os.path.normpath(local_display_path) if local_display_path else '',
            'serverPath': os.path.normpath(server_display_path) if server_display_path else '',
            'reason': str(reason or '').strip() or 'blocked',
            'error': str(error or '').strip(),
        }

    def _copy_local_project_manager_file(self, local_file_path, server_file_path):
        local_display_path = os.path.normpath(str(local_file_path or '').strip())
        server_display_path = os.path.normpath(str(server_file_path or '').strip())
        server_parent_path = os.path.dirname(server_display_path)
        if server_parent_path:
            os.makedirs(self._to_windows_extended_path(server_parent_path), exist_ok=True)

        shutil.copy2(
            self._to_windows_extended_path(local_display_path),
            self._to_windows_extended_path(server_display_path),
        )

    def _get_local_project_manager_excluded_root_names(self):
        return {'archive', '0 archive'}

    def _get_local_project_manager_excluded_file_extensions(self):
        return {'.dwl', '.dwl2'}

    def _get_local_project_manager_managed_root_names(self, settings):
        return self._normalize_copy_project_folder_names(
            [*self._get_copy_project_disciplines(settings), 'Xrefs']
        )

    def _split_local_project_manager_relative_path(self, relative_path):
        normalized_relative_path = os.path.normpath(str(relative_path or '').strip())
        if not normalized_relative_path or normalized_relative_path in {'.', os.curdir}:
            return []
        return [
            part
            for part in re.split(r'[\\/]+', normalized_relative_path)
            if part and part not in {'.', os.curdir}
        ]

    def _describe_local_project_manager_relative_path(self, relative_path, managed_root_names=None):
        normalized_relative_path = os.path.normpath(str(relative_path or '').strip())
        path_parts = self._split_local_project_manager_relative_path(normalized_relative_path)
        root_folder = path_parts[0] if path_parts else ''
        managed_root_name_keys = {
            str(name or '').strip().lower()
            for name in (managed_root_names or [])
            if str(name or '').strip()
        }
        scope_type = (
            'managed'
            if root_folder and root_folder.lower() in managed_root_name_keys
            else 'additive_only'
        )
        return {
            'relativePath': normalized_relative_path,
            'parts': path_parts,
            'rootFolder': root_folder,
            'scopeType': scope_type,
        }

    def _get_local_project_manager_target_conflict(self, root_path, relative_path):
        normalized_root_path = os.path.normpath(str(root_path or '').strip())
        normalized_relative_path = os.path.normpath(str(relative_path or '').strip())
        if not normalized_root_path or not normalized_relative_path:
            return ''

        parent_collision_path = self._get_local_project_manager_path_collision(
            normalized_root_path,
            normalized_relative_path,
        )
        if parent_collision_path:
            return os.path.normpath(parent_collision_path)

        target_path = os.path.normpath(
            os.path.join(normalized_root_path, normalized_relative_path)
        )
        target_copy_path = self._to_windows_extended_path(target_path)
        if not os.path.exists(target_copy_path):
            return ''
        if self._is_copy_project_path_directory(target_path):
            return target_path
        if not self._is_copy_project_path_file(target_path):
            return target_path
        return ''

    def _build_local_project_manager_candidate_entry(
        self,
        local_project_path,
        server_project_path,
        relative_path,
        reason,
        local_file=None,
        server_file=None,
        direction='to_server',
        root_folder='',
        scope_type='',
        change_type='',
        direction_label='',
    ):
        normalized_relative_path = os.path.normpath(str(relative_path or '').strip())
        local_file = local_file or {}
        server_file = server_file or {}
        direction_key = str(direction or '').strip().lower()
        path_details = self._describe_local_project_manager_relative_path(
            normalized_relative_path
        )
        resolved_root_folder = str(root_folder or path_details.get('rootFolder') or '').strip()
        resolved_scope_type = str(scope_type or path_details.get('scopeType') or '').strip().lower()
        if resolved_scope_type not in {'managed', 'additive_only'}:
            resolved_scope_type = 'additive_only'
        source_file = local_file if direction_key == 'to_server' else server_file
        source_size_bytes = source_file.get('sizeBytes')
        source_size_label = str(source_file.get('sizeLabel') or '').strip()
        if not source_size_label:
            source_size_label = (
                self._format_copy_project_size_label(source_size_bytes)
                if source_size_bytes is not None
                else 'Size unavailable'
            )

        resolved_change_type = str(change_type or '').strip().lower()
        if resolved_change_type not in {'newer', 'missing'}:
            resolved_change_type = (
                'missing'
                if str(reason or '').strip().lower() in {'server_missing', 'local_missing'}
                else 'newer'
            )

        resolved_direction_label = str(direction_label or '').strip()
        if not resolved_direction_label:
            if direction_key == 'to_local':
                resolved_direction_label = (
                    'Missing locally'
                    if str(reason or '').strip().lower() == 'local_missing'
                    else 'Newer than local'
                )
            else:
                resolved_direction_label = (
                    'Missing on server'
                    if str(reason or '').strip().lower() == 'server_missing'
                    else 'Newer than server'
                )

        should_auto_select = resolved_change_type in {'newer', 'missing'}

        return {
            'relativePath': normalized_relative_path,
            'path': normalized_relative_path,
            'rootFolder': resolved_root_folder,
            'scopeType': resolved_scope_type,
            'changeType': resolved_change_type,
            'directionLabel': resolved_direction_label,
            'localPath': os.path.normpath(
                str(local_file.get('path') or os.path.join(local_project_path, normalized_relative_path))
            ) if normalized_relative_path else '',
            'serverPath': os.path.normpath(
                str(server_file.get('path') or os.path.join(server_project_path, normalized_relative_path))
            ) if normalized_relative_path else '',
            'reason': str(reason or '').strip(),
            'localModifiedAt': str(local_file.get('modifiedAt') or '').strip(),
            'serverModifiedAt': str(server_file.get('modifiedAt') or '').strip(),
            'sizeBytes': source_size_bytes,
            'sizeLabel': source_size_label,
            'selectedByDefault': should_auto_select,
        }

    def _resolve_local_project_manager_pair(
        self,
        local_project_path,
        server_project_path=None,
        launch_context=None,
    ):
        normalized_local_project_path = os.path.normpath(str(local_project_path or '').strip())
        if not normalized_local_project_path:
            return {
                'status': 'error',
                'code': 'local_project_path_required',
                'message': 'Local project path is required.',
            }
        if not os.path.isdir(self._to_windows_extended_path(normalized_local_project_path)):
            return {
                'status': 'error',
                'message': 'Local project path does not exist.',
            }

        settings = self.get_user_settings()
        source_resolution = self._resolve_copy_project_source_path(
            server_project_path,
            launch_context,
            settings,
        )
        if source_resolution.get('status') != 'success':
            return source_resolution

        normalized_server_path = source_resolution.get('path') or ''
        resolved_from_workroom = bool(source_resolution.get('resolvedFromWorkroom'))
        resolution_mode = str(source_resolution.get('resolutionMode') or '').strip()
        workroom_project_path = str(source_resolution.get('workroomProjectPath') or '').strip()

        if not os.path.isdir(self._to_windows_extended_path(normalized_server_path)):
            if resolved_from_workroom:
                return {
                    'status': 'error',
                    'code': 'manual_selection_required',
                    'message': 'Could not auto-resolve project folder from Project Workroom. Please select it manually.',
                    'resolvedFromWorkroom': True,
                    'resolvedServerProjectPath': normalized_server_path,
                    'resolutionMode': 'project_id_ancestor_not_accessible',
                    'workroomProjectPath': workroom_project_path,
                }
            return {'status': 'error', 'message': 'Server project path does not exist.'}

        return {
            'status': 'success',
            'localProjectPath': normalized_local_project_path,
            'serverProjectPath': normalized_server_path,
            'resolvedServerProjectPath': normalized_server_path,
            'resolvedFromWorkroom': resolved_from_workroom,
            'resolutionMode': resolution_mode or 'manual_selection',
            'workroomProjectPath': workroom_project_path,
        }

    def _compare_local_project_manager_files(self, local_project_path, server_project_path):
        normalized_local_project_path = os.path.normpath(str(local_project_path or '').strip())
        normalized_server_project_path = os.path.normpath(str(server_project_path or '').strip())
        if not normalized_local_project_path or not normalized_server_project_path:
            return {
                'status': 'error',
                'message': 'Local and server project paths are required.',
            }

        try:
            settings = self.get_user_settings()
        except Exception:
            settings = {}
        managed_root_names = self._get_local_project_manager_managed_root_names(settings)
        excluded_root_names = self._get_local_project_manager_excluded_root_names()
        excluded_file_exts = self._get_local_project_manager_excluded_file_extensions()
        local_scan = self._scan_copy_project_files(
            normalized_local_project_path,
            excluded_root_names=excluded_root_names,
            excluded_file_extensions=excluded_file_exts,
        )
        server_scan = self._scan_copy_project_files(
            normalized_server_project_path,
            excluded_root_names=excluded_root_names,
            excluded_file_extensions=excluded_file_exts,
        )

        local_files_by_relative_path = {
            str(entry.get('relativePath') or '').strip().lower(): entry
            for entry in local_scan.get('files', [])
            if str(entry.get('relativePath') or '').strip()
        }
        server_files_by_relative_path = {
            str(entry.get('relativePath') or '').strip().lower(): entry
            for entry in server_scan.get('files', [])
            if str(entry.get('relativePath') or '').strip()
        }

        local_to_server_candidates = []
        server_to_local_candidates = []
        local_to_server_blocked_entries = []
        server_to_local_blocked_entries = []
        equal_files = []
        comparison_tolerance_seconds = 60.0
        seen_local_blocked = set()
        seen_server_blocked = set()

        def add_local_blocked_entry(entry):
            dedupe_key = (
                f"{str(entry.get('relativePath') or '').strip().lower()}|"
                f"{str(entry.get('serverPath') or '').strip().lower()}|"
                f"{str(entry.get('reason') or '').strip().lower()}|"
                f"{str(entry.get('error') or '').strip().lower()}"
            )
            if dedupe_key in seen_local_blocked:
                return
            seen_local_blocked.add(dedupe_key)
            local_to_server_blocked_entries.append(entry)

        def add_server_blocked_entry(entry):
            dedupe_key = (
                f"{str(entry.get('relativePath') or '').strip().lower()}|"
                f"{str(entry.get('localPath') or '').strip().lower()}|"
                f"{str(entry.get('reason') or '').strip().lower()}|"
                f"{str(entry.get('error') or '').strip().lower()}"
            )
            if dedupe_key in seen_server_blocked:
                return
            seen_server_blocked.add(dedupe_key)
            server_to_local_blocked_entries.append(entry)

        for scan_error in local_scan.get('scanErrors', []):
            add_local_blocked_entry(
                self._build_local_project_manager_blocked_entry(
                    normalized_local_project_path,
                    normalized_server_project_path,
                    scan_error.get('relativePath') or '',
                    'scan_error',
                    local_path=scan_error.get('path') or '',
                    error=scan_error.get('error') or '',
                )
            )

        for scan_error in server_scan.get('scanErrors', []):
            add_server_blocked_entry(
                self._build_local_project_manager_blocked_entry(
                    normalized_local_project_path,
                    normalized_server_project_path,
                    scan_error.get('relativePath') or '',
                    'scan_error',
                    server_path=scan_error.get('path') or '',
                    error=scan_error.get('error') or '',
                )
            )

        all_relative_path_keys = sorted(
            set(local_files_by_relative_path.keys()) | set(server_files_by_relative_path.keys())
        )
        for relative_path_key in all_relative_path_keys:
            local_file = local_files_by_relative_path.get(relative_path_key)
            server_file = server_files_by_relative_path.get(relative_path_key)
            relative_path = str(
                (local_file or server_file or {}).get('relativePath') or ''
            ).strip()
            if not relative_path:
                continue
            path_details = self._describe_local_project_manager_relative_path(
                relative_path,
                managed_root_names=managed_root_names,
            )
            root_folder = str(path_details.get('rootFolder') or '').strip()
            scope_type = str(path_details.get('scopeType') or 'additive_only').strip().lower()

            if local_file and server_file:
                local_modified_timestamp = float(local_file.get('modifiedTimestamp') or 0.0)
                server_modified_timestamp = float(server_file.get('modifiedTimestamp') or 0.0)
                time_difference = local_modified_timestamp - server_modified_timestamp
                if time_difference > comparison_tolerance_seconds:
                    conflict_path = self._get_local_project_manager_target_conflict(
                        normalized_server_project_path,
                        relative_path,
                    )
                    if conflict_path:
                        add_local_blocked_entry(
                            self._build_local_project_manager_blocked_entry(
                                normalized_local_project_path,
                                normalized_server_project_path,
                                relative_path,
                                'file_directory_conflict',
                                local_path=local_file.get('path') or '',
                                server_path=conflict_path,
                            )
                        )
                        continue
                    local_to_server_candidates.append(
                        self._build_local_project_manager_candidate_entry(
                            normalized_local_project_path,
                            normalized_server_project_path,
                            relative_path,
                            'local_newer',
                            local_file=local_file,
                            server_file=server_file,
                            direction='to_server',
                            root_folder=root_folder,
                            scope_type=scope_type,
                            change_type='newer',
                            direction_label='Newer than server',
                        )
                    )
                    continue

                if time_difference < -comparison_tolerance_seconds:
                    conflict_path = self._get_local_project_manager_target_conflict(
                        normalized_local_project_path,
                        relative_path,
                    )
                    if conflict_path:
                        add_server_blocked_entry(
                            self._build_local_project_manager_blocked_entry(
                                normalized_local_project_path,
                                normalized_server_project_path,
                                relative_path,
                                'file_directory_conflict',
                                local_path=conflict_path,
                                server_path=server_file.get('path') or '',
                            )
                        )
                        continue
                    server_to_local_candidates.append(
                        self._build_local_project_manager_candidate_entry(
                            normalized_local_project_path,
                            normalized_server_project_path,
                            relative_path,
                            'server_newer',
                            local_file=local_file,
                            server_file=server_file,
                            direction='to_local',
                            root_folder=root_folder,
                            scope_type=scope_type,
                            change_type='newer',
                            direction_label='Newer than local',
                        )
                    )
                    continue

                equal_files.append({
                    'relativePath': relative_path,
                    'path': relative_path,
                    'rootFolder': root_folder,
                    'scopeType': scope_type,
                })
                continue

            if local_file:
                conflict_path = self._get_local_project_manager_target_conflict(
                    normalized_server_project_path,
                    relative_path,
                )
                if conflict_path:
                    add_local_blocked_entry(
                        self._build_local_project_manager_blocked_entry(
                            normalized_local_project_path,
                            normalized_server_project_path,
                            relative_path,
                            'file_directory_conflict',
                            local_path=local_file.get('path') or '',
                            server_path=conflict_path,
                        )
                    )
                    continue
                local_to_server_candidates.append(
                    self._build_local_project_manager_candidate_entry(
                        normalized_local_project_path,
                        normalized_server_project_path,
                        relative_path,
                        'server_missing',
                        local_file=local_file,
                        direction='to_server',
                        root_folder=root_folder,
                        scope_type=scope_type,
                        change_type='missing',
                        direction_label='Missing on server',
                    )
                )
                continue

            if server_file:
                conflict_path = self._get_local_project_manager_target_conflict(
                    normalized_local_project_path,
                    relative_path,
                )
                if conflict_path:
                    add_server_blocked_entry(
                        self._build_local_project_manager_blocked_entry(
                            normalized_local_project_path,
                            normalized_server_project_path,
                            relative_path,
                            'file_directory_conflict',
                            local_path=conflict_path,
                            server_path=server_file.get('path') or '',
                        )
                    )
                    continue
                server_to_local_candidates.append(
                    self._build_local_project_manager_candidate_entry(
                        normalized_local_project_path,
                        normalized_server_project_path,
                        relative_path,
                        'local_missing',
                        server_file=server_file,
                        direction='to_local',
                        root_folder=root_folder,
                        scope_type=scope_type,
                        change_type='missing',
                        direction_label='Missing locally',
                    )
                )

        local_to_server_candidates.sort(
            key=lambda item: str(item.get('relativePath') or '').lower()
        )
        server_to_local_candidates.sort(
            key=lambda item: str(item.get('relativePath') or '').lower()
        )
        local_to_server_blocked_entries.sort(
            key=lambda item: str(item.get('relativePath') or '').lower()
        )
        server_to_local_blocked_entries.sort(
            key=lambda item: str(item.get('relativePath') or '').lower()
        )
        equal_files.sort(key=lambda item: str(item.get('relativePath') or '').lower())

        return {
            'status': 'success',
            'managedRootNames': managed_root_names,
            'localToServerCandidates': local_to_server_candidates,
            'serverToLocalCandidates': server_to_local_candidates,
            'localToServerBlockedEntries': local_to_server_blocked_entries,
            'serverToLocalBlockedEntries': server_to_local_blocked_entries,
            'equalFiles': equal_files,
        }

    def _create_local_project_manager_backup(
        self,
        project_root_path,
        relative_paths,
        direction='to_server',
    ):
        normalized_project_root_path = os.path.normpath(str(project_root_path or '').strip())
        normalized_relative_paths = self._normalize_local_project_manager_relative_paths(
            relative_paths
        )
        if not normalized_project_root_path or not normalized_relative_paths:
            return {
                'status': 'success',
                'backupCreated': False,
                'backupPath': '',
                'copiedFileCount': 0,
                'failedFiles': [],
            }

        existing_file_entries = []
        for relative_path in normalized_relative_paths:
            source_path = os.path.normpath(
                os.path.join(normalized_project_root_path, relative_path)
            )
            if not self._is_copy_project_path_file(source_path):
                continue
            existing_file_entries.append((relative_path, source_path))

        if not existing_file_entries:
            return {
                'status': 'success',
                'backupCreated': False,
                'backupPath': '',
                'copiedFileCount': 0,
                'failedFiles': [],
            }

        direction_key = str(direction or '').strip().lower()
        if direction_key == 'to_server':
            archive_root = os.path.join(
                normalized_project_root_path,
                'Archive',
                'Local Project Manager',
            )
        else:
            project_name = os.path.basename(normalized_project_root_path.rstrip('\\/'))
            archive_root = os.path.join(
                _get_windows_documents_dir(),
                'Local Projects',
                '0 Archive',
                project_name,
            )

        backup_path = self._reserve_unique_archive_folder(archive_root)
        copied_file_count = 0
        failed_files = []
        for relative_path, source_path in existing_file_entries:
            backup_target_path = os.path.normpath(os.path.join(backup_path, relative_path))
            backup_target_parent = os.path.dirname(backup_target_path)
            if backup_target_parent:
                os.makedirs(self._to_windows_extended_path(backup_target_parent), exist_ok=True)
            try:
                shutil.copy2(
                    self._to_windows_extended_path(source_path),
                    self._to_windows_extended_path(backup_target_path),
                )
                copied_file_count += 1
            except Exception as e:
                failed_files.append({
                    'relativePath': relative_path,
                    'source': source_path,
                    'destination': backup_target_path,
                    'error': str(e),
                })

        self._cleanup_old_backups(archive_root, max_backups=5)
        return {
            'status': 'success',
            'backupCreated': copied_file_count > 0 or bool(failed_files),
            'backupPath': backup_path,
            'copiedFileCount': copied_file_count,
            'failedFiles': failed_files,
        }

    def list_local_project_manager_projects(self, local_root=None):
        """List local project folders available under Documents\\Local Projects."""
        try:
            resolved_local_root = os.path.normpath(
                str(local_root or os.path.join(_get_windows_documents_dir(), 'Local Projects')).strip()
            )
            local_projects = []
            if os.path.isdir(self._to_windows_extended_path(resolved_local_root)):
                with os.scandir(self._to_windows_extended_path(resolved_local_root)) as entries:
                    for entry in entries:
                        if not self._is_copy_project_entry_directory(entry):
                            continue
                        if str(entry.name or '').strip().lower() == '0 archive':
                            continue

                        local_project_path = os.path.normpath(
                            os.path.join(resolved_local_root, str(entry.name or '').strip())
                        )
                        modified_info = self._get_copy_project_latest_modified_info(local_project_path)
                        local_projects.append({
                            'name': str(entry.name or '').strip(),
                            'localProjectPath': local_project_path,
                            'lastModifiedAt': modified_info.get('modifiedAt') or '',
                            'lastModifiedLabel': modified_info.get('modifiedLabel') or '',
                        })

            local_projects.sort(key=lambda item: str(item.get('name') or '').lower())
            return {
                'status': 'success',
                'localRootPath': resolved_local_root,
                'projects': local_projects,
            }
        except Exception as e:
            logging.error(f"Error listing local project manager projects: {e}")
            return {'status': 'error', 'message': str(e)}

    def _summarize_local_project_manager_comparison(self, comparison_result):
        local_to_server_candidates = list(
            comparison_result.get('localToServerCandidates', []) or []
        )
        server_to_local_candidates = list(
            comparison_result.get('serverToLocalCandidates', []) or []
        )
        equal_files = list(comparison_result.get('equalFiles', []) or [])

        local_to_server_count = len(local_to_server_candidates)
        server_to_local_count = len(server_to_local_candidates)
        managed_local_newer_count = sum(
            1
            for item in local_to_server_candidates
            if str(item.get('changeType') or '').strip().lower() == 'newer'
        )
        managed_server_newer_count = sum(
            1
            for item in server_to_local_candidates
            if str(item.get('changeType') or '').strip().lower() == 'newer'
        )

        summary = 'equal'
        recommendation = 'none'
        if managed_local_newer_count and managed_server_newer_count:
            summary = 'mixed'
            recommendation = 'sync-to-server'
        elif managed_local_newer_count:
            summary = 'local-newer'
            recommendation = 'sync-to-server'
        elif managed_server_newer_count:
            summary = 'server-newer'
            recommendation = 'copy-to-local'
        elif local_to_server_count and server_to_local_count:
            summary = 'additions-both'
            recommendation = 'sync-to-server'
        elif local_to_server_count:
            summary = 'local-additions'
            recommendation = 'sync-to-server'
        elif server_to_local_count:
            summary = 'server-additions'
            recommendation = 'copy-to-local'

        return {
            'summary': summary,
            'recommendation': recommendation,
            'localToServerCandidateCount': local_to_server_count,
            'serverToLocalCandidateCount': server_to_local_count,
            'managedLocalToServerNewerCount': managed_local_newer_count,
            'managedServerToLocalNewerCount': managed_server_newer_count,
            'equalFileCount': len(equal_files),
        }

    def _preview_local_project_manager_direction(
        self,
        local_project_path,
        server_project_path=None,
        launch_context=None,
        direction='to_server',
    ):
        pair_resolution = self._resolve_local_project_manager_pair(
            local_project_path,
            server_project_path,
            launch_context,
        )
        if pair_resolution.get('status') != 'success':
            return pair_resolution

        comparison_result = self._compare_local_project_manager_files(
            pair_resolution.get('localProjectPath') or '',
            pair_resolution.get('serverProjectPath') or '',
        )
        if comparison_result.get('status') != 'success':
            return comparison_result

        direction_key = str(direction or '').strip().lower()
        if direction_key == 'to_local':
            candidate_key = 'serverToLocalCandidates'
            blocked_key = 'serverToLocalBlockedEntries'
        else:
            candidate_key = 'localToServerCandidates'
            blocked_key = 'localToServerBlockedEntries'

        candidate_files = list(comparison_result.get(candidate_key, []) or [])
        blocked_entries = list(comparison_result.get(blocked_key, []) or [])
        comparison_summary = self._summarize_local_project_manager_comparison(
            comparison_result
        )

        return {
            'status': 'success',
            'localProjectPath': pair_resolution.get('localProjectPath') or '',
            'serverProjectPath': pair_resolution.get('serverProjectPath') or '',
            'resolvedServerProjectPath': pair_resolution.get('resolvedServerProjectPath') or '',
            'resolvedFromWorkroom': bool(pair_resolution.get('resolvedFromWorkroom')),
            'resolutionMode': str(pair_resolution.get('resolutionMode') or '').strip() or 'manual_selection',
            'workroomProjectPath': str(pair_resolution.get('workroomProjectPath') or '').strip(),
            'candidateFiles': candidate_files,
            'blockedEntries': blocked_entries,
            'candidateFileCount': len(candidate_files),
            'blockedEntryCount': len(blocked_entries),
            'selectedCandidateCount': sum(
                1 for item in candidate_files if item.get('selectedByDefault') is True
            ),
            'managedRootNames': list(comparison_result.get('managedRootNames', []) or []),
            'summary': comparison_summary.get('summary') or 'equal',
            'recommendation': comparison_summary.get('recommendation') or 'none',
            'localToServerCandidateCount': comparison_summary.get('localToServerCandidateCount', 0),
            'serverToLocalCandidateCount': comparison_summary.get('serverToLocalCandidateCount', 0),
            'managedLocalToServerNewerCount': comparison_summary.get('managedLocalToServerNewerCount', 0),
            'managedServerToLocalNewerCount': comparison_summary.get('managedServerToLocalNewerCount', 0),
            'equalFileCount': comparison_summary.get('equalFileCount', 0),
            'hasMixedUpdates': comparison_summary.get('summary') == 'mixed',
        }

    def _apply_local_project_manager_direction(
        self,
        local_project_path,
        server_project_path=None,
        selected_relative_paths=None,
        launch_context=None,
        direction='to_server',
    ):
        pair_resolution = self._resolve_local_project_manager_pair(
            local_project_path,
            server_project_path,
            launch_context,
        )
        if pair_resolution.get('status') != 'success':
            return pair_resolution

        normalized_relative_paths = self._normalize_local_project_manager_relative_paths(
            selected_relative_paths
        )
        direction_key = str(direction or '').strip().lower()
        if not normalized_relative_paths:
            return {
                'status': 'error',
                'code': 'no_files_selected',
                'message': (
                    'Select at least one file to copy to local.'
                    if direction_key == 'to_local'
                    else 'Select at least one file to sync to the server.'
                ),
            }

        normalized_local_project_path = pair_resolution.get('localProjectPath') or ''
        normalized_server_project_path = pair_resolution.get('serverProjectPath') or ''
        comparison_result = self._compare_local_project_manager_files(
            normalized_local_project_path,
            normalized_server_project_path,
        )
        if comparison_result.get('status') != 'success':
            return comparison_result

        if direction_key == 'to_local':
            candidate_key = 'serverToLocalCandidates'
            blocked_key = 'serverToLocalBlockedEntries'
            source_path_key = 'serverPath'
            destination_path_key = 'localPath'
            backup_project_root = normalized_local_project_path
            message_success = 'Selected files copied to local project.'
            message_empty = 'No files were copied to local project.'
        else:
            candidate_key = 'localToServerCandidates'
            blocked_key = 'localToServerBlockedEntries'
            source_path_key = 'localPath'
            destination_path_key = 'serverPath'
            backup_project_root = normalized_server_project_path
            message_success = 'Selected files synced to server.'
            message_empty = 'No files were synced.'

        candidate_lookup = {}
        for entry in comparison_result.get(candidate_key, []) or []:
            relative_path = str(entry.get('relativePath') or '').strip()
            if relative_path:
                candidate_lookup[relative_path.lower()] = entry

        blocked_lookup = {}
        for entry in comparison_result.get(blocked_key, []) or []:
            relative_path = str(entry.get('relativePath') or '').strip()
            if not relative_path:
                continue
            blocked_lookup.setdefault(relative_path.lower(), []).append(entry)

        pending_candidates = []
        failed_files = []
        blocked_entries = []
        seen_blocked_keys = set()
        for relative_path in normalized_relative_paths:
            candidate_entry = candidate_lookup.get(relative_path.lower())
            if candidate_entry:
                pending_candidates.append(candidate_entry)
                continue

            matching_blocked_entries = blocked_lookup.get(relative_path.lower(), [])
            if matching_blocked_entries:
                for blocked_entry in matching_blocked_entries:
                    blocked_key_value = (
                        f"{str(blocked_entry.get('relativePath') or '').strip().lower()}|"
                        f"{str(blocked_entry.get('localPath') or '').strip().lower()}|"
                        f"{str(blocked_entry.get('serverPath') or '').strip().lower()}|"
                        f"{str(blocked_entry.get('reason') or '').strip().lower()}|"
                        f"{str(blocked_entry.get('error') or '').strip().lower()}"
                    )
                    if blocked_key_value in seen_blocked_keys:
                        continue
                    seen_blocked_keys.add(blocked_key_value)
                    blocked_entries.append(blocked_entry)
                continue

            source_fallback_path = os.path.normpath(
                os.path.join(
                    normalized_server_project_path if direction_key == 'to_local' else normalized_local_project_path,
                    relative_path,
                )
            )
            destination_fallback_path = os.path.normpath(
                os.path.join(
                    normalized_local_project_path if direction_key == 'to_local' else normalized_server_project_path,
                    relative_path,
                )
            )
            failed_files.append({
                'relativePath': relative_path,
                'source': source_fallback_path,
                'destination': destination_fallback_path,
                'error': 'Selected file is no longer available for copying.',
            })

        backup_result = {
            'status': 'success',
            'backupCreated': False,
            'backupPath': '',
            'copiedFileCount': 0,
            'failedFiles': [],
        }
        if pending_candidates:
            backup_result = self._create_local_project_manager_backup(
                backup_project_root,
                [entry.get('relativePath') for entry in pending_candidates],
                direction=direction_key,
            )
            if backup_result.get('status') != 'success':
                return backup_result
            if backup_result.get('failedFiles'):
                return {
                    'status': 'error',
                    'code': 'backup_failed',
                    'message': 'Failed to create a backup before copying files.',
                    'localProjectPath': normalized_local_project_path,
                    'serverProjectPath': normalized_server_project_path,
                    'backupPath': backup_result.get('backupPath') or '',
                    'backupResult': backup_result,
                }

        copied_files = []
        for candidate_entry in pending_candidates:
            relative_path = str(candidate_entry.get('relativePath') or '').strip()
            if not relative_path:
                continue

            source_path = os.path.normpath(
                str(candidate_entry.get(source_path_key) or '').strip()
            )
            destination_path = os.path.normpath(
                str(candidate_entry.get(destination_path_key) or '').strip()
            )

            if not self._is_copy_project_path_file(source_path):
                failed_files.append({
                    'relativePath': relative_path,
                    'source': source_path,
                    'destination': destination_path,
                    'error': 'Source file was not found.',
                })
                continue

            conflict_path = self._get_local_project_manager_target_conflict(
                backup_project_root,
                relative_path,
            )
            if conflict_path:
                blocked_entry = self._build_local_project_manager_blocked_entry(
                    normalized_local_project_path,
                    normalized_server_project_path,
                    relative_path,
                    'file_directory_conflict',
                    local_path=destination_path if direction_key == 'to_local' else source_path,
                    server_path=destination_path if direction_key != 'to_local' else source_path,
                )
                blocked_key_value = (
                    f"{str(blocked_entry.get('relativePath') or '').strip().lower()}|"
                    f"{str(blocked_entry.get('localPath') or '').strip().lower()}|"
                    f"{str(blocked_entry.get('serverPath') or '').strip().lower()}|"
                    f"{str(blocked_entry.get('reason') or '').strip().lower()}|"
                    f"{str(blocked_entry.get('error') or '').strip().lower()}"
                )
                if blocked_key_value not in seen_blocked_keys:
                    seen_blocked_keys.add(blocked_key_value)
                    blocked_entries.append(blocked_entry)
                continue

            try:
                self._copy_local_project_manager_file(source_path, destination_path)
                copied_files.append({
                    'relativePath': relative_path,
                    'source': source_path,
                    'destination': destination_path,
                })
            except Exception as e:
                failed_files.append({
                    'relativePath': relative_path,
                    'source': source_path,
                    'destination': destination_path,
                    'error': str(e),
                })

        return {
            'status': 'success',
            'message': message_success if copied_files else message_empty,
            'localProjectPath': normalized_local_project_path,
            'serverProjectPath': normalized_server_project_path,
            'resolvedServerProjectPath': pair_resolution.get('resolvedServerProjectPath') or '',
            'resolvedFromWorkroom': bool(pair_resolution.get('resolvedFromWorkroom')),
            'resolutionMode': str(pair_resolution.get('resolutionMode') or '').strip() or 'manual_selection',
            'workroomProjectPath': str(pair_resolution.get('workroomProjectPath') or '').strip(),
            'copiedFiles': copied_files,
            'copiedFileCount': len(copied_files),
            'failedFiles': failed_files,
            'failedFileCount': len(failed_files),
            'blockedEntries': blocked_entries,
            'blockedEntryCount': len(blocked_entries),
            'backupCreated': bool(backup_result.get('backupCreated')),
            'backupPath': str(backup_result.get('backupPath') or '').strip(),
            'backupResult': backup_result,
        }

    def preview_local_project_manager_copy_to_local(
        self,
        local_project_path,
        server_project_path=None,
        launch_context=None,
    ):
        """Compare a server project against an existing local copy and recommend overwrite-to-local files."""
        try:
            return self._preview_local_project_manager_direction(
                local_project_path,
                server_project_path,
                launch_context,
                direction='to_local',
            )
        except Exception as e:
            logging.error(f"Error previewing Local Project Manager copy-to-local: {e}")
            return {'status': 'error', 'message': str(e)}

    def preview_local_project_manager_sync(self, local_project_path, server_project_path=None, launch_context=None):
        """Compare a local project copy against the server project and recommend sync-up files."""
        try:
            return self._preview_local_project_manager_direction(
                local_project_path,
                server_project_path,
                launch_context,
                direction='to_server',
            )
        except Exception as e:
            logging.error(f"Error previewing Local Project Manager sync: {e}")
            return {'status': 'error', 'message': str(e)}

    def apply_local_project_manager_copy_to_local(
        self,
        local_project_path,
        server_project_path,
        selected_relative_paths,
        launch_context=None,
    ):
        """Copy selected server files into the existing local project using relative paths."""
        try:
            return self._apply_local_project_manager_direction(
                local_project_path,
                server_project_path,
                selected_relative_paths,
                launch_context,
                direction='to_local',
            )
        except Exception as e:
            logging.error(f"Error applying Local Project Manager copy-to-local: {e}")
            return {'status': 'error', 'message': str(e)}

    def apply_local_project_manager_sync(self, local_project_path, server_project_path, selected_relative_paths, launch_context=None):
        """Copy selected local files back to the server project using relative paths."""
        try:
            return self._apply_local_project_manager_direction(
                local_project_path,
                server_project_path,
                selected_relative_paths,
                launch_context,
                direction='to_server',
            )
        except Exception as e:
            logging.error(f"Error applying Local Project Manager sync: {e}")
            return {'status': 'error', 'message': str(e)}

    def _copy_folder_contents(self, source_folder, destination_folder):
        """Recursively copy a folder with per-file failure tracking."""
        source_display_root = os.path.normpath(source_folder)
        destination_display_root = os.path.normpath(destination_folder)
        source_copy_root = self._to_windows_extended_path(source_display_root)
        destination_copy_root = self._to_windows_extended_path(destination_display_root)

        copied_file_count = 0
        failed_files = []

        os.makedirs(destination_copy_root, exist_ok=True)

        for current_root, child_dirs, child_files in os.walk(source_copy_root):
            relative_root = os.path.relpath(current_root, source_copy_root)
            if relative_root in ('.', os.curdir):
                relative_root = ''

            destination_root = destination_copy_root if not relative_root else os.path.join(
                destination_copy_root, relative_root
            )
            source_display = source_display_root if not relative_root else os.path.join(
                source_display_root, relative_root
            )
            destination_display = destination_display_root if not relative_root else os.path.join(
                destination_display_root, relative_root
            )

            try:
                os.makedirs(destination_root, exist_ok=True)
            except Exception as e:
                failed_files.append({
                    'source': source_display,
                    'destination': destination_display,
                    'error': str(e),
                })
                continue

            for child_dir in child_dirs:
                source_dir = os.path.join(source_display, child_dir)
                destination_dir = os.path.join(destination_display, child_dir)
                try:
                    os.makedirs(os.path.join(destination_root, child_dir), exist_ok=True)
                except Exception as e:
                    failed_files.append({
                        'source': source_dir,
                        'destination': destination_dir,
                        'error': str(e),
                    })

            for child_file in child_files:
                source_path = os.path.join(current_root, child_file)
                destination_path = os.path.join(destination_root, child_file)
                source_display_path = os.path.join(source_display, child_file)
                destination_display_path = os.path.join(destination_display, child_file)
                try:
                    shutil.copy2(source_path, destination_path)
                    copied_file_count += 1
                except Exception as e:
                    failed_files.append({
                        'source': source_display_path,
                        'destination': destination_display_path,
                        'error': str(e),
                    })

        return {
            'copiedFileCount': copied_file_count,
            'failedFiles': failed_files,
        }

    def _copy_folder_direct_files(self, source_folder, destination_folder):
        source_display_root = os.path.normpath(source_folder)
        destination_display_root = os.path.normpath(destination_folder)
        source_copy_root = self._to_windows_extended_path(source_display_root)
        destination_copy_root = self._to_windows_extended_path(destination_display_root)

        copied_file_count = 0
        failed_files = []
        source_file_count = 0

        os.makedirs(destination_copy_root, exist_ok=True)

        try:
            with os.scandir(source_copy_root) as entries:
                for entry in entries:
                    try:
                        entry_stat = entry.stat(follow_symlinks=False)
                    except Exception as e:
                        failed_files.append({
                            'source': os.path.join(source_display_root, str(entry.name or '')),
                            'destination': os.path.join(destination_display_root, str(entry.name or '')),
                            'error': str(e),
                        })
                        continue

                    if not self._is_copy_project_entry_file(entry, entry_stat):
                        continue

                    source_file_count += 1
                    source_path = entry.path
                    destination_path = os.path.join(destination_copy_root, entry.name)
                    source_display_path = os.path.join(source_display_root, entry.name)
                    destination_display_path = os.path.join(destination_display_root, entry.name)

                    try:
                        shutil.copy2(source_path, destination_path)
                        copied_file_count += 1
                    except Exception as e:
                        failed_files.append({
                            'source': source_display_path,
                            'destination': destination_display_path,
                            'error': str(e),
                        })
        except Exception as e:
            failed_files.append({
                'source': source_display_root,
                'destination': destination_display_root,
                'error': str(e),
            })

        return {
            'copiedFileCount': copied_file_count,
            'failedFiles': failed_files,
            'sourceFileCount': source_file_count,
        }

    def _copy_folder_selected_children(self, source_folder, destination_folder, selected_child_names, include_parent_root_files):
        copied_file_count = 0
        failed_files = []
        missing_child_folders = []
        copied_any_source_content = False

        if include_parent_root_files:
            direct_files_result = self._copy_folder_direct_files(source_folder, destination_folder)
            copied_file_count += int(direct_files_result.get('copiedFileCount', 0) or 0)
            failed_files.extend(direct_files_result.get('failedFiles', []))
            copied_any_source_content = int(direct_files_result.get('sourceFileCount', 0) or 0) > 0

        for child_name in self._normalize_copy_project_folder_names(selected_child_names):
            source_child_folder = os.path.join(source_folder, child_name)
            destination_child_folder = os.path.join(destination_folder, child_name)
            if os.path.isdir(self._to_windows_extended_path(source_child_folder)):
                copy_result = self._copy_folder_contents(source_child_folder, destination_child_folder)
                copied_file_count += int(copy_result.get('copiedFileCount', 0) or 0)
                failed_files.extend(copy_result.get('failedFiles', []))
                copied_any_source_content = True
            else:
                missing_child_folders.append(os.path.join(os.path.basename(source_folder), child_name))

        return {
            'copiedFileCount': copied_file_count,
            'failedFiles': failed_files,
            'missingChildFolders': missing_child_folders,
            'copiedAnySourceContent': copied_any_source_content,
        }

    def _resolve_copy_project_source_path(self, server_project_path, launch_context, settings):
        raw_server_path = str(server_project_path or '').strip()
        if raw_server_path:
            return {
                'status': 'success',
                'path': os.path.normpath(raw_server_path),
                'resolvedFromWorkroom': False,
                'resolutionMode': 'manual_selection',
                'workroomProjectPath': '',
            }

        context = self._resolve_workroom_context(settings, launch_context)
        source = str(context.get('source') or '').strip().lower()
        context_project_path = str(context.get('project_path') or '').strip()

        if source != 'workroom':
            if context_project_path:
                return {
                    'status': 'success',
                    'path': os.path.normpath(context_project_path),
                    'resolvedFromWorkroom': False,
                    'resolutionMode': 'launch_context_project_path',
                    'workroomProjectPath': '',
                }
            return {
                'status': 'error',
                'code': 'server_project_path_required',
                'message': 'Server project path is required.',
                'resolvedFromWorkroom': False,
                'resolutionMode': 'missing_server_path',
                'workroomProjectPath': '',
            }

        if not context_project_path:
            return {
                'status': 'error',
                'code': 'manual_selection_required',
                'message': 'Could not auto-resolve project folder from Project Workroom. Please select it manually.',
                'resolvedFromWorkroom': True,
                'resolutionMode': 'workroom_missing_project_path',
                'workroomProjectPath': '',
            }

        if os.path.isdir(self._to_windows_extended_path(context_project_path)):
            return {
                'status': 'success',
                'path': os.path.normpath(context_project_path),
                'resolvedFromWorkroom': True,
                'resolutionMode': 'workroom_project_path',
                'workroomProjectPath': context_project_path,
            }

        resolved_root = self._find_workroom_project_root_by_id(context_project_path)
        if not resolved_root:
            return {
                'status': 'error',
                'code': 'manual_selection_required',
                'message': 'Could not auto-resolve project folder from Project Workroom. Please select it manually.',
                'resolvedFromWorkroom': True,
                'resolutionMode': 'project_id_ancestor_not_found',
                'workroomProjectPath': context_project_path,
            }

        return {
            'status': 'success',
            'path': os.path.normpath(resolved_root),
            'resolvedFromWorkroom': True,
            'resolutionMode': 'project_id_ancestor',
            'workroomProjectPath': context_project_path,
        }

    def _get_backup_drawings_folder_names(self, settings):
        return self._normalize_copy_project_folder_names(
            [*self._get_copy_project_disciplines(settings), 'Xrefs']
        )

    def _build_backup_drawings_timestamp(self, now=None):
        return (now or datetime.datetime.now()).strftime("%Y%m%d_%H%M%S")

    def _reserve_unique_archive_folder(self, archive_root, timestamp=None):
        archive_root = os.path.normpath(str(archive_root or '').strip())
        if not archive_root:
            raise ValueError("Archive root path is required.")

        archive_root_copy_path = self._to_windows_extended_path(archive_root)
        os.makedirs(archive_root_copy_path, exist_ok=True)

        base_name = str(timestamp or self._build_backup_drawings_timestamp()).strip()
        if not base_name:
            raise ValueError("Archive folder name is required.")

        candidate_index = 1
        while True:
            suffix = '' if candidate_index == 1 else f" ({candidate_index})"
            candidate_path = os.path.join(archive_root, f"{base_name}{suffix}")
            candidate_copy_path = self._to_windows_extended_path(candidate_path)
            try:
                os.makedirs(candidate_copy_path, exist_ok=False)
                return os.path.normpath(candidate_path)
            except FileExistsError:
                candidate_index += 1

    def backup_project_drawings(self, project_root_path=None, launch_context=None):
        """Copy configured discipline folders and Xrefs into Archive\\<timestamp>."""
        try:
            settings = self.get_user_settings()
            source_resolution = self._resolve_copy_project_source_path(
                project_root_path, launch_context, settings
            )
            if source_resolution.get('status') != 'success':
                return source_resolution

            normalized_project_root = source_resolution.get('path') or ''
            resolved_from_workroom = bool(source_resolution.get('resolvedFromWorkroom'))
            resolution_mode = str(source_resolution.get('resolutionMode') or '').strip()
            workroom_project_path = str(source_resolution.get('workroomProjectPath') or '').strip()

            if not os.path.isdir(self._to_windows_extended_path(normalized_project_root)):
                if resolved_from_workroom:
                    return {
                        'status': 'error',
                        'code': 'manual_selection_required',
                        'message': 'Could not auto-resolve project folder from Project Workroom. Please select it manually.',
                        'resolvedFromWorkroom': True,
                        'resolvedProjectRootPath': normalized_project_root,
                        'resolutionMode': 'project_id_ancestor_not_accessible',
                        'workroomProjectPath': workroom_project_path,
                    }
                return {'status': 'error', 'message': 'Project root path does not exist.'}

            archive_root_path = os.path.normpath(os.path.join(normalized_project_root, 'Archive'))
            archive_path = self._reserve_unique_archive_folder(archive_root_path)
            required_folders = self._get_backup_drawings_folder_names(settings)

            copied_folders = []
            missing_source_folders = []
            copied_file_count = 0
            failed_files = []

            for folder_name in required_folders:
                source_folder = os.path.join(normalized_project_root, folder_name)
                destination_folder = os.path.join(archive_path, folder_name)
                if os.path.isdir(self._to_windows_extended_path(source_folder)):
                    copy_result = self._copy_folder_contents(source_folder, destination_folder)
                    copied_folders.append(folder_name)
                    copied_file_count += int(copy_result.get('copiedFileCount', 0) or 0)
                    failed_files.extend(copy_result.get('failedFiles', []))
                else:
                    missing_source_folders.append(folder_name)

            if not copied_folders:
                shutil.rmtree(self._to_windows_extended_path(archive_path), ignore_errors=True)
                return {
                    'status': 'error',
                    'code': 'nothing_to_backup',
                    'message': 'No configured discipline folders or Xrefs folder were found to back up.',
                    'projectRootPath': normalized_project_root,
                    'resolvedProjectRootPath': normalized_project_root,
                    'resolvedFromWorkroom': resolved_from_workroom,
                    'resolutionMode': resolution_mode or 'manual_selection',
                    'workroomProjectPath': workroom_project_path,
                    'archiveRootPath': archive_root_path,
                    'archivePath': archive_path,
                    'disciplines': self._get_copy_project_disciplines(settings),
                    'copiedFolders': [],
                    'missingSourceFolders': missing_source_folders,
                    'copiedFileCount': 0,
                    'failedFileCount': 0,
                    'failedFiles': [],
                }

            failed_file_count = len(failed_files)
            has_warnings = bool(missing_source_folders or failed_file_count)
            message = 'Drawing backup created.'
            if has_warnings:
                message = 'Drawing backup created with warnings.'

            return {
                'status': 'success',
                'message': message,
                'projectRootPath': normalized_project_root,
                'resolvedProjectRootPath': normalized_project_root,
                'resolvedFromWorkroom': resolved_from_workroom,
                'resolutionMode': resolution_mode or 'manual_selection',
                'workroomProjectPath': workroom_project_path,
                'archiveRootPath': archive_root_path,
                'archivePath': archive_path,
                'disciplines': self._get_copy_project_disciplines(settings),
                'copiedFolders': copied_folders,
                'missingSourceFolders': missing_source_folders,
                'copiedFileCount': copied_file_count,
                'failedFileCount': failed_file_count,
                'failedFiles': failed_files,
            }
        except Exception as e:
            logging.error(f"Error backing up project drawings: {e}")
            return {'status': 'error', 'message': str(e)}

    def preview_copy_project_locally(self, server_project_path=None, launch_context=None):
        """Return subfolder options and sizes before copying a project locally."""
        try:
            settings = self.get_user_settings()
            source_resolution = self._resolve_copy_project_source_path(
                server_project_path, launch_context, settings
            )
            if source_resolution.get('status') != 'success':
                return source_resolution

            normalized_server_path = source_resolution.get('path') or ''
            resolved_from_workroom = bool(source_resolution.get('resolvedFromWorkroom'))
            resolution_mode = str(source_resolution.get('resolutionMode') or '').strip()
            workroom_project_path = str(source_resolution.get('workroomProjectPath') or '').strip()

            if not os.path.isdir(self._to_windows_extended_path(normalized_server_path)):
                if resolved_from_workroom:
                    return {
                        'status': 'error',
                        'code': 'manual_selection_required',
                        'message': 'Could not auto-resolve project folder from Project Workroom. Please select it manually.',
                        'resolvedFromWorkroom': True,
                        'resolvedServerProjectPath': normalized_server_path,
                        'resolutionMode': 'project_id_ancestor_not_accessible',
                        'workroomProjectPath': workroom_project_path,
                    }
                return {'status': 'error', 'message': 'Server project path does not exist.'}

            target_info = self._get_copy_project_target_info(normalized_server_path)
            if target_info.get('status') != 'success':
                return target_info

            local_project_path = target_info.get('localProjectPath') or ''
            local_project_copy_path = self._to_windows_extended_path(local_project_path)
            if os.path.exists(local_project_copy_path) and not os.path.isdir(local_project_copy_path):
                return {
                    'status': 'error',
                    'message': f'Local project path is unavailable: {local_project_path}'
                }

            default_folder_names = self._get_copy_project_default_folder_names(settings)
            default_folder_keys = {folder_name.lower() for folder_name in default_folder_names}
            excluded_root_name_keys = self._get_local_project_manager_excluded_root_names()
            folder_options = []

            with os.scandir(self._to_windows_extended_path(normalized_server_path)) as entries:
                for entry in entries:
                    if not self._is_copy_project_entry_directory(entry):
                        continue

                    folder_name = str(entry.name or '').strip()
                    if not folder_name:
                        continue
                    if folder_name.lower() in excluded_root_name_keys:
                        continue

                    folder_path = os.path.normpath(os.path.join(normalized_server_path, folder_name))
                    size_info = self._get_copy_project_folder_size_info(folder_path)
                    folder_options.append({
                        'name': folder_name,
                        'path': folder_path,
                        'sizeBytes': size_info.get('sizeBytes'),
                        'sizeLabel': size_info.get('sizeLabel') or 'Size unavailable',
                        'selectedByDefault': folder_name.lower() in default_folder_keys,
                        'sizeStatus': size_info.get('sizeStatus') or 'unavailable',
                        'hasChildFolders': self._get_copy_project_has_child_folders(folder_path),
                        'childrenLoaded': False,
                    })

            folder_options.sort(
                key=lambda item: (
                    not bool(item.get('selectedByDefault')),
                    str(item.get('name') or '').lower(),
                )
            )

            return {
                'status': 'success',
                'serverProjectPath': normalized_server_path,
                'resolvedServerProjectPath': normalized_server_path,
                'resolvedFromWorkroom': resolved_from_workroom,
                'resolutionMode': resolution_mode or 'manual_selection',
                'workroomProjectPath': workroom_project_path,
                'localProjectPath': local_project_path,
                'projectName': target_info.get('projectName') or '',
                'localProjectExists': bool(target_info.get('localProjectExists')),
                'folderOptions': folder_options,
            }
        except Exception as e:
            logging.error(f"Error previewing local project copy: {e}")
            return {'status': 'error', 'message': str(e)}

    def preview_copy_project_locally_child_folders(self, server_project_path=None, parent_folder_name=None, launch_context=None):
        """Return one additional preview level for a selected top-level project folder."""
        try:
            settings = self.get_user_settings()
            source_resolution = self._resolve_copy_project_source_path(
                server_project_path, launch_context, settings
            )
            if source_resolution.get('status') != 'success':
                return source_resolution

            normalized_server_path = source_resolution.get('path') or ''
            resolved_from_workroom = bool(source_resolution.get('resolvedFromWorkroom'))
            resolution_mode = str(source_resolution.get('resolutionMode') or '').strip()
            workroom_project_path = str(source_resolution.get('workroomProjectPath') or '').strip()

            if not os.path.isdir(self._to_windows_extended_path(normalized_server_path)):
                if resolved_from_workroom:
                    return {
                        'status': 'error',
                        'code': 'manual_selection_required',
                        'message': 'Could not auto-resolve project folder from Project Workroom. Please select it manually.',
                        'resolvedFromWorkroom': True,
                        'resolvedServerProjectPath': normalized_server_path,
                        'resolutionMode': 'project_id_ancestor_not_accessible',
                        'workroomProjectPath': workroom_project_path,
                    }
                return {'status': 'error', 'message': 'Server project path does not exist.'}

            target_info = self._get_copy_project_target_info(normalized_server_path)
            if target_info.get('status') != 'success':
                return target_info

            parent_folder_names = self._normalize_copy_project_folder_names([parent_folder_name])
            if not parent_folder_names:
                return {
                    'status': 'error',
                    'code': 'parent_folder_name_required',
                    'message': 'Parent folder name is required.',
                }

            normalized_parent_folder_name = parent_folder_names[0]
            parent_folder_path = os.path.normpath(
                os.path.join(normalized_server_path, normalized_parent_folder_name)
            )
            if not os.path.isdir(self._to_windows_extended_path(parent_folder_path)):
                return {
                    'status': 'error',
                    'code': 'parent_folder_not_found',
                    'message': f'Parent folder was not found: {normalized_parent_folder_name}',
                    'serverProjectPath': normalized_server_path,
                    'resolvedServerProjectPath': normalized_server_path,
                    'resolvedFromWorkroom': resolved_from_workroom,
                    'resolutionMode': resolution_mode or 'manual_selection',
                    'workroomProjectPath': workroom_project_path,
                }

            parent_direct_files_size_info = self._get_copy_project_direct_files_size_info(parent_folder_path)
            child_folder_options = []

            with os.scandir(self._to_windows_extended_path(parent_folder_path)) as entries:
                for entry in entries:
                    if not self._is_copy_project_entry_directory(entry):
                        continue

                    child_folder_name = str(entry.name or '').strip()
                    if not child_folder_name:
                        continue

                    child_folder_path = os.path.normpath(os.path.join(parent_folder_path, child_folder_name))
                    size_info = self._get_copy_project_folder_size_info(child_folder_path)
                    child_folder_options.append({
                        'name': child_folder_name,
                        'path': child_folder_path,
                        'relativePath': os.path.join(normalized_parent_folder_name, child_folder_name),
                        'sizeBytes': size_info.get('sizeBytes'),
                        'sizeLabel': size_info.get('sizeLabel') or 'Size unavailable',
                        'sizeStatus': size_info.get('sizeStatus') or 'unavailable',
                    })

            child_folder_options.sort(key=lambda item: str(item.get('name') or '').lower())

            return {
                'status': 'success',
                'serverProjectPath': normalized_server_path,
                'resolvedServerProjectPath': normalized_server_path,
                'resolvedFromWorkroom': resolved_from_workroom,
                'resolutionMode': resolution_mode or 'manual_selection',
                'workroomProjectPath': workroom_project_path,
                'localProjectPath': target_info.get('localProjectPath') or '',
                'projectName': target_info.get('projectName') or '',
                'parentFolderName': normalized_parent_folder_name,
                'parentFolderPath': parent_folder_path,
                'parentDirectFilesSizeBytes': parent_direct_files_size_info.get('sizeBytes'),
                'parentDirectFilesSizeLabel': parent_direct_files_size_info.get('sizeLabel') or 'Size unavailable',
                'parentDirectFilesSizeStatus': parent_direct_files_size_info.get('sizeStatus') or 'unavailable',
                'childFolderOptions': child_folder_options,
            }
        except Exception as e:
            logging.error(f"Error previewing child folders for local project copy: {e}")
            return {'status': 'error', 'message': str(e)}

    def _build_copy_project_folder_requests(self, settings, selected_folder_names=None, selected_folder_requests=None):
        if selected_folder_requests is not None:
            return self._normalize_copy_project_folder_requests(selected_folder_requests)
        if selected_folder_names is None:
            return [
                {
                    'name': folder_name,
                    'mode': 'all',
                    'selectedChildNames': [],
                    'includeParentRootFiles': False,
                }
                for folder_name in self._get_copy_project_default_folder_names(settings)
            ]
        return [
            {
                'name': folder_name,
                'mode': 'all',
                'selectedChildNames': [],
                'includeParentRootFiles': False,
            }
            for folder_name in self._normalize_copy_project_folder_names(selected_folder_names)
        ]

    def _scan_copy_project_direct_server_files(self, folder_path, relative_prefix):
        normalized_folder_path = os.path.normpath(str(folder_path or '').strip())
        normalized_prefix = os.path.normpath(str(relative_prefix or '').strip())
        files = []
        scan_errors = []

        try:
            with os.scandir(self._to_windows_extended_path(normalized_folder_path)) as entries:
                for entry in entries:
                    entry_name = str(entry.name or '').strip()
                    if not entry_name:
                        continue
                    relative_path = os.path.normpath(os.path.join(normalized_prefix, entry_name))
                    try:
                        entry_stat = entry.stat(follow_symlinks=False)
                    except Exception as e:
                        scan_errors.append({
                            'relativePath': relative_path,
                            'path': os.path.join(normalized_folder_path, entry_name),
                            'error': str(e),
                            'reason': 'scan_error',
                        })
                        continue

                    if not self._is_copy_project_entry_file(entry, entry_stat):
                        continue

                    file_size = max(0, int(getattr(entry_stat, 'st_size', 0) or 0))
                    file_path = os.path.join(normalized_folder_path, entry_name)
                    files.append({
                        'relativePath': relative_path,
                        'path': file_path,
                        'sizeBytes': file_size,
                        'sizeLabel': self._format_copy_project_size_label(file_size),
                        'modifiedAt': _get_file_modified_iso(file_path),
                        'modifiedTimestamp': float(getattr(entry_stat, 'st_mtime', 0.0) or 0.0),
                    })
        except Exception as e:
            scan_errors.append({
                'relativePath': normalized_prefix,
                'path': normalized_folder_path,
                'error': str(e),
                'reason': 'scan_error',
            })

        return {
            'files': files,
            'scanErrors': scan_errors,
        }

    def _prefix_copy_project_scan_results(self, scan_result, relative_prefix):
        normalized_prefix = os.path.normpath(str(relative_prefix or '').strip())
        prefixed_files = []
        prefixed_errors = []

        for entry in scan_result.get('files', []) or []:
            entry_relative_path = str(entry.get('relativePath') or '').strip()
            relative_path = (
                os.path.normpath(os.path.join(normalized_prefix, entry_relative_path))
                if entry_relative_path
                else normalized_prefix
            )
            prefixed_files.append({
                **entry,
                'relativePath': relative_path,
            })

        for entry in scan_result.get('scanErrors', []) or []:
            entry_relative_path = str(entry.get('relativePath') or '').strip()
            relative_path = (
                os.path.normpath(os.path.join(normalized_prefix, entry_relative_path))
                if entry_relative_path
                else normalized_prefix
            )
            prefixed_errors.append({
                **entry,
                'relativePath': relative_path,
            })

        return {
            'files': prefixed_files,
            'scanErrors': prefixed_errors,
        }

    def _scan_copy_project_selected_server_payload(self, server_project_path, folder_requests):
        normalized_server_path = os.path.normpath(str(server_project_path or '').strip())
        files = []
        scan_errors = []
        missing_server_folders = []

        for folder_request in folder_requests or []:
            folder_name = str(folder_request.get('name') or '').strip()
            if not folder_name:
                continue

            source_folder = os.path.normpath(os.path.join(normalized_server_path, folder_name))
            if not os.path.isdir(self._to_windows_extended_path(source_folder)):
                missing_server_folders.append(folder_name)
                continue

            if str(folder_request.get('mode') or '').strip().lower() == 'subset':
                if bool(folder_request.get('includeParentRootFiles')):
                    direct_result = self._scan_copy_project_direct_server_files(
                        source_folder,
                        folder_name,
                    )
                    files.extend(direct_result.get('files', []))
                    scan_errors.extend(direct_result.get('scanErrors', []))

                for child_name in self._normalize_copy_project_folder_names(
                    folder_request.get('selectedChildNames')
                ):
                    source_child_folder = os.path.normpath(os.path.join(source_folder, child_name))
                    relative_prefix = os.path.normpath(os.path.join(folder_name, child_name))
                    if os.path.isdir(self._to_windows_extended_path(source_child_folder)):
                        child_scan = self._prefix_copy_project_scan_results(
                            self._scan_copy_project_files(source_child_folder),
                            relative_prefix,
                        )
                        files.extend(child_scan.get('files', []))
                        scan_errors.extend(child_scan.get('scanErrors', []))
                    else:
                        missing_server_folders.append(relative_prefix)
                continue

            folder_scan = self._prefix_copy_project_scan_results(
                self._scan_copy_project_files(source_folder),
                folder_name,
            )
            files.extend(folder_scan.get('files', []))
            scan_errors.extend(folder_scan.get('scanErrors', []))

        deduped_files = {}
        for entry in files:
            relative_path = str(entry.get('relativePath') or '').strip()
            if not relative_path:
                continue
            deduped_files[relative_path.lower()] = entry

        return {
            'files': sorted(
                deduped_files.values(),
                key=lambda item: str(item.get('relativePath') or '').lower(),
            ),
            'scanErrors': sorted(
                scan_errors,
                key=lambda item: str(item.get('relativePath') or '').lower(),
            ),
            'missingServerFolders': missing_server_folders,
        }

    def _build_copy_project_replacement_risk_file(self, local_file, server_file=None, reason='local_only'):
        local_file = local_file or {}
        server_file = server_file or {}
        local_size = local_file.get('sizeBytes')
        return {
            'relativePath': os.path.normpath(str(local_file.get('relativePath') or '').strip()),
            'localPath': os.path.normpath(str(local_file.get('path') or '').strip()),
            'serverPath': os.path.normpath(str(server_file.get('path') or '').strip()) if server_file else '',
            'reason': str(reason or '').strip() or 'local_only',
            'localModifiedAt': str(local_file.get('modifiedAt') or '').strip(),
            'serverModifiedAt': str(server_file.get('modifiedAt') or '').strip() if server_file else '',
            'sizeBytes': local_size,
            'sizeLabel': (
                str(local_file.get('sizeLabel') or '').strip()
                or (
                    self._format_copy_project_size_label(local_size)
                    if local_size is not None
                    else 'Size unavailable'
                )
            ),
        }

    def preview_copy_project_locally_replacement(
        self,
        server_project_path=None,
        launch_context=None,
        selected_folder_names=None,
        selected_folder_requests=None,
    ):
        """Preview local files that would be deleted before replacing an existing local project."""
        try:
            settings = self.get_user_settings()
            source_resolution = self._resolve_copy_project_source_path(
                server_project_path, launch_context, settings
            )
            if source_resolution.get('status') != 'success':
                return source_resolution

            normalized_server_path = source_resolution.get('path') or ''
            resolved_from_workroom = bool(source_resolution.get('resolvedFromWorkroom'))
            resolution_mode = str(source_resolution.get('resolutionMode') or '').strip()
            workroom_project_path = str(source_resolution.get('workroomProjectPath') or '').strip()

            if not os.path.isdir(self._to_windows_extended_path(normalized_server_path)):
                return {'status': 'error', 'message': 'Server project path does not exist.'}

            target_info = self._get_copy_project_target_info(normalized_server_path)
            if target_info.get('status') != 'success':
                return target_info

            local_project_path = target_info.get('localProjectPath') or ''
            local_project_exists = os.path.isdir(self._to_windows_extended_path(local_project_path))
            folder_requests = self._build_copy_project_folder_requests(
                settings,
                selected_folder_names=selected_folder_names,
                selected_folder_requests=selected_folder_requests,
            )
            if not folder_requests:
                return {
                    'status': 'error',
                    'code': 'no_folders_selected',
                    'message': 'Select at least one folder to copy locally.',
                }

            selected_server_scan = self._scan_copy_project_selected_server_payload(
                normalized_server_path,
                folder_requests,
            )
            server_files_by_relative_path = {
                str(entry.get('relativePath') or '').strip().lower(): entry
                for entry in selected_server_scan.get('files', []) or []
                if str(entry.get('relativePath') or '').strip()
            }

            newer_local_files = []
            local_only_files = []
            blocked_entries = []
            comparison_tolerance_seconds = 60.0

            for scan_error in selected_server_scan.get('scanErrors', []) or []:
                blocked_entries.append(
                    self._build_local_project_manager_blocked_entry(
                        local_project_path,
                        normalized_server_path,
                        scan_error.get('relativePath') or '',
                        'server_scan_error',
                        server_path=scan_error.get('path') or '',
                        error=scan_error.get('error') or '',
                    )
                )

            if local_project_exists:
                local_scan = self._scan_copy_project_files(local_project_path)
                for scan_error in local_scan.get('scanErrors', []) or []:
                    blocked_entries.append(
                        self._build_local_project_manager_blocked_entry(
                            local_project_path,
                            normalized_server_path,
                            scan_error.get('relativePath') or '',
                            'local_scan_error',
                            local_path=scan_error.get('path') or '',
                            error=scan_error.get('error') or '',
                        )
                    )

                for local_file in local_scan.get('files', []) or []:
                    relative_path = str(local_file.get('relativePath') or '').strip()
                    if not relative_path:
                        continue
                    server_file = server_files_by_relative_path.get(relative_path.lower())
                    if not server_file:
                        local_only_files.append(
                            self._build_copy_project_replacement_risk_file(
                                local_file,
                                None,
                                reason='local_only',
                            )
                        )
                        continue

                    local_modified_timestamp = float(local_file.get('modifiedTimestamp') or 0.0)
                    server_modified_timestamp = float(server_file.get('modifiedTimestamp') or 0.0)
                    if local_modified_timestamp - server_modified_timestamp > comparison_tolerance_seconds:
                        newer_local_files.append(
                            self._build_copy_project_replacement_risk_file(
                                local_file,
                                server_file,
                                reason='local_newer',
                            )
                        )

            return {
                'status': 'success',
                'serverProjectPath': normalized_server_path,
                'resolvedServerProjectPath': normalized_server_path,
                'resolvedFromWorkroom': resolved_from_workroom,
                'resolutionMode': resolution_mode or 'manual_selection',
                'workroomProjectPath': workroom_project_path,
                'localProjectPath': local_project_path,
                'projectName': target_info.get('projectName') or '',
                'localProjectExists': local_project_exists,
                'folderRequests': folder_requests,
                'selectedServerFileCount': len(selected_server_scan.get('files', []) or []),
                'missingServerFolders': selected_server_scan.get('missingServerFolders', []),
                'newerLocalFiles': newer_local_files,
                'newerLocalFileCount': len(newer_local_files),
                'localOnlyFiles': local_only_files,
                'localOnlyFileCount': len(local_only_files),
                'blockedEntries': blocked_entries,
                'blockedEntryCount': len(blocked_entries),
                'requiresConfirmation': local_project_exists,
            }
        except Exception as e:
            logging.error(f"Error previewing local project replacement: {e}")
            return {'status': 'error', 'message': str(e)}

    def _backup_existing_local_project_before_replace(self, local_project_path):
        normalized_local_project_path = os.path.normpath(str(local_project_path or '').strip())
        if not normalized_local_project_path:
            return {'status': 'error', 'message': 'Local project path is required.'}
        if not os.path.isdir(self._to_windows_extended_path(normalized_local_project_path)):
            return {
                'status': 'success',
                'backupCreated': False,
                'backupPath': '',
                'copiedFileCount': 0,
                'failedFiles': [],
                'failedFileCount': 0,
            }

        project_name = os.path.basename(normalized_local_project_path.rstrip('\\/'))
        backup_root = os.path.join(
            _get_windows_documents_dir(),
            'Local Projects',
            '0 Archive',
            project_name,
        )
        backup_path = self._reserve_unique_archive_folder(backup_root)
        copy_result = self._copy_folder_contents(normalized_local_project_path, backup_path)
        failed_files = list(copy_result.get('failedFiles', []) or [])
        self._cleanup_old_backups(backup_root, max_backups=5)
        return {
            'status': 'success',
            'backupCreated': True,
            'backupPath': backup_path,
            'copiedFileCount': int(copy_result.get('copiedFileCount', 0) or 0),
            'failedFiles': failed_files,
            'failedFileCount': len(failed_files),
        }

    def copy_project_locally(
        self,
        server_project_path=None,
        launch_context=None,
        selected_folder_names=None,
        selected_folder_requests=None,
        replace_existing_local=False,
    ):
        """Copy key project folders from server to local Documents\\Local Projects."""
        try:
            settings = self.get_user_settings()
            source_resolution = self._resolve_copy_project_source_path(
                server_project_path, launch_context, settings
            )
            if source_resolution.get('status') != 'success':
                return source_resolution

            normalized_server_path = source_resolution.get('path') or ''
            resolved_from_workroom = bool(source_resolution.get('resolvedFromWorkroom'))
            resolution_mode = str(source_resolution.get('resolutionMode') or '').strip()
            workroom_project_path = str(source_resolution.get('workroomProjectPath') or '').strip()

            if not os.path.isdir(self._to_windows_extended_path(normalized_server_path)):
                if resolved_from_workroom:
                    return {
                        'status': 'error',
                        'code': 'manual_selection_required',
                        'message': 'Could not auto-resolve project folder from Project Workroom. Please select it manually.',
                        'resolvedFromWorkroom': True,
                        'resolvedServerProjectPath': normalized_server_path,
                        'resolutionMode': 'project_id_ancestor_not_accessible',
                        'workroomProjectPath': workroom_project_path,
                    }
                return {'status': 'error', 'message': 'Server project path does not exist.'}

            target_info = self._get_copy_project_target_info(normalized_server_path)
            if target_info.get('status') != 'success':
                return target_info

            project_name = target_info.get('projectName') or ''
            local_root = target_info.get('localRootPath') or ''
            local_project_path = target_info.get('localProjectPath') or ''
            os.makedirs(self._to_windows_extended_path(local_root), exist_ok=True)
            local_project_copy_path = self._to_windows_extended_path(local_project_path)
            replace_existing_local = bool(replace_existing_local)

            if os.path.exists(local_project_copy_path):
                if os.path.isdir(local_project_copy_path):
                    if not replace_existing_local:
                        return {
                            'status': 'error',
                            'code': 'local_project_exists',
                            'message': f'Local project already exists: {local_project_path}',
                            'serverProjectPath': normalized_server_path,
                            'resolvedServerProjectPath': normalized_server_path,
                            'resolvedFromWorkroom': resolved_from_workroom,
                            'resolutionMode': resolution_mode or 'manual_selection',
                            'workroomProjectPath': workroom_project_path,
                            'localProjectPath': local_project_path,
                            'projectName': project_name,
                        }
                else:
                    return {
                        'status': 'error',
                        'message': f'Local project path is unavailable: {local_project_path}'
                    }

            disciplines = self._get_copy_project_disciplines(settings)
            folder_requests = self._build_copy_project_folder_requests(
                settings,
                selected_folder_names=selected_folder_names,
                selected_folder_requests=selected_folder_requests,
            )

            if not folder_requests:
                return {
                    'status': 'error',
                    'code': 'no_folders_selected',
                    'message': 'Select at least one folder to copy locally.',
                }

            backup_result = {
                'status': 'success',
                'backupCreated': False,
                'backupPath': '',
                'copiedFileCount': 0,
                'failedFiles': [],
                'failedFileCount': 0,
            }
            if os.path.isdir(local_project_copy_path) and replace_existing_local:
                backup_result = self._backup_existing_local_project_before_replace(local_project_path)
                if backup_result.get('status') != 'success':
                    return backup_result
                if backup_result.get('failedFiles'):
                    return {
                        'status': 'error',
                        'code': 'local_backup_failed',
                        'message': 'Failed to back up the existing local project before replacing it.',
                        'serverProjectPath': normalized_server_path,
                        'resolvedServerProjectPath': normalized_server_path,
                        'resolvedFromWorkroom': resolved_from_workroom,
                        'resolutionMode': resolution_mode or 'manual_selection',
                        'workroomProjectPath': workroom_project_path,
                        'localProjectPath': local_project_path,
                        'projectName': project_name,
                        'backupResult': backup_result,
                        'backupPath': backup_result.get('backupPath') or '',
                    }
                try:
                    shutil.rmtree(local_project_copy_path)
                except Exception as e:
                    return {
                        'status': 'error',
                        'code': 'local_project_delete_failed',
                        'message': f'Failed to delete existing local project: {e}',
                        'serverProjectPath': normalized_server_path,
                        'resolvedServerProjectPath': normalized_server_path,
                        'resolvedFromWorkroom': resolved_from_workroom,
                        'resolutionMode': resolution_mode or 'manual_selection',
                        'workroomProjectPath': workroom_project_path,
                        'localProjectPath': local_project_path,
                        'projectName': project_name,
                        'backupResult': backup_result,
                        'backupPath': backup_result.get('backupPath') or '',
                    }

            required_folders = [request.get('name') for request in folder_requests]
            os.makedirs(local_project_copy_path, exist_ok=False)
            for folder_name in required_folders:
                folder_path = os.path.join(local_project_path, folder_name)
                os.makedirs(self._to_windows_extended_path(folder_path), exist_ok=True)

            copied_folders = []
            missing_server_folders = []
            copied_file_count = 0
            failed_files = []
            for folder_request in folder_requests:
                folder_name = folder_request.get('name') or ''
                if not folder_name:
                    continue

                source_folder = os.path.join(normalized_server_path, folder_name)
                destination_folder = os.path.join(local_project_path, folder_name)
                if os.path.isdir(self._to_windows_extended_path(source_folder)):
                    if str(folder_request.get('mode') or '').strip().lower() == 'subset':
                        copy_result = self._copy_folder_selected_children(
                            source_folder,
                            destination_folder,
                            folder_request.get('selectedChildNames'),
                            bool(folder_request.get('includeParentRootFiles')),
                        )
                        if copy_result.get('copiedAnySourceContent'):
                            copied_folders.append(folder_name)
                        copied_file_count += int(copy_result.get('copiedFileCount', 0) or 0)
                        failed_files.extend(copy_result.get('failedFiles', []))
                        missing_server_folders.extend(copy_result.get('missingChildFolders', []))
                    else:
                        copy_result = self._copy_folder_contents(source_folder, destination_folder)
                        copied_folders.append(folder_name)
                        copied_file_count += int(copy_result.get('copiedFileCount', 0) or 0)
                        failed_files.extend(copy_result.get('failedFiles', []))
                else:
                    missing_server_folders.append(folder_name)

            failed_file_count = len(failed_files)
            copy_warnings = []
            message = 'Project copied locally.'
            if failed_file_count:
                message = f'Project copied locally with warnings: {failed_file_count} file(s) failed.'
                copy_warnings.append(message)

            return {
                'status': 'success',
                'message': message,
                'serverProjectPath': normalized_server_path,
                'resolvedServerProjectPath': normalized_server_path,
                'resolvedFromWorkroom': resolved_from_workroom,
                'resolutionMode': resolution_mode or 'manual_selection',
                'workroomProjectPath': workroom_project_path,
                'localProjectPath': local_project_path,
                'projectName': project_name,
                'disciplines': disciplines,
                'copiedFolders': copied_folders,
                'missingServerFolders': missing_server_folders,
                'copyWarnings': copy_warnings,
                'copiedFileCount': copied_file_count,
                'failedFileCount': failed_file_count,
                'failedFiles': failed_files,
                'replacedExistingLocal': replace_existing_local,
                'backupCreated': bool(backup_result.get('backupCreated')),
                'backupPath': str(backup_result.get('backupPath') or '').strip(),
                'backupResult': backup_result,
            }
        except Exception as e:
            logging.error(f"Error copying project locally: {e}")
            return {'status': 'error', 'message': str(e)}

    def compare_project_timestamps(self, local_project_path, server_project_path):
        """Compare file timestamps between local and server project folders."""
        try:
            normalized_local = os.path.normpath(str(local_project_path or '').strip())
            normalized_server = os.path.normpath(str(server_project_path or '').strip())

            if not normalized_local:
                return {'status': 'error', 'message': 'Local project path is required.'}
            if not normalized_server:
                return {'status': 'error', 'message': 'Server project path is required.'}

            local_ext = self._to_windows_extended_path(normalized_local)
            server_ext = self._to_windows_extended_path(normalized_server)

            if not os.path.isdir(local_ext):
                return {'status': 'error', 'message': 'Local project path does not exist.'}
            if not os.path.isdir(server_ext):
                return {'status': 'error', 'message': 'Server project path does not exist.'}

            comparison_result = self._compare_local_project_manager_files(
                normalized_local,
                normalized_server,
            )
            if comparison_result.get('status') != 'success':
                return comparison_result

            comparison_summary = self._summarize_local_project_manager_comparison(
                comparison_result
            )
            newer_local_files = list(
                entry
                for entry in (comparison_result.get('localToServerCandidates', []) or [])
                if str(entry.get('changeType') or '').strip().lower() == 'newer'
            )
            newer_server_files = list(
                entry
                for entry in (comparison_result.get('serverToLocalCandidates', []) or [])
                if str(entry.get('changeType') or '').strip().lower() == 'newer'
            )
            equal_files = list(comparison_result.get('equalFiles', []) or [])

            return {
                'status': 'success',
                'summary': comparison_summary.get('summary') or 'equal',
                'recommendation': comparison_summary.get('recommendation') or 'none',
                'newerLocalFiles': newer_local_files,
                'newerServerFiles': newer_server_files,
                'equalFiles': equal_files,
                'localToServerCandidateCount': comparison_summary.get('localToServerCandidateCount', 0),
                'serverToLocalCandidateCount': comparison_summary.get('serverToLocalCandidateCount', 0),
                'managedLocalToServerNewerCount': comparison_summary.get('managedLocalToServerNewerCount', 0),
                'managedServerToLocalNewerCount': comparison_summary.get('managedServerToLocalNewerCount', 0),
                'localToServerBlockedEntries': list(
                    comparison_result.get('localToServerBlockedEntries', []) or []
                ),
                'serverToLocalBlockedEntries': list(
                    comparison_result.get('serverToLocalBlockedEntries', []) or []
                ),
                'localProjectPath': normalized_local,
                'serverProjectPath': normalized_server,
            }
        except Exception as e:
            logging.error(f"Error comparing project timestamps: {e}")
            return {'status': 'error', 'message': str(e)}

    def _create_local_project_backup(self, source_path, backup_type='server'):
        """Create a timestamped backup before overwriting files."""
        try:
            source_ext = self._to_windows_extended_path(os.path.normpath(source_path))
            if not os.path.exists(source_ext):
                return {'status': 'error', 'message': 'Source path does not exist.'}

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            if backup_type == 'server':
                backup_root = os.path.join(os.path.dirname(source_path), 'Archive')
            else:
                local_root = os.path.join(_get_windows_documents_dir(), 'Local Projects')
                backup_root = os.path.join(local_root, '0 Archive')

            os.makedirs(self._to_windows_extended_path(backup_root), exist_ok=True)

            folder_name = os.path.basename(source_path.rstrip('\\/'))
            backup_path = os.path.join(backup_root, f'{folder_name}_{timestamp}')

            copy_result = self._copy_folder_contents(source_path, backup_path)

            return {
                'status': 'success',
                'backupPath': backup_path,
                'backupType': backup_type,
                'copiedFileCount': copy_result.get('copiedFileCount', 0),
                'failedFiles': copy_result.get('failedFiles', []),
            }
        except Exception as e:
            logging.error(f"Error creating backup: {e}")
            return {'status': 'error', 'message': str(e)}

    def _cleanup_old_backups(self, backup_root, max_backups=5):
        """Keep only the most recent backups, delete oldest ones."""
        try:
            backup_ext = self._to_windows_extended_path(os.path.normpath(backup_root))
            if not os.path.isdir(backup_ext):
                return

            existing_backups = []
            for entry in os.scandir(backup_ext):
                if entry.is_dir():
                    existing_backups.append({
                        'name': entry.name,
                        'path': entry.path,
                        'created': entry.stat().st_ctime,
                    })

            if len(existing_backups) <= max_backups:
                return

            existing_backups.sort(key=lambda x: x['created'])
            backups_to_delete = existing_backups[:-max_backups]

            for backup in backups_to_delete:
                try:
                    shutil.rmtree(backup['path'], ignore_errors=True)
                except Exception:
                    pass
        except Exception as e:
            logging.warning(f"Error cleaning up old backups: {e}")

    def copy_project_to_server(self, local_project_path=None, launch_context=None, selected_folder_names=None, selected_folder_requests=None):
        """Copy key project folders from local Documents\\Local Projects to server."""
        try:
            normalized_local_path = os.path.normpath(str(local_project_path or '').strip())
            if not normalized_local_path:
                return {'status': 'error', 'message': 'Local project path is required.'}

            local_ext = self._to_windows_extended_path(normalized_local_path)
            if not os.path.isdir(local_ext):
                return {'status': 'error', 'message': 'Local project path does not exist.'}

            project_name = os.path.basename(normalized_local_path.rstrip('\\/'))
            if not project_name:
                return {'status': 'error', 'message': 'Invalid local project path.'}

            source = str(launch_context.get('source') or '').strip().lower() if launch_context else ''
            server_path = ''
            if source == 'workroom':
                server_path = str(launch_context.get('rootProjectPath') or launch_context.get('projectPath') or '').strip()
            elif launch_context and launch_context.get('projectPath'):
                server_path = str(launch_context.get('projectPath')).strip()

            if not server_path:
                return {'status': 'error', 'message': 'Server project path not provided. Open from a deliverable to specify the server location.'}

            normalized_server_path = os.path.normpath(server_path)
            server_ext = self._to_windows_extended_path(normalized_server_path)
            if not os.path.isdir(server_ext):
                return {'status': 'error', 'message': 'Server project path does not exist.'}

            settings = self.get_user_settings()
            disciplines = self._get_copy_project_disciplines(settings)

            if selected_folder_requests is not None:
                folder_requests = self._normalize_copy_project_folder_requests(selected_folder_requests)
            elif selected_folder_names is None:
                folder_requests = [
                    {
                        'name': folder_name,
                        'mode': 'all',
                        'selectedChildNames': [],
                        'includeParentRootFiles': False,
                    }
                    for folder_name in self._get_copy_project_default_folder_names(settings)
                ]
            else:
                folder_requests = [
                    {
                        'name': folder_name,
                        'mode': 'all',
                        'selectedChildNames': [],
                        'includeParentRootFiles': False,
                    }
                    for folder_name in self._normalize_copy_project_folder_names(selected_folder_names)
                ]

            if not folder_requests:
                return {
                    'status': 'error',
                    'code': 'no_folders_selected',
                    'message': 'Select at least one folder to copy to server.',
                }

            backup_results = []
            for folder_request in folder_requests:
                folder_name = folder_request.get('name') or ''
                if not folder_name:
                    continue

                server_folder = os.path.join(normalized_server_path, folder_name)
                if os.path.isdir(self._to_windows_extended_path(server_folder)):
                    backup_result = self._create_local_project_backup(server_folder, backup_type='server')
                    backup_results.append(backup_result)

                    backup_root = os.path.join(normalized_server_path, 'Archive')
                    self._cleanup_old_backups(backup_root, max_backups=5)

            local_backup_root = os.path.join(_get_windows_documents_dir(), 'Local Projects', '0 Archive')
            self._cleanup_old_backups(local_backup_root, max_backups=5)

            copied_folders = []
            missing_local_folders = []
            copied_file_count = 0
            failed_files = []

            for folder_request in folder_requests:
                folder_name = folder_request.get('name') or ''
                if not folder_name:
                    continue

                source_folder = os.path.join(normalized_local_path, folder_name)
                destination_folder = os.path.join(normalized_server_path, folder_name)

                if os.path.isdir(self._to_windows_extended_path(source_folder)):
                    local_backup = self._create_local_project_backup(destination_folder, backup_type='local')
                    backup_results.append(local_backup)

                    if str(folder_request.get('mode') or '').strip().lower() == 'subset':
                        copy_result = self._copy_folder_selected_children(
                            source_folder,
                            destination_folder,
                            folder_request.get('selectedChildNames'),
                            bool(folder_request.get('includeParentRootFiles')),
                        )
                        if copy_result.get('copiedAnySourceContent'):
                            copied_folders.append(folder_name)
                        copied_file_count += int(copy_result.get('copiedFileCount', 0) or 0)
                        failed_files.extend(copy_result.get('failedFiles', []))
                    else:
                        copy_result = self._copy_folder_contents(source_folder, destination_folder)
                        copied_folders.append(folder_name)
                        copied_file_count += int(copy_result.get('copiedFileCount', 0) or 0)
                        failed_files.extend(copy_result.get('failedFiles', []))
                else:
                    missing_local_folders.append(folder_name)

            failed_file_count = len(failed_files)
            message = 'Project copied to server.'
            if failed_file_count:
                message = f'Project copied to server with warnings: {failed_file_count} file(s) failed.'

            return {
                'status': 'success',
                'message': message,
                'localProjectPath': normalized_local_path,
                'serverProjectPath': normalized_server_path,
                'projectName': project_name,
                'disciplines': disciplines,
                'copiedFolders': copied_folders,
                'missingLocalFolders': missing_local_folders,
                'copiedFileCount': copied_file_count,
                'failedFileCount': failed_file_count,
                'failedFiles': failed_files,
                'backupResults': [r for r in backup_results if r.get('status') == 'success'],
            }
        except Exception as e:
            logging.error(f"Error copying project to server: {e}")
            return {'status': 'error', 'message': str(e)}

    def get_wire_sizer_url(self):
        """Return a URL or relative path to the Wire Sizer build output."""
        candidate = BASE_DIR / "WireSizerApplication" / "dist" / "index.html"
        try:
            if candidate.exists():
                try:
                    rel_path = candidate.relative_to(BASE_DIR)
                    return {
                        'status': 'success',
                        'url': rel_path.as_posix()
                    }
                except ValueError:
                    return {
                        'status': 'success',
                        'url': candidate.resolve().as_uri()
                    }
        except Exception as e:
            logging.error(f"Error resolving Wire Sizer path: {e}")
            return {
                'status': 'error',
                'message': 'Failed to resolve Wire Sizer path.'
            }
        return {
            'status': 'error',
            'message': 'Wire Sizer build not found. Run npm install and npm run build in WireSizerApplication.'
        }

    # --- Circuit Breaker AI (Panel Schedule) ---

    def _find_free_port(self):
        """Find an available TCP port."""
        import socket
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.bind(('127.0.0.1', 0))
            return s.getsockname()[1]

    def start_circuit_breaker_server(self):
        """Start the CircuitBreakerAI FastAPI server as a subprocess."""
        # If already running, return the URL
        if hasattr(self, '_cb_process') and self._cb_process and self._cb_process.poll() is None:
            return {'status': 'success', 'url': f'http://127.0.0.1:{self._cb_port}'}

        server_script = BASE_DIR / "CircuitBreakerAI" / "server.py"
        if not server_script.exists():
            return {'status': 'error', 'message': 'CircuitBreakerAI server.py not found.'}

        port = self._find_free_port()
        self._cb_port = port

        python_exe = sys.executable

        # Build environment with API key - prefer user settings, then env vars
        env = os.environ.copy()
        settings = self.get_user_settings()
        api_key = settings.get('apiKey', '').strip()
        if not api_key:
            api_key = os.environ.get('GEMINI_API_KEY', os.environ.get('GOOGLE_API_KEY', ''))
        if api_key:
            env['GEMINI_API_KEY'] = api_key

        try:
            startupinfo = None
            creationflags = 0
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = 0  # SW_HIDE
                creationflags = subprocess.CREATE_NO_WINDOW

            self._cb_process = subprocess.Popen(
                [python_exe, str(server_script), "--port", str(port)],
                cwd=str(server_script.parent),
                startupinfo=startupinfo,
                creationflags=creationflags,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                env=env,
            )

            # Wait for server to be ready (poll the port)
            import socket
            for _ in range(100):  # 10 seconds max
                try:
                    with socket.create_connection(('127.0.0.1', port), timeout=0.1):
                        return {'status': 'success', 'url': f'http://127.0.0.1:{port}'}
                except (ConnectionRefusedError, OSError):
                    time.sleep(0.1)

            return {'status': 'error', 'message': 'CircuitBreakerAI server failed to start.'}
        except Exception as e:
            logging.error(f"Error starting CircuitBreakerAI server: {e}")
            return {'status': 'error', 'message': str(e)}

    def stop_circuit_breaker_server(self):
        """Stop the CircuitBreakerAI server subprocess."""
        if hasattr(self, '_cb_process') and self._cb_process:
            try:
                self._cb_process.terminate()
                self._cb_process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                self._cb_process.kill()
            except Exception:
                pass
            self._cb_process = None
        return {'status': 'success'}

    def save_circuit_breaker_schedule(self, suggested_name=None):
        """Save the CircuitBreakerAI schedule to a user-selected path."""
        try:
            if isinstance(suggested_name, dict):
                suggested_name = suggested_name.get('suggestedName') or suggested_name.get('suggested_name')

            output_path = BASE_DIR / "CircuitBreakerAI" / "ElectricalPanels" / "Filled_Panel_Schedules.xlsx"
            if not output_path.exists():
                return {'status': 'error', 'message': 'No schedule file found. Please generate a panel first.'}

            default_name = str(suggested_name or 'Panel_Schedule').strip() or 'Panel_Schedule'
            selection = self.select_template_save_location(default_name=default_name, file_type='xlsx')
            if selection.get('status') != 'success':
                return selection

            dest_path = selection.get('path')
            if not dest_path:
                return {'status': 'cancelled', 'path': None}

            if not dest_path.lower().endswith('.xlsx'):
                dest_path = f"{dest_path}.xlsx"

            dest_dir = os.path.dirname(dest_path)
            if dest_dir and not os.path.isdir(dest_dir):
                os.makedirs(dest_dir, exist_ok=True)

            shutil.copy2(str(output_path), dest_path)

            try:
                os.remove(output_path)
            except Exception as cleanup_error:
                logging.warning(f"Could not remove CircuitBreakerAI temp schedule: {cleanup_error}")

            return {'status': 'success', 'path': dest_path}
        except Exception as e:
            logging.error(f"Error saving CircuitBreakerAI schedule: {e}")
            return {'status': 'error', 'message': str(e)}

    def _ensure_panel_schedule_job_state(self):
        if not hasattr(self, "_panel_schedule_job_lock") or self._panel_schedule_job_lock is None:
            self._panel_schedule_job_lock = threading.Lock()
        if not hasattr(self, "_panel_schedule_job_record"):
            self._panel_schedule_job_record = None

    def _normalize_panel_schedule_payload(self, payload):
        data = payload or {}
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except Exception:
                data = {}
        return data if isinstance(data, dict) else {}

    def _count_panel_schedule_payload_panels(self, payload):
        data = self._normalize_panel_schedule_payload(payload)
        panels_payload = data.get('panels')
        if isinstance(panels_payload, list):
            count = sum(1 for item in panels_payload if isinstance(item, dict))
            if count > 0:
                return count
        return 1

    def _extract_panel_schedule_output_path(self, payload):
        data = self._normalize_panel_schedule_payload(payload)
        output_mode = str(
            data.get('outputMode') or data.get('output_mode') or 'new'
        ).strip().lower()
        if output_mode not in ("new", "existing"):
            output_mode = "new"
        output_path = (
            data.get('outputPath')
            or data.get('output_path')
            or (data.get('newOutputPath') if output_mode == "new" else data.get('existingOutputPath'))
            or (data.get('new_output_path') if output_mode == "new" else data.get('existing_output_path'))
        )
        output_path = str(output_path or '').strip()
        return os.path.normpath(output_path) if output_path else ""

    def _build_panel_schedule_job_record(self, job_id, payload, panel_count):
        output_path = self._extract_panel_schedule_output_path(payload)
        activity_id = str(
            (payload or {}).get('activityId') or (payload or {}).get('activity_id') or ''
        ).strip()
        return {
            'jobId': str(job_id or '').strip(),
            'toolId': 'toolCircuitBreaker',
            'activityId': activity_id,
            'status': 'running',
            'message': f"Running {panel_count} panel{'s' if panel_count != 1 else ''} in background...",
            'panelCount': max(int(panel_count or 0), 1),
            'completedCount': 0,
            'outputPath': output_path,
            'outputFolder': os.path.dirname(output_path) if output_path else '',
            'successCount': 0,
            'failureCount': 0,
            'results': [],
            'startedAt': utc_now_iso(),
            'finishedAt': '',
        }

    def _store_panel_schedule_job_record(self, record):
        self._ensure_panel_schedule_job_state()
        with self._panel_schedule_job_lock:
            self._panel_schedule_job_record = deepcopy(record) if isinstance(record, dict) else None

    def _get_panel_schedule_job_record(self, job_id=None):
        self._ensure_panel_schedule_job_state()
        with self._panel_schedule_job_lock:
            record = deepcopy(self._panel_schedule_job_record)
        if not isinstance(record, dict):
            return None
        if job_id is not None and str(record.get('jobId') or '') != str(job_id or '').strip():
            return None
        return record

    def _build_panel_schedule_terminal_record(self, job_id, result):
        existing = self._get_panel_schedule_job_record(job_id) or {}
        payload = result if isinstance(result, dict) else {}
        status = str(payload.get('status') or 'error').strip().lower()
        if status not in ('success', 'error'):
            status = 'error'

        def _coerce_non_negative_int(value, default=0):
            try:
                return max(int(value), 0)
            except Exception:
                return default

        output_path = str(
            payload.get('outputPath') or existing.get('outputPath') or ''
        ).strip()
        output_path = os.path.normpath(output_path) if output_path else ''
        panel_count = _coerce_non_negative_int(
            payload.get('panelCount'),
            _coerce_non_negative_int(existing.get('panelCount'), 1),
        ) or 1
        return {
            'jobId': str(job_id or '').strip(),
            'toolId': str(existing.get('toolId') or 'toolCircuitBreaker').strip() or 'toolCircuitBreaker',
            'activityId': str(
                payload.get('activityId') or existing.get('activityId') or ''
            ).strip(),
            'status': status,
            'message': str(payload.get('message') or '').strip(),
            'panelCount': panel_count,
            'completedCount': _coerce_non_negative_int(
                payload.get('completedCount'),
                _coerce_non_negative_int(
                    payload.get('successCount'),
                    0,
                ) + _coerce_non_negative_int(payload.get('failureCount'), 0),
            ),
            'outputPath': output_path,
            'outputFolder': str(
                payload.get('outputFolder') or existing.get('outputFolder') or (
                    os.path.dirname(output_path) if output_path else ''
                )
            ).strip(),
            'successCount': _coerce_non_negative_int(payload.get('successCount')),
            'failureCount': _coerce_non_negative_int(payload.get('failureCount')),
            'results': deepcopy(payload.get('results') or []) if isinstance(payload.get('results'), list) else [],
            'startedAt': str(existing.get('startedAt') or utc_now_iso()).strip(),
            'finishedAt': utc_now_iso(),
        }

    def get_panel_schedule_background_status(self, job_id):
        normalized_job_id = str(job_id or '').strip()
        if not normalized_job_id:
            return {'status': 'not_found', 'jobId': ''}
        record = self._get_panel_schedule_job_record(normalized_job_id)
        if not record:
            return {'status': 'not_found', 'jobId': normalized_job_id}
        return record

    def run_panel_schedule_background(self, payload):
        """Runs Panel Schedule AI in a background thread and notifies the UI."""
        if getattr(self, "_panel_schedule_running", False):
            return {'status': 'error', 'message': 'Panel Schedule AI is already running.'}

        normalized_payload = dict(payload or {}) if isinstance(payload, dict) else {}
        panel_count = self._count_panel_schedule_payload_panels(payload)
        job_id = uuid.uuid4().hex
        self._store_panel_schedule_job_record(
            self._build_panel_schedule_job_record(job_id, normalized_payload, panel_count)
        )
        self._panel_schedule_running = True
        try:
            thread = threading.Thread(
                target=self._panel_schedule_worker,
                args=(job_id, normalized_payload),
                daemon=True
            )
            self._panel_schedule_thread = thread
            thread.start()
        except Exception as e:
            self._panel_schedule_running = False
            self._store_panel_schedule_job_record(
                self._build_panel_schedule_terminal_record(
                    job_id,
                    {
                        'status': 'error',
                        'message': str(e),
                        'panelCount': panel_count,
                    },
                )
            )
            return {'status': 'error', 'message': str(e)}
        return {
            'status': 'started',
            'jobId': job_id,
            'panelCount': panel_count,
            'activityId': str(
                normalized_payload.get('activityId') or normalized_payload.get('activity_id') or ''
            ).strip(),
        }

    def _panel_schedule_worker(self, job_id, payload):
        window = webview.windows[0] if webview.windows else None
        result = {'status': 'error', 'message': 'Panel Schedule AI failed.'}
        try:
            result = self._process_panel_schedule_payload(payload, job_id=job_id)
        except Exception as e:
            logging.error(f"Panel Schedule AI error: {e}")
            result = {'status': 'error', 'message': str(e)}
        finally:
            self._panel_schedule_running = False

        terminal_record = self._build_panel_schedule_terminal_record(job_id, result)
        self._store_panel_schedule_job_record(terminal_record)

        if not window:
            return

        try:
            js_payload = json.dumps(terminal_record)
            window.evaluate_js(
                f"window.handlePanelScheduleResult({js_payload})"
            )
        except Exception as e:
            logging.error(f"Failed to notify Panel Schedule AI result: {e}")

    def _normalize_panel_schedule_paths(self, value):
        values = value if isinstance(value, (list, tuple)) else [value]
        normalized = []
        for raw in values:
            path = str(raw or '').strip()
            if path:
                normalized.append(path)
        return normalized

    def _get_panel_schedule_path_value(self, data, plural_keys, singular_keys):
        if not isinstance(data, dict):
            return None
        for key in plural_keys:
            if key in data:
                return data.get(key)
        for key in singular_keys:
            if key in data:
                return data.get(key)
        return []

    def _decode_data_url(self, data_url):
        if not data_url or "base64," not in data_url:
            raise ValueError("Invalid image data.")
        header, encoded = data_url.split("base64,", 1)
        mime = ""
        if header.startswith("data:"):
            mime = header.split(";", 1)[0].replace("data:", "").strip()
        return base64.b64decode(encoded), mime

    def _get_email_links_root(self):
        root = os.path.join(os.path.dirname(TASKS_FILE), "email-links")
        os.makedirs(root, exist_ok=True)
        return os.path.abspath(root)

    def _is_within_directory(self, base_path, target_path):
        try:
            base = os.path.abspath(base_path)
            target = os.path.abspath(target_path)
            return os.path.commonpath([base, target]) == base
        except Exception:
            return False

    def _sanitize_email_path_component(self, value, fallback):
        text = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", str(value or "").strip())
        text = re.sub(r"\s+", " ", text).strip().strip(".")
        text = text[:80]
        return text or fallback

    def _coerce_local_file_path(self, raw):
        value = str(raw or "").strip()
        if not value:
            return ""
        if value.lower().startswith("file://"):
            parsed = urlparse(value)
            path = unquote(parsed.path or "")
            if parsed.netloc:
                unc_path = path.replace("/", "\\")
                if unc_path and not unc_path.startswith("\\"):
                    unc_path = "\\" + unc_path
                return os.path.normpath(f"\\\\{parsed.netloc}{unc_path}")
            if sys.platform == "win32" and re.match(r"^/[A-Za-z]:", path):
                path = path[1:]
            return os.path.normpath(path)
        return os.path.normpath(value)

    def _coerce_local_email_path(self, raw):
        return self._coerce_local_file_path(raw)

    def _panel_schedule_save_uploads(self, uploads, label):
        temp_paths = []
        resolved_paths = []
        for upload in uploads or []:
            if not isinstance(upload, dict):
                continue
            data_url = upload.get("dataUrl") or upload.get("data_url")
            if not data_url:
                continue
            file_name = str(upload.get("name") or f"{label}.jpg")
            try:
                raw, mime = self._decode_data_url(data_url)
            except Exception:
                continue
            suffix = os.path.splitext(file_name)[1] or ""
            if not suffix:
                if mime.endswith("png"):
                    suffix = ".png"
                elif mime.endswith("jpeg") or mime.endswith("jpg"):
                    suffix = ".jpg"
                elif mime.endswith("heic"):
                    suffix = ".heic"
                else:
                    suffix = ".img"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(raw)
            tmp.close()
            temp_paths.append(tmp.name)
            resolved_paths.append(tmp.name)
        return resolved_paths, temp_paths

    def _resolve_panel_schedule_api_key(self):
        settings = self.get_user_settings()
        api_key = settings.get('apiKey', '').strip()
        if not api_key:
            api_key = os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY') or ''
        return (api_key or '').strip()

    def _format_panel_schedule_ai_error(self, exc):
        msg = str(exc)
        lower = msg.lower()
        if ("api key expired" in lower or
                "api_key_invalid" in lower or
                "invalid api key" in lower):
            return ('Your Google API key is expired/invalid. '
                    'Create a new key in Google AI Studio, update your settings, then try again.')
        if "deadline" in lower or "unavailable" in lower or "503" in lower:
            return ('Gemini took too long to respond (the request timed out). '
                    'This usually means the AI server is under heavy load or the '
                    'images are very large. Click Rerun to try again, or split the '
                    'panel into smaller image groups (one breaker section per run).')
        if "model" in lower and ("not found" in lower or "does not exist" in lower):
            return 'AI model not available. The Gemini 3 Flash model may not be accessible with your API key.'
        if "quota" in lower or "rate limit" in lower:
            return 'API rate limit exceeded. Please wait a moment and try again.'
        return msg

    def _normalize_panel_schedule_extension(self, ext_value):
        ext = str(ext_value or "").strip().lower()
        if not ext:
            return ""
        if not ext.startswith("."):
            ext = f".{ext}"
        return ext

    def _panel_schedule_excel_requirement_message(self):
        return "Editing .xls requires Microsoft Excel installed on this machine."

    def _convert_excel_workbook(self, source_path, dest_path, target_extension):
        target_extension = self._normalize_panel_schedule_extension(target_extension)
        if target_extension not in (".xlsx", ".xls"):
            raise ValueError("Excel conversion target must be .xlsx or .xls.")

        source_abs = os.path.abspath(str(source_path))
        dest_abs = os.path.abspath(str(dest_path))
        if not os.path.exists(source_abs):
            raise FileNotFoundError(f"Workbook not found: {source_abs}")

        try:
            import win32com.client
        except Exception as exc:
            raise RuntimeError(self._panel_schedule_excel_requirement_message()) from exc

        excel = None
        workbook = None
        file_format = 56 if target_extension == ".xls" else 51

        try:
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False

            workbook = excel.Workbooks.Open(source_abs)
            workbook.SaveAs(dest_abs, FileFormat=file_format)
        except Exception as exc:
            raise RuntimeError(
                f"{self._panel_schedule_excel_requirement_message()} "
                f"Excel conversion failed: {exc}"
            ) from exc
        finally:
            if workbook is not None:
                try:
                    workbook.Close(False)
                except Exception:
                    pass
            if excel is not None:
                try:
                    excel.Quit()
                except Exception:
                    pass

    def _update_panel_schedule_workbook(self, panel_data, output_path, use_extracted_kva=False):
        output_extension = self._normalize_panel_schedule_extension(
            os.path.splitext(str(output_path))[1]
        )
        if output_extension == ".xlsx":
            return cb_update_excel_workbook(
                panel_data,
                output_path,
                use_extracted_kva=use_extracted_kva,
            )
        if output_extension != ".xls":
            raise ValueError("Panel schedule must be an .xlsx or .xls file.")

        temp_xlsx = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        temp_xlsx_path = temp_xlsx.name
        temp_xlsx.close()

        try:
            self._convert_excel_workbook(output_path, temp_xlsx_path, ".xlsx")
            sheet_name = cb_update_excel_workbook(
                panel_data,
                temp_xlsx_path,
                use_extracted_kva=use_extracted_kva,
            )
            self._convert_excel_workbook(temp_xlsx_path, output_path, ".xls")
            return sheet_name
        finally:
            try:
                os.remove(temp_xlsx_path)
            except Exception:
                pass

    def _build_panel_schedule_prompt(self, panel_name, num_breaker_imgs, num_dir_imgs):
        return f"""
Analyze these electrical panel photos for Panel: {panel_name}.

You are provided with {num_breaker_imgs} images of the CIRCUIT BREAKERS (first {num_breaker_imgs} images)
and {num_dir_imgs} images of the CIRCUIT DIRECTORY (last {num_dir_imgs} images).

IMAGE INTERPRETATION RULES
- All breaker images belong to the SAME panel.
- All directory images belong to the SAME panel.
- Breaker coverage may be split across multiple photos of the same panel, such as upper half, middle, and bottom half images.
- Directory coverage may be split across multiple photos, such as circuits 1-42 on one image and 43-84 on another.
- Merge partial and overlapping views into one complete understanding of the panel.
- Do not treat split photos as separate panels.
- Reconcile overlapping information by using visible circuit numbers, breaker positions, and the clearest label text.
- Breaker images come first. Directory images come last.
- Do not assume each image shows the entire panel or a complete directory by itself.

TASK 1: HEADER
- Extract Voltage, Bus Rating, Wire, Phase, Mounting, Enclosure.
- Look at the directory images or labels on the panel across all images.
- Extract the panel AIC rating into JSON field "aic_rating" when visibly shown (examples: "10 KAIC", "22,000 AIC").
- If AIC is blank, hidden, unreadable, or not present, set "aic_rating" to "".
- Do not confuse AIC rating with Bus Rating, breaker amperage, voltage, or main requirement.

TASK 2: CIRCUITS & POLES
- Identify every breaker visible across the Breaker Images.
- Use visible breaker positions and circuit numbering to merge split sections from multiple photos.
- CRITICAL: Determine the 'poles' (1, 2, or 3).
    - A 1-pole breaker takes up 1 circuit space.
    - A 2-pole breaker has a tied handle and takes up 2 vertical circuit spaces (e.g. 1 & 3).
    - A 3-pole breaker has a tied handle and takes up 3 vertical circuit spaces (e.g. 1, 3, & 5).
- Provide the circuit_number as the TOP-most circuit number the breaker occupies.
- Extract Amperage and Load Description from the labels (Breaker Images) or circuit directory (Directory Images).
- Resolve ditto marks (") if seen in the directory.
- If the same circuit appears in overlapping photos, combine the evidence and keep one final record.

TASK 3: LOAD TYPES
- LIGHTING -> 'C', RECEPTACLES -> 'G', MOTORS/HVAC -> 'M', KITCHEN -> 'K', DEDICATED -> 'D'
""".strip()

    def _build_existing_directory_prompt(self, panel_name, num_dir_imgs):
        return f"""
Analyze these electrical panel schedule directory document photos for Panel: {panel_name}.

You are provided with {num_dir_imgs} images of the existing panel directory document. This document contains the circuit information (circuit numbers, load descriptions, breaker ratings/amperages, poles, and any visible kVA/load values) needed to recreate the panel schedule.

IMAGE INTERPRETATION RULES
- All directory images belong to the SAME panel directory document.
- Directory coverage may be split across multiple photos, such as circuits 1-42 on one image and 43-84 on another.
- Merge partial and overlapping views into one complete understanding of the panel.
- Reconcile overlapping information by using visible circuit numbers, breaker positions, and the clearest label text.

TASK 1: HEADER
- Extract Voltage, Bus Rating, Wire, Phase, Mounting, Enclosure, and panel AIC rating from any headers, tables, or title block details.
- Set the JSON field "aic_rating" to the visible AIC value (examples: "10 KAIC", "22,000 AIC").
- If AIC is blank, hidden, unreadable, or not present, set "aic_rating" to "".
- Do not confuse AIC rating with Bus Rating, breaker amperage, voltage, or main requirement.

TASK 2: CIRCUITS & POLES
- Identify every circuit entry in the directory document.
- CRITICAL: Determine the 'poles' (1, 2, or 3) for each active circuit breaker.
    - Look at how descriptions or circuit lines are grouped or span multiple spaces to represent multi-pole breakers.
    - A 2-pole breaker has tied circuit spaces (e.g., descriptions spanning circuits 1 and 3, or a single breaker amperage serving a 2-circuit load).
    - Provide the circuit_number as the TOP-most circuit number the breaker occupies.
- Extract Amperage (e.g. 20A, 30A, 50A), Load Description, and the visible kVA/load values from the directory table.
- For each circuit, set the JSON field "phase_kva" to an ordered array of visible kVA decimals for the occupied circuit rows, top-to-bottom (examples: ["0.72"] for 1-pole, ["1.20", "1.10"] for 2-pole, ["2.00", "2.10", "2.05"] for 3-pole).
- If a value is shown in VA, convert it to kVA. If an occupied row's kVA/load value is blank, hidden, unreadable, or not present, use "" for that row in "phase_kva".
- Do not add, total, or sum phase loads together. Do not repeat a 2-pole or 3-pole total on every occupied row.
- Keep the legacy JSON field "kva" blank unless only one single visible total is available and there are no row-by-row phase values.
- Do not estimate kVA and do not infer it from breaker amperage, voltage, load type, or description.
- Resolve ditto marks (") if seen in the directory.
- If the same circuit appears in overlapping photos, combine the evidence and keep one final record.

TASK 3: LOAD TYPES
- LIGHTING -> 'C', RECEPTACLES -> 'G', MOTORS/HVAC -> 'M', KITCHEN -> 'K', DEDICATED -> 'D'
""".strip()

    def _analyze_panel_schedule_images(self, panel_name, breaker_paths, directory_paths, input_mode="field_photos"):
        api_key = self._resolve_panel_schedule_api_key()
        if not api_key:
            raise RuntimeError(
                'AI API key is not configured. Please add it in Settings or set GOOGLE_API_KEY in your .env.'
            )

        self._ensure_aiohttp()
        client = genai.Client(
            api_key=api_key,
            http_options=types.HttpOptions(timeout=600000),
        )

        num_breaker_imgs = len(breaker_paths)
        num_dir_imgs = len(directory_paths)
        if input_mode == "existing_directory":
            prompt = self._build_existing_directory_prompt(panel_name, num_dir_imgs)
        else:
            prompt = self._build_panel_schedule_prompt(
                panel_name, num_breaker_imgs, num_dir_imgs
            )

        gemini_images = []
        try:
            for path in breaker_paths + directory_paths:
                if not os.path.exists(path):
                    raise ValueError(f"Image not found: {path}")
                gemini_images.append(_open_panel_schedule_image(path))

            cb_enforce_rate_limit()

            response = None
            for attempt in range(3):
                try:
                    response = client.models.generate_content(
                        model="gemini-3.5-flash",
                        contents=[prompt, *gemini_images],
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            response_schema=PanelData
                        ),
                    )
                    break
                except Exception as exc:
                    msg = str(exc).lower()
                    transient = (
                        "unavailable" in msg
                        or "deadline" in msg
                        or "503" in msg
                        or "504" in msg
                    )
                    if not transient or attempt == 2:
                        raise
                    time.sleep(5 * (3 ** attempt))
        finally:
            for img in gemini_images:
                try:
                    img.close()
                except Exception:
                    pass

        panel_data = response.parsed
        if panel_data is None:
            raw_text = response.text or ""
            if not raw_text.strip():
                raise RuntimeError("AI returned an empty response. Please try again.")
            data = json.loads(raw_text)
            if hasattr(PanelData, "model_validate"):
                panel_data = PanelData.model_validate(data)
            else:
                panel_data = PanelData.parse_obj(data)

        panel_data.panel_name = panel_name
        return panel_data

    def _process_panel_schedule_payload(self, payload, job_id=None):
        data = self._normalize_panel_schedule_payload(payload)
        activity_id = str(
            data.get('activityId') or data.get('activity_id') or ''
        ).strip()

        output_mode = str(
            data.get('outputMode') or data.get('output_mode') or 'new'
        ).strip().lower()
        if output_mode not in ("new", "existing"):
            output_mode = "new"

        output_path = (
            data.get('outputPath')
            or data.get('output_path')
            or (data.get('newOutputPath') if output_mode == "new" else data.get('existingOutputPath'))
        )
        output_path = str(output_path or '').strip()
        if not output_path:
            raise ValueError("Panel schedule file is required.")

        output_path = os.path.normpath(output_path)
        output_extension_hint = self._normalize_panel_schedule_extension(
            data.get('outputExtension') or data.get('output_extension')
        )
        output_extension = self._normalize_panel_schedule_extension(
            os.path.splitext(output_path)[1]
        )

        created_workbook = False

        if output_mode == "new":
            if not output_extension:
                output_extension = output_extension_hint if output_extension_hint in (".xlsx", ".xls") else ".xlsx"
                output_path = f"{output_path}{output_extension}"
            if output_extension not in (".xlsx", ".xls"):
                raise ValueError("Panel schedule must be an .xlsx or .xls file.")
            if os.path.exists(output_path):
                raise ValueError("The selected file already exists. Choose a new name or add to the existing schedule.")
            if not CB_TEMPLATE_PATH.exists():
                raise ValueError("Panel schedule template not found.")
            dest_dir = os.path.dirname(output_path)
            if dest_dir:
                os.makedirs(dest_dir, exist_ok=True)
            if output_extension == ".xlsx":
                shutil.copy2(str(CB_TEMPLATE_PATH), output_path)
            else:
                self._convert_excel_workbook(str(CB_TEMPLATE_PATH), output_path, ".xls")
            created_workbook = True
        else:
            if output_extension not in (".xlsx", ".xls"):
                raise ValueError("Panel schedule must be an .xlsx or .xls file.")
            if not os.path.exists(output_path):
                raise ValueError("Selected panel schedule was not found.")

        panels_payload = data.get('panels')
        panel_requests = []
        if isinstance(panels_payload, list):
            for index, raw_panel in enumerate(panels_payload):
                if not isinstance(raw_panel, dict):
                    continue
                panel_name = str(
                    raw_panel.get('panelName') or raw_panel.get('panel_name') or ''
                ).strip()
                if not panel_name:
                    panel_name = f"PANEL {index + 1}"
                panel_id = str(
                    raw_panel.get('panelId') or raw_panel.get('panel_id') or f"panel_{index + 1}"
                ).strip() or f"panel_{index + 1}"
                panel_requests.append({
                    'panel_id': panel_id,
                    'panel_name': panel_name,
                    'input_mode': str(raw_panel.get('inputMode') or raw_panel.get('input_mode') or 'field_photos').strip().lower(),
                    'breaker_paths': self._normalize_panel_schedule_paths(
                        self._get_panel_schedule_path_value(
                            raw_panel,
                            ('breakerPaths', 'breaker_paths'),
                            ('breakerPath', 'breaker_path'),
                        )
                    ),
                    'directory_paths': self._normalize_panel_schedule_paths(
                        self._get_panel_schedule_path_value(
                            raw_panel,
                            ('directoryPaths', 'directory_paths'),
                            ('directoryPath', 'directory_path'),
                        )
                    ),
                    'breaker_uploads': raw_panel.get('breakerUploads') or raw_panel.get('breaker_uploads') or [],
                    'directory_uploads': raw_panel.get('directoryUploads') or raw_panel.get('directory_uploads') or [],
                })

        if not panel_requests:
            panel_name = str(data.get('panelName') or data.get('panel_name') or '').strip()
            if not panel_name:
                panel_name = "PANEL"
            panel_requests.append({
                'panel_id': 'panel_1',
                'panel_name': panel_name,
                'input_mode': str(data.get('inputMode') or data.get('input_mode') or 'field_photos').strip().lower(),
                'breaker_paths': self._normalize_panel_schedule_paths(
                    self._get_panel_schedule_path_value(
                        data,
                        ('breakerPaths', 'breaker_paths'),
                        ('breakerPath', 'breaker_path'),
                    )
                ),
                'directory_paths': self._normalize_panel_schedule_paths(
                    self._get_panel_schedule_path_value(
                        data,
                        ('directoryPaths', 'directory_paths'),
                        ('directoryPath', 'directory_path'),
                    )
                ),
                'breaker_uploads': data.get('breakerUploads') or data.get('breaker_uploads') or [],
                'directory_uploads': data.get('directoryUploads') or data.get('directory_uploads') or [],
            })

        results = []
        success_count = 0
        failure_count = 0
        panel_count = len(panel_requests)

        def _store_running_record(message, completed_count):
            if not job_id:
                return
            existing = self._get_panel_schedule_job_record(job_id) or {}
            self._store_panel_schedule_job_record({
                'jobId': str(job_id or '').strip(),
                'toolId': 'toolCircuitBreaker',
                'activityId': activity_id or str(existing.get('activityId') or '').strip(),
                'status': 'running',
                'message': str(message or '').strip(),
                'panelCount': max(int(panel_count or 0), 1),
                'completedCount': max(int(completed_count or 0), 0),
                'outputPath': output_path,
                'outputFolder': os.path.dirname(output_path) if output_path else '',
                'successCount': success_count,
                'failureCount': failure_count,
                'results': deepcopy(results),
                'startedAt': str(existing.get('startedAt') or utc_now_iso()).strip(),
                'finishedAt': '',
            })

        for index, panel_request in enumerate(panel_requests):
            panel_id = str(panel_request.get('panel_id') or f"panel_{index + 1}").strip() or f"panel_{index + 1}"
            panel_name = str(panel_request.get('panel_name') or '').strip() or f"PANEL {index + 1}"
            input_mode = str(panel_request.get('input_mode') or 'field_photos').strip().lower()
            breaker_paths = list(panel_request.get('breaker_paths') or [])
            directory_paths = list(panel_request.get('directory_paths') or [])
            breaker_uploads = panel_request.get('breaker_uploads') or []
            directory_uploads = panel_request.get('directory_uploads') or []
            temp_paths = []
            _store_running_record(
                f"Processing panel {index + 1} of {panel_count}: {panel_name}",
                success_count + failure_count,
            )

            try:
                if breaker_uploads:
                    uploaded_paths, created = self._panel_schedule_save_uploads(
                        breaker_uploads, "breaker"
                    )
                    breaker_paths.extend(uploaded_paths)
                    temp_paths.extend(created)
                if directory_uploads:
                    uploaded_paths, created = self._panel_schedule_save_uploads(
                        directory_uploads, "directory"
                    )
                    directory_paths.extend(uploaded_paths)
                    temp_paths.extend(created)

                if input_mode == "existing_directory":
                    if not directory_paths:
                        raise ValueError(
                            "At least one directory photo is required for existing directory mode."
                        )
                else:
                    if not breaker_paths or not directory_paths:
                        raise ValueError(
                            "At least one breaker photo and at least one directory photo are required."
                        )

                try:
                    panel_data = self._analyze_panel_schedule_images(
                        panel_name, breaker_paths, directory_paths, input_mode=input_mode
                    )
                except Exception as e:
                    raise RuntimeError(self._format_panel_schedule_ai_error(e)) from e

                sheet_name = self._update_panel_schedule_workbook(
                    panel_data,
                    output_path,
                    use_extracted_kva=input_mode == "existing_directory",
                )
                success_count += 1
                results.append({
                    'panelId': panel_id,
                    'panelName': panel_name,
                    'status': 'success',
                    'sheetName': sheet_name,
                    'message': f"Panel '{sheet_name}' added to schedule."
                })
            except Exception as e:
                failure_count += 1
                results.append({
                    'panelId': panel_id,
                    'panelName': panel_name,
                    'status': 'error',
                    'message': str(e)
                })
            finally:
                for temp_path in temp_paths:
                    try:
                        os.remove(temp_path)
                    except Exception:
                        pass
            _store_running_record(
                f"Processed {success_count + failure_count} of {panel_count} panels...",
                success_count + failure_count,
            )

        output_folder = os.path.dirname(output_path)
        first_success = next(
            (item for item in results if item.get('status') == 'success'),
            None
        )

        if success_count == 0:
            if created_workbook:
                try:
                    os.remove(output_path)
                except Exception:
                    pass
            first_error_message = next(
                (item.get('message') for item in results if item.get('message')),
                "Panel Schedule AI failed."
            )
            return {
                'status': 'error',
                'message': first_error_message,
                'toolId': 'toolCircuitBreaker',
                'activityId': activity_id,
                'outputPath': output_path,
                'outputFolder': output_folder,
                'panelCount': panel_count,
                'completedCount': success_count + failure_count,
                'successCount': success_count,
                'failureCount': failure_count,
                'results': results
            }

        if failure_count > 0:
            message = (
                f"Added {success_count} panel sheet{'s' if success_count != 1 else ''}; "
                f"{failure_count} panel{'s' if failure_count != 1 else ''} failed."
            )
        else:
            message = (
                f"Added {success_count} panel sheet{'s' if success_count != 1 else ''} to schedule."
            )

        return {
            'status': 'success',
            'message': message,
            'toolId': 'toolCircuitBreaker',
            'activityId': activity_id,
            'outputPath': output_path,
            'outputFolder': output_folder,
            'sheetName': first_success.get('sheetName') if first_success else '',
            'panelCount': panel_count,
            'completedCount': success_count + failure_count,
            'successCount': success_count,
            'failureCount': failure_count,
            'results': results
        }

    def _normalize_launch_context(self, launch_context):
        if isinstance(launch_context, dict):
            return launch_context
        if isinstance(launch_context, str):
            raw_text = launch_context.strip()
            if not raw_text:
                return {}
            try:
                parsed = json.loads(raw_text)
                if isinstance(parsed, dict):
                    return parsed
                logging.info(
                    "_normalize_launch_context: Ignoring JSON launch context because it is not an object "
                    f"(type={type(parsed).__name__}).")
                return {}
            except Exception as e:
                preview = raw_text[:120]
                logging.warning(
                    "_normalize_launch_context: Failed to parse launch context JSON string "
                    f"(preview={preview!r}): {e}")
                return {}
        if launch_context is not None:
            logging.info(
                "_normalize_launch_context: Ignoring unsupported launch context type "
                f"(type={type(launch_context).__name__}).")
        return {}

    def _primary_discipline_from_settings(self, settings):
        discipline = settings.get('discipline', 'Electrical')
        if isinstance(discipline, list):
            discipline = next((d for d in discipline if d), 'Electrical')
        discipline = str(discipline or 'Electrical').strip()
        discipline_map = {
            'electrical': 'Electrical',
            'mechanical': 'Mechanical',
            'plumbing': 'Plumbing',
            'general': 'Electrical'
        }
        return discipline_map.get(discipline.lower(), 'Electrical')

    def _get_settings_workroom_disciplines(self, settings):
        discipline_map = {
            'electrical': 'Electrical',
            'mechanical': 'Mechanical',
            'plumbing': 'Plumbing',
        }
        raw_value = settings.get('discipline', ['Electrical'])
        if isinstance(raw_value, str):
            candidates = [raw_value]
        elif isinstance(raw_value, list):
            candidates = raw_value
        else:
            candidates = ['Electrical']

        result = []
        seen = set()
        for item in candidates:
            normalized = discipline_map.get(str(item or '').strip().lower())
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            result.append(normalized)
        return result or ['Electrical']

    def _resolve_workroom_discipline(self, settings, launch_context):
        settings_disciplines = self._get_settings_workroom_disciplines(settings)
        context = self._normalize_launch_context(launch_context)
        requested = str(context.get('discipline') or '').strip().lower()
        mapped_requested = {
            'electrical': 'Electrical',
            'mechanical': 'Mechanical',
            'plumbing': 'Plumbing',
        }.get(requested)
        # Workroom selection should drive CAD routing when explicitly provided.
        if mapped_requested:
            return mapped_requested, 'launch_context'
        if len(settings_disciplines) == 1:
            return settings_disciplines[0], 'settings_single'
        return settings_disciplines[0], 'settings_fallback'

    def _resolve_workroom_context(self, settings, launch_context):
        context = self._normalize_launch_context(launch_context)
        source = str(context.get('source') or '').strip().lower()
        project_path = str(
            context.get('projectPath')
            or context.get('project_path')
            or ''
        ).strip()
        root_project_path = str(context.get('rootProjectPath') or '').strip()
        effective_project_path = root_project_path or project_path
        if source == 'workroom':
            discipline, discipline_source = self._resolve_workroom_discipline(
                settings, launch_context)
        else:
            discipline = self._primary_discipline_from_settings(settings)
            discipline_source = 'settings_primary'
        normalized_project_path = os.path.normpath(
            effective_project_path) if effective_project_path else ''
        logging.info(
            f"_resolve_workroom_context: source={source or 'none'} "
            f"discipline={discipline} discipline_source={discipline_source} "
            f"project_path={normalized_project_path or '<empty>'} "
            f"(projectPath={project_path or '<empty>'}, rootProjectPath={root_project_path or '<empty>'})")
        return {
            'source': source,
            'project_path': normalized_project_path,
            'discipline': discipline,
            'discipline_source': discipline_source,
        }

    def _resolve_launch_context_default_directory(self, launch_context, allow_workroom=False):
        context = self._normalize_launch_context(launch_context)
        source = str(context.get('source') or '').strip().lower()
        if source == 'workroom' and not allow_workroom:
            return ''

        raw_path = str(
            context.get('rootProjectPath')
            or context.get('projectPath')
            or context.get('project_path')
            or ''
        ).strip()
        if not raw_path:
            return ''

        normalized_path = os.path.normpath(raw_path)
        normalized_copy_path = self._to_windows_extended_path(normalized_path)
        if os.path.isdir(normalized_copy_path):
            return normalized_path

        parent_dir = os.path.dirname(normalized_path)
        if parent_dir and os.path.isdir(self._to_windows_extended_path(parent_dir)):
            return os.path.normpath(parent_dir)
        return ''

    def _is_workroom_auto_select_enabled(self, settings, launch_context):
        context = self._resolve_workroom_context(settings, launch_context)
        source = str(context.get('source') or '').strip().lower()
        auto_select_setting = bool(settings.get('workroomAutoSelectCadFiles', True))
        if source != 'workroom':
            logging.info(
                "_is_workroom_auto_select_enabled: disabled because launch source is not workroom "
                f"(source={source or 'none'}, launch_context_type={type(launch_context).__name__}).")
            return False
        if not auto_select_setting:
            logging.info(
                "_is_workroom_auto_select_enabled: disabled because workroomAutoSelectCadFiles is off.")
            return False
        return True

    def _should_force_workroom_clean_xrefs_manual_selection(self, context):
        source = str((context or {}).get('source') or '').strip().lower()
        return source == 'workroom'

    def _list_base_level_dwgs(self, folder_path):
        if not folder_path or not os.path.isdir(folder_path):
            return []
        file_paths = []
        try:
            with os.scandir(folder_path) as entries:
                for entry in entries:
                    if entry.is_file() and entry.name.lower().endswith('.dwg'):
                        file_paths.append(entry.path)
        except Exception as e:
            logging.warning(f"Failed to list DWGs in {folder_path}: {e}")
            return []
        file_paths.sort(key=lambda path: os.path.basename(path).lower())
        return file_paths

    def _write_files_list_temp(self, file_paths):
        temp_file = tempfile.NamedTemporaryFile(
            mode='w', suffix='.txt', delete=False, encoding='utf-8')
        for path in file_paths:
            temp_file.write(path + '\n')
        temp_file.close()
        return temp_file.name

    def _source_display_path(self, source):
        if not isinstance(source, dict):
            return ''
        display_path = str(source.get('displayPath') or '').strip()
        if display_path:
            return display_path
        if str(source.get('kind') or '').strip() == 'zipEntry':
            zip_path = str(source.get('zipPath') or '').strip()
            entry_name = str(source.get('entryName') or '').strip()
            if zip_path and entry_name:
                return f'{zip_path}::{entry_name}'
        return str(source.get('path') or '').strip()

    def _normalize_dwg_file_sources(self, sources, require_exists=True):
        if not isinstance(sources, (list, tuple)):
            return []
        normalized_sources = []
        seen = set()
        for raw_source in sources:
            if isinstance(raw_source, str):
                raw_source = {'kind': 'file', 'path': raw_source}
            if not isinstance(raw_source, dict):
                continue
            kind = str(raw_source.get('kind') or 'file').strip()
            if kind == 'zipEntry':
                zip_path = os.path.normpath(str(raw_source.get('zipPath') or '').strip())
                entry_name = str(raw_source.get('entryName') or '').strip().replace('\\', '/')
                if not zip_path or not entry_name:
                    continue
                if os.path.splitext(entry_name)[1].lower() != '.dwg':
                    continue
                if require_exists and not os.path.isfile(zip_path):
                    continue
                project_root = str(raw_source.get('projectRoot') or '').strip()
                source = {
                    'kind': 'zipEntry',
                    'zipPath': zip_path,
                    'entryName': entry_name,
                    'projectRoot': os.path.normpath(project_root) if project_root else '',
                    'displayPath': str(raw_source.get('displayPath') or f'{zip_path}::{entry_name}').strip(),
                }
                key = ('zipEntry', os.path.normcase(zip_path), entry_name.lower())
            else:
                path = os.path.normpath(str(raw_source.get('path') or '').strip())
                if not path or os.path.splitext(path)[1].lower() != '.dwg':
                    continue
                if require_exists and not os.path.isfile(path):
                    continue
                source = {
                    'kind': 'file',
                    'path': path,
                    'displayPath': str(raw_source.get('displayPath') or path).strip(),
                }
                project_root = str(raw_source.get('projectRoot') or '').strip()
                if project_root:
                    source['projectRoot'] = os.path.normpath(project_root)
                key = ('file', os.path.normcase(path))
            if key in seen:
                continue
            seen.add(key)
            normalized_sources.append(source)
        return normalized_sources

    def _write_dwg_source_manifest_temp(self, sources):
        normalized_sources = self._normalize_dwg_file_sources(
            sources, require_exists=True)
        if not normalized_sources:
            return ''
        temp_file = tempfile.NamedTemporaryFile(
            mode='w', suffix='.json', delete=False, encoding='utf-8')
        json.dump(normalized_sources, temp_file, ensure_ascii=True)
        temp_file.close()
        return temp_file.name

    def _get_launch_context_dwg_file_sources(self, launch_context):
        context = self._normalize_launch_context(launch_context)
        return self._normalize_dwg_file_sources(
            context.get('dwgFileSources'), require_exists=True)

    def _get_workflow_project_path(self, settings, launch_context):
        context = self._resolve_workroom_context(settings, launch_context)
        project_path = str(context.get('project_path') or '').strip()
        if project_path:
            return project_path
        launch_payload = self._normalize_launch_context(launch_context)
        return os.path.normpath(str(
            launch_payload.get('rootProjectPath')
            or launch_payload.get('projectPath')
            or ''
        ).strip()) if launch_payload else ''

    def _path_has_archive_part(self, path):
        parts = re.split(r'[\\/]+', str(path or ''))
        return any(
            str(part or '').strip().lower() in {'archive', '0 archive'}
            for part in parts
        )

    def _resolve_workflow_electrical_dwg_files(self, settings, launch_context):
        project_path = self._get_workflow_project_path(settings, launch_context)
        if not project_path:
            return []
        folder_resolution = self._resolve_workroom_discipline_folder(
            project_path, 'Electrical')
        electrical_folder = folder_resolution.get('resolved_folder') or ''
        if not electrical_folder:
            return []
        return self._list_base_level_dwgs(electrical_folder)

    def _read_workflow_xref_scan_results(self, settings, dwg_files):
        dwg_paths = self._normalize_dwg_file_paths(dwg_files, require_exists=True)
        if not dwg_paths:
            return []
        acad_path = str((settings or {}).get('autocadPath') or '').strip()
        script_path = os.path.join(BASE_DIR, "scripts", "ListDwgXrefs.ps1")
        if not acad_path or not os.path.isfile(acad_path) or not os.path.exists(script_path):
            return []
        files_list_path = self._write_files_list_temp(dwg_paths)
        try:
            command = self._build_powershell_script_command(
                script_path,
                '-AcadCore',
                acad_path,
                '-FilesListPath',
                files_list_path,
            )
            timeout = max(60, min(600, len(dwg_paths) * 45))
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=timeout,
            )
            output = (result.stdout or '') + '\n' + (result.stderr or '')
            for line in reversed(output.splitlines()):
                line = line.strip()
                if line.startswith('XREF_JSON:'):
                    payload = line[len('XREF_JSON:'):].strip()
                    parsed = json.loads(payload)
                    return parsed if isinstance(parsed, list) else []
            if result.returncode != 0:
                logging.warning(
                    "ListDwgXrefs.ps1 failed with code %s: %s",
                    result.returncode,
                    output[-1000:],
                )
        except Exception as exc:
            logging.warning("Workflow XREF scan failed: %s", exc)
        finally:
            try:
                os.unlink(files_list_path)
            except Exception:
                pass
        return []

    def _extract_xref_target_basenames(self, scan_results):
        seen = set()
        basenames = []

        def _add_candidate(value):
            raw = str(value or '').strip()
            if not raw:
                return
            candidate = raw.replace('/', os.sep).replace('\\', os.sep)
            base_name = os.path.basename(candidate)
            if not base_name:
                base_name = raw
            if base_name.lower().endswith('.dwg'):
                base_name = os.path.splitext(base_name)[0]
            base_name = base_name.strip()
            if not base_name:
                return
            key = base_name.lower()
            if key in seen:
                return
            seen.add(key)
            basenames.append(base_name)

        for item in scan_results or []:
            if isinstance(item, str):
                _add_candidate(item)
                continue
            if not isinstance(item, dict):
                continue
            refs = item.get('references')
            if not isinstance(refs, list):
                refs = [item]
            for ref in refs:
                if isinstance(ref, str):
                    _add_candidate(ref)
                elif isinstance(ref, dict):
                    _add_candidate(
                        ref.get('path')
                        or ref.get('fileName')
                        or ref.get('baseName')
                        or ref.get('name')
                    )
        return basenames

    def _zipinfo_timestamp(self, info):
        try:
            return datetime.datetime(*info.date_time).timestamp()
        except Exception:
            return 0

    def _iter_arch_dwg_source_candidates(self, arch_folder, project_root, include_zips=True):
        if not arch_folder or not os.path.isdir(arch_folder):
            return []
        candidates = []
        for root, dirs, files in os.walk(arch_folder):
            dirs[:] = [
                name for name in dirs
                if not self._path_has_archive_part(os.path.join(root, name))
            ]
            if self._path_has_archive_part(root):
                continue
            try:
                folder_mtime = os.path.getmtime(root)
            except Exception:
                folder_mtime = 0
            for filename in files:
                full_path = os.path.join(root, filename)
                lower_name = filename.lower()
                if lower_name.endswith('.dwg'):
                    try:
                        file_mtime = os.path.getmtime(full_path)
                    except Exception:
                        file_mtime = 0
                    candidates.append({
                        'matchKey': os.path.splitext(filename)[0].lower(),
                        'folderMtime': folder_mtime,
                        'fileMtime': file_mtime,
                        'source': {
                            'kind': 'file',
                            'path': os.path.normpath(full_path),
                            'displayPath': os.path.normpath(full_path),
                            'projectRoot': os.path.normpath(project_root) if project_root else '',
                        },
                    })
                    continue
                if not include_zips or not lower_name.endswith('.zip'):
                    continue
                try:
                    zip_mtime = os.path.getmtime(full_path)
                    with zipfile.ZipFile(full_path) as archive:
                        for info in archive.infolist():
                            if info.is_dir():
                                continue
                            entry_name = str(info.filename or '').replace('\\', '/')
                            if not entry_name.lower().endswith('.dwg'):
                                continue
                            if self._path_has_archive_part(entry_name):
                                continue
                            entry_base = os.path.basename(entry_name)
                            candidates.append({
                                'matchKey': os.path.splitext(entry_base)[0].lower(),
                                'folderMtime': folder_mtime,
                                'fileMtime': max(zip_mtime, self._zipinfo_timestamp(info)),
                                'source': {
                                    'kind': 'zipEntry',
                                    'zipPath': os.path.normpath(full_path),
                                    'entryName': entry_name,
                                    'projectRoot': os.path.normpath(project_root) if project_root else '',
                                    'displayPath': f'{os.path.normpath(full_path)}::{entry_name}',
                                },
                            })
                except Exception as exc:
                    logging.info("Skipping unreadable Arch ZIP %s: %s", full_path, exc)
        candidates.sort(
            key=lambda item: (
                -float(item.get('folderMtime') or 0),
                -float(item.get('fileMtime') or 0),
                self._source_display_path(item.get('source') or {}).lower(),
            )
        )
        return candidates

    def _resolve_workflow_clean_xref_sources(self, settings, launch_context):
        electrical_dwgs = self._resolve_workflow_electrical_dwg_files(
            settings, launch_context)
        if not electrical_dwgs:
            return []
        scan_results = self._read_workflow_xref_scan_results(settings, electrical_dwgs)
        target_basenames = self._extract_xref_target_basenames(scan_results)
        if not target_basenames:
            return []
        project_path = self._get_workflow_project_path(settings, launch_context)
        if not project_path:
            return []
        project_root = (
            self._find_workroom_project_root_by_id(project_path)
            or os.path.normpath(project_path)
        )
        arch_resolution = self._resolve_workroom_discipline_folder(project_path, 'Arch')
        arch_folder = arch_resolution.get('resolved_folder') or ''
        include_zips = _normalize_workflow_cad_defaults(
            (settings or {}).get('workflowCadDefaults')
        ).get('cleanXrefsSearchZipArchives', True)
        candidates = self._iter_arch_dwg_source_candidates(
            arch_folder, project_root, include_zips=include_zips)
        if not candidates:
            return []
        by_key = {}
        for candidate in candidates:
            by_key.setdefault(candidate.get('matchKey'), []).append(candidate)
        selected = []
        seen_sources = set()
        for basename in target_basenames:
            matches = by_key.get(str(basename or '').strip().lower()) or []
            if not matches:
                continue
            source = matches[0].get('source') or {}
            if source.get('kind') == 'zipEntry':
                source_key = (
                    'zipEntry',
                    os.path.normcase(source.get('zipPath') or ''),
                    str(source.get('entryName') or '').lower(),
                )
            else:
                source_key = ('file', os.path.normcase(source.get('path') or ''))
            if source_key in seen_sources:
                continue
            seen_sources.add(source_key)
            selected.append(source)
        return self._normalize_dwg_file_sources(selected, require_exists=True)

    def _trace_cad_auto_select(self, event, **fields):
        try:
            payload = {
                'timestamp': datetime.datetime.now(
                    datetime.timezone.utc).isoformat().replace('+00:00', 'Z'),
                'event': str(event or '').strip() or 'unknown',
            }
            for key, value in (fields or {}).items():
                payload[str(key)] = value
            with open(CAD_AUTO_SELECT_TRACE_FILE, 'a', encoding='utf-8') as trace_file:
                trace_file.write(
                    json.dumps(payload, ensure_ascii=True, default=str) + '\n')
        except Exception:
            pass

    def trace_cad_auto_select_event(self, event, payload=None):
        payload_data = dict(payload) if isinstance(payload, dict) else {
            'value': payload,
        }
        payload_data.setdefault('trace_source', 'frontend')
        self._trace_cad_auto_select(
            str(event or '').strip() or 'trace_event',
            **payload_data,
        )
        return {'status': 'success'}

    def clear_cad_auto_select_trace(self):
        try:
            with open(CAD_AUTO_SELECT_TRACE_FILE, 'w', encoding='utf-8'):
                pass
            return {
                'status': 'success',
                'path': CAD_AUTO_SELECT_TRACE_FILE,
            }
        except Exception as e:
            logging.error(f"clear_cad_auto_select_trace failed: {e}")
            return {
                'status': 'error',
                'message': str(e),
                'path': CAD_AUTO_SELECT_TRACE_FILE,
            }

    def get_cad_auto_select_trace(self, line_limit=200):
        try:
            try:
                limit = int(line_limit)
            except (TypeError, ValueError):
                limit = 200
            limit = max(1, min(limit, 1000))

            if not os.path.exists(CAD_AUTO_SELECT_TRACE_FILE):
                return {
                    'status': 'success',
                    'path': CAD_AUTO_SELECT_TRACE_FILE,
                    'entries': [],
                    'lines': [],
                    'lineCount': 0,
                }

            with open(CAD_AUTO_SELECT_TRACE_FILE, 'r', encoding='utf-8') as trace_file:
                lines = trace_file.read().splitlines()
            selected_lines = lines[-limit:]
            entries = []
            for line in selected_lines:
                try:
                    entries.append(json.loads(line))
                except json.JSONDecodeError:
                    entries.append({'raw': line})
            return {
                'status': 'success',
                'path': CAD_AUTO_SELECT_TRACE_FILE,
                'entries': entries,
                'lines': selected_lines,
                'lineCount': len(selected_lines),
            }
        except Exception as e:
            logging.error(f"get_cad_auto_select_trace failed: {e}")
            return {
                'status': 'error',
                'message': str(e),
                'path': CAD_AUTO_SELECT_TRACE_FILE,
                'entries': [],
                'lines': [],
                'lineCount': 0,
            }

    def _normalize_dwg_file_paths(self, file_paths, require_exists=True):
        if not isinstance(file_paths, (list, tuple, set)):
            return []

        resolved_paths = []
        seen = set()
        for raw_path in file_paths:
            raw_text = str(raw_path or '').strip()
            if not raw_text:
                continue
            normalized = os.path.normpath(raw_text)
            if os.path.splitext(normalized)[1].lower() != '.dwg':
                continue
            if require_exists and not os.path.isfile(normalized):
                continue
            key = os.path.normcase(normalized)
            if key in seen:
                continue
            seen.add(key)
            resolved_paths.append(normalized)
        return resolved_paths

    def _get_workroom_cad_file_cache_store(self):
        cache = getattr(self, '_workroom_cad_file_cache', None)
        if not isinstance(cache, dict):
            cache = {}
            self._workroom_cad_file_cache = cache
        return cache

    def _get_workroom_cad_file_cache_key(self, project_path, discipline):
        normalized_project_path = os.path.normpath(
            str(project_path or '').strip()) if project_path else ''
        normalized_discipline = str(discipline or '').strip()
        if not normalized_project_path or not normalized_discipline:
            return None
        return (os.path.normcase(normalized_project_path), normalized_discipline.lower())

    def _set_workroom_cad_file_cache_entry(self, project_path, discipline, folder_path, file_paths, resolution_mode=''):
        cache_key = self._get_workroom_cad_file_cache_key(project_path, discipline)
        if cache_key is None:
            return None

        normalized_project_path = os.path.normpath(str(project_path or '').strip())
        normalized_discipline = str(discipline or '').strip()
        normalized_folder_path = os.path.normpath(
            str(folder_path or '').strip()) if folder_path else ''
        normalized_files = self._normalize_dwg_file_paths(
            file_paths, require_exists=False)

        entry = {
            'project_path': normalized_project_path,
            'discipline': normalized_discipline,
            'folder_path': normalized_folder_path,
            'files': normalized_files,
            'resolution_mode': str(resolution_mode or '').strip(),
        }
        self._get_workroom_cad_file_cache_store()[cache_key] = entry
        return entry

    def _get_workroom_cad_file_cache_entry(self, project_path, discipline):
        cache_key = self._get_workroom_cad_file_cache_key(project_path, discipline)
        if cache_key is None:
            return None
        entry = self._get_workroom_cad_file_cache_store().get(cache_key)
        return entry if isinstance(entry, dict) else None

    def _get_launch_context_cad_file_paths(self, launch_context):
        context = self._normalize_launch_context(launch_context)
        return self._normalize_dwg_file_paths(
            context.get('cadFilePaths'), require_exists=True)

    def _get_shared_parent_folder(self, file_paths):
        parent_dirs = []
        seen = set()
        for path in file_paths:
            parent = os.path.dirname(os.path.normpath(str(path or '').strip()))
            if not parent:
                continue
            key = os.path.normcase(parent)
            if key in seen:
                continue
            seen.add(key)
            parent_dirs.append(parent)
        if len(parent_dirs) == 1:
            return parent_dirs[0]
        return ''

    def _find_project_root_by_id(self, path):
        normalized_path = os.path.normpath(str(path or '').strip()) if path else ''
        if not normalized_path:
            return ''

        current = normalized_path.rstrip("\\/")
        if not current:
            return ''

        basename = os.path.basename(current).strip()
        if os.path.isfile(current) or (not os.path.isdir(current) and os.path.splitext(basename)[1]):
            parent = os.path.dirname(current)
            current = parent.rstrip("\\/") or parent

        while current:
            folder_name = os.path.basename(current).strip()
            if re.match(r'^\d{6}(?!\d)', folder_name):
                logging.info(
                    "_find_project_root_by_id: matched %s from %s",
                    current,
                    normalized_path,
                )
                return current

            parent = os.path.dirname(current)
            if not parent:
                break
            if os.path.normcase(parent) == os.path.normcase(current):
                break
            current = parent.rstrip("\\/") or parent

        return ''

    def _find_workroom_project_root_by_id(self, project_path):
        return self._find_project_root_by_id(project_path)

    def _resolve_workroom_discipline_folder(self, project_path, discipline):
        requested_discipline = str(discipline or '').strip() or 'Electrical'
        discipline_map = {
            'electrical': 'Electrical',
            'mechanical': 'Mechanical',
            'plumbing': 'Plumbing',
            'arch': 'Arch',
        }
        resolved_discipline = discipline_map.get(
            requested_discipline.lower(), requested_discipline)

        known_folder_names = {'electrical', 'mechanical', 'plumbing', 'arch'}
        candidate_entries = []
        seen = set()

        def _add_candidate(path_value, mode):
            if not path_value:
                return
            normalized = os.path.normpath(path_value)
            key = os.path.normcase(normalized)
            if key in seen:
                return
            seen.add(key)
            candidate_entries.append({
                'mode': mode,
                'path': normalized,
            })

        normalized_project_path = os.path.normpath(
            project_path) if project_path else ''
        project_base = os.path.basename(
            normalized_project_path.rstrip("\\/")
        ).strip() if normalized_project_path else ''
        project_root_by_id = self._find_workroom_project_root_by_id(
            normalized_project_path)

        if project_root_by_id:
            _add_candidate(
                os.path.join(project_root_by_id, resolved_discipline),
                'project_id_ancestor_child_folder',
            )

        if normalized_project_path:
            if project_base.lower() == resolved_discipline.lower():
                _add_candidate(
                    normalized_project_path, 'project_path_is_discipline_folder')

            _add_candidate(
                os.path.join(normalized_project_path, resolved_discipline),
                'project_path_child_folder',
            )

            if project_base.lower() in known_folder_names:
                parent_path = os.path.dirname(
                    normalized_project_path.rstrip("\\/"))
                if parent_path:
                    _add_candidate(
                        os.path.join(parent_path, resolved_discipline),
                        'discipline_sibling_folder',
                    )

        for entry in candidate_entries:
            folder_path = entry['path']
            if os.path.isdir(folder_path):
                return {
                    'resolved_folder': folder_path,
                    'mode': entry['mode'],
                    'discipline': resolved_discipline,
                    'candidates': [item['path'] for item in candidate_entries],
                }

        return {
            'resolved_folder': '',
            'mode': 'not_found',
            'discipline': resolved_discipline,
            'candidates': [item['path'] for item in candidate_entries],
        }

    def _is_workroom_auto_select_tool_allowed(self, tool_name):
        normalized_name = str(tool_name or '').strip()
        if normalized_name.endswith('_test_mode'):
            normalized_name = normalized_name[:-10]
        return normalized_name in {
            'run_publish_script',
            'run_manage_layers_script',
        }

    def _resolve_workroom_auto_file_selection(self, settings, launch_context, tool_name):
        if not self._is_workroom_auto_select_tool_allowed(tool_name):
            logging.info(
                f"{tool_name}: Workroom DWG auto-select skipped (tool not allowlisted).")
            self._trace_cad_auto_select(
                'auto_select_skipped_tool_not_allowed',
                tool_name=tool_name,
            )
            return None
        context = self._resolve_workroom_context(settings, launch_context)
        launch_payload = self._normalize_launch_context(launch_context)
        project_path = context.get('project_path') or ''
        discipline = context.get('discipline') or 'Electrical'
        explicit_dwg_files = self._get_launch_context_cad_file_paths(launch_context)
        workflow_preselected = launch_payload.get('workflowPreselectedDwgFiles') is True

        def _use_explicit_dwg_files(selection_source='launch_context_explicit_files'):
            files_list_path = self._write_files_list_temp(explicit_dwg_files)
            shared_parent = self._get_shared_parent_folder(explicit_dwg_files)
            logging.info(
                f"{tool_name}: Auto-selected {len(explicit_dwg_files)} explicit DWG(s) from launch context "
                f"(folder={shared_parent or '<multiple>'}).")
            self._trace_cad_auto_select(
                'auto_select_selected',
                tool_name=tool_name,
                selection_source=selection_source,
                project_path=project_path,
                discipline=discipline,
                folder_path=shared_parent,
                files_list_path=files_list_path,
                count=len(explicit_dwg_files),
                file_paths=explicit_dwg_files,
            )
            return {
                'files_list_path': files_list_path,
                'project_path': project_path,
                'discipline': discipline,
                'folder_path': shared_parent,
                'count': len(explicit_dwg_files),
                'resolution_mode': selection_source,
            }

        self._trace_cad_auto_select(
            'auto_select_request',
            tool_name=tool_name,
            source=context.get('source') or '',
            project_path=project_path,
            discipline=discipline,
            discipline_source=context.get('discipline_source') or '',
            launch_context=launch_payload,
        )
        if workflow_preselected and explicit_dwg_files:
            return _use_explicit_dwg_files('workflow_preselected_files')
        if workflow_preselected and 'cadFilePaths' in launch_payload:
            self._trace_cad_auto_select(
                'auto_select_explicit_invalid',
                tool_name=tool_name,
                project_path=project_path,
                discipline=discipline,
                cad_file_paths=launch_payload.get('cadFilePaths'),
                workflow_preselected=True,
            )
            return None

        if not self._is_workroom_auto_select_enabled(settings, launch_context):
            if explicit_dwg_files:
                return _use_explicit_dwg_files()
            self._trace_cad_auto_select(
                'auto_select_skipped_disabled',
                tool_name=tool_name,
                launch_context=launch_payload,
            )
            return None

        cached_entry = self._get_workroom_cad_file_cache_entry(
            project_path, discipline)
        if cached_entry:
            cached_dwg_files = self._normalize_dwg_file_paths(
                cached_entry.get('files'), require_exists=True)
            if cached_dwg_files:
                files_list_path = self._write_files_list_temp(cached_dwg_files)
                cached_folder_path = str(cached_entry.get('folder_path') or '').strip()
                folder_path = cached_folder_path or self._get_shared_parent_folder(
                    cached_dwg_files)
                logging.info(
                    f"{tool_name}: Auto-selected {len(cached_dwg_files)} DWG(s) from cached Workroom detection "
                    f"(folder={folder_path or '<multiple>'}; cached_mode={cached_entry.get('resolution_mode') or 'unknown'}).")
                self._trace_cad_auto_select(
                    'auto_select_selected',
                    tool_name=tool_name,
                    selection_source='workroom_cached_detection',
                    project_path=project_path,
                    discipline=str(cached_entry.get('discipline') or discipline),
                    folder_path=folder_path,
                    files_list_path=files_list_path,
                    count=len(cached_dwg_files),
                    file_paths=cached_dwg_files,
                    cached_resolution_mode=cached_entry.get('resolution_mode') or '',
                )
                return {
                    'files_list_path': files_list_path,
                    'project_path': project_path,
                    'discipline': str(cached_entry.get('discipline') or discipline),
                    'folder_path': folder_path,
                    'count': len(cached_dwg_files),
                    'resolution_mode': 'workroom_cached_detection',
                }
            cached_file_count = len(cached_entry.get('files') or [])
            if cached_file_count:
                logging.info(
                    f"{tool_name}: Cached Workroom detection had {cached_file_count} DWG(s), but none still exist. "
                    "Falling back.")
                self._trace_cad_auto_select(
                    'auto_select_cached_stale',
                    tool_name=tool_name,
                    project_path=project_path,
                    discipline=discipline,
                    cached_file_count=cached_file_count,
                    cached_files=cached_entry.get('files') or [],
                )
            else:
                logging.info(
                    f"{tool_name}: Cached Workroom detection is empty for project_path={project_path} "
                    f"discipline={discipline}. Falling back.")
                self._trace_cad_auto_select(
                    'auto_select_cached_empty',
                    tool_name=tool_name,
                    project_path=project_path,
                    discipline=discipline,
                )
        if explicit_dwg_files:
            return _use_explicit_dwg_files()
        if 'cadFilePaths' in launch_payload:
            logging.info(
                f"{tool_name}: Launch context cadFilePaths were provided but no valid DWG files remained. "
                "Falling back to folder scan.")
            self._trace_cad_auto_select(
                'auto_select_explicit_invalid',
                tool_name=tool_name,
                project_path=project_path,
                discipline=discipline,
                cad_file_paths=launch_payload.get('cadFilePaths'),
            )
        if not project_path:
            logging.info(
                f"{tool_name}: Workroom auto-select fallback to manual file picker (missing project path).")
            self._trace_cad_auto_select(
                'auto_select_fallback_manual',
                tool_name=tool_name,
                reason='missing_project_path',
                discipline=discipline,
            )
            return None
        folder_resolution = self._resolve_workroom_discipline_folder(
            project_path, discipline)
        discipline_folder = folder_resolution.get('resolved_folder') or ''
        if not discipline_folder:
            candidates = folder_resolution.get('candidates') or []
            candidate_text = '; '.join(candidates) if candidates else 'none'
            logging.info(
                f"{tool_name}: Workroom auto-select fallback to manual file picker "
                f"(project_path={project_path}; discipline={discipline}; checked={candidate_text}).")
            self._trace_cad_auto_select(
                'auto_select_fallback_manual',
                tool_name=tool_name,
                reason='discipline_folder_not_found',
                project_path=project_path,
                discipline=discipline,
                checked=candidates,
            )
            return None
        dwg_files = self._list_base_level_dwgs(discipline_folder)
        if not dwg_files:
            logging.info(
                f"{tool_name}: Workroom auto-select fallback to manual file picker (no DWGs in {discipline_folder}).")
            self._trace_cad_auto_select(
                'auto_select_fallback_manual',
                tool_name=tool_name,
                reason='no_dwgs_in_folder',
                project_path=project_path,
                discipline=discipline,
                folder_path=discipline_folder,
            )
            return None
        files_list_path = self._write_files_list_temp(dwg_files)
        logging.info(
            f"{tool_name}: Auto-selected {len(dwg_files)} DWG(s) from {discipline_folder} "
            f"(mode={folder_resolution.get('mode')}).")
        self._trace_cad_auto_select(
            'auto_select_selected',
            tool_name=tool_name,
            selection_source=folder_resolution.get('mode', '') or 'folder_scan',
            project_path=project_path,
            discipline=folder_resolution.get('discipline') or discipline,
            folder_path=discipline_folder,
            files_list_path=files_list_path,
            count=len(dwg_files),
            file_paths=dwg_files,
        )
        return {
            'files_list_path': files_list_path,
            'project_path': project_path,
            'discipline': folder_resolution.get('discipline') or discipline,
            'folder_path': discipline_folder,
            'count': len(dwg_files),
            'resolution_mode': folder_resolution.get('mode', ''),
        }

    def _resolve_workroom_arch_folder(self, settings, launch_context):
        context = self._resolve_workroom_context(settings, launch_context)
        source = str(context.get('source') or '').strip().lower()
        if source != 'workroom':
            return ''
        project_path = context.get('project_path') or ''
        if not project_path:
            logging.info(
                "_resolve_workroom_arch_folder: Missing project path for workroom launch.")
            return ''
        arch_resolution = self._resolve_workroom_discipline_folder(
            project_path, 'Arch')
        arch_folder = arch_resolution.get('resolved_folder') or ''
        if os.path.isdir(arch_folder):
            logging.info(
                f"_resolve_workroom_arch_folder: Resolved Arch folder at {arch_folder} "
                f"(mode={arch_resolution.get('mode')}).")
            return arch_folder
        candidates = arch_resolution.get('candidates') or []
        if candidates:
            logging.info(
                "_resolve_workroom_arch_folder: Arch folder not found. "
                f"Checked: {'; '.join(candidates)}")
        return ''

    def get_workroom_cad_files(self, launch_context=None):
        try:
            settings = self.get_user_settings()
            context = self._resolve_workroom_context(settings, launch_context)
            project_path = context.get('project_path') or ''
            discipline = context.get('discipline') or self._primary_discipline_from_settings(
                settings)
            normalized_launch_context = self._normalize_launch_context(launch_context)
            self._trace_cad_auto_select(
                'get_workroom_cad_files_request',
                project_path=project_path,
                discipline=discipline,
                source=context.get('source') or '',
                discipline_source=context.get('discipline_source') or '',
                launch_context=normalized_launch_context,
            )

            if not project_path:
                self._trace_cad_auto_select(
                    'get_workroom_cad_files_missing_project_path',
                    discipline=discipline,
                    launch_context=normalized_launch_context,
                )
                return {
                    'status': 'success',
                    'projectPath': '',
                    'discipline': discipline,
                    'folderPath': '',
                    'resolutionMode': 'missing_project_path',
                    'files': [],
                    'message': 'Project folder path is missing.',
                }

            folder_resolution = self._resolve_workroom_discipline_folder(
                project_path, discipline)
            resolved_discipline = folder_resolution.get(
                'discipline') or discipline
            folder_path = folder_resolution.get('resolved_folder') or ''
            resolution_mode = folder_resolution.get('mode') or ''

            if not folder_path:
                candidates = folder_resolution.get('candidates') or []
                if candidates:
                    logging.info(
                        "get_workroom_cad_files: Discipline folder not found "
                        f"(project_path={project_path}; discipline={resolved_discipline}; "
                        f"checked={'; '.join(candidates)})."
                    )
                else:
                    logging.info(
                        "get_workroom_cad_files: Discipline folder not found "
                        f"(project_path={project_path}; discipline={resolved_discipline})."
                    )
                self._set_workroom_cad_file_cache_entry(
                    project_path,
                    resolved_discipline,
                    '',
                    [],
                    resolution_mode or 'not_found',
                )
                self._trace_cad_auto_select(
                    'get_workroom_cad_files_result',
                    project_path=project_path,
                    discipline=resolved_discipline,
                    resolution_mode=resolution_mode or 'not_found',
                    folder_path='',
                    candidate_paths=candidates,
                    count=0,
                    file_paths=[],
                )
                return {
                    'status': 'success',
                    'projectPath': project_path,
                    'discipline': resolved_discipline,
                    'folderPath': '',
                    'resolutionMode': resolution_mode or 'not_found',
                    'files': [],
                    'message': f'{resolved_discipline} folder not found.',
                }

            dwg_paths = self._list_base_level_dwgs(folder_path)
            files = [
                {
                    'name': os.path.basename(path),
                    'path': path,
                }
                for path in dwg_paths
            ]
            self._set_workroom_cad_file_cache_entry(
                project_path,
                resolved_discipline,
                folder_path,
                dwg_paths,
                resolution_mode,
            )
            self._trace_cad_auto_select(
                'get_workroom_cad_files_result',
                project_path=project_path,
                discipline=resolved_discipline,
                resolution_mode=resolution_mode,
                folder_path=folder_path,
                count=len(dwg_paths),
                file_paths=dwg_paths,
            )

            return {
                'status': 'success',
                'projectPath': project_path,
                'discipline': resolved_discipline,
                'folderPath': folder_path,
                'resolutionMode': resolution_mode,
                'files': files,
                'message': (
                    f'Found {len(files)} DWG file{"s" if len(files) != 1 else ""}.'
                ),
            }
        except Exception as e:
            logging.error(f"get_workroom_cad_files failed: {e}")
            self._trace_cad_auto_select(
                'get_workroom_cad_files_error',
                error=str(e),
                launch_context=self._normalize_launch_context(launch_context),
            )
            return {'status': 'error', 'message': str(e), 'files': []}

    def _notify_activity_status(self, payload):
        try:
            if not webview.windows:
                return
            safe_payload = payload if isinstance(payload, dict) else {}
            webview.windows[0].evaluate_js(
                f"window.updateActivityStatus({json.dumps(safe_payload)})"
            )
        except Exception as e:
            logging.debug(f"_notify_activity_status failed: {e}")

    def _notify_tool_status(self, tool_id, message, activity_id=None, **payload):
        try:
            if not webview.windows:
                return
            normalized_tool_id = str(tool_id or '').strip()
            normalized_activity_id = str(activity_id or '').strip()
            normalized_message = str(message or '').strip()
            extra_payload = dict(payload or {})
            if normalized_activity_id or extra_payload:
                activity_payload = {
                    'toolId': normalized_tool_id,
                    'activityId': normalized_activity_id,
                    'message': normalized_message,
                }
                activity_payload.update(extra_payload)
                self._notify_activity_status(activity_payload)
                return
            js_message = json.dumps(normalized_message)
            webview.windows[0].evaluate_js(
                f'window.updateToolStatus("{normalized_tool_id}", {js_message})')
        except Exception as e:
            logging.debug(
                f"_notify_tool_status failed for {tool_id}: {e}")

    def _record_test_mode_event(self, event):
        if not self.test_mode:
            return
        event_data = dict(event or {})
        event_data['timestamp'] = datetime.datetime.utcnow().isoformat() + 'Z'
        self._test_mode_records.append(event_data)

    def get_test_mode_records(self):
        return {
            'status': 'success',
            'test_mode': self.test_mode,
            'records': list(self._test_mode_records),
        }

    def _run_workroom_cad_tool_in_test_mode(self, settings, launch_context, tool_id, tool_name, activity_id=None):
        context = self._resolve_workroom_context(settings, launch_context)
        source = context.get('source') or 'none'
        project_path = context.get('project_path') or ''
        discipline = context.get('discipline') or self._primary_discipline_from_settings(settings)
        discipline_source = context.get('discipline_source') or 'unknown'

        if not self._is_workroom_auto_select_enabled(settings, launch_context):
            message = (
                "TEST MODE: Workroom auto-select is disabled for this run. "
                f"(source={source}, project_path={project_path or '<empty>'}, "
                f"discipline={discipline}, discipline_source={discipline_source})"
            )
            self._record_test_mode_event({
                'tool_id': tool_id,
                'tool_name': tool_name,
                'status': 'error',
                'reason': 'auto_select_disabled',
                'source': source,
                'project_path': project_path,
                'discipline': discipline,
                'discipline_source': discipline_source,
            })
            self._notify_tool_status(
                tool_id,
                f"ERROR: {message}",
                activity_id=activity_id,
            )
            return {'status': 'error', 'test_mode': True, 'message': message}

        auto_selection = self._resolve_workroom_auto_file_selection(
            settings, launch_context, f'{tool_name}_test_mode'
        )
        if not auto_selection:
            message = (
                "TEST MODE: Expected auto-selected DWG files but none were resolved. "
                f"(source={source}, project_path={project_path or '<empty>'}, "
                f"discipline={discipline}, discipline_source={discipline_source})"
            )
            self._record_test_mode_event({
                'tool_id': tool_id,
                'tool_name': tool_name,
                'status': 'error',
                'reason': 'no_auto_selected_files',
                'source': source,
                'project_path': project_path,
                'discipline': discipline,
                'discipline_source': discipline_source,
            })
            self._notify_tool_status(
                tool_id,
                f"ERROR: {message}",
                activity_id=activity_id,
            )
            return {'status': 'error', 'test_mode': True, 'message': message}

        count = int(auto_selection.get('count') or 0)
        folder_path = auto_selection.get('folder_path') or ''
        self._record_test_mode_event({
            'tool_id': tool_id,
            'tool_name': tool_name,
            'status': 'success',
            'source': source,
            'project_path': project_path,
            'discipline': discipline,
            'discipline_source': discipline_source,
            'folder_path': folder_path,
            'count': count,
            'resolution_mode': auto_selection.get('resolution_mode') or '',
        })
        self._notify_tool_status(
            tool_id,
            f"TEST MODE: Auto-selected {count} DWG(s) from {folder_path or 'unknown folder'}.",
            activity_id=activity_id,
        )
        # Match normal tool completion signal so UI state is identical to production runs.
        self._notify_tool_status(tool_id, "DONE", activity_id=activity_id)
        return {
            'status': 'success',
            'test_mode': True,
            'tool_id': tool_id,
            'tool_name': tool_name,
            'source': source,
            'project_path': project_path,
            'discipline': discipline,
            'discipline_source': discipline_source,
            'folder_path': folder_path,
            'count': count,
            'activityId': str(activity_id or '').strip(),
        }

    def _run_workroom_clean_xrefs_tool_in_test_mode(self, settings, launch_context, activity_id=None):
        context = self._resolve_workroom_context(settings, launch_context)
        source = context.get('source') or 'none'
        project_path = context.get('project_path') or ''
        discipline = context.get('discipline') or self._primary_discipline_from_settings(settings)
        discipline_source = context.get('discipline_source') or 'unknown'

        if not self._should_force_workroom_clean_xrefs_manual_selection(context):
            message = (
                "TEST MODE: Non-workroom launch. "
                "Clean-xrefs uses manual file selection with the default picker location "
                f"(source={source}, project_path={project_path or '<empty>'}, "
                f"discipline={discipline}, discipline_source={discipline_source})."
            )
            self._record_test_mode_event({
                'tool_id': 'toolCleanXrefs',
                'tool_name': 'run_clean_xrefs_script',
                'status': 'success',
                'reason': 'non_workroom_manual_picker',
                'source': source,
                'project_path': project_path,
                'discipline': discipline,
                'discipline_source': discipline_source,
                'manual_selection_enforced': False,
                'arch_folder': '',
            })
            self._notify_tool_status(
                'toolCleanXrefs',
                message,
                activity_id=activity_id,
            )
            self._notify_tool_status(
                'toolCleanXrefs',
                "DONE",
                activity_id=activity_id,
            )
            return {
                'status': 'success',
                'test_mode': True,
                'tool_id': 'toolCleanXrefs',
                'tool_name': 'run_clean_xrefs_script',
                'source': source,
                'project_path': project_path,
                'discipline': discipline,
                'discipline_source': discipline_source,
                'manual_selection_enforced': False,
                'arch_folder': '',
                'activityId': str(activity_id or '').strip(),
            }

        arch_folder = self._resolve_workroom_arch_folder(settings, launch_context)
        if arch_folder:
            message = (
                "TEST MODE: Workroom manual selection enforced. "
                f"Opening Arch folder picker at {arch_folder}."
            )
        else:
            message = (
                "TEST MODE: Workroom manual selection enforced. "
                "Arch folder not found; opening default file picker."
            )

        self._record_test_mode_event({
            'tool_id': 'toolCleanXrefs',
            'tool_name': 'run_clean_xrefs_script',
            'status': 'success',
            'source': source,
            'project_path': project_path,
            'discipline': discipline,
            'discipline_source': discipline_source,
            'manual_selection_enforced': True,
            'arch_folder': arch_folder,
        })
        self._notify_tool_status(
            'toolCleanXrefs',
            message,
            activity_id=activity_id,
        )
        # Match normal tool completion signal so UI state is identical to production runs.
        self._notify_tool_status(
            'toolCleanXrefs',
            "DONE",
            activity_id=activity_id,
        )
        return {
            'status': 'success',
            'test_mode': True,
            'tool_id': 'toolCleanXrefs',
            'tool_name': 'run_clean_xrefs_script',
            'source': source,
            'project_path': project_path,
            'discipline': discipline,
            'discipline_source': discipline_source,
            'manual_selection_enforced': True,
            'arch_folder': arch_folder,
            'activityId': str(activity_id or '').strip(),
        }

    def run_publish_script(self, launch_context=None, activity_id=None, params_override=None):
        """Runs the PlotDWGs.ps1 PowerShell script with progress updates.

        params_override (dict, optional): per-invocation overrides for publishDwgOptions
        (autoDetectPaperSize, shrinkPercent, stripPdfLayers). Used by the workflow runner.
        """
        script_path = os.path.join(BASE_DIR, "scripts", "PlotDWGs.ps1")
        if not os.path.exists(script_path):
            raise Exception("PlotDWGs.ps1 not found in scripts directory.")
        settings = self.get_user_settings()
        launch_payload = self._normalize_launch_context(launch_context)
        workflow_blocking = launch_payload.get('workflowBlocking') is True
        if self.test_mode:
            return self._run_workroom_cad_tool_in_test_mode(
                settings,
                launch_context,
                'toolPublishDwgs',
                'run_publish_script',
                activity_id=activity_id,
            )
        acad_path = settings.get('autocadPath', '')
        if not acad_path:
            raise Exception("No AutoCAD version selected in settings.")
        publish_options = dict(settings.get('publishDwgOptions') or {})
        if isinstance(params_override, dict):
            publish_options.update(params_override)
        auto_detect = publish_options.get('autoDetectPaperSize', True)
        shrink_percent = publish_options.get('shrinkPercent', 100)
        strip_pdf_layers = publish_options.get('stripPdfLayers', True)

        def _ps_bool(value):
            return "1" if value else "0"

        auto_selection = self._resolve_workroom_auto_file_selection(
            settings, launch_context, 'run_publish_script')
        default_directory = self._resolve_launch_context_default_directory(
            launch_context
        )
        if self._is_workroom_auto_select_enabled(settings, launch_context) and not auto_selection:
            fallback_context = self._resolve_workroom_context(
                settings, launch_context)
            logging.info(
                "run_publish_script: Workroom auto-select unavailable; opening file picker "
                f"(source={fallback_context.get('source') or 'none'}, "
                f"project_path={fallback_context.get('project_path') or '<empty>'}, "
                f"discipline={fallback_context.get('discipline')}, "
                f"discipline_source={fallback_context.get('discipline_source')}).")
            self._notify_tool_status(
                'toolPublishDwgs',
                "Workroom auto-select unavailable. Opening file picker...",
                activity_id=activity_id,
            )
            self._trace_cad_auto_select(
                'tool_manual_fallback',
                tool_id='toolPublishDwgs',
                tool_name='run_publish_script',
                reason='auto_selection_unavailable',
                source=fallback_context.get('source') or 'none',
                project_path=fallback_context.get('project_path') or '',
                discipline=fallback_context.get('discipline') or '',
            )
        command = self._build_powershell_script_command(
            script_path,
            '-AcadCore',
            acad_path,
            '-AutoDetectPaperSize',
            _ps_bool(auto_detect),
            '-ShrinkPercent',
            shrink_percent,
            '-StripPdfLayers',
            _ps_bool(strip_pdf_layers),
        )
        if auto_selection:
            command.extend([
                '-FilesListPath',
                auto_selection['files_list_path'],
            ])
            selected_count = auto_selection.get('count')
            if isinstance(selected_count, int) and selected_count > 0:
                self._notify_tool_status(
                    'toolPublishDwgs',
                    f"Using auto-selected DWGs ({selected_count}) via {auto_selection.get('resolution_mode') or 'workroom'}...",
                    activity_id=activity_id,
                )
        elif default_directory:
            command.extend([
                '-DefaultDirectory',
                default_directory,
            ])
        self._trace_cad_auto_select(
            'tool_command_launch',
            tool_id='toolPublishDwgs',
            tool_name='run_publish_script',
            command=command,
            command_type='argv',
            has_files_list_path='-FilesListPath' in command,
            files_list_path=auto_selection.get('files_list_path') if auto_selection else '',
            auto_selection=auto_selection,
        )
        run_kwargs = {'activity_id': activity_id}
        if workflow_blocking:
            run_kwargs['wait'] = True
        script_result = self._run_script_with_progress(
            command,
            'toolPublishDwgs',
            **run_kwargs,
        )
        if workflow_blocking and isinstance(script_result, dict) and script_result.get('status') == 'error':
            return {
                'status': 'error',
                'message': script_result.get('message') or 'Publish DWGs failed.',
                'activityId': str(activity_id or '').strip(),
                'scriptResult': script_result,
            }
        return {'status': 'success', 'activityId': str(activity_id or '').strip()}

    def run_manage_layers_script(self, launch_context=None, activity_id=None, params_override=None):
        """Runs the ManageLayersDWGs.ps1 PowerShell script with progress updates.

        params_override (dict, optional): per-invocation overrides for manageLayersOptions
        keys (e.g., scanAllLayers, freezePatterns, thawPatterns). Used by the workflow runner
        so a single workflow step can supply its own freeze/thaw patterns without mutating
        the user's global manageLayersOptions.
        """
        script_path = os.path.join(BASE_DIR, "scripts", "ManageLayersDWGs.ps1")
        if not os.path.exists(script_path):
            raise Exception(
                "ManageLayersDWGs.ps1 not found in scripts directory.")
        settings = self.get_user_settings()
        launch_payload = self._normalize_launch_context(launch_context)
        workflow_blocking = launch_payload.get('workflowBlocking') is True
        if self.test_mode:
            return self._run_workroom_cad_tool_in_test_mode(
                settings,
                launch_context,
                'toolManageLayers',
                'run_manage_layers_script',
                activity_id=activity_id,
            )
        acad_path = settings.get('autocadPath', '')
        if not acad_path:
            raise Exception("No AutoCAD version selected in settings.")
        manage_options = dict(settings.get('manageLayersOptions') or {})
        if isinstance(params_override, dict):
            manage_options.update(params_override)
        scan_all = manage_options.get('scanAllLayers', True)
        freeze_patterns = _normalize_layer_pattern_list(
            manage_options.get('freezePatterns'))
        thaw_patterns = _normalize_layer_pattern_list(
            manage_options.get('thawPatterns'))

        def _ps_bool(value):
            return "1" if value else "0"

        auto_selection = self._resolve_workroom_auto_file_selection(
            settings, launch_context, 'run_manage_layers_script')
        default_directory = self._resolve_launch_context_default_directory(
            launch_context
        )
        if self._is_workroom_auto_select_enabled(settings, launch_context) and not auto_selection:
            fallback_context = self._resolve_workroom_context(
                settings, launch_context)
            logging.info(
                "run_manage_layers_script: Workroom auto-select unavailable; opening file picker "
                f"(source={fallback_context.get('source') or 'none'}, "
                f"project_path={fallback_context.get('project_path') or '<empty>'}, "
                f"discipline={fallback_context.get('discipline')}, "
                f"discipline_source={fallback_context.get('discipline_source')}).")
            self._notify_tool_status(
                'toolManageLayers',
                "Workroom auto-select unavailable. Opening file picker...",
                activity_id=activity_id,
            )
            self._trace_cad_auto_select(
                'tool_manual_fallback',
                tool_id='toolManageLayers',
                tool_name='run_manage_layers_script',
                reason='auto_selection_unavailable',
                source=fallback_context.get('source') or 'none',
                project_path=fallback_context.get('project_path') or '',
                discipline=fallback_context.get('discipline') or '',
            )
        command = self._build_powershell_script_command(
            script_path,
            '-AcadCore',
            acad_path,
            '-ScanAllLayers',
            _ps_bool(scan_all),
        )
        if freeze_patterns:
            command.extend(['-FreezePatterns', ';'.join(freeze_patterns)])
        if thaw_patterns:
            command.extend(['-ThawPatterns', ';'.join(thaw_patterns)])
        if auto_selection:
            command.extend([
                '-FilesListPath',
                auto_selection['files_list_path'],
            ])
            selected_count = auto_selection.get('count')
            if isinstance(selected_count, int) and selected_count > 0:
                self._notify_tool_status(
                    'toolManageLayers',
                    f"Using auto-selected DWGs ({selected_count}) via {auto_selection.get('resolution_mode') or 'workroom'}...",
                    activity_id=activity_id,
                )
        elif default_directory:
            command.extend([
                '-DefaultDirectory',
                default_directory,
            ])
        self._trace_cad_auto_select(
            'tool_command_launch',
            tool_id='toolManageLayers',
            tool_name='run_manage_layers_script',
            command=command,
            command_type='argv',
            has_files_list_path='-FilesListPath' in command,
            files_list_path=auto_selection.get('files_list_path') if auto_selection else '',
            auto_selection=auto_selection,
        )
        run_kwargs = {'activity_id': activity_id}
        if workflow_blocking:
            run_kwargs['wait'] = True
        script_result = self._run_script_with_progress(
            command,
            'toolManageLayers',
            **run_kwargs,
        )
        if workflow_blocking and isinstance(script_result, dict) and script_result.get('status') == 'error':
            return {
                'status': 'error',
                'message': script_result.get('message') or 'Freeze/Thaw Layers failed.',
                'activityId': str(activity_id or '').strip(),
                'scriptResult': script_result,
            }
        return {'status': 'success', 'activityId': str(activity_id or '').strip()}

    def run_clean_xrefs_script(self, launch_context=None, activity_id=None, params_override=None):
        """Runs the removeXREFPaths.ps1 PowerShell script with progress updates.

        params_override (dict, optional): per-invocation overrides for cleanDwgOptions keys
        (stripXrefs, setByLayer, purge, audit, hatchColor). Used by the workflow runner.
        """
        try:
            script_path = os.path.join(BASE_DIR, "scripts", "removeXREFPaths.ps1")
            if not os.path.exists(script_path):
                raise Exception(
                    "removeXREFPaths.ps1 not found in scripts directory.")
            settings = self.get_user_settings()
            launch_payload = self._normalize_launch_context(launch_context)
            workflow_blocking = launch_payload.get('workflowBlocking') is True
            if self.test_mode:
                return self._run_workroom_clean_xrefs_tool_in_test_mode(
                    settings,
                    launch_context,
                    activity_id=activity_id,
                )
            acad_path = settings.get('autocadPath', '')
            if not acad_path:
                raise Exception("No AutoCAD version selected in settings.")
            context = self._resolve_workroom_context(settings, launch_context)
            discipline = context.get('discipline') or self._primary_discipline_from_settings(settings)
            discipline_short_map = {
                'Electrical': 'E',
                'Mechanical': 'M',
                'Plumbing': 'P',
                'General': 'G'
            }
            discipline_short = discipline_short_map.get(
                discipline, (discipline[:1] or 'E').upper())

            clean_options = dict(settings.get('cleanDwgOptions') or {})
            if isinstance(params_override, dict):
                clean_options.update(params_override)
            def _bool_setting(key, default=True):
                value = clean_options.get(key, default)
                return bool(value) if value is not None else default

            strip_xrefs = _bool_setting('stripXrefs', True)
            set_by_layer = _bool_setting('setByLayer', True)
            purge = _bool_setting('purge', True)
            audit = _bool_setting('audit', True)
            hatch_color = _bool_setting('hatchColor', True)

            def _ps_bool(value):
                return "1" if value else "0"

            source = context.get('source') or 'none'
            project_path = context.get('project_path') or ''
            discipline_source = context.get('discipline_source') or 'unknown'
            source_manifest_path = ''
            workflow_dwg_sources = self._get_launch_context_dwg_file_sources(
                launch_context)
            if workflow_dwg_sources:
                source_manifest_path = self._write_dwg_source_manifest_temp(
                    workflow_dwg_sources)
            force_manual_selection = self._should_force_workroom_clean_xrefs_manual_selection(
                context) and not source_manifest_path
            default_directory = self._resolve_launch_context_default_directory(
                launch_context
            )

            auto_selection = None
            arch_folder = self._resolve_workroom_arch_folder(
                settings, launch_context) if force_manual_selection else ''
            if force_manual_selection:
                if arch_folder:
                    self._notify_tool_status(
                        'toolCleanXrefs',
                        "Opening Arch folder for file selection...",
                        activity_id=activity_id,
                    )
                    logging.info(
                        "run_clean_xrefs_script: Workroom manual selection enforced. "
                        f"(source={source}, project_path={project_path or '<empty>'}, "
                        f"discipline={discipline}, discipline_source={discipline_source}, "
                        f"arch_folder={arch_folder})")
                else:
                    self._notify_tool_status(
                        'toolCleanXrefs',
                        "Arch folder not found. Opening file picker...",
                        activity_id=activity_id,
                    )
                    logging.info(
                        "run_clean_xrefs_script: Workroom manual selection enforced, Arch folder not found. "
                        f"(source={source}, project_path={project_path or '<empty>'}, "
                        f"discipline={discipline}, discipline_source={discipline_source})")
            else:
                auto_selection = self._resolve_workroom_auto_file_selection(
                    settings, launch_context, 'run_clean_xrefs_script')
                if self._is_workroom_auto_select_enabled(settings, launch_context) and not auto_selection:
                    fallback_context = self._resolve_workroom_context(
                        settings, launch_context)
                    logging.info(
                        "run_clean_xrefs_script: Workroom auto-select unavailable; opening file picker "
                        f"(source={fallback_context.get('source') or 'none'}, "
                        f"project_path={fallback_context.get('project_path') or '<empty>'}, "
                        f"discipline={fallback_context.get('discipline')}, "
                        f"discipline_source={fallback_context.get('discipline_source')}).")
                    self._notify_tool_status(
                        'toolCleanXrefs',
                        "Workroom auto-select unavailable. Opening file picker...",
                        activity_id=activity_id,
                    )
            if source_manifest_path:
                self._notify_tool_status(
                    'toolCleanXrefs',
                    f"Using workflow-selected DWGs ({len(workflow_dwg_sources)})...",
                    activity_id=activity_id,
                )

            command = self._build_powershell_script_command(
                script_path,
                '-AcadCore',
                acad_path,
                '-DisciplineShort',
                discipline_short,
                '-StripXrefs',
                _ps_bool(strip_xrefs),
                '-SetByLayer',
                _ps_bool(set_by_layer),
                '-Purge',
                _ps_bool(purge),
                '-Audit',
                _ps_bool(audit),
                '-HatchColor',
                _ps_bool(hatch_color),
            )
            if source_manifest_path:
                command.extend(['-SourceManifestPath', source_manifest_path])
            if auto_selection:
                command.extend(['-FilesListPath', auto_selection["files_list_path"]])
            if arch_folder:
                command.extend(['-DefaultDirectory', arch_folder])
            elif default_directory:
                command.extend(['-DefaultDirectory', default_directory])
            run_kwargs = {'activity_id': activity_id}
            if workflow_blocking:
                run_kwargs['wait'] = True
            script_result = self._run_script_with_progress(
                command,
                'toolCleanXrefs',
                **run_kwargs,
            )
            if workflow_blocking and isinstance(script_result, dict) and script_result.get('status') == 'error':
                return {
                    'status': 'error',
                    'message': script_result.get('message') or 'Prepare CAD for XREF failed.',
                    'activityId': str(activity_id or '').strip(),
                    'scriptResult': script_result,
                }
            return {'status': 'success', 'activityId': str(activity_id or '').strip()}
        except Exception as e:
            logging.error(f"run_clean_xrefs_script failed: {e}")
            self._notify_tool_status(
                'toolCleanXrefs',
                f"ERROR: {str(e)}",
                activity_id=activity_id,
            )
            return {'status': 'error', 'message': str(e)}

    def get_workflow_tools(self):
        """Returns the JSON-safe registry of tools available as workflow steps."""
        return {'status': 'success', 'tools': get_workflow_tool_descriptors()}

    def _build_workflow_step_launch_context(self, base_launch_context, inputs, required):
        """Construct a per-step launch_context that injects pre-flight user inputs.

        Does NOT mutate `base_launch_context`. Maps the schema-declared input keys onto
        the launch_context fields the existing tool methods already consume:
          - 'projectFolder' -> ctx['project_path'], ctx['projectPath'], ctx['rootProjectPath']
          - 'dwgFiles'      -> ctx['cadFilePaths']
        For projectFolder, when no `source` is set on the base context, marks the
        derived context as 'manual' so downstream code treats it as an explicit path
        instead of an unresolved Workroom reference.
        """
        ctx = dict(base_launch_context or {})
        if not isinstance(inputs, dict):
            return ctx
        required_keys = {
            str(spec.get('key') or '').strip()
            for spec in (required or [])
            if isinstance(spec, dict)
        }

        if 'projectFolder' in required_keys and inputs.get('projectFolder'):
            folder = str(inputs['projectFolder']).strip()
            if folder:
                ctx['project_path'] = folder
                ctx['projectPath'] = folder
                ctx['rootProjectPath'] = folder
                if not str(ctx.get('source') or '').strip():
                    ctx['source'] = 'manual'

        if 'dwgFiles' in required_keys:
            raw_files = inputs.get('dwgFiles')
            if isinstance(raw_files, (list, tuple)):
                cleaned = [str(p).strip() for p in raw_files if str(p or '').strip()]
                if cleaned:
                    ctx['cadFilePaths'] = cleaned
                    raw_sources = inputs.get('dwgFileSources')
                    sources = self._normalize_dwg_file_sources(
                        raw_sources, require_exists=False)
                    if not sources:
                        sources = self._normalize_dwg_file_sources(
                            [{'kind': 'file', 'path': path} for path in cleaned],
                            require_exists=False,
                        )
                    if sources:
                        ctx['dwgFileSources'] = sources
                    ctx['workflowPreselectedDwgFiles'] = True
        return ctx

    def resolve_workflow_defaults(self, workflow_id, launch_context=None):
        """Read-only: suggest pre-flight input values for each step that needs them.

        Returns {'status': 'success', 'defaults': {<stepIndex>: {<key>: value}}}
        Uses the same Workroom-resolution helpers individual tools already use so the
        user sees the same auto-resolved values they'd get if they ran each tool alone.
        Does not fire any PowerShell scripts or write activity events.
        """
        try:
            workflow_id_str = str(workflow_id or '').strip()
            if not workflow_id_str:
                return {'status': 'error', 'message': 'Workflow id is required.'}
            settings = self.get_user_settings()
            workflows = settings.get('workflows') or []
            workflow = next(
                (w for w in workflows
                 if isinstance(w, dict) and str(w.get('id') or '').strip() == workflow_id_str),
                None,
            )
            if not workflow:
                return {'status': 'error', 'message': f'Workflow not found: {workflow_id_str}'}

            steps = workflow.get('steps') or []
            defaults = {}
            workflow_cad_defaults = _normalize_workflow_cad_defaults(
                settings.get('workflowCadDefaults'))

            workroom_ctx = None
            workroom_files_cache = None  # cached resolution; reused for every dwgFiles step
            electrical_files_cache = None
            clean_xref_sources_cache = None

            def _get_workroom_ctx():
                nonlocal workroom_ctx
                if workroom_ctx is None:
                    try:
                        workroom_ctx = self._resolve_workroom_context(settings, launch_context) or {}
                    except Exception:
                        workroom_ctx = {}
                return workroom_ctx

            def _get_workroom_files():
                nonlocal workroom_files_cache
                if workroom_files_cache is not None:
                    return workroom_files_cache
                workroom_files_cache = []
                try:
                    selection = self._resolve_workroom_auto_file_selection(
                        settings, launch_context, 'run_publish_script')
                except Exception:
                    selection = None
                if selection and selection.get('files_list_path'):
                    try:
                        with open(selection['files_list_path'], 'r', encoding='utf-8') as fh:
                            workroom_files_cache = [
                                line.strip() for line in fh if line.strip()
                            ]
                    except Exception:
                            workroom_files_cache = []
                return workroom_files_cache

            def _get_electrical_files():
                nonlocal electrical_files_cache
                if electrical_files_cache is None:
                    try:
                        electrical_files_cache = self._resolve_workflow_electrical_dwg_files(
                            settings, launch_context)
                    except Exception:
                        electrical_files_cache = []
                return electrical_files_cache

            def _get_clean_xref_sources():
                nonlocal clean_xref_sources_cache
                if clean_xref_sources_cache is None:
                    try:
                        clean_xref_sources_cache = self._resolve_workflow_clean_xref_sources(
                            settings, launch_context)
                    except Exception:
                        logging.exception("Failed to resolve workflow clean-xrefs defaults.")
                        clean_xref_sources_cache = []
                return clean_xref_sources_cache

            for index, step in enumerate(steps):
                if not isinstance(step, dict):
                    continue
                tool_id = str(step.get('toolId') or '').strip()
                entry = WORKFLOW_TOOL_REGISTRY.get(tool_id)
                if not entry:
                    continue
                required = entry.get('requiredInputs') or []
                if not required:
                    continue

                step_defaults = {}
                for spec in required:
                    key = str(spec.get('key') or '').strip()
                    if key == 'projectFolder':
                        ctx = _get_workroom_ctx()
                        candidate = (
                            str(ctx.get('project_path') or '').strip()
                            or str((launch_context or {}).get('project_path') or '').strip()
                        )
                        if candidate:
                            step_defaults['projectFolder'] = candidate
                    elif key == 'dwgFiles':
                        if (
                            tool_id == 'manageLayers'
                            and workflow_cad_defaults.get('manageLayersDwgSource') == 'electricalTopLevel'
                        ):
                            files = _get_electrical_files()
                        elif (
                            tool_id == 'cleanXrefs'
                            and workflow_cad_defaults.get('cleanXrefsDwgSource') == 'electricalXrefsToNewestArch'
                        ):
                            sources = _get_clean_xref_sources()
                            files = [self._source_display_path(source) for source in sources]
                            if files:
                                step_defaults['dwgFileSources'] = list(sources)
                        else:
                            files = _get_workroom_files()
                        if files:
                            step_defaults['dwgFiles'] = list(files)
                if step_defaults:
                    defaults[str(index)] = step_defaults

            return {'status': 'success', 'defaults': defaults}
        except Exception as exc:
            logging.exception("resolve_workflow_defaults failed.")
            return {'status': 'error', 'message': str(exc)}

    def run_workflow(self, workflow_id, launch_context=None, activity_id=None, step_inputs=None):
        """Sequentially run the steps of a saved workflow. Stops on first failure.

        Each step dispatches to a callable in WORKFLOW_TOOL_REGISTRY. The parent
        activity_id receives high-level progress messages; each step is given its own
        sub-activity id so per-step progress shows as a distinct row in the activity tray.

        step_inputs (dict, optional): JSON-safe per-step pre-flight inputs collected
        by the workflow UI before dispatch. Keys are stringified step indices (0-based);
        values are dicts like {'projectFolder': '...', 'dwgFiles': [...]}. When provided,
        each step's launch_context is augmented with the matching input fields so the
        tool methods don't need to prompt the user again.
        """
        try:
            workflow_id_str = str(workflow_id or '').strip()
            if not workflow_id_str:
                return {'status': 'error', 'message': 'Workflow id is required.'}
            settings = self.get_user_settings()
            workflows = settings.get('workflows') or []
            workflow = next(
                (w for w in workflows
                 if isinstance(w, dict) and str(w.get('id') or '').strip() == workflow_id_str),
                None,
            )
            if not workflow:
                return {'status': 'error', 'message': f'Workflow not found: {workflow_id_str}'}

            name = str(workflow.get('name') or 'Unnamed workflow').strip() or 'Unnamed workflow'
            steps = workflow.get('steps') or []
            if not isinstance(steps, list) or not steps:
                return {'status': 'error', 'message': 'Workflow has no steps.'}

            parent_activity_id = str(activity_id or '').strip()
            total = len(steps)

            def _post_parent(message, *, status=None, progress=None):
                if not parent_activity_id:
                    return
                payload = {
                    'toolId': 'toolWorkflow',
                    'activityId': parent_activity_id,
                    'message': message,
                    'label': f'Workflow: {name}',
                    'workflowTitle': name,
                }
                if progress is not None:
                    payload['progress'] = progress
                if status:
                    payload['status'] = status
                self._notify_activity_status(payload)

            _post_parent(f'Starting workflow: {name}', progress=3)

            for index, step in enumerate(steps, start=1):
                if not isinstance(step, dict):
                    msg = f'Step {index} is malformed.'
                    _post_parent(f'ERROR: {msg}', status='error')
                    return {'status': 'error', 'failedStep': index, 'message': msg}

                tool_id = str(step.get('toolId') or '').strip()
                entry = WORKFLOW_TOOL_REGISTRY.get(tool_id)
                if not entry:
                    msg = f'Unknown tool "{tool_id}" in step {index}.'
                    _post_parent(f'ERROR: {msg}', status='error')
                    return {'status': 'error', 'failedStep': index, 'message': msg}

                step_params = step.get('params')
                if not isinstance(step_params, dict):
                    step_params = {}
                display = entry.get('displayName') or tool_id
                progress = int(5 + (index - 1) * 90 / max(total, 1))
                _post_parent(f'Step {index}/{total}: {display}…', progress=progress)
                step_activity_id = (
                    f'{parent_activity_id}-step{index}' if parent_activity_id else ''
                )
                inputs_for_step = {}
                if isinstance(step_inputs, dict):
                    candidate = step_inputs.get(str(index - 1))
                    if isinstance(candidate, dict):
                        inputs_for_step = candidate
                per_step_context = self._build_workflow_step_launch_context(
                    base_launch_context=launch_context,
                    inputs=inputs_for_step,
                    required=entry.get('requiredInputs') or [],
                )
                per_step_context['workflowBlocking'] = True
                try:
                    result = entry['invoke'](self, per_step_context, step_activity_id, step_params)
                except Exception as exc:
                    logging.exception(
                        f"run_workflow: step {index} ({tool_id}) raised an exception.")
                    msg = str(exc) or 'Unknown error.'
                    _post_parent(
                        f'ERROR: workflow halted at step {index} ({display}): {msg}',
                        status='error',
                    )
                    return {
                        'status': 'error',
                        'failedStep': index,
                        'tool': tool_id,
                        'message': msg,
                    }

                if isinstance(result, dict) and result.get('status') == 'error':
                    msg = str(result.get('message') or 'Step reported an error.').strip()
                    _post_parent(
                        f'ERROR: workflow halted at step {index} ({display}): {msg}',
                        status='error',
                    )
                    return {
                        'status': 'error',
                        'failedStep': index,
                        'tool': tool_id,
                        'message': msg,
                        'stepResult': result,
                    }

            _post_parent(
                f'Workflow "{name}" complete ({total}/{total} steps).',
                status='success',
                progress=100,
            )
            return {
                'status': 'success',
                'workflowId': workflow_id_str,
                'stepCount': total,
                'activityId': parent_activity_id,
            }
        except Exception as exc:
            logging.exception("run_workflow failed unexpectedly.")
            if activity_id:
                try:
                    workflow_title = str(locals().get('name') or '').strip()
                    payload = {
                        'toolId': 'toolWorkflow',
                        'activityId': str(activity_id).strip(),
                        'message': f'ERROR: {exc}',
                        'status': 'error',
                    }
                    if workflow_title:
                        payload['label'] = f'Workflow: {workflow_title}'
                        payload['workflowTitle'] = workflow_title
                    self._notify_activity_status(payload)
                except Exception:
                    pass
            return {'status': 'error', 'message': str(exc)}

    def select_files(self, options):
        """Shows a file dialog and returns selected paths."""
        try:
            options = options or {}
            window = webview.windows[0]
            default_directory = options.get('default_directory')
            if default_directory in (None, ''):
                default_directory = options.get('default_dir')
            if default_directory in (None, ''):
                default_directory = options.get('defaultDirectory')
            directory = self._resolve_dialog_directory(default_directory)
            file_paths = window.create_file_dialog(
                webview.FileDialog.OPEN,  # DEPRECATION FIX
                allow_multiple=options.get('allow_multiple', False),
                file_types=tuple(options.get('file_types', ())),
                directory=directory,
            )
            if not file_paths:
                return {'status': 'cancelled', 'paths': []}
            return {'status': 'success', 'paths': file_paths}
        except TypeError:
            try:
                options = options or {}
                window = webview.windows[0]
                file_paths = window.create_file_dialog(
                    webview.FileDialog.OPEN,  # DEPRECATION FIX
                    allow_multiple=options.get('allow_multiple', False),
                    file_types=tuple(options.get('file_types', ()))
                )
                if not file_paths:
                    return {'status': 'cancelled', 'paths': []}
                return {'status': 'success', 'paths': file_paths}
            except Exception as e:
                logging.error(f"Error in file dialog fallback: {e}")
                return {'status': 'error', 'message': str(e)}
        except Exception as e:
            logging.error(f"Error in file dialog: {e}")
            return {'status': 'error', 'message': str(e)}







    def import_and_process_csv(self):
        """Handles CSV file import dialog."""
        try:
            window = webview.windows[0]
            file_paths = window.create_file_dialog(
                webview.FileDialog.OPEN,  # DEPRECATION FIX
                allow_multiple=False,
                file_types=('CSV Files (*.csv)',)
            )
            if not file_paths:
                return {'status': 'cancelled', 'data': []}
            with open(file_paths[0], 'r', encoding='utf-8-sig') as f:
                csv_content = f.read()
            projects = self._process_csv_rows(csv_content, file_paths[0])
            return {'status': 'success', 'data': projects}
        except Exception as e:
            logging.error(f"Error during CSV import: {e}")
            return {'status': 'error', 'message': str(e), 'data': []}

    def _process_csv_rows(self, csv_content, csv_path):
        """Helper to process CSV data."""
        new_projects = []
        fallback_paths = {}
        csv_dir = os.path.dirname(csv_path)
        fallback_json_path = os.path.join(csv_dir, 'tasks.json')
        if os.path.exists(fallback_json_path):
            try:
                with open(fallback_json_path, 'r', encoding='utf-8') as f:
                    fallback_data = json.load(f)
                    for p in fallback_data:
                        if p.get('id') and p.get('path'):
                            fallback_paths[str(p['id']).strip()] = p['path']
            except Exception as e:
                logging.warning(f"Could not load fallback tasks.json: {e}")
        f = io.StringIO(csv_content)
        reader = csv.reader(f)
        try:
            header = next(reader)
            if 'project name' not in ''.join(header).lower():
                f.seek(0)
        except StopIteration:
            return []
        for row in reader:
            if not any(field.strip() for field in row):
                continue
            row.extend([''] * 12)
            proj_id, name, nick, notes, due, tasks_str, path, *_ = row
            if not any([proj_id, name, path]):
                continue
            final_path = (path or '').strip()
            if final_path and not os.path.exists(os.path.normpath(final_path)):
                fallback_path = fallback_paths.get((proj_id or '').strip())
                if fallback_path:
                    final_path = fallback_path
            tasks_list = [
                {'text': t.strip(), 'done': False, 'links': []}
                for t in tasks_str.replace('\r', '\n').replace(';', '\n').replace('\u2022', '\n').split('\n')
                if t.strip()
            ]
            project = {
                'id': (proj_id or '').strip(),
                'name': (name or '').strip(),
                'nick': (nick or '').strip(),
                'notes': (notes or '').strip(),
                'due': (due or '').strip(),
                'path': final_path,
                'tasks': tasks_list,
                'refs': [],
                'statuses': []
            }
            new_projects.append(project)
        return new_projects


# --- Main Application Setup ---
if __name__ == '__main__':
    api = Api()
    window = webview.create_window(
        'ACIES Desktop Application',
        'index.html',
        js_api=api,
        width=1400,
        height=900,
        resizable=True,
        min_size=(1024, 768)
    )
    # --- FIX: Removed the line that caused the circular reference ---
    webview.start()
    # Cleanup subprocesses on exit
    api.stop_circuit_breaker_server()
